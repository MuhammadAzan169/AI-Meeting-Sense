# REPSession.py - RAG-Powered Critical Evaluation Engine for Real Estate Sales
import json
import os
import re
import traceback
import pickle
from datetime import datetime
from typing import Optional, List, Dict, Any, Tuple
from pathlib import Path
from enum import Enum

import requests
import configparser
import numpy as np

# RAG and Embeddings
try:
    from sentence_transformers import SentenceTransformer
    EMBEDDINGS_AVAILABLE = True
except ImportError:
    EMBEDDINGS_AVAILABLE = False
    print("⚠️ sentence-transformers not available. Install: pip install sentence-transformers")

try:
    import faiss
    FAISS_AVAILABLE = True
except ImportError:
    FAISS_AVAILABLE = False
    print("⚠️ FAISS not available. Install: pip install faiss-cpu")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not available. Install: pip install python-docx")

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.pdfgen import canvas
from dotenv import load_dotenv
from pydantic import BaseModel, Field, field_validator, ConfigDict


# ===================== PYDANTIC MODELS V2 =====================

class LLMProvider(str, Enum):
    """Supported LLM providers"""
    OPENAI = "openai"
    OPENROUTER = "openrouter"
    CUSTOM = "custom"

class LLMConfig(BaseModel):
    """LLM Configuration Model"""
    model_config = ConfigDict(protected_namespaces=())
    
    provider: str = Field(..., description="LLM provider name")
    api_key: str = Field(..., description="API key for the LLM provider")
    model: str = Field(default="gpt-4o-mini", description="Model name to use")
    base_url: Optional[str] = Field(default="", description="Base URL for API calls")
    temperature: float = Field(default=0.2, ge=0.0, le=2.0, description="Temperature for generation")
    max_tokens: int = Field(default=5000, ge=1, le=32000, description="Maximum tokens to generate")
    timeout: int = Field(default=10000, ge=1, description="Timeout in seconds for API calls")
    
    @field_validator('base_url', mode='before')
    @classmethod
    def set_base_url_based_on_provider(cls, v, info):
        """Set default base URL based on provider if not provided"""
        if v:
            return v
        
        provider = info.data.get('provider', 'openai').lower().strip()
        if provider == "openai":
            return "https://api.openai.com/v1"
        elif provider == "openrouter":
            return "https://openrouter.ai/api/v1"
        else:
            return v

class TranscriptSegment(BaseModel):
    """Transcript segment model"""
    speaker_id: Optional[str] = None
    speaker_name: Optional[str] = None
    transcript: str
    segment_id: Optional[int] = None
    start: Optional[float] = 0.0
    end: Optional[float] = 0.0
    duration: Optional[float] = 0.0

class TranscriptData(BaseModel):
    """Complete transcript data model"""
    transcripts: List[TranscriptSegment]
    agenda: str
    metadata: Dict[str, Any] = {}
    summary_info: Dict[str, Any]

class AgendaAnalysis(BaseModel):
    """Agenda analysis results"""
    agenda_words: List[str]
    relevant_words: List[str]
    relevance_percentage: float

class MeetingAnalysis(BaseModel):
    """Meeting analysis results"""
    content: str
    meeting_id: str
    transcript_data: TranscriptData
    agenda_analysis: AgendaAnalysis
    enhanced_analysis: Dict[str, Any] = {}


# ===================== VALIDATION MODELS =====================

class ActionItem(BaseModel):
    """Action item with strict date validation"""
    task: str
    responsible: str
    deadline: str = Field(default="No due date was mentioned")
    priority: str = Field(default="Medium")  # High, Medium, Low
    status: str = Field(default="Pending")
    
    @field_validator('deadline', mode='before')
    @classmethod
    def validate_deadline(cls, v):
        """Ensure deadline is properly handled"""
        if not v or v.strip() == "":
            return "No due date was mentioned"
        return v

class SpeakerInsight(BaseModel):
    """Speaker insight analysis"""
    communication_style: str
    strengths: List[str]
    weaknesses: List[str]

class ClientInsight(BaseModel):
    """Client insight analysis"""
    client_understanding: str
    buying_intent: str  # Explicit, Implicit, or Unclear

class MeetingAnalysisOutput(BaseModel):
    """Validated meeting analysis output"""
    meeting_overview: Dict[str, str]
    executive_summary: str
    participant_summary: str
    key_discussion_points: List[str]
    action_items: List[ActionItem]
    speaker_insights: Dict[str, Any]  # Contains employee and client insights
    objections: List[str]
    opportunities: List[str]
    decisions: List[str]
    commitments: List[str]
    deal_status: str
    intelligence_status: str
    closing_probability: str


class KnowledgeChunk(BaseModel):
    """Knowledge base chunk with embedding"""
    chunk_id: int
    text: str
    embedding: Optional[List[float]] = None
    metadata: Dict[str, Any] = {}


class CoachingEvaluation(BaseModel):
    """RAG-powered coaching evaluation output"""
    executive_evaluation: str
    critical_analysis: str
    strengths: List[str]
    improvement_areas: List[str]
    coaching_recommendations: List[str]
    knowledge_references: List[str]
    performance_scores: Dict[str, Any] = {}


# ===================== MAIN CLASS =====================

class RealEstateSalesMeetingSummarizer:
    def __init__(self, transcript_file: Optional[str] = None, knowledge_base_path: Optional[str] = None):
        self.config_file = "meeting_config.ini"
        self.transcript_file = transcript_file
        self.summaries_folder = "RealEstateMeetingRecords"
        self.company_name = "Real Estate Sales Coaching & Evaluation"
        
        # RAG components
        self.knowledge_base_path = knowledge_base_path or "knowledge_base.docx"
        self.faiss_index_path = "knowledge_base.faiss"
        self.chunks_path = "knowledge_base_chunks.pkl"
        self.embedding_model_name = "all-MiniLM-L6-v2"
        self.embedding_model = None
        self.faiss_index = None
        self.knowledge_chunks = []
        
        # Load configuration
        self._load_environment_variables()
        self._setup_llm_config()
        self._setup_paths_configuration()
        self._setup_folders()
        self._setup_professional_styles()
        
        # Initialize RAG system
        self._initialize_rag_system()

    def _load_environment_variables(self):
        """Load configuration from .env file"""
        # Use relative path - look for .env in the same directory as this script
        script_dir = os.path.dirname(os.path.abspath(__file__))
        env_path = os.path.join(script_dir, '.env')
        
        print(f"🔧 DEBUG - Looking for .env file at: {env_path}")
        
        if os.path.exists(env_path):
            print(f"✅ Found .env file at: {env_path}")
            
            # Read and display the .env file content (without API key)
            try:
                with open(env_path, 'r') as f:
                    lines = f.readlines()
                    print(f"🔧 DEBUG - .env file content (excluding sensitive data):")
                    for line in lines:
                        if 'API_KEY' not in line.upper() and 'KEY' not in line.upper():  # Don't print API keys
                            print(f"   {line.strip()}")
                        else:
                            print(f"   [REDACTED: Contains API key]")
            except Exception as e:
                print(f"⚠️ Could not read .env file: {e}")
            
            load_dotenv(env_path, override=True)
            print(f"✅ Loaded environment variables from: {env_path}")
        else:
            print(f"⚠️ .env file not found at: {env_path}")
            print(f"   Looking for .env in current working directory...")
            
            # Try current working directory
            cwd_env_path = os.path.join(os.getcwd(), '.env')
            if os.path.exists(cwd_env_path):
                print(f"✅ Found .env file in current directory: {cwd_env_path}")
                load_dotenv(cwd_env_path, override=True)
            else:
                print(f"⚠️ .env file not found in current directory either")
                print(f"   Loading from system environment variables...")
                load_dotenv(override=True)
        
        # Debug: Check what environment variables are set
        print(f"🔧 DEBUG - Current environment variables:")
        print(f"   LLM_PROVIDER: {os.getenv('LLM_PROVIDER', 'NOT SET')}")
        print(f"   LLM_MODEL: {os.getenv('LLM_MODEL', 'NOT SET')}")
        print(f"   LLM_BASE_URL: {os.getenv('LLM_BASE_URL', 'NOT SET')}")
        print(f"   LLM_API_KEY present: {'Yes' if os.getenv('LLM_API_KEY') else 'No'}")
        print(f"   LLM_TEMPERATURE: {os.getenv('LLM_TEMPERATURE', 'NOT SET')}")
        print(f"   LLM_MAX_TOKENS: {os.getenv('LLM_MAX_TOKENS', 'NOT SET')}")
        print(f"   LLM_TIMEOUT: {os.getenv('LLM_TIMEOUT', 'NOT SET')}")

    def _setup_llm_config(self):
        """Setup LLM configuration using Pydantic model"""
        try:
            # Read from environment variables with exact names
            provider = os.getenv('LLM_PROVIDER', 'openai')
            api_key = os.getenv('LLM_API_KEY', '')
            model = os.getenv('LLM_MODEL', 'gpt-4o-mini')
            base_url = os.getenv('LLM_BASE_URL', '')
            temperature = float(os.getenv('LLM_TEMPERATURE', '0.2'))
            max_tokens = int(os.getenv('LLM_MAX_TOKENS', '5000'))
            timeout = int(os.getenv('LLM_TIMEOUT', '10000'))
            
            # DEBUG: Print what we're reading from environment
            print(f"🔧 DEBUG - Environment variables loaded:")
            print(f"   LLM_PROVIDER: {provider}")
            print(f"   LLM_MODEL: {model}")
            print(f"   LLM_BASE_URL: {base_url}")
            print(f"   LLM_API_KEY present: {'Yes' if api_key else 'No'}")
            print(f"   LLM_TEMPERATURE: {temperature}")
            print(f"   LLM_MAX_TOKENS: {max_tokens}")
            print(f"   LLM_TIMEOUT: {timeout}")
            
            # Check if we have the default value that shouldn't be there
            if 'llama' in model.lower() or 'meta-llama' in model:
                print(f"⚠️ WARNING: Found LLaMA model in environment: {model}")
                print(f"   Overriding with gpt-4o-mini as specified in .env")
                model = 'gpt-4o-mini'
            
            # Create LLMConfig using Pydantic model
            self.llm_config = LLMConfig(
                provider=provider,
                api_key=api_key,
                model=model,
                base_url=base_url,
                temperature=temperature,
                max_tokens=max_tokens,
                timeout=timeout
            )
            
            print(f"✅ LLM Configuration Final:")
            print(f"   Provider: {self.llm_config.provider}")
            print(f"   Model: {self.llm_config.model}")
            print(f"   Base URL: {self.llm_config.base_url}")
            
            # Check if API key is configured
            if not self.llm_config.api_key or self.llm_config.api_key == '':
                print("❌ API key not configured. Please set LLM_API_KEY in .env file")
            
        except Exception as e:
            print(f"❌ Error setting up LLM configuration: {e}")
            traceback.print_exc()
            # Set default configuration
            self.llm_config = LLMConfig(
                provider="openai",
                api_key="",
                model="gpt-4o-mini",
                base_url="https://api.openai.com/v1"
            )

    def _setup_paths_configuration(self):
        """Initialize paths configuration"""
        self.config = configparser.ConfigParser()
        if os.path.exists(self.config_file):
            self.config.read(self.config_file)
        else:
            self.config['Paths'] = {
                'transcript_path': self.transcript_file or 'conversation.json',
                'summaries_folder': self.summaries_folder
            }
            self._save_config()

    def _save_config(self):
        """Save configuration to file"""
        with open(self.config_file, 'w') as f:
            self.config.write(f)

    def _setup_folders(self):
        """Create necessary folders"""
        Path(self.summaries_folder).mkdir(exist_ok=True)

    def _setup_professional_styles(self):
        """Setup professional balanced styles"""
        self.styles = getSampleStyleSheet()

        # Professional but clean styles
        self.pro_styles = {}

        # Main Title Style - DARK BLUE
        self.pro_styles['MainTitle'] = ParagraphStyle(
            name='MainTitle',
            parent=self.styles['Title'],
            fontSize=24,
            textColor=colors.HexColor("#0D47A1"),
            spaceAfter=6,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
            leading=28
        )

        # Meeting Title Style - LIGHT BLUE
        self.pro_styles['MeetingTitle'] = ParagraphStyle(
            name='MeetingTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor("#1976D2"),
            spaceAfter=10,
            alignment=TA_CENTER,
            fontName='Helvetica',
            leading=18
        )

        # Section Header - DARK BLUE
        self.pro_styles['SectionHeader'] = ParagraphStyle(
            name='SectionHeader',
            parent=self.styles['Heading1'],
            fontSize=14,
            textColor=colors.HexColor("#0D47A1"),
            spaceAfter=6,
            spaceBefore=12,
            fontName='Helvetica-Bold',
            alignment=TA_LEFT
        )

        # Regular Text
        self.pro_styles['Regular'] = ParagraphStyle(
            name='Regular',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.black,
            spaceAfter=4,
            fontName='Helvetica',
            leading=12,
            wordWrap='LTR'
        )

        # Table Header
        self.pro_styles['TableHeader'] = ParagraphStyle(
            name='TableHeader',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.white,
            fontName='Helvetica-Bold',
            alignment=TA_CENTER,
            leading=12,
            wordWrap='LTR'
        )

        # Table Text
        self.pro_styles['TableText'] = ParagraphStyle(
            name='TableText',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.black,
            leading=12,
            fontName='Helvetica',
            wordWrap='LTR'
        )

        # Key Point Style
        self.pro_styles['KeyPoint'] = ParagraphStyle(
            name='KeyPoint',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor("#333333"),
            leftIndent=15,
            spaceAfter=3,
            fontName='Helvetica',
            leading=12,
            wordWrap='LTR'
        )

        # Confidential Style
        self.pro_styles['Confidential'] = ParagraphStyle(
            name='Confidential',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor("#D32F2F"),
            borderPadding=6,
            borderColor=colors.HexColor("#D32F2F"),
            borderWidth=1,
            alignment=TA_CENTER
        )

        # First Page Title Style - DARK BLUE
        self.pro_styles['FirstPageTitle'] = ParagraphStyle(
            name='FirstPageTitle',
            parent=self.styles['Title'],
            fontSize=28,
            textColor=colors.HexColor("#0D47A1"),
            spaceAfter=15,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
            leading=32
        )

        # First Page Subtitle
        self.pro_styles['FirstPageSubtitle'] = ParagraphStyle(
            name='FirstPageSubtitle',
            parent=self.styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor("#666666"),
            spaceAfter=20,
            alignment=TA_CENTER,
            fontName='Helvetica',
            leading=14
        )

        # High Risk Highlight - RED (NO BACKGROUND COLOR)
        self.pro_styles['HighRiskHighlight'] = ParagraphStyle(
            name='HighRiskHighlight',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.black,
            fontName='Helvetica',
            leftIndent=15,
            spaceAfter=3
        )

        # Medium Risk Highlight - YELLOW/ORANGE (NO BACKGROUND COLOR)
        self.pro_styles['MediumRiskHighlight'] = ParagraphStyle(
            name='MediumRiskHighlight',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.black,
            fontName='Helvetica',
            leftIndent=15,
            spaceAfter=3
        )

        # Low Risk Highlight - GREEN (NO BACKGROUND COLOR)
        self.pro_styles['LowRiskHighlight'] = ParagraphStyle(
            name='LowRiskHighlight',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.black,
            fontName='Helvetica',
            leftIndent=15,
            spaceAfter=3
        )

        # Success Highlight - GREEN (NO BACKGROUND COLOR)
        self.pro_styles['SuccessHighlight'] = ParagraphStyle(
            name='SuccessHighlight',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.black,
            borderPadding=4,
            borderColor=colors.HexColor("#2E7D32"),
            borderWidth=0.5,
            fontName='Helvetica-Bold'
        )

    def _get_api_headers(self) -> Dict[str, str]:
        """Get appropriate headers for the LLM provider"""
        provider = self.llm_config.provider.lower().strip()
        base_headers = {
            "Authorization": f"Bearer {self.llm_config.api_key}",
            "Content-Type": "application/json"
        }
        
        if provider == "openrouter":
            base_headers.update({
                "HTTP-Referer": "https://github.com/legal-audit-summarizer",
                "X-Title": "Real Estate Sales Meeting Record"
            })
        
        return base_headers

    # ============================================================
    # RAG SYSTEM METHODS
    # ============================================================

    def _initialize_rag_system(self):
        """Initialize RAG system: load or create FAISS index"""
        print("\n🧠 Initializing RAG System...")
        
        if not EMBEDDINGS_AVAILABLE or not FAISS_AVAILABLE:
            print("⚠️ RAG components not available. System will work without knowledge base.")
            return
        
        # Load embedding model
        try:
            print(f"   Loading embedding model: {self.embedding_model_name}...")
            self.embedding_model = SentenceTransformer(self.embedding_model_name)
            print("   ✅ Embedding model loaded")
        except Exception as e:
            print(f"   ❌ Failed to load embedding model: {e}")
            return
        
        # Check if index exists
        if os.path.exists(self.faiss_index_path) and os.path.exists(self.chunks_path):
            print("   📂 Loading existing FAISS index...")
            try:
                self._load_faiss_index()
                print(f"   ✅ Loaded {len(self.knowledge_chunks)} knowledge chunks")
                return
            except Exception as e:
                print(f"   ⚠️ Failed to load index: {e}. Will rebuild...")
        
        # Build new index
        if not os.path.exists(self.knowledge_base_path):
            print(f"   ⚠️ Knowledge base not found: {self.knowledge_base_path}")
            print(f"   System will work without knowledge base grounding.")
            return
        
        print(f"   📚 Building FAISS index from {self.knowledge_base_path}...")
        self._build_faiss_index()

    def _load_knowledge_base(self) -> str:
        """Load knowledge base from .docx file"""
        if not DOCX_AVAILABLE:
            print("   ❌ python-docx not available")
            return ""
        
        try:
            doc = Document(self.knowledge_base_path)
            full_text = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            full_text.append(text)
            
            content = "\n".join(full_text)
            print(f"   ✅ Loaded {len(content)} characters from knowledge base")
            return content
            
        except Exception as e:
            print(f"   ❌ Error loading knowledge base: {e}")
            traceback.print_exc()
            return ""

    def _semantic_chunk(self, text: str, chunk_size: int = 500, overlap: int = 100) -> List[str]:
        """Chunk text semantically with overlap"""
        # Split by paragraphs first
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        chunks = []
        current_chunk = []
        current_length = 0
        
        for para in paragraphs:
            para_length = len(para)
            
            # If single paragraph exceeds chunk_size, split it
            if para_length > chunk_size:
                sentences = re.split(r'(?<=[.!?])\s+', para)
                for sent in sentences:
                    if current_length + len(sent) > chunk_size and current_chunk:
                        chunks.append(' '.join(current_chunk))
                        # Keep overlap
                        overlap_text = ' '.join(current_chunk[-2:]) if len(current_chunk) > 1 else current_chunk[-1] if current_chunk else ""
                        current_chunk = [overlap_text, sent] if overlap_text else [sent]
                        current_length = len(overlap_text) + len(sent)
                    else:
                        current_chunk.append(sent)
                        current_length += len(sent)
            else:
                if current_length + para_length > chunk_size and current_chunk:
                    chunks.append(' '.join(current_chunk))
                    # Keep overlap
                    overlap_text = ' '.join(current_chunk[-2:]) if len(current_chunk) > 1 else current_chunk[-1] if current_chunk else ""
                    current_chunk = [overlap_text, para] if overlap_text else [para]
                    current_length = len(overlap_text) + para_length
                else:
                    current_chunk.append(para)
                    current_length += para_length
        
        # Add remaining chunk
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks

    def _build_faiss_index(self):
        """Build FAISS index from knowledge base"""
        # Load knowledge base
        kb_content = self._load_knowledge_base()
        if not kb_content:
            print("   ❌ No knowledge base content to index")
            return
        
        # Chunk the content
        print("   🔪 Chunking knowledge base...")
        text_chunks = self._semantic_chunk(kb_content, chunk_size=500, overlap=100)
        print(f"   ✅ Created {len(text_chunks)} chunks")
        
        # Generate embeddings
        print("   🧠 Generating embeddings...")
        embeddings = self.embedding_model.encode(text_chunks, show_progress_bar=True, convert_to_numpy=True)
        
        # Create FAISS index
        dimension = embeddings.shape[1]
        index = faiss.IndexFlatL2(dimension)
        index.add(embeddings.astype('float32'))
        
        self.faiss_index = index
        self.knowledge_chunks = [
            KnowledgeChunk(chunk_id=i, text=chunk, embedding=emb.tolist())
            for i, (chunk, emb) in enumerate(zip(text_chunks, embeddings))
        ]
        
        # Save index and chunks
        print("   💾 Saving FAISS index...")
        faiss.write_index(index, self.faiss_index_path)
        
        with open(self.chunks_path, 'wb') as f:
            pickle.dump(self.knowledge_chunks, f)
        
        print(f"   ✅ FAISS index built with {len(self.knowledge_chunks)} chunks")

    def _load_faiss_index(self):
        """Load existing FAISS index"""
        self.faiss_index = faiss.read_index(self.faiss_index_path)
        
        with open(self.chunks_path, 'rb') as f:
            self.knowledge_chunks = pickle.load(f)

    def retrieve_relevant_knowledge(self, query: str, top_k: int = 5) -> List[KnowledgeChunk]:
        """Retrieve most relevant knowledge chunks for a query"""
        if not self.faiss_index or not self.knowledge_chunks:
            print("   ⚠️ FAISS index not available")
            return []
        
        # Generate query embedding
        query_embedding = self.embedding_model.encode([query], convert_to_numpy=True)
        
        # Search FAISS index
        distances, indices = self.faiss_index.search(query_embedding.astype('float32'), top_k)
        
        # Retrieve chunks
        retrieved_chunks = []
        for idx in indices[0]:
            if idx < len(self.knowledge_chunks):
                retrieved_chunks.append(self.knowledge_chunks[idx])
        
        return retrieved_chunks

    # ============================================================
    # END RAG SYSTEM METHODS
    # ============================================================

    def create_professional_header_footer(self, canvas, doc):
        """Create professional header/footer"""
        canvas.saveState()

        # Header - DARK BLUE
        header_height = 0.7 * inch
        canvas.setFillColor(colors.HexColor("#0D47A1"))
        canvas.rect(0, doc.pagesize[1] - header_height, doc.pagesize[0], header_height, fill=1, stroke=0)

        # Company name
        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica-Bold", 12)
        canvas.drawString(0.7 * inch, doc.pagesize[1] - 0.45 * inch, self.company_name)
        canvas.setFont("Helvetica", 8)
        canvas.drawString(0.7 * inch, doc.pagesize[1] - 0.65 * inch, "Sales Meeting Summary")

        # Document title
        canvas.setFont("Helvetica-Bold", 11)
        canvas.drawRightString(doc.pagesize[0] - 0.7 * inch, doc.pagesize[1] - 0.45 * inch, "SALES MEETING RECORD")
        canvas.setFont("Helvetica", 8)
        canvas.drawRightString(doc.pagesize[0] - 0.7 * inch, doc.pagesize[1] - 0.65 * inch, "Professional Summary")

        # Footer - LIGHT BLUE
        footer_height = 0.5 * inch
        canvas.setFillColor(colors.HexColor("#1976D2"))
        canvas.rect(0, 0, doc.pagesize[0], footer_height, fill=1, stroke=0)

        # Page number
        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica", 9)
        canvas.drawCentredString(doc.pagesize[0] / 2, 0.25 * inch, f"Page {doc.page}")

        canvas.setFont("Helvetica", 8)
        timestamp = datetime.now().strftime("%Y-%m%d %H:%M")
        footer_left = f"© {datetime.now().year} | Generated: {timestamp}"
        canvas.drawString(0.7 * inch, 0.15 * inch, footer_left)
        canvas.drawRightString(doc.pagesize[0] - 0.7 * inch, 0.15 * inch, "CONFIDENTIAL")

        canvas.restoreState()

    def create_first_page_header_footer(self, canvas, doc):
        """Create special header/footer for first page"""
        canvas.saveState()
        
        # No header on first page
        
        # Simple footer only - LIGHT BLUE
        canvas.setFillColor(colors.HexColor("#1976D2"))
        canvas.rect(0, 0, doc.pagesize[0], 0.4 * inch, fill=1, stroke=0)
        
        # Page number
        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica", 9)
        canvas.drawCentredString(doc.pagesize[0] / 2, 0.15 * inch, f"Page {doc.page}")
        
        canvas.setFont("Helvetica", 8)
        canvas.drawString(0.7 * inch, 0.1 * inch, f"© {datetime.now().year}")
        canvas.drawRightString(doc.pagesize[0] - 0.7 * inch, 0.1 * inch, "CONFIDENTIAL")
        
        canvas.restoreState()

    def load_transcript_data(self) -> Optional[TranscriptData]:
        """Load transcript data - Updated to handle your JSON structure"""
        try:
            if not os.path.exists(self.transcript_file):
                print(f"❌ Transcript file not found: {self.transcript_file}")
                return None

            with open(self.transcript_file, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # Your JSON is an array of conversation segments
            # Extract agenda from metadata or set default
            agenda = "Real Estate Sales Consultation"
            
            # Create speaker mapping
            speaker_mapping = {}
            employee_names = []
            client_names = []
            all_speakers = set()
            
            # First pass: collect all unique speakers
            for segment in data:
                speaker = segment.get('speaker', '')
                if speaker:
                    all_speakers.add(speaker)
            
            # Identify employee: any speaker that doesn't contain "Speaker" in the name
            # Agent name will never be Speaker_n
            for speaker in all_speakers:
                if "Speaker" not in speaker and speaker:
                    # This is the agent/employee
                    employee_names.append(speaker)
                    speaker_mapping[speaker] = f"{speaker} (Sales Consultant)"
                else:
                    # These are clients (Speaker_1, Speaker_2, etc.)
                    speaker_mapping[speaker] = "Client"
                    client_names.append("Client")

            print(f"📝 Found {len(data)} conversation segments")
            print(f"📋 Agenda: {agenda}")
            print(f"👥 Company Employee(s): {', '.join(employee_names)}")
            print(f"👥 Client(s): {', '.join(set(client_names))}")

            # Process transcripts
            enhanced_transcripts = []
            total_words = 0
            speakers = set()
            total_seconds = 0

            for i, segment in enumerate(data):
                try:
                    speaker = segment.get('speaker', f"Speaker_{i+1}")
                    transcript_text = segment.get('transcript', '')
                    
                    # Get speaker name from mapping
                    if speaker in speaker_mapping:
                        speaker_name = speaker_mapping[speaker]
                        speakers.add(speaker_name)
                    else:
                        # Default to "Client" for unknown speakers
                        speaker_name = "Client"
                        speakers.add(speaker_name)

                    if transcript_text and len(transcript_text.strip()) > 3:
                        enhanced_transcripts.append(TranscriptSegment(
                            speaker_id=speaker,
                            speaker_name=speaker_name,
                            transcript=transcript_text.strip(),
                            segment_id=i + 1,
                            start=segment.get('start', 0),
                            end=segment.get('end', 0),
                            duration=segment.get('duration', 0)
                        ))
                        total_words += len(transcript_text.split())
                        total_seconds += segment.get('duration', 0)

                except Exception as e:
                    print(f"⚠ Error processing segment {i}: {e}")
                    continue

            print(f"✅ Processed {len(enhanced_transcripts)} valid segments")

            # Calculate total duration
            minutes = int(total_seconds // 60)
            seconds = int(total_seconds % 60)
            if minutes > 0:
                duration_formatted = f"{minutes} minutes {seconds} seconds"
            else:
                duration_formatted = f"{seconds} seconds"

            return TranscriptData(
                transcripts=enhanced_transcripts,
                agenda=agenda,
                metadata={},
                summary_info={
                    'total_segments': len(enhanced_transcripts),
                    'total_words': total_words,
                    'total_duration': total_seconds,
                    'duration_formatted': duration_formatted,
                    'unique_speakers': len(speakers),
                    'speakers': list(speakers),
                    'employee_names': employee_names,
                    'client_names': ["Client"]  # Use generic "Client" name
                }
            )

        except Exception as e:
            print(f"❌ Error loading transcript: {e}")
            traceback.print_exc()
            return None

    def _extract_agenda_analysis(self, transcripts: List[TranscriptSegment], agenda: str) -> AgendaAnalysis:
        """Extract agenda analysis from transcript"""
        all_text = " ".join([t.transcript for t in transcripts]).lower()
        agenda_words = set(re.findall(r'\b\w+\b', agenda.lower()))

        relevant_words = [word for word in agenda_words if word in all_text and len(word) > 3]
        relevance_percentage = (len(relevant_words) / len(agenda_words) * 100) if agenda_words else 0

        return AgendaAnalysis(
            agenda_words=list(agenda_words),
            relevant_words=relevant_words,
            relevance_percentage=relevance_percentage
        )

    def _call_llm(self, prompt: str, max_tokens: Optional[int] = None, temperature: Optional[float] = None) -> Optional[str]:
        """Make API call to LLM using Pydantic configuration"""
        
        if not self.llm_config.api_key or self.llm_config.api_key == '':
            print("❌ API key not configured. Please set LLM_API_KEY in .env file")
            return None
        
        headers = self._get_api_headers()
        
        # Use base_url from LLMConfig Pydantic model
        endpoint = f"{self.llm_config.base_url}/chat/completions"
        
        # Prepare payload based on provider
        provider = self.llm_config.provider.lower().strip()
        
        # DEBUG: Show what we're sending
        print(f"🔧 DEBUG - API Request Details:")
        print(f"   Endpoint: {endpoint}")
        print(f"   Provider: {provider}")
        print(f"   Model: {self.llm_config.model}")
        
        if provider in ["openai", "openrouter"]:
            payload = {
                "model": self.llm_config.model,
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                "max_tokens": max_tokens or self.llm_config.max_tokens,
                "temperature": temperature or self.llm_config.temperature
            }
        else:
            # Custom provider - use standard OpenAI format
            payload = {
                "model": self.llm_config.model,
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                "max_tokens": max_tokens or self.llm_config.max_tokens,
                "temperature": temperature or self.llm_config.temperature
            }
        
        try:
            print(f"🔗 Calling {self.llm_config.provider} API...")
            print(f"   Model: {self.llm_config.model}")
            
            response = requests.post(
                endpoint,
                headers=headers,
                json=payload,
                timeout=self.llm_config.timeout
            )
            
            # DEBUG: Show response status
            print(f"🔧 DEBUG - API Response Status: {response.status_code}")
            
            response.raise_for_status()
            result = response.json()
            
            # DEBUG: Show response structure
            print(f"🔧 DEBUG - API Response Keys: {list(result.keys())}")
            if 'choices' in result:
                print(f"🔧 DEBUG - Number of choices: {len(result['choices'])}")
            
            # Extract response based on provider
            if provider in ["openai", "openrouter"]:
                return result['choices'][0]['message']['content']
            else:
                # Try standard OpenAI format
                if 'choices' in result and len(result['choices']) > 0:
                    if 'message' in result['choices'][0]:
                        return result['choices'][0]['message']['content']
                    elif 'text' in result['choices'][0]:
                        return result['choices'][0]['text']
                
                return str(result)
                
        except requests.exceptions.RequestException as e:
            print(f"❌ API request failed: {e}")
            if hasattr(e, 'response') and e.response is not None:
                print(f"   Status Code: {e.response.status_code}")
                print(f"   Response: {e.response.text[:500]}")
            return None
        except KeyError as e:
            print(f"❌ Unexpected API response format: {e}")
            print(f"🔧 DEBUG - Full response: {result}")
            return None
        except Exception as e:
            print(f"❌ Error calling LLM: {e}")
            traceback.print_exc()
            return None

    def _format_transcript_for_analysis(self, transcripts: List[TranscriptSegment]) -> str:
        """Format transcript for analysis"""
        formatted = ""
        for segment in transcripts:
            speaker = segment.speaker_name or segment.speaker_id or f"Speaker_{segment.segment_id}"
            formatted += f"{speaker}: {segment.transcript}\n\n"
        return formatted

    def _parse_and_validate_sections(self, analysis: str) -> Dict[str, Any]:
        """Parse and validate coaching evaluation sections"""
        # First parse into sections
        sections = {}
        current_section = None
        current_content = []
        
        lines = analysis.split('\n')
        
        for line in lines:
            line_stripped = line.strip()
            
            # Main section detection (##)
            if line_stripped.startswith('## '):
                if current_section:
                    sections[current_section] = '\n'.join(current_content).strip()
                
                current_section = line_stripped[3:].strip()
                current_content = []
            
            # Content accumulation - keep original line for formatting
            elif current_section is not None:
                current_content.append(line)
        
        # Save last section
        if current_section:
            sections[current_section] = '\n'.join(current_content).strip()
        
        # For coaching evaluation, we keep content mostly as-is for proper formatting
        # The PDF generation methods will handle specific parsing
        return sections

    def _wrap_table_text(self, text: str, max_length: int = 80, column_type: str = "task") -> str:
        """Wrap text for table cells to prevent overflow"""
        # Use smaller max_length for deadline and responsible columns
        if column_type == "deadline":
            max_length = 40
        elif column_type == "responsible":
            max_length = 25
        elif column_type == "task":
            max_length = 60
        
        if len(text) <= max_length:
            return text
        
        words = text.split()
        lines = []
        current_line = []
        
        for word in words:
            if len(' '.join(current_line + [word])) <= max_length:
                current_line.append(word)
            else:
                if current_line:
                    lines.append(' '.join(current_line))
                current_line = [word]
        
        if current_line:
            lines.append(' '.join(current_line))
        
        return '\n'.join(lines)

    # ============================================================
    # UNIFIED COACHING EVALUATION METHODS
    # ============================================================

    def generate_unified_coaching_evaluation(self) -> Optional[str]:
        """Generate unified coaching evaluation combining RAG and coaching insights"""
        print("🎯 Starting Unified Coaching Evaluation...")
        print("="*60)

        # Ask for JSON file path if not already set
        if not self.transcript_file or not os.path.exists(self.transcript_file):
            self.transcript_file = input("📂 Enter the path to your JSON transcript file: ").strip().strip('"')
        
        if not os.path.exists(self.transcript_file):
            print(f"❌ Transcript file not found: {self.transcript_file}")
            return None

        # Load transcript data
        transcript_data = self.load_transcript_data()
        if not transcript_data:
            return None
        
        # Generate meeting ID
        meeting_id = f"COACH_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        # Generate unified analysis
        print("🧠 Generating unified coaching evaluation...")
        analysis_content = self._generate_unified_llm_analysis(transcript_data, meeting_id)
        
        if not analysis_content:
            print("❌ Failed to generate unified analysis")
            return None
        
        # Create MeetingAnalysis object
        analysis = MeetingAnalysis(
            content=analysis_content,
            meeting_id=meeting_id,
            transcript_data=transcript_data,
            agenda_analysis=self._extract_agenda_analysis(transcript_data.transcripts, transcript_data.agenda),
            enhanced_analysis=self._parse_unified_sections(analysis_content)
        )
        
        # Generate PDF
        pdf_path = self._create_unified_pdf(analysis)
        
        if pdf_path:
            print(f"\n✅ UNIFIED COACHING EVALUATION COMPLETED!")
            print(f"   📄 Report: {pdf_path}")
            print(f"   🔐 Meeting ID: {analysis.meeting_id}")
            print(f"   👤 Agent: {', '.join(analysis.transcript_data.summary_info['employee_names'])}")
            print(f"   ⏱️  Duration: {analysis.transcript_data.summary_info['duration_formatted']}")
            print(f"   🧠 RAG System: {'Active' if self.faiss_index else 'Not Available'}")
            print("\n" + "="*60)
            print("🎯 REPORT STRUCTURE:")
            print("   ✓ Executive Evaluation Summary")
            print("   ✓ Performance Evaluation (Based on Best Practices)")
            print("   ✓ Coaching Snapshot (60-Second Read)")
            print("   ✓ Coaching Playbook")
            print("   ✓ Next 7-Day Focus")
            print("   ✓ Deal & Intelligence Status")
            return pdf_path
        
        return None

    def _generate_unified_llm_analysis(self, transcript_data: TranscriptData, meeting_id: str) -> Optional[str]:
        """Generate unified coaching evaluation using RAG and coaching insights"""
        
        formatted_transcript = self._format_transcript_for_analysis(transcript_data.transcripts)
        
        # Get employee name from transcript data
        employee_name = "Employee"
        if transcript_data.summary_info['employee_names']:
            employee_name = transcript_data.summary_info['employee_names'][0]
        
        # RAG RETRIEVAL
        print("🔍 Retrieving relevant knowledge from knowledge base...")
        retrieved_context = ""
        
        if self.faiss_index and self.knowledge_chunks:
            query = f"Sales coaching evaluation for: {transcript_data.agenda}. {formatted_transcript[:500]}"
            relevant_chunks = self.retrieve_relevant_knowledge(query, top_k=8)
            
            if relevant_chunks:
                print(f"   ✅ Retrieved {len(relevant_chunks)} relevant knowledge chunks")
                retrieved_context = "\n\n".join([
                    f"[Knowledge Reference {i+1}]\n{chunk.text}"
                    for i, chunk in enumerate(relevant_chunks)
                ])
        else:
            print("   ⚠️ RAG system not initialized")
        
        # UNIFIED PROMPT
        prompt = f"""You are a SENIOR SALES COACH with 20+ years of experience. You combine deep sales expertise with practical coaching insights.

Your task: Create a UNIFIED COACHING EVALUATION that blends strategic analysis with actionable coaching.

==================================================
KNOWLEDGE CONTEXT (Best Practices Reference)
==================================================

{retrieved_context if retrieved_context else "[Using general sales best practices]"}

==================================================
MEETING INFORMATION
==================================================

Meeting ID: {meeting_id}
Agenda: {transcript_data.agenda}
Participants: {', '.join(transcript_data.summary_info['speakers'])}
Agent: {employee_name}
Duration: {transcript_data.summary_info['duration_formatted']}

==================================================
TRANSCRIPT
==================================================

{formatted_transcript}

==================================================
YOUR TASK: UNIFIED COACHING EVALUATION
==================================================

Generate your evaluation in THIS EXACT STRUCTURE:

## 1. EXECUTIVE EVALUATION SUMMARY
[2-3 paragraphs: What happened? What went well? What needs immediate attention? Be direct and coaching-oriented.]

## 2. PARTICIPANT SUMMARY
[Clear participant roles. Format: "Company Representative: {employee_name} (Sales Consultant), Client: [Client Name if mentioned, otherwise 'Client']"]

## 3. KEY DISCUSSION POINTS
[List 5-7 key points discussed. Use bullet points for clarity]
- [Point 1]
- [Point 2]

## 4. ACTION ITEMS
[Create a table with these columns]
| Task Description | Responsible Person | Deadline | Priority | Status |
|---|---|---|---|---|
| [Clear, specific task] | [Name] | [Only if date mentioned, otherwise "No due date was mentioned"] | [High/Medium/Low] | Pending |

## 5. PERFORMANCE EVALUATION (Based on Best Practices)
[Deep analysis comparing agent performance to best practices. Use bullet points for clarity when possible]
**Communication Quality:** [Analyze with specific examples]
**Client Engagement:** [Analyze with specific examples]
**Objection Handling:** [Analyze with specific examples]
**Deal Progression:** [Analyze with specific examples]

## 6. STRENGTHS (Evidence-Based)
[List 4-6 specific strengths with evidence. Use bullet points]
- **[Strength 1]**: [Evidence from transcript] → [Why this matters]
- **[Strength 2]**: [Evidence from transcript] → [Impact on deal]

## 7. COACHING SNAPSHOT (60-SECOND READ)

**What went well:**
- [2-3 bullet points of successes]

**What cost leverage:**
- [2-3 bullet points of missed opportunities]

**What to fix in the next meeting:**
- [2-3 specific, actionable fixes]

**Overall Deal Health:**
[🟢 Positive / 🟡 Caution / 🔴 Negative] — [One-sentence assessment]

## 8. OBJECTIONS AND OPPORTUNITIES

**Objections:**
- [Objection 1]
- [Objection 2]

**Opportunities:**
- [Opportunity 1]
- [Opportunity 2]

## 9. COACHING PLAYBOOK — WHAT TO DO NEXT TIME

**IMPORTANT: For commission-related objections, use this exact format:**

**When client says:** "I’ve seen other agents offering lower commission rates."
> **Say this instead:** "I understand that lower commissions can be appealing, but my higher commission allows us to invest more in marketing your property, which often results in a faster sale and a higher final price."
> **Why this works:** This response directly addresses the client's concern while emphasizing the value of the services provided.

**Rule for Objection Handling:**
> • Always acknowledge the client's concern first, then provide a well-rounded response that highlights your unique value proposition.
> • If you go longer, stop and ask: "Does that make sense, or should I simplify it?"
> **Why this works:** This ensures the client feels heard and allows for clarification if needed.

**Before closing, always ask:**
> "Are you ready to move forward with the listing agreement today?"
> **Why this works:** This direct question prompts the client to make a decision and reinforces the urgency of the situation.

**For other objections from the transcript:**

**When client says:** "[Exact client objection from transcript]"
> **Say this instead:** "[Exact script agent should use]"
> **Why this works:** [Brief explanation]

**Rule for [Specific Situation]:**
> [Specific process] → [Time limit] max
> **If you go longer, stop and ask:** "Does that make sense, or should I simplify it?"
> **Why this works:** [Brief explanation]

**Before closing, always ask:**
> "[Exact closing question]"
> **Why this works:** [Brief explanation]

## 10. NEXT 7-DAY COACHING FOCUS
For the next week, focus on ONLY these three things:

1. [Focus 1: Specific behavior to practice]
2. [Focus 2: Specific phrase to use]
3. [Focus 3: Specific closing technique]

## 11. DECISIONS AND COMMITMENTS

**Decisions Made:**
- [Decision 1]
- [Decision 2]

**Commitments Given:**
- [Commitment 1]
- [Commitment 2]

## 12. PERFORMANCE SCORES

**Communication Effectiveness:** [X/10]
**Justification:** [Based on transcript evidence]

**Client Engagement:** [X/10]
**Justification:** [Based on transcript evidence]

**Objection Handling:** [X/10]
**Justification:** [Based on transcript evidence]

**Deal Progression:** [X/10]
**Justification:** [Based on transcript evidence]

**Overall Performance:** [X/10]
**Summary:** [One-sentence bottom line]

## 13. DEAL AND INTELLIGENCE STATUS

**Deal Status Evaluation:**
[Current stage, momentum, next milestone]

**Intelligence Quality:**
[Quality of information gathered]

**Risk Assessment:**
[Potential risks and concerns]

**Recommended Next Steps:**
1. [Action 1]
2. [Action 2]
3. [Action 3]

**Deal Probability:** [Percentage]
**Justification:** [2-3 sentences explaining probability]

==================================================
CRITICAL RULES
==================================================

1. **USE BULLET POINTS** when listing items - not paragraphs
2. **BE SPECIFIC** - reference exact transcript moments
3. **GROUND IN BEST PRACTICES** - reference knowledge context when relevant
4. **ACTIONABLE COACHING** - every recommendation should be implementable
5. **NO ACADEMIC LANGUAGE** - sound like a human coach
6. **LLM DECIDES FORMAT** - use points for lists, paragraphs for analysis
7. **REMOVE "KNOWLEDGE BASE REFERENCES USED"** - don't include this section
8. **COMBINE INSIGHTS** - blend strategic analysis with tactical coaching

Tone: Professional, direct, coaching-focused, actionable

Now provide the UNIFIED COACHING EVALUATION:
"""

        print("🧠 Generating unified coaching evaluation...")
        response = self._call_llm(prompt, max_tokens=12000, temperature=0.3)
        
        return response

    def _parse_unified_sections(self, analysis: str) -> Dict[str, Any]:
        """Parse unified evaluation sections"""
        sections = {}
        current_section = None
        current_content = []
        
        lines = analysis.split('\n')
        
        for line in lines:
            line_stripped = line.strip()
            
            # Section detection (##)
            if line_stripped.startswith('## '):
                if current_section:
                    sections[current_section] = '\n'.join(current_content).strip()
                
                current_section = line_stripped[3:].strip()
                current_content = []
            elif current_section is not None:
                current_content.append(line)
        
        if current_section:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections

    def _create_unified_pdf(self, analysis: MeetingAnalysis) -> Optional[str]:
        """Create unified coaching evaluation PDF"""
        output_path = os.path.join(self.summaries_folder, f"{analysis.meeting_id}_unified_coaching.pdf")

        print(f"📊 Creating unified coaching PDF for {analysis.meeting_id}...")

        try:
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=0.7 * inch,
                leftMargin=0.7 * inch,
                topMargin=1.0 * inch,
                bottomMargin=0.8 * inch
            )

            story = []
            
            # Add first page
            self._create_unified_first_page(story, analysis)
            
            # Parse and add sections
            sections = self._parse_unified_sections(analysis.content)
            self._add_unified_sections(story, sections, analysis)
            
            def on_first_page(canvas, doc):
                self.create_unified_first_page_header_footer(canvas, doc)
            
            def on_later_pages(canvas, doc):
                self.create_unified_header_footer(canvas, doc)
            
            doc.build(story, onFirstPage=on_first_page, onLaterPages=on_later_pages)
            
            print(f"✅ Unified PDF created: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"❌ Error creating PDF: {e}")
            traceback.print_exc()
            return None

    def _create_unified_first_page(self, story, analysis: MeetingAnalysis):
        """Create first page for unified report"""
        story.append(Spacer(1, 0.8 * inch))
        
        # Main title
        story.append(Paragraph("COACHING EVALUATION REPORT", self.pro_styles['FirstPageTitle']))
        story.append(Spacer(1, 0.1 * inch))
        story.append(Paragraph("Unified RAG-Powered Performance Analysis", ParagraphStyle(
            name='UnifiedSubtitle',
            fontName='Helvetica',
            fontSize=14,
            textColor=colors.HexColor("#666666"),
            alignment=TA_CENTER,
            spaceAfter=20
        )))
        
        # Meeting details
        agenda = analysis.transcript_data.agenda
        duration = analysis.transcript_data.summary_info['duration_formatted']
        employee = ', '.join(analysis.transcript_data.summary_info['employee_names']) if analysis.transcript_data.summary_info['employee_names'] else "Not specified"
        
        details_data = [
            ['Evaluation Details', ''],
            ['Meeting ID:', analysis.meeting_id],
            ['Date:', datetime.now().strftime('%B %d, %Y')],
            ['Agenda:', agenda],
            ['Duration:', duration],
            ['Agent Evaluated:', employee],
            ['Report Type:', 'Unified Coaching Evaluation'],
        ]
        
        details_table = Table(details_data, colWidths=[2.0 * inch, 4.5 * inch])
        details_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (1, 0), colors.HexColor("#0D47A1")),
            ('TEXTCOLOR', (0, 0), (1, 0), colors.white),
            ('FONTNAME', (0, 0), (1, 0), "Helvetica-Bold"),
            ('FONTSIZE', (0, 0), (1, 0), 12),
            ('ALIGN', (0, 0), (1, 0), 'CENTER'),
            ('PADDING', (0, 0), (1, 0), 10),
            ('BACKGROUND', (0, 1), (0, -1), colors.HexColor("#E3F2FD")),
            ('FONTNAME', (0, 1), (0, -1), "Helvetica-Bold"),
            ('FONTSIZE', (0, 1), (0, -1), 10),
            ('ALIGN', (0, 1), (0, -1), 'RIGHT'),
            ('PADDING', (0, 1), (0, -1), 8),
            ('FONTNAME', (1, 1), (1, -1), "Helvetica"),
            ('FONTSIZE', (1, 1), (1, -1), 10),
            ('ALIGN', (1, 1), (1, -1), 'LEFT'),
            ('PADDING', (1, 1), (1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#BBDEFB")),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        
        story.append(details_table)
        story.append(Spacer(1, 0.4 * inch))
        
        # Confidential notice
        story.append(Table(
            [[Paragraph("CONFIDENTIAL — FOR COACHING PURPOSES ONLY", ParagraphStyle(
                name='ConfidentialUnified',
                fontName='Helvetica-Bold',
                fontSize=11,
                textColor=colors.HexColor("#D32F2F"),
                alignment=TA_CENTER
            ))]],
            colWidths=[6.5 * inch],
            style=TableStyle([
                ('BACKGROUND', (0, 0), (0, 0), colors.white),
                ('BOX', (0, 0), (0, 0), 1, colors.HexColor("#D32F2F")),
                ('PADDING', (0, 0), (0, 0), 10),
            ])
        ))
        
        story.append(PageBreak())

    def create_unified_first_page_header_footer(self, canvas, doc):
        """Create unified first page header/footer"""
        canvas.saveState()
        
        # Simple footer
        canvas.setFillColor(colors.HexColor("#1976D2"))
        canvas.rect(0, 0, doc.pagesize[0], 0.4 * inch, fill=1, stroke=0)
        
        # Page number
        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica", 9)
        canvas.drawCentredString(doc.pagesize[0] / 2, 0.15 * inch, f"Page {doc.page}")
        
        canvas.setFont("Helvetica", 8)
        canvas.drawString(0.7 * inch, 0.1 * inch, f"© {datetime.now().year}")
        canvas.drawRightString(doc.pagesize[0] - 0.7 * inch, 0.1 * inch, "UNIFIED EVALUATION")
        
        canvas.restoreState()

    def create_unified_header_footer(self, canvas, doc):
        """Create unified header/footer"""
        canvas.saveState()

        # Header - DARK BLUE
        header_height = 0.7 * inch
        canvas.setFillColor(colors.HexColor("#0D47A1"))
        canvas.rect(0, doc.pagesize[1] - header_height, doc.pagesize[0], header_height, fill=1, stroke=0)

        # Company name
        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica-Bold", 12)
        canvas.drawString(0.7 * inch, doc.pagesize[1] - 0.45 * inch, "REAL ESTATE COACHING")
        canvas.setFont("Helvetica", 8)
        canvas.drawString(0.7 * inch, doc.pagesize[1] - 0.65 * inch, "Unified Performance Evaluation")

        # Document title
        canvas.setFont("Helvetica-Bold", 11)
        canvas.drawRightString(doc.pagesize[0] - 0.7 * inch, doc.pagesize[1] - 0.45 * inch, "COACHING REPORT")
        canvas.setFont("Helvetica", 8)
        canvas.drawRightString(doc.pagesize[0] - 0.7 * inch, doc.pagesize[1] - 0.65 * inch, "RAG-Powered Insights")

        # Footer - LIGHT BLUE
        footer_height = 0.5 * inch
        canvas.setFillColor(colors.HexColor("#1976D2"))
        canvas.rect(0, 0, doc.pagesize[0], footer_height, fill=1, stroke=0)

        # Page number
        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica", 9)
        canvas.drawCentredString(doc.pagesize[0] / 2, 0.25 * inch, f"Page {doc.page}")

        canvas.setFont("Helvetica", 8)
        timestamp = datetime.now().strftime("%Y-%m%d %H:%M")
        footer_left = f"© {datetime.now().year} | Generated: {timestamp}"
        canvas.drawString(0.7 * inch, 0.15 * inch, footer_left)
        canvas.drawRightString(doc.pagesize[0] - 0.7 * inch, 0.15 * inch, "CONFIDENTIAL")

        canvas.restoreState()

    def _add_unified_sections(self, story, sections: Dict[str, Any], analysis: MeetingAnalysis):
        """Add unified sections to PDF"""
        
        # Define EXACT section order as specified
        section_order = [
            "1. EXECUTIVE EVALUATION SUMMARY",
            "2. PARTICIPANT SUMMARY",
            "3. KEY DISCUSSION POINTS",
            "4. ACTION ITEMS",
            "5. PERFORMANCE EVALUATION (Based on Best Practices)",
            "6. STRENGTHS (Evidence-Based)",
            "7. COACHING SNAPSHOT (60-SECOND READ)",
            "8. OBJECTIONS AND OPPORTUNITIES",
            "9. COACHING PLAYBOOK — WHAT TO DO NEXT TIME",
            "10. NEXT 7-DAY COACHING FOCUS",
            "11. DECISIONS AND COMMITMENTS",
            "12. PERFORMANCE SCORES",
            "13. DEAL AND INTELLIGENCE STATUS"
        ]
        
        for section_title in section_order:
            if section_title in sections:
                content = sections[section_title]
                
                # Add section header
                story.append(Paragraph(section_title, self.pro_styles['SectionHeader']))
                story.append(Spacer(1, 8))
                
                # Process each section type
                if "EXECUTIVE EVALUATION SUMMARY" in section_title:
                    self._create_text_section(story, content)
                elif "PARTICIPANT SUMMARY" in section_title:
                    self._create_text_section(story, content)
                elif "KEY DISCUSSION POINTS" in section_title:
                    self._create_bullet_list(story, content)
                elif "ACTION ITEMS" in section_title:
                    parsed_items = self._parse_action_items_table(content)
                    self._create_professional_action_items(story, parsed_items)
                elif "PERFORMANCE EVALUATION" in section_title:
                    self._create_performance_evaluation(story, content)
                elif "STRENGTHS" in section_title:
                    self._create_strengths_section(story, content)
                elif "COACHING SNAPSHOT" in section_title:
                    self._create_coaching_snapshot(story, content)
                elif "OBJECTIONS AND OPPORTUNITIES" in section_title:
                    self._create_objections_opportunities(story, content)
                elif "COACHING PLAYBOOK" in section_title:
                    self._create_coaching_playbook(story, content)
                elif "NEXT 7-DAY" in section_title:
                    self._create_7day_focus(story, content)
                elif "DECISIONS AND COMMITMENTS" in section_title:
                    self._create_decisions_commitments(story, content)
                elif "PERFORMANCE SCORES" in section_title:
                    self._create_performance_scores(story, content)
                elif "DEAL AND INTELLIGENCE STATUS" in section_title:
                    self._create_deal_intelligence_status(story, content)
                
                story.append(Spacer(1, 12))

    def _create_performance_evaluation(self, story, content: str):
        """Create performance evaluation section"""
        if not content or content.strip() == "":
            # If content is empty, add a placeholder
            story.append(Paragraph("Performance evaluation analysis was not generated by the LLM.", 
                                 ParagraphStyle(
                                     name='EmptySection',
                                     fontName='Helvetica-Italic',
                                     fontSize=10,
                                     textColor=colors.HexColor("#666666"),
                                     spaceAfter=12
                                 )))
            return
        
        # Clean up markdown formatting
        content = content.replace('**', '')
        
        # Parse the content to extract subheadings
        # The content format is: "Communication Quality: [content] Client Engagement: [content] etc."
        
        # Define the subheadings we're looking for
        subheadings = [
            "Communication Quality:",
            "Client Engagement:",
            "Objection Handling:",
            "Deal Progression:"
        ]
        
        # Find and process each subheading
        current_section = None
        sections = {}
        
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            
            # Check if this line starts a new subheading
            found_subheading = False
            for subheading in subheadings:
                if line.startswith(subheading):
                    current_section = subheading.replace(':', '')
                    sections[current_section] = line.replace(subheading, '').strip()
                    found_subheading = True
                    break
            
            # If not a subheading, append to current section
            if not found_subheading and current_section and line:
                if current_section in sections:
                    sections[current_section] += ' ' + line
                else:
                    sections[current_section] = line
        
        # Create proper subheadings with content
        for subheading, content_text in sections.items():
            if content_text:
                # Add subheading
                story.append(Paragraph(f"<b>{subheading}:</b>", ParagraphStyle(
                    name='EvalSubheader',
                    parent=self.pro_styles['Regular'],
                    fontSize=11,
                    fontName='Helvetica-Bold',
                    textColor=colors.HexColor("#0D47A1"),
                    spaceAfter=4,
                    spaceBefore=12 if subheading == list(sections.keys())[0] else 8
                )))
                
                # Add content
                story.append(Paragraph(content_text.strip(), self.pro_styles['Regular']))
                story.append(Spacer(1, 6))
        
        # If we couldn't parse subheadings, show as-is
        if not sections:
            story.append(Paragraph(content, self.pro_styles['Regular']))

    def _create_coaching_snapshot(self, story, content: str):
        """Create coaching snapshot section"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if line.startswith('**What went well:**'):
                story.append(Paragraph("<b>✓ What went well:</b>", ParagraphStyle(
                    name='SnapshotHeader',
                    fontName='Helvetica-Bold',
                    fontSize=11,
                    textColor=colors.HexColor("#2E7D32"),
                    spaceAfter=4,
                    spaceBefore=0
                )))
            elif line.startswith('**What cost leverage:**'):
                story.append(Paragraph("<b>⚠ What cost leverage:</b>", ParagraphStyle(
                    name='SnapshotHeader',
                    fontName='Helvetica-Bold',
                    fontSize=11,
                    textColor=colors.HexColor("#FF8F00"),
                    spaceAfter=4,
                    spaceBefore=12
                )))
            elif line.startswith('**What to fix in the next meeting:**'):
                story.append(Paragraph("<b>🔧 What to fix in the next meeting:</b>", ParagraphStyle(
                    name='SnapshotHeader',
                    fontName='Helvetica-Bold',
                    fontSize=11,
                    textColor=colors.HexColor("#0D47A1"),
                    spaceAfter=4,
                    spaceBefore=12
                )))
            elif line.startswith('**Overall Deal Health:**'):
                # Extract the emoji and text
                health_text = line.replace('**Overall Deal Health:**', '').strip()
                story.append(Paragraph(f"<b>📊 Overall Deal Health:</b> {health_text}", ParagraphStyle(
                    name='SnapshotHeader',
                    fontName='Helvetica-Bold',
                    fontSize=11,
                    textColor=colors.HexColor("#424242"),
                    spaceAfter=4,
                    spaceBefore=12
                )))
            elif line.startswith('-') or line.startswith('•'):
                text = line[1:].strip()
                story.append(Paragraph(f"• {text}", self.pro_styles['KeyPoint']))
                story.append(Spacer(1, 2))
            elif line and not line.startswith('**'):
                story.append(Paragraph(line, self.pro_styles['Regular']))
                story.append(Spacer(1, 4))

    def _create_coaching_playbook(self, story, content: str):
        """Create coaching playbook section - IMPROVED VERSION WITH BETTER PARSING"""
        if not content or content.strip() == "" or content.strip() == "No specific coaching opportunities identified in this conversation.":
            story.append(Paragraph("No specific coaching opportunities identified in this conversation.", 
                                 ParagraphStyle(
                                     name='NoCoaching',
                                     fontName='Helvetica-Italic',
                                     fontSize=10,
                                     textColor=colors.HexColor("#666666"),
                                     spaceAfter=12
                                 )))
            return
        
        # Clean up the content - remove redundant markdown formatting
        content = content.strip()
        
        # Split into lines for processing
        lines = content.split('\n')
        
        # Track if we've parsed any structured content
        parsed_structured_content = False
        
        # Process each section
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Handle "When client says:" section
            if 'When client says:' in line:
                parsed_structured_content = True
                # Extract the client quote
                quote_match = re.search(r'[""](.*?)[""]', line)
                if quote_match:
                    client_quote = quote_match.group(1)
                else:
                    # Try to extract from the line content
                    client_quote = line.split('When client says:')[-1].strip().strip('"')
                
                # Add to story
                story.append(Paragraph("<b>When client says:</b>", ParagraphStyle(
                    name='PlaybookClientSays',
                    fontName='Helvetica-Bold',
                    fontSize=10,
                    textColor=colors.HexColor("#0D47A1"),
                    spaceAfter=4,
                    spaceBefore=0
                )))
                
                story.append(Paragraph(f"\"{client_quote}\"", ParagraphStyle(
                    name='PlaybookQuote',
                    fontName='Helvetica-Oblique',
                    fontSize=10,
                    textColor=colors.HexColor("#424242"),
                    leftIndent=15,
                    spaceAfter=8
                )))
                
                # Look for "Say this instead:" - check current and next lines
                say_line = line
                if 'Say this instead:' not in line and i + 1 < len(lines):
                    # Check if next line contains it
                    if 'Say this instead:' in lines[i + 1]:
                        i += 1
                        say_line = lines[i]
                
                # Extract agent response
                if 'Say this instead:' in say_line:
                    response_match = re.search(r'[""](.*?)[""]', say_line)
                    if response_match:
                        agent_response = response_match.group(1)
                    else:
                        agent_response = say_line.split('Say this instead:')[-1].strip().strip('"')
                    
                    # Clean up markdown formatting
                    agent_response = agent_response.replace('**', '').replace('>', '').strip()
                    
                    story.append(Paragraph("<b>Say this instead:</b>", ParagraphStyle(
                        name='PlaybookResponse',
                        fontName='Helvetica-Bold',
                        fontSize=10,
                        textColor=colors.HexColor("#2E7D32"),
                        spaceAfter=4,
                        spaceBefore=8
                    )))
                    
                    story.append(Paragraph(f"\"{agent_response}\"", ParagraphStyle(
                        name='PlaybookQuote',
                        fontName='Helvetica-Oblique',
                        fontSize=10,
                        textColor=colors.HexColor("#424242"),
                        leftIndent=15,
                        spaceAfter=8
                    )))
                
                # Look for "Why this works:" - check current and next lines
                why_line = ""
                if i + 1 < len(lines):
                    next_lines_to_check = min(3, len(lines) - i - 1)
                    for j in range(1, next_lines_to_check + 1):
                        if 'Why this works:' in lines[i + j]:
                            why_line = lines[i + j]
                            i += j
                            break
                
                # Extract explanation
                if why_line and 'Why this works:' in why_line:
                    explanation = why_line.split('Why this works:')[-1].strip()
                    explanation = explanation.replace('**', '').replace('>', '').strip()
                    
                    story.append(Paragraph("<b>Why this works:</b>", ParagraphStyle(
                        name='PlaybookWhy',
                        fontName='Helvetica-Bold',
                        fontSize=10,
                        textColor=colors.HexColor("#666666"),
                        spaceAfter=4,
                        spaceBefore=8
                    )))
                    
                    story.append(Paragraph(explanation, ParagraphStyle(
                        name='PlaybookExplanation',
                        fontName='Helvetica',
                        fontSize=10,
                        textColor=colors.HexColor("#424242"),
                        leftIndent=15,
                        spaceAfter=12
                    )))
            
            # Handle "Rule for Objection Handling:" section
            elif line.startswith('**Rule for') or 'Rule for Objection Handling:' in line:
                parsed_structured_content = True
                # Extract the rule text
                rule_text = line.replace('**Rule for', '').replace('**', '').replace('Rule for Objection Handling:', '').strip()
                if ':**' in rule_text:
                    rule_text = rule_text.split(':**')[1].strip()
                
                story.append(Paragraph(f"<b>Rule for Objection Handling:</b>", ParagraphStyle(
                    name='PlaybookRule',
                    fontName='Helvetica-Bold',
                    fontSize=10,
                    textColor=colors.HexColor("#1976D2"),
                    spaceAfter=4,
                    spaceBefore=12
                )))
                
                # Look for the rule details in next lines
                rule_details = []
                j = i + 1
                while j < len(lines) and j < i + 5:  # Check next 5 lines
                    next_line = lines[j].strip()
                    if next_line and not next_line.startswith('**'):  # Stop at next section
                        rule_details.append(next_line.replace('>', '').replace('**', '').strip())
                        j += 1
                    else:
                        break
                
                if rule_details:
                    for detail in rule_details:
                        story.append(Paragraph(f"• {detail}", ParagraphStyle(
                            name='PlaybookDetail',
                            fontName='Helvetica',
                            fontSize=10,
                            textColor=colors.HexColor("#424242"),
                            leftIndent=15,
                            spaceAfter=2
                        )))
                
                story.append(Spacer(1, 8))
            
            # Handle "Before closing, always ask:" section
            elif 'Before closing, always ask:' in line:
                parsed_structured_content = True
                # Extract the closing question
                question_match = re.search(r'[""](.*?)[""]', line)
                if question_match:
                    closing_question = question_match.group(1)
                else:
                    # Try to extract from the line
                    closing_question = line.split('Before closing, always ask:')[-1].strip()
                    closing_question = closing_question.replace('**', '').replace('>', '').strip()
                
                # Skip if empty
                if not closing_question or closing_question in ['""', "''"]:
                    i += 1
                    continue
                
                story.append(Paragraph("<b>Before closing, always ask:</b>", ParagraphStyle(
                    name='PlaybookClosing',
                    fontName='Helvetica-Bold',
                    fontSize=10,
                    textColor=colors.HexColor("#D32F2F"),
                    spaceAfter=4,
                    spaceBefore=12
                )))
                
                story.append(Paragraph(f"\"{closing_question}\"", ParagraphStyle(
                    name='PlaybookQuote',
                    fontName='Helvetica-Oblique',
                    fontSize=10,
                    textColor=colors.HexColor("#424242"),
                    leftIndent=15,
                    spaceAfter=8
                )))
                
                # Look for "Why this works:" explanation
                why_explanation = ""
                if i + 1 < len(lines):
                    next_lines_to_check = min(3, len(lines) - i - 1)
                    for j in range(1, next_lines_to_check + 1):
                        if 'Why this works:' in lines[i + j]:
                            why_explanation = lines[i + j].split('Why this works:')[-1].strip()
                            why_explanation = why_explanation.replace('**', '').replace('>', '').strip()
                            i += j
                            break
                
                if why_explanation:
                    story.append(Paragraph("<b>Why this works:</b>", ParagraphStyle(
                        name='PlaybookWhy',
                        fontName='Helvetica-Bold',
                        fontSize=10,
                        textColor=colors.HexColor("#666666"),
                        spaceAfter=4,
                        spaceBefore=8
                    )))
                    
                    story.append(Paragraph(why_explanation, ParagraphStyle(
                        name='PlaybookExplanation',
                        fontName='Helvetica',
                        fontSize=10,
                        textColor=colors.HexColor("#424242"),
                        leftIndent=15,
                        spaceAfter=12
                    )))
            
            i += 1
        
        # Only show raw content if we didn't parse any structured content
        if not parsed_structured_content:
            # Clean up the content first
            cleaned_content = content.replace('**', '').replace('>', '').strip()
            story.append(Paragraph(cleaned_content, self.pro_styles['Regular']))

    def _create_7day_focus(self, story, content: str):
        """Create next 7-day coaching focus"""
        story.append(Paragraph("<b>For the next week, focus on ONLY these three things:</b>", 
                             ParagraphStyle(
                                 name='FocusHeader',
                                 fontName='Helvetica-Bold',
                                 fontSize=11,
                                 textColor=colors.HexColor("#D32F2F"),
                                 spaceAfter=8
                             )))
        
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('1.') or line.startswith('2.') or line.startswith('3.'):
                text = line[2:].strip() if len(line) > 2 else line
                story.append(Paragraph(f"<b>{line[:2]}</b> {text}", ParagraphStyle(
                    name='FocusItem',
                    fontName='Helvetica-Bold',
                    fontSize=10,
                    textColor=colors.HexColor("#424242"),
                    leftIndent=15,
                    spaceAfter=6
                )))

    # ============================================================
    # HELPER METHODS FOR UNIFIED REPORT
    # ============================================================

    def _create_text_section(self, story, content):
        """Create a simple text section"""
        if isinstance(content, str):
            paragraphs = content.split('\n\n') if '\n\n' in content else [content]
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), self.pro_styles['Regular']))
                    story.append(Spacer(1, 4))

    def _create_bullet_list(self, story, content):
        """Create a bullet point list section"""
        if isinstance(content, str):
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('-') or line.startswith('•'):
                    text = line[1:].strip()
                    story.append(Paragraph(f"• {text}", self.pro_styles['KeyPoint']))
                    story.append(Spacer(1, 3))
        elif isinstance(content, list):
            for item in content:
                story.append(Paragraph(f"• {item}", self.pro_styles['KeyPoint']))
                story.append(Spacer(1, 3))

    def _parse_action_items_table(self, content: str) -> List[Dict]:
        """Parse markdown table from action items section into list of dictionaries"""
        items = []
        
        if not content or not isinstance(content, str):
            return items
        
        lines = content.strip().split('\n')
        
        # Find table rows (skip header and separator)
        in_table = False
        for line in lines:
            line = line.strip()
            
            # Detect table rows (start with |)
            if line.startswith('|'):
                # Skip header row
                if 'Task Description' in line or 'Responsible Person' in line:
                    in_table = True
                    continue
                
                # Skip separator row (---|---|---)
                if '---' in line or '===' in line:
                    continue
                
                # Parse data row
                if in_table:
                    parts = [p.strip() for p in line.split('|')]
                    # Remove empty first/last elements from split
                    parts = [p for p in parts if p]
                    
                    if len(parts) >= 5:
                        items.append({
                            'task': parts[0],
                            'responsible': parts[1],
                            'deadline': parts[2],
                            'priority': parts[3],
                            'status': parts[4]
                        })
        
        return items

    def _create_professional_action_items(self, story, content: List[Dict]):
        """Create professional action items table with proper text wrapping"""
        if not content:
            return
            
        # Header style for centered alignment
        header_style = ParagraphStyle(
            name='TableHeader',
            fontName='Helvetica-Bold',
            fontSize=10,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#0D47A1")
        )
        
        table_data = [[
            Paragraph('Task Description', header_style),
            Paragraph('Responsible Person', header_style),
            Paragraph('Deadline', header_style),
            Paragraph('Priority', header_style),
            Paragraph('Status', header_style)
        ]]

        for item in content:
            task_desc = self._wrap_table_text(item.get('task', ''), column_type="task")
            responsible = self._wrap_table_text(item.get('responsible', ''), column_type="responsible")
            deadline = self._wrap_table_text(item.get('deadline', 'No due date was mentioned'), column_type="deadline")
            priority = item.get('priority', 'Medium')
            status = item.get('status', 'Pending')
            
            # Color coding for priority
            if priority.lower() == 'high':
                priority_color = colors.HexColor("#D32F2F")  # Red
            elif priority.lower() == 'medium':
                priority_color = colors.HexColor("#FF8F00")  # Orange
            else:
                priority_color = colors.HexColor("#2E7D32")  # Green
            
            table_data.append([
                Paragraph(task_desc, ParagraphStyle(
                    name='TableTask',
                    fontName='Helvetica',
                    fontSize=10,
                    textColor=colors.black,
                    wordWrap='LTR',
                    leading=12
                )),
                Paragraph(responsible, ParagraphStyle(
                    name='TableResponsible',
                    fontName='Helvetica',
                    fontSize=10,
                    alignment=TA_CENTER,
                    wordWrap='LTR'
                )),
                Paragraph(deadline, ParagraphStyle(
                    name='TableDeadline',
                    fontName='Helvetica',
                    fontSize=10,
                    alignment=TA_CENTER
                )),
                Paragraph(priority, ParagraphStyle(
                    name='TablePriority',
                    fontName='Helvetica-Bold',
                    fontSize=10,
                    alignment=TA_CENTER,
                    textColor=priority_color
                )),
                Paragraph(status, ParagraphStyle(
                    name='TableStatus',
                    fontName='Helvetica',
                    fontSize=10,
                    alignment=TA_CENTER
                ))
            ])

        if len(table_data) > 1:
            # Calculate column widths - Adjusted to prevent overlapping
            col_widths = [2.5 * inch, 1.8 * inch, 1.6 * inch, 0.7 * inch, 0.7 * inch]
            
            # Create table
            table = Table(table_data, colWidths=col_widths, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.white),  # White header
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),  # Black grid only
                ('PADDING', (0, 0), (-1, -1), 6),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),  # White rows
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            
            story.append(KeepTogether(table))
        else:
            # Fallback if no table found
            story.append(Paragraph("No action items found", self.pro_styles['Regular']))

    def _create_strengths_section(self, story, content):
        """Create strengths with evidence section"""
        if isinstance(content, str):
            # Parse bullet points
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('-') or line.startswith('•'):
                    # Parse strength with evidence
                    text = line[1:].strip()
                    
                    # Replace ** with bold tags
                    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
                    
                    story.append(Paragraph(f"• {text}", ParagraphStyle(
                        name='StrengthPoint',
                        parent=self.pro_styles['KeyPoint'],
                        textColor=colors.black,
                        fontName='Helvetica',
                        leftIndent=15
                    )))
                    story.append(Spacer(1, 4))
        elif isinstance(content, list):
            for item in content:
                story.append(Paragraph(f"• {item}", ParagraphStyle(
                    name='StrengthPoint',
                    parent=self.pro_styles['KeyPoint'],
                    textColor=colors.black
                )))
                story.append(Spacer(1, 4))

    def _create_objections_opportunities(self, story, content):
        """Create objections and opportunities section"""
        if isinstance(content, str):
            lines = content.split('\n')
            current_subsection = None
            
            for line in lines:
                line = line.strip()
                if 'Objections:' in line:
                    story.append(Paragraph("<b>Objections:</b>", ParagraphStyle(
                        name='ObjSubheader',
                        parent=self.pro_styles['Regular'],
                        fontSize=10,
                        fontName='Helvetica-Bold',
                        spaceAfter=4,
                        spaceBefore=6
                    )))
                    current_subsection = 'objections'
                elif 'Opportunities:' in line:
                    story.append(Paragraph("<b>Opportunities:</b>", ParagraphStyle(
                        name='OppSubheader',
                        parent=self.pro_styles['Regular'],
                        fontSize=10,
                        fontName='Helvetica-Bold',
                        spaceAfter=4,
                        spaceBefore=6
                    )))
                    current_subsection = 'opportunities'
                elif line.startswith('-') or line.startswith('•'):
                    text = line[1:].strip()
                    color = colors.HexColor("#D32F2F") if current_subsection == 'objections' else colors.HexColor("#2E7D32")
                    story.append(Paragraph(f"• {text}", ParagraphStyle(
                        name=f'{current_subsection}Point',
                        parent=self.pro_styles['KeyPoint'],
                        textColor=color,
                        leftIndent=15
                    )))
                    story.append(Spacer(1, 3))

    def _create_decisions_commitments(self, story, content):
        """Create decisions and commitments section"""
        if isinstance(content, str):
            lines = content.split('\n')
            current_subsection = None
            
            for line in lines:
                line = line.strip()
                if 'Decisions Made:' in line or 'Decisions:' in line:
                    story.append(Paragraph("<b>Decisions Made:</b>", ParagraphStyle(
                        name='DecisionsSubheader',
                        parent=self.pro_styles['Regular'],
                        fontSize=10,
                        fontName='Helvetica-Bold',
                        spaceAfter=4,
                        spaceBefore=6
                    )))
                    current_subsection = 'decisions'
                elif 'Commitments Given:' in line or 'Commitments:' in line:
                    story.append(Paragraph("<b>Commitments Given:</b>", ParagraphStyle(
                        name='CommitmentsSubheader',
                        parent=self.pro_styles['Regular'],
                        fontSize=10,
                        fontName='Helvetica-Bold',
                        spaceAfter=4,
                        spaceBefore=6
                    )))
                    current_subsection = 'commitments'
                elif line.startswith('-') or line.startswith('•'):
                    text = line[1:].strip()
                    story.append(Paragraph(f"• {text}", self.pro_styles['KeyPoint']))
                    story.append(Spacer(1, 3))

    def _create_performance_scores(self, story, content: str):
        """Create performance scores table with proper parsing"""
        table_data = [['Performance Dimension', 'Score', 'Justification']]
        
        # Parse the content for scores
        lines = content.split('\n')
        current_dimension = None
        current_score = None
        
        for line in lines:
            line = line.strip()
            
            # Look for dimension lines (like "**Communication Effectiveness:** 8/10")
            if line.startswith('**') and ':' in line and '/' in line:
                parts = line.replace('**', '').split(':', 1)
                if len(parts) == 2:
                    current_dimension = parts[0].strip()
                    current_score = parts[1].strip()
            
            # Look for justification lines
            elif line.startswith('**Justification:**') and current_dimension and current_score:
                justification = line.replace('**Justification:**', '').strip()
                table_data.append([
                    Paragraph(current_dimension, self.pro_styles['TableText']),
                    Paragraph(current_score, ParagraphStyle(
                        name='ScoreText',
                        parent=self.pro_styles['TableText'],
                        fontName='Helvetica-Bold',
                        alignment=TA_CENTER,
                        fontSize=11
                    )),
                    Paragraph(justification, self.pro_styles['TableText'])
                ])
                current_dimension = None
                current_score = None
            
            # Look for overall assessment
            elif line.startswith('**Overall Performance:**'):
                parts = line.replace('**Overall Performance:**', '').split('**Summary:**')
                if len(parts) == 2:
                    overall_score = parts[0].strip()
                    summary = parts[1].replace('**', '').strip()
                    table_data.append([
                        Paragraph('<b>Overall Assessment</b>', self.pro_styles['TableText']),
                        Paragraph(overall_score, ParagraphStyle(
                            name='ScoreText',
                            parent=self.pro_styles['TableText'],
                            fontName='Helvetica-Bold',
                            alignment=TA_CENTER,
                            fontSize=11
                        )),
                        Paragraph(summary, self.pro_styles['TableText'])
                    ])
        
        # If we couldn't parse specific dimensions but have content
        if len(table_data) == 1 and content.strip():
            # Fallback: show content as-is
            story.append(Paragraph(content, self.pro_styles['Regular']))
            return
        
        # Create the table
        if len(table_data) > 1:
            table = Table(table_data, colWidths=[1.8*inch, 0.8*inch, 3.9*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1976D2")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), "Helvetica-Bold"),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#BBDEFB")),
                ('PADDING', (0, 0), (-1, -1), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(table)
        else:
            story.append(Paragraph("No performance scores available", self.pro_styles['Regular']))

    def _create_deal_intelligence_status(self, story, content: str):
        """Create comprehensive deal and intelligence status evaluation section"""
        if isinstance(content, str):
            lines = content.split('\n')
            current_subsection = None
            
            for line in lines:
                line_stripped = line.strip()
                
                # Detect subsections
                if line_stripped.startswith('**') and line_stripped.endswith('**'):
                    # Subsection header
                    subsection_title = line_stripped.replace('**', '').replace(':', '').strip()
                    current_subsection = subsection_title
                    
                    # Color-coded subsection headers
                    if 'Deal Status' in subsection_title:
                        color = colors.HexColor("#0D47A1")  # Blue
                    elif 'Intelligence' in subsection_title:
                        color = colors.HexColor("#1976D2")  # Light Blue
                    elif 'Risk' in subsection_title:
                        color = colors.HexColor("#D32F2F")  # Red
                    elif 'Next Steps' in subsection_title or 'Recommended' in subsection_title:
                        color = colors.HexColor("#2E7D32")  # Green
                    elif 'Probability' in subsection_title:
                        color = colors.HexColor("#FF8F00")  # Orange
                    else:
                        color = colors.HexColor("#424242")  # Dark Gray
                    
                    story.append(Paragraph(subsection_title, ParagraphStyle(
                        name=f'DealSubsection_{subsection_title[:10]}',
                        fontName='Helvetica-Bold',
                        fontSize=11,
                        textColor=color,
                        spaceBefore=8,
                        spaceAfter=4
                    )))
                    
                elif line_stripped.startswith('-') or line_stripped.startswith('•'):
                    # Bullet point
                    text = line_stripped[1:].strip()
                    story.append(Paragraph(f"• {text}", self.pro_styles['KeyPoint']))
                    story.append(Spacer(1, 3))
                    
                elif line_stripped and line_stripped[0].isdigit() and '.' in line_stripped[:3]:
                    # Numbered list item
                    story.append(Paragraph(line_stripped, ParagraphStyle(
                        name='NumberedItem',
                        fontName='Helvetica',
                        fontSize=10,
                        leftIndent=12,
                        spaceBefore=3,
                        spaceAfter=3,
                        textColor=colors.black
                    )))
                    
                elif ':' in line_stripped and len(line_stripped.split(':')[0]) < 40:
                    # Label: Value format
                    parts = line_stripped.split(':', 1)
                    label = parts[0].strip()
                    value = parts[1].strip() if len(parts) > 1 else ''
                    
                    # Convert markdown bold to HTML bold
                    label = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', label)
                    
                    # Highlight probability percentage
                    if 'Probability' in label or '%' in value:
                        story.append(Paragraph(f"<b>{label}:</b> <font color='#FF8F00'><b>{value}</b></font>", 
                                             self.pro_styles['Regular']))
                    else:
                        story.append(Paragraph(f"<b>{label}:</b> {value}", self.pro_styles['Regular']))
                    story.append(Spacer(1, 4))
                    
                elif line_stripped and current_subsection:
                    # Regular paragraph content under a subsection
                    story.append(Paragraph(line_stripped, self.pro_styles['Regular']))
                    story.append(Spacer(1, 4))

    # ============================================================
    # MAIN FUNCTION
    # ============================================================

def main():
    """Main execution function"""
    print("🏠 REAL ESTATE COACHING EVALUATION SYSTEM")
    print("="*60)
    
    # Ask for transcript file
    default_file = "conversation.json"
    transcript_file = input(f"📂 Enter JSON transcript file path (press Enter for '{default_file}'): ").strip().strip('"')
    
    if not transcript_file:
        transcript_file = default_file
    
    # Create evaluator
    evaluator = RealEstateSalesMeetingSummarizer(transcript_file=transcript_file)
    
    # Generate unified report
    print(f"\n🚀 Generating Unified Coaching Evaluation...")
    evaluator.generate_unified_coaching_evaluation()


if __name__ == "__main__":
    main()