# rizwan.py - RAG-Powered Critical Evaluation Engine for Real Estate Sales
# Combined Best-of-Both: Enterprise Visual Design + Robust Parsing
# Output: 4 PDFs:
#   1. coaching_summary.pdf           — Coaching Summary (Version 1)
#   2. agent_summary.pdf              — Agent Profile Summary (Version 2)
#   3. coaching_insights_viz.pdf      — Visualization PDF for Coaching Insights
#   4. agent_performance_viz.pdf      — Visualization PDF for Agent Performance
# + 8 JSON files:
#   1. coaching_insights.json         — Raw data for coaching visualization PDF
#   2. agent_performance.json         — Raw data for agent performance visualization PDF
#   3. action_items.json              — All action items extracted from the summary
#   4. all_visualizations.json        — Master combined visualization data (Pydantic-validated)
#   5. coaching_recommendations.json  — Structured coaching recommendations (strengths, failures, playbook, 7-day plan)
#   6. agent_tier_calculation.json    — Agent tier classification methodology and calculation
#   7. agent_performance_aggregate.json — Aggregated performance metrics for dashboard (new)
#   8. deal_outcome.json              — Final deal outcome / client decision status
# Brutally honest, zero-repetition coaching evaluation

import json
import os
import re
import traceback
import pickle
import math
import logging
from datetime import datetime
from typing import Optional, List, Dict, Any, Tuple
from pathlib import Path
from enum import Enum

import requests
import configparser
import numpy as np

# RAG and Embeddings
try:
    from sentence_transformers import SentenceTransformer
    EMBEDDINGS_AVAILABLE = True
except ImportError:
    EMBEDDINGS_AVAILABLE = False
    print("sentence-transformers not available. Install: pip install sentence-transformers")

try:
    import faiss
    FAISS_AVAILABLE = True
except ImportError:
    FAISS_AVAILABLE = False
    print("FAISS not available. Install: pip install faiss-cpu")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available. Install: pip install python-docx")

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, KeepTogether, Flowable, HRFlowable, CondPageBreak
)
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.pdfgen import canvas
from reportlab.graphics.shapes import Drawing, Rect, String, Circle, Line, Polygon
from reportlab.graphics import renderPDF
from dotenv import load_dotenv
from pydantic import BaseModel, Field, field_validator, ConfigDict

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


# ════════════════════════════════════════════════════════
# PYDANTIC V2 MODELS
# ════════════════════════════════════════════════════════

class LLMProvider(str, Enum):
    OPENAI = "openai"
    OPENROUTER = "openrouter"
    CUSTOM = "custom"

class LLMConfig(BaseModel):
    model_config = ConfigDict(protected_namespaces=())
    provider: str = Field(..., description="LLM provider name")
    api_key: str = Field(..., description="API key for the LLM provider")
    model: str = Field(..., description="Model name to use")
    base_url: Optional[str] = Field(..., description="Base URL for API calls")
    temperature: float = Field(..., ge=0.0, le=2.0)
    max_tokens: int = Field(..., ge=1, le=32000)
    timeout: int = Field(..., ge=1)

    @field_validator('base_url', mode='before')
    @classmethod
    def set_base_url_based_on_provider(cls, v, info):
        if v:
            return v
        provider = info.data.get('provider', 'openai').lower().strip()
        if provider == "openai":
            return "https://api.openai.com/v1"
        elif provider == "openrouter":
            return "https://openrouter.ai/api/v1"
        return v

class TranscriptSegment(BaseModel):
    speaker_id: Optional[str] = None
    speaker_name: Optional[str] = None
    transcript: str
    segment_id: Optional[int] = None
    start: Optional[float] = 0.0
    end: Optional[float] = 0.0
    duration: Optional[float] = 0.0

class TranscriptData(BaseModel):
    transcripts: List[TranscriptSegment]
    agenda: str
    metadata: Dict[str, Any] = {}
    summary_info: Dict[str, Any]

class AgendaAnalysis(BaseModel):
    agenda_words: List[str]
    relevant_words: List[str]
    relevance_percentage: float

class MeetingAnalysis(BaseModel):
    content: str
    meeting_id: str
    transcript_data: TranscriptData
    agenda_analysis: AgendaAnalysis
    enhanced_analysis: Dict[str, Any] = {}

class ActionItem(BaseModel):
    task: str
    responsible: str
    deadline: str = Field(default="No due date was mentioned")
    priority: str = Field(default="Medium")
    status: str = Field(default="Pending")

    @field_validator('deadline', mode='before')
    @classmethod
    def validate_deadline(cls, v):
        if not v or v.strip() == "":
            return "No due date was mentioned"
        return v

class KnowledgeChunk(BaseModel):
    chunk_id: int
    text: str
    embedding: Optional[List[float]] = None
    metadata: Dict[str, Any] = {}

class CoachingEvaluation(BaseModel):
    executive_evaluation: str
    critical_analysis: str
    strengths: List[str]
    improvement_areas: List[str]
    coaching_recommendations: List[str]
    knowledge_references: List[str]
    performance_scores: Dict[str, Any] = {}


# ═══════════════════════════════════════════════════
# ENHANCED PYDANTIC MODELS — NEW JSON OUTPUTS
# ═══════════════════════════════════════════════════

class VisualizationEntry(BaseModel):
    """Single visualization entry with enhanced schema for frontend rendering."""
    model_config = ConfigDict(protected_namespaces=())
    visualization_id: str = Field(..., description="Unique chart identifier")
    visualization_name: str = Field(..., description="Display title")
    category: str = Field(default="general", description="coaching_insights | agent_performance | combined")
    related_pdf: str = Field(default="", description="Which PDF this chart appears in")
    chart_type: str = Field(..., description="bar | horizontalBar | line | donut | radar | gauge | funnel | scatter | stackedBar | heatmap | quadrant | scoreCard | table")
    description: str = Field(default="")
    labels: List[str] = Field(default_factory=list)
    datasets: List[Dict[str, Any]] = Field(default_factory=list)
    metric_source: str = Field(default="transcript_analysis", description="Where the metric comes from")
    metric_value: Optional[Any] = Field(default=None, description="Primary metric value if applicable")
    explanation: str = Field(default="", description="What this visualization reveals")
    visualization_priority: int = Field(default=5, ge=1, le=10, description="1=highest priority, 10=lowest")
    display_style: Dict[str, Any] = Field(default_factory=dict, description="Frontend rendering hints")
    metadata: Dict[str, Any] = Field(default_factory=dict)
    timestamp: str = Field(default="")

class AllVisualizationsOutput(BaseModel):
    """Master JSON combining all visualization data."""
    model_config = ConfigDict(protected_namespaces=())
    generated_at: str
    report_type: str = "all_visualizations"
    schema_version: str = "2.0"
    metadata: Dict[str, Any] = {}
    total_visualizations: int = 0
    coaching_insights_count: int = 0
    agent_performance_count: int = 0
    visualizations: List[VisualizationEntry] = Field(default_factory=list)

class StrengthEntry(BaseModel):
    """A single strength finding."""
    strength: str
    timestamp: str = ""
    evidence: str = ""
    impact: str = ""

class FailureEntry(BaseModel):
    """A single failure/weakness finding."""
    failure_title: str
    timestamp: str = ""
    quote: str = ""
    what_happened: str = ""
    what_top_agent_would_do: str = ""
    revenue_impact: str = ""

class PlaybookEntry(BaseModel):
    """A single coaching playbook scenario."""
    scenario: str
    client_trigger: str = ""
    risk: str = ""
    assertive_approach: str = ""
    consultative_approach: str = ""
    why_it_works: str = ""

class CoachingDayPlan(BaseModel):
    """A single day's coaching plan."""
    day: int
    focus: str
    drill: str = ""
    metric: str = ""

class CoachingRecommendationsOutput(BaseModel):
    """Structured coaching recommendations extracted from LLM analysis."""
    model_config = ConfigDict(protected_namespaces=())
    generated_at: str
    report_type: str = "coaching_recommendations"
    schema_version: str = "2.0"
    agent_name: str = ""
    metadata: Dict[str, Any] = {}
    strengths: List[StrengthEntry] = Field(default_factory=list)
    failures: List[FailureEntry] = Field(default_factory=list)
    tactical_playbook: List[PlaybookEntry] = Field(default_factory=list)
    seven_day_plan: List[CoachingDayPlan] = Field(default_factory=list)
    summary: Dict[str, Any] = Field(default_factory=dict)

class AgentTierCalculationOutput(BaseModel):
    """Agent tier classification methodology and calculation."""
    model_config = ConfigDict(protected_namespaces=())
    generated_at: str
    report_type: str = "agent_tier_calculation"
    schema_version: str = "2.0"
    agent_name: str = ""
    metrics_used: Dict[str, float] = Field(default_factory=dict, description="Score label → value")
    weight_distribution: Dict[str, float] = Field(default_factory=dict, description="Category → weight %")
    scoring_formula: str = Field(default="weighted_average", description="How the final score was derived")
    overall_score: float = 0.0
    tier_thresholds: Dict[str, str] = Field(default_factory=lambda: {
        "Elite Performer": "9.0-10.0",
        "Strong Performer": "7.0-8.9",
        "Developing Agent": "5.0-6.9",
        "Needs Improvement": "3.0-4.9",
        "Critical Development": "0.0-2.9",
    })
    final_agent_tier: str = ""
    tier_justification: str = ""
    promotion_criteria: List[str] = Field(default_factory=list)
    comparative_metrics: List[Dict[str, str]] = Field(default_factory=list)
    metadata: Dict[str, Any] = Field(default_factory=dict)


# ════════════════════════════════════════════════════════
# NEW PYDANTIC MODEL FOR AGGREGATED PERFORMANCE DATA (JSON 7)
# ════════════════════════════════════════════════════════
class AgentPerformanceAggregate(BaseModel):
    """Aggregated performance data for dashboard across meetings."""
    model_config = ConfigDict(protected_namespaces=())
    generated_at: str
    agent_name: str
    transcript_file: str
    meeting_date: str  # can be derived from transcript or current date
    overall_performance_score: float
    overall_performance_score_out_of: int = 10
    overall_deal_probability: Optional[float] = None
    deal_probability_out_of: int = 100
    talk_ratio: Dict[str, Optional[float]]  # e.g., {"agent": 45, "client": 55}
    performance_dimensions: List[Dict[str, Any]]  # label, score, max, justification, confidence
    discovery_pillars: List[Dict[str, Any]]    # pillar, status, evidence, confidence
    emotional_intelligence: Dict[str, Any]     # dimensions scores and overall
    tone_scores: Dict[str, float]              # individual tone scores (excluding overall)
    negotiation_scores: Dict[str, float]       # individual negotiation scores
    ethics_risk: Optional[str] = None
    coaching_priority_areas: List[Dict[str, Any]]  # sorted by lowest scores
    # Conversation metrics
    conversation_metrics: Dict[str, Any] = Field(default_factory=dict)
    # Detailed counts for visualizations
    question_type_counts: Dict[str, int] = Field(default_factory=dict)
    objection_type_counts: Dict[str, int] = Field(default_factory=dict)
    objection_handling_counts: Dict[str, Dict[str, int]] = Field(default_factory=dict)
    language_phrase_counts: Dict[str, int] = Field(default_factory=dict)  # confident, hedging, filler
    response_length_distribution: Dict[str, int] = Field(default_factory=dict)
    response_delay_distribution: Dict[str, int] = Field(default_factory=dict)
    sentiment_progression: List[Dict[str, Any]] = Field(default_factory=list)  # phase, score
    trust_vs_pressure: Optional[Dict[str, float]] = None  # trust, pressure
    deal_outcome: Optional[Dict[str, Any]] = Field(default_factory=lambda: None, description="Final deal outcome from transcript")


# ════════════════════════════════════════════════════════
# DEAL OUTCOME PYDANTIC MODEL
# ════════════════════════════════════════════════════════

class DealOutcomeResult(BaseModel):
    """Final deal outcome / client decision extracted from transcript."""
    model_config = ConfigDict(protected_namespaces=())
    generated_at: str = Field(default="")
    report_type: str = Field(default="deal_outcome")
    schema_version: str = Field(default="1.0")
    agent_name: str = Field(default="")
    status: str = Field(
        default="Decision Pending / Not Clearly Stated",
        description=(
            "Final deal status. One of: 'Deal Closed', 'Client Declined the Deal', "
            "'Client Requested Time to Decide', 'Decision Pending / Not Clearly Stated'"
        ),
    )
    confidence: str = Field(
        default="low",
        description="Confidence in the determination: high, medium, or low",
    )
    evidence: str = Field(
        default="Insufficient transcript evidence to determine deal outcome.",
        description="Direct transcript evidence supporting the status determination",
    )
    supporting_quotes: List[str] = Field(
        default_factory=list,
        description="Exact quotes from the transcript supporting the determination",
    )
    metadata: Dict[str, Any] = Field(default_factory=dict)


# ════════════════════════════════════════════════════════
# TRANSCRIPT VALIDATION — ANTI-HALLUCINATION LAYER
# ════════════════════════════════════════════════════════

class TranscriptValidationResult(BaseModel):
    """Result of pre-analysis transcript validation."""
    model_config = ConfigDict(protected_namespaces=())
    is_valid: bool = Field(default=False, description="Whether transcript passed all validation checks")
    has_sufficient_length: bool = Field(default=False, description="Whether transcript has enough exchanges")
    is_real_estate_related: bool = Field(default=False, description="Whether transcript is a real estate conversation")
    total_segments: int = Field(default=0, description="Total number of transcript segments")
    meaningful_exchanges: int = Field(default=0, description="Number of meaningful back-and-forth exchanges")
    real_estate_relevance_score: float = Field(default=0.0, description="0.0 to 1.0 relevance score")
    failure_reasons: List[str] = Field(default_factory=list, description="List of validation failure reasons")
    validation_message: str = Field(default="", description="Human-readable validation summary")

    # Validation thresholds
    MIN_MEANINGFUL_EXCHANGES: int = 10
    MIN_RELEVANCE_SCORE: float = 0.25


class TranscriptValidator:
    """Pre-analysis validation layer to prevent hallucinated outputs.
    
    Validates:
    1. Transcript length (minimum 10 meaningful exchanges)
    2. Domain relevance (must be a real estate conversation)
    
    All downstream generation (PDFs, JSONs, visualizations) must check
    validation results before producing content.
    """

    # Real estate domain keywords — expanded list for robust detection
    REAL_ESTATE_KEYWORDS = {
        # Property types
        'property', 'home', 'house', 'condo', 'apartment', 'townhouse', 'duplex',
        'land', 'lot', 'estate', 'mansion', 'bungalow', 'villa', 'flat',
        'residential', 'commercial', 'industrial', 'rental', 'investment',
        # Transaction terms
        'listing', 'list', 'sell', 'sale', 'buy', 'purchase', 'offer', 'closing',
        'escrow', 'contract', 'deed', 'title', 'settlement', 'transaction',
        'mortgage', 'loan', 'financing', 'down payment', 'pre-approval',
        'pre-qualified', 'lender', 'appraisal', 'inspection', 'contingency',
        # Pricing & valuation
        'price', 'pricing', 'value', 'valuation', 'comps', 'comparable',
        'market analysis', 'cma', 'asking price', 'list price', 'sale price',
        'assessed', 'equity', 'appreciation', 'depreciation',
        # Agents & roles
        'agent', 'realtor', 'broker', 'buyer', 'seller', 'homeowner',
        'listing agent', 'buyer agent', 'commission', 'mls',
        # Property features
        'bedroom', 'bathroom', 'kitchen', 'garage', 'yard', 'pool',
        'square feet', 'sqft', 'acres', 'lot size', 'renovation',
        'remodel', 'upgrade', 'staging', 'curb appeal',
        # Market terms
        'market', 'inventory', 'days on market', 'dom', 'sold',
        'pending', 'active', 'under contract', 'contingent',
        'foreclosure', 'short sale', 'bank owned', 'reo',
        # Location & neighborhood
        'neighborhood', 'location', 'school district', 'zoning',
        'hoa', 'homeowners association', 'subdivision',
        # Showing & open house
        'showing', 'open house', 'walkthrough', 'tour',
        'viewing', 'schedule', 'appointment',
        # Negotiation
        'negotiate', 'negotiation', 'counteroffer', 'counter offer',
        'concession', 'contingency', 'earnest money',
        # Moving & timeline
        'move', 'moving', 'relocate', 'relocation', 'timeline',
        'move-in', 'possession', 'occupancy',
    }

    # Minimum word count for a segment to be considered "meaningful"
    MIN_WORDS_PER_SEGMENT = 3

    @classmethod
    def validate(cls, transcript_data: 'TranscriptData') -> TranscriptValidationResult:
        """Run all validation checks on a transcript before analysis begins.
        
        Returns a TranscriptValidationResult with pass/fail status and detailed reasons.
        """
        result = TranscriptValidationResult()
        result.total_segments = len(transcript_data.transcripts)

        # ── CHECK 1: Transcript Length ──
        meaningful_count = cls._count_meaningful_exchanges(transcript_data)
        result.meaningful_exchanges = meaningful_count
        result.has_sufficient_length = meaningful_count >= result.MIN_MEANINGFUL_EXCHANGES

        if not result.has_sufficient_length:
            result.failure_reasons.append(
                f"Insufficient transcript data: found {meaningful_count} meaningful exchanges, "
                f"but at least {result.MIN_MEANINGFUL_EXCHANGES} are required for reliable analysis."
            )

        # ── CHECK 2: Domain Relevance ──
        relevance_score = cls._compute_real_estate_relevance(transcript_data)
        result.real_estate_relevance_score = relevance_score
        result.is_real_estate_related = relevance_score >= result.MIN_RELEVANCE_SCORE

        if not result.is_real_estate_related:
            result.failure_reasons.append(
                "The transcript does not contain a real estate related discussion between an agent "
                "and a client. Domain relevance score: {:.1%} (minimum required: {:.0%}).".format(
                    relevance_score, result.MIN_RELEVANCE_SCORE
                )
            )

        # ── FINAL VERDICT ──
        result.is_valid = result.has_sufficient_length and result.is_real_estate_related

        if result.is_valid:
            result.validation_message = (
                f"Transcript validation passed. {meaningful_count} meaningful exchanges detected. "
                f"Real estate relevance: {relevance_score:.1%}."
            )
        elif not result.has_sufficient_length and not result.is_real_estate_related:
            result.validation_message = (
                "The transcript could not be analyzed because it either contains insufficient dialogue "
                "or does not represent a real estate conversation between an agent and a client."
            )
        elif not result.has_sufficient_length:
            result.validation_message = (
                "This meeting transcript contains too little conversation to produce reliable analysis. "
                f"At least {result.MIN_MEANINGFUL_EXCHANGES} meaningful exchanges are required. "
                f"Only {meaningful_count} were found."
            )
        else:
            result.validation_message = (
                "The transcript does not contain a real estate related discussion between an agent "
                "and a client. Therefore, meaningful analysis cannot be generated."
            )

        logger.info(f"Transcript validation: valid={result.is_valid}, "
                     f"exchanges={meaningful_count}, relevance={relevance_score:.2%}")
        return result

    @classmethod
    def _count_meaningful_exchanges(cls, transcript_data: 'TranscriptData') -> int:
        """Count segments with meaningful content (not just filler, greetings, or noise).
        
        A meaningful exchange is a segment with at least MIN_WORDS_PER_SEGMENT words
        that contains substantive content (not just 'yeah', 'okay', 'mm-hmm').
        """
        FILLER_ONLY = {
            'yeah', 'yes', 'no', 'okay', 'ok', 'right', 'uh', 'um', 'hmm',
            'mm-hmm', 'mhm', 'uh-huh', 'sure', 'yep', 'nope', 'alright',
            'great', 'good', 'wow', 'oh', 'ah', 'hey', 'hi', 'hello',
            'thanks', 'thank you', 'bye', 'goodbye', 'definitely',
        }
        count = 0
        for seg in transcript_data.transcripts:
            text = seg.transcript.strip()
            if not text:
                continue
            words = text.lower().split()
            # Filter out pure filler responses
            non_filler_words = [w for w in words if w.strip('.,!?') not in FILLER_ONLY]
            if len(non_filler_words) >= cls.MIN_WORDS_PER_SEGMENT:
                count += 1
        return count

    @classmethod
    def _compute_real_estate_relevance(cls, transcript_data: 'TranscriptData') -> float:
        """Compute what fraction of transcript segments contain real estate keywords.
        
        Returns a score between 0.0 and 1.0.
        """
        if not transcript_data.transcripts:
            return 0.0

        full_text = ' '.join(seg.transcript.lower() for seg in transcript_data.transcripts)
        total_words = full_text.split()
        if not total_words:
            return 0.0

        # Count how many unique RE keywords appear in the transcript
        keywords_found = set()
        for keyword in cls.REAL_ESTATE_KEYWORDS:
            if keyword in full_text:
                keywords_found.add(keyword)

        # Also count segments that contain at least one RE keyword
        segments_with_re = 0
        for seg in transcript_data.transcripts:
            seg_lower = seg.transcript.lower()
            if any(kw in seg_lower for kw in cls.REAL_ESTATE_KEYWORDS):
                segments_with_re += 1

        total_segs = len(transcript_data.transcripts)
        segment_ratio = segments_with_re / total_segs if total_segs > 0 else 0.0
        keyword_density = len(keywords_found) / 30.0  # Normalize: 30 unique keywords = 1.0

        # Combined score: 60% segment coverage + 40% keyword density
        relevance = 0.6 * segment_ratio + 0.4 * min(keyword_density, 1.0)
        return round(relevance, 4)


# ════════════════════════════════════════════════════════
# VALIDATION MESSAGE CONSTANTS
# ════════════════════════════════════════════════════════

VALIDATION_MESSAGES = {
    "insufficient_data": "Not enough transcript data to generate meaningful insights.",
    "not_real_estate": (
        "The transcript does not contain a real estate related discussion between an agent "
        "and a client. Therefore, meaningful analysis cannot be generated."
    ),
    "global_failure": (
        "The transcript could not be analyzed because it either contains insufficient dialogue "
        "or does not represent a real estate conversation between an agent and a client."
    ),
    "cannot_determine": "Information cannot be determined from the transcript.",
    "viz_insufficient": "Visualization cannot be generated due to insufficient transcript data.",
}


# ════════════════════════════════════════════════════════
# VISUALIZATION LOGGER — JSON BACKUP FOR FRONTEND
# ════════════════════════════════════════════════════════

class VisualizationLogger:
    """Centralized visualization data storage system.
    
    Captures raw data from every chart/graph generated during evaluation
    and exports to a structured JSON file optimized for frontend libraries
    (Chart.js, D3.js, ECharts, Recharts).
    
    Usage:
        logger = VisualizationLogger(report_type="real_estate_sales_coaching")
        logger.log_chart(
            chart_id="confidence_score_chart",
            title="Agent Confidence Score",
            chart_type="bar",
            labels=["Opening", "Discovery"],
            datasets=[{"label": "Score", "values": [7, 8]}]
        )
        logger.save_to_json("visualizations_backup.json")
    """

    def __init__(self, report_type: str = "real_estate_sales_coaching"):
        self._visualizations: List[Dict[str, Any]] = []
        self._report_type = report_type
        self._created_at = datetime.now().isoformat()
        self._metadata: Dict[str, Any] = {}

    def set_metadata(self, **kwargs):
        """Set report-level metadata (agent name, call ID, duration, etc.)."""
        self._metadata.update(kwargs)

    def log_chart(
        self,
        chart_id: str,
        title: str,
        chart_type: str,
        labels: Optional[List[str]] = None,
        datasets: Optional[List[Dict[str, Any]]] = None,
        description: str = "",
        metadata: Optional[Dict[str, Any]] = None,
        options: Optional[Dict[str, Any]] = None,
    ):
        """Log a single visualization's raw data for frontend reconstruction.
        
        Args:
            chart_id: Unique identifier (e.g., 'talk_time_donut').
            title: Display title for the chart.
            chart_type: One of 'bar', 'horizontalBar', 'line', 'doughnut', 'radar',
                        'funnel', 'gauge', 'scatter', 'stackedBar', 'heatmap',
                        'quadrant', 'scoreCard', 'table'.
            labels: Category / axis labels.
            datasets: List of dataset dicts, each with 'label', 'values',
                      and optional 'colors', 'borderColors'.
            description: Human-readable description of what the chart shows.
            metadata: Chart-specific metadata (units, thresholds, etc.).
            options: Frontend rendering hints (y_range, stacked, etc.).
        """
        entry: Dict[str, Any] = {
            "id": chart_id,
            "title": title,
            "type": chart_type,
            "description": description,
            "labels": labels or [],
            "datasets": datasets or [],
            "metadata": metadata or {},
            "timestamp": datetime.now().isoformat(),
        }
        if options:
            entry["options"] = options
        self._visualizations.append(entry)

    def get_all(self) -> List[Dict[str, Any]]:
        """Return all logged visualizations."""
        return list(self._visualizations)

    def to_dict(self) -> Dict[str, Any]:
        """Build the complete JSON-serializable output dict."""
        return {
            "generated_at": self._created_at,
            "report_type": self._report_type,
            "schema_version": "2.0",
            "metadata": self._metadata,
            "total_visualizations": len(self._visualizations),
            "visualizations": self._visualizations,
        }

    def save_to_json(self, filepath: str):
        """Write the full visualization backup to a JSON file."""
        payload = self.to_dict()
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(payload, f, indent=2, ensure_ascii=False, default=str)
            logger.info(f"Saved visualization JSON: {filepath}")
        except Exception as e:
            logger.error(f"Failed to save visualization JSON: {e}")


# ════════════════════════════════════════════════════════
# CUSTOM FLOWABLES (from Enterprise v2.0)
# ════════════════════════════════════════════════════════

class ScoreBar(Flowable):
    """Horizontal color-coded score bar: green >= 75%, yellow >= 50%, red < 50%."""
    def __init__(self, label, score, max_score=10, width=400, height=22):
        super().__init__()
        self.label = label
        self.score = score
        self.max_score = max_score
        self.width = width
        self.height = height

    def wrap(self, availWidth, availHeight):
        return self.width, self.height + 4

    def draw(self):
        pct = self.score / self.max_score if self.max_score else 0
        bar_w = self.width * 0.55
        bar_x = self.width * 0.40
        # Background
        self.canv.setFillColor(colors.HexColor("#ECEFF1"))
        self.canv.roundRect(bar_x, 4, bar_w, self.height - 4, 3, fill=1, stroke=0)
        # Filled portion
        fill_color = self._get_color_from_pct(pct)
        self.canv.setFillColor(colors.HexColor(fill_color))
        self.canv.roundRect(bar_x, 4, bar_w * pct, self.height - 4, 3, fill=1, stroke=0)
        # Label
        self.canv.setFont("Helvetica", 8)
        self.canv.setFillColor(colors.HexColor("#424242"))
        self.canv.drawString(4, 8, self.label)
        # Score text
        self.canv.setFont("Helvetica-Bold", 9)
        self.canv.setFillColor(colors.HexColor(fill_color))
        self.canv.drawString(bar_x + bar_w + 6, 8, f"{self.score}/{self.max_score}")

    def _get_color_from_pct(self, pct):
        if pct >= 0.75:
            return "#43A047"
        elif pct >= 0.50:
            return "#FF8F00"
        else:
            return "#E53935"


class SectionDivider(Flowable):
    """Thin horizontal rule divider."""
    def __init__(self, width=460, color="#1565C0", thickness=1.5):
        super().__init__()
        self.width = width
        self.color = color
        self.thickness = thickness

    def wrap(self, availWidth, availHeight):
        return self.width, self.thickness + 6

    def draw(self):
        self.canv.setStrokeColor(colors.HexColor(self.color))
        self.canv.setLineWidth(self.thickness)
        self.canv.line(0, 3, self.width, 3)


class ShadedBox(Flowable):
    """Metric card with label + value in a shaded container."""
    def __init__(self, label, value, width=140, bg="#F5F7FA", border="#1565C0"):
        super().__init__()
        self.label = label
        self.value = value
        self.bw = width
        self.bg = bg
        self.border = border

    def wrap(self, availWidth, availHeight):
        return self.bw, 48

    def draw(self):
        self.canv.setFillColor(colors.HexColor(self.bg))
        self.canv.setStrokeColor(colors.HexColor(self.border))
        self.canv.roundRect(0, 0, self.bw, 44, 4, fill=1, stroke=1)
        self.canv.setFont("Helvetica-Bold", 14)
        self.canv.setFillColor(colors.HexColor(self.border))
        self.canv.drawCentredString(self.bw / 2, 24, str(self.value))
        self.canv.setFont("Helvetica", 7)
        self.canv.setFillColor(colors.HexColor("#757575"))
        self.canv.drawCentredString(self.bw / 2, 10, self.label)


class RadarChartFlowable(Flowable):
    """Spider/radar chart for multi-dimensional score visualization."""
    def __init__(self, scores_dict, width=280, height=280):
        super().__init__()
        self.scores = scores_dict
        self.chart_w = width
        self.chart_h = height

    def wrap(self, availWidth, availHeight):
        return self.chart_w, self.chart_h

    def draw(self):
        labels = list(self.scores.keys())
        values = list(self.scores.values())
        n = len(labels)
        if n < 3:
            return
        cx, cy = self.chart_w / 2, self.chart_h / 2
        r = min(cx, cy) - 30
        angles = [2 * math.pi * i / n - math.pi / 2 for i in range(n)]

        # Grid circles
        self.canv.setStrokeColor(colors.HexColor("#E0E0E0"))
        self.canv.setLineWidth(0.3)
        for frac in [0.25, 0.5, 0.75, 1.0]:
            self.canv.circle(cx, cy, r * frac, fill=0)

        # Axes
        for angle in angles:
            x2 = cx + r * math.cos(angle)
            y2 = cy + r * math.sin(angle)
            self.canv.line(cx, cy, x2, y2)

        # Data polygon
        max_val = 10
        pts = []
        for i, val in enumerate(values):
            frac = val / max_val
            x = cx + r * frac * math.cos(angles[i])
            y = cy + r * frac * math.sin(angles[i])
            pts.append((x, y))

        if pts:
            self.canv.setFillColor(colors.Color(0.08, 0.35, 0.69, alpha=0.15))
            self.canv.setStrokeColor(colors.HexColor("#1565C0"))
            self.canv.setLineWidth(1.5)
            path = self.canv.beginPath()
            path.moveTo(pts[0][0], pts[0][1])
            for px, py in pts[1:]:
                path.lineTo(px, py)
            path.close()
            self.canv.drawPath(path, fill=1, stroke=1)
            # Dots
            self.canv.setFillColor(colors.HexColor("#1565C0"))
            for px, py in pts:
                self.canv.circle(px, py, 3, fill=1, stroke=0)

        # Labels
        self.canv.setFont("Helvetica", 7)
        self.canv.setFillColor(colors.HexColor("#424242"))
        for i, label in enumerate(labels):
            lx = cx + (r + 24) * math.cos(angles[i])
            ly = cy + (r + 24) * math.sin(angles[i])
            short = label[:20] + "..." if len(label) > 20 else label
            self.canv.drawCentredString(lx, ly - 3, short)


class DonutChartFlowable(Flowable):
    """Donut chart showing percentage splits (e.g., talk-time distribution)."""
    def __init__(self, segments, width=220, height=220):
        """segments: list of (label, value, hex_color)"""
        super().__init__()
        self.segments = segments
        self.chart_w = width
        self.chart_h = height

    def wrap(self, availWidth, availHeight):
        return self.chart_w + 120, self.chart_h

    def draw(self):
        cx, cy = self.chart_w / 2, self.chart_h / 2
        r_outer = min(cx, cy) - 10
        r_inner = r_outer * 0.55
        total = sum(v for _, v, _ in self.segments) or 1
        start_angle = 90

        for label, value, hex_color in self.segments:
            sweep = (value / total) * 360
            self.canv.setFillColor(colors.HexColor(hex_color))
            self.canv.setStrokeColor(colors.white)
            self.canv.setLineWidth(2)
            # Draw wedge
            path = self.canv.beginPath()
            path.moveTo(cx, cy)
            path.arcTo(cx - r_outer, cy - r_outer, cx + r_outer, cy + r_outer,
                       start_angle, sweep)
            path.close()
            self.canv.drawPath(path, fill=1, stroke=1)
            start_angle += sweep

        # Inner circle (donut hole)
        self.canv.setFillColor(colors.white)
        self.canv.setStrokeColor(colors.white)
        self.canv.circle(cx, cy, r_inner, fill=1, stroke=0)

        # Center text
        self.canv.setFont("Helvetica-Bold", 9)
        self.canv.setFillColor(colors.HexColor("#424242"))
        self.canv.drawCentredString(cx, cy + 2, "Talk Time")
        self.canv.setFont("Helvetica", 7)
        self.canv.drawCentredString(cx, cy - 10, "Distribution")

        # Legend
        legend_x = self.chart_w + 10
        legend_y = self.chart_h - 30
        for label, value, hex_color in self.segments:
            pct = (value / total) * 100
            self.canv.setFillColor(colors.HexColor(hex_color))
            self.canv.rect(legend_x, legend_y, 10, 10, fill=1, stroke=0)
            self.canv.setFont("Helvetica", 7.5)
            self.canv.setFillColor(colors.HexColor("#424242"))
            self.canv.drawString(legend_x + 14, legend_y + 1, f"{label}: {pct:.0f}%")
            legend_y -= 16


class HBarChartFlowable(Flowable):
    """Horizontal bar chart with labels and color-coded bars."""
    def __init__(self, data, width=440, bar_height=18, spacing=6, max_val=10):
        """data: list of (label, value, hex_color)"""
        super().__init__()
        self.data = data
        self.chart_w = width
        self.bar_h = bar_height
        self.spacing = spacing
        self.max_val = max_val

    def wrap(self, availWidth, availHeight):
        h = len(self.data) * (self.bar_h + self.spacing) + 10
        return self.chart_w, h

    def draw(self):
        label_w = self.chart_w * 0.30
        bar_w = self.chart_w * 0.55
        bar_x = label_w + 10
        y = len(self.data) * (self.bar_h + self.spacing)

        for label, value, hex_color in self.data:
            y -= (self.bar_h + self.spacing)
            # Background bar
            self.canv.setFillColor(colors.HexColor("#ECEFF1"))
            self.canv.roundRect(bar_x, y, bar_w, self.bar_h, 3, fill=1, stroke=0)
            # Filled bar
            pct = min(value / self.max_val, 1.0) if self.max_val else 0
            self.canv.setFillColor(colors.HexColor(hex_color))
            self.canv.roundRect(bar_x, y, bar_w * pct, self.bar_h, 3, fill=1, stroke=0)
            # Label
            self.canv.setFont("Helvetica", 7.5)
            self.canv.setFillColor(colors.HexColor("#424242"))
            short = label[:28] + ".." if len(label) > 28 else label
            self.canv.drawRightString(label_w, y + 5, short)
            # Value
            self.canv.setFont("Helvetica-Bold", 8)
            self.canv.setFillColor(colors.HexColor(hex_color))
            self.canv.drawString(bar_x + bar_w + 6, y + 5, f"{value:.1f}")


class LineChartFlowable(Flowable):
    """Line chart for sentiment over time or skill trends."""
    def __init__(self, data_points, width=420, height=180, y_label="Score", x_labels=None,
                 line_color="#1565C0", fill_color=None, y_range=None):
        """data_points: list of numeric values. x_labels: list of str labels for x-axis."""
        super().__init__()
        self.data = data_points
        self.chart_w = width
        self.chart_h = height
        self.y_label = y_label
        self.x_labels = x_labels
        self.line_color = line_color
        self.fill_color = fill_color
        self.y_range = y_range

    def wrap(self, availWidth, availHeight):
        return self.chart_w, self.chart_h + 30

    def draw(self):
        if not self.data or len(self.data) < 2:
            return
        margin_l, margin_b, margin_t, margin_r = 40, 28, 10, 10
        plot_w = self.chart_w - margin_l - margin_r
        plot_h = self.chart_h - margin_b - margin_t
        ox, oy = margin_l, margin_b

        y_min = self.y_range[0] if self.y_range else min(self.data)
        y_max = self.y_range[1] if self.y_range else max(self.data)
        if y_max == y_min:
            y_max = y_min + 1

        # Grid lines
        self.canv.setStrokeColor(colors.HexColor("#E0E0E0"))
        self.canv.setLineWidth(0.3)
        for i in range(5):
            gy = oy + plot_h * i / 4
            self.canv.line(ox, gy, ox + plot_w, gy)
            val = y_min + (y_max - y_min) * i / 4
            self.canv.setFont("Helvetica", 6)
            self.canv.setFillColor(colors.HexColor("#9E9E9E"))
            self.canv.drawRightString(ox - 4, gy - 3, f"{val:.1f}")

        # Axes
        self.canv.setStrokeColor(colors.HexColor("#BDBDBD"))
        self.canv.setLineWidth(0.5)
        self.canv.line(ox, oy, ox + plot_w, oy)
        self.canv.line(ox, oy, ox, oy + plot_h)

        # Plot points
        n = len(self.data)
        pts = []
        for i, val in enumerate(self.data):
            x = ox + plot_w * i / (n - 1)
            y = oy + plot_h * (val - y_min) / (y_max - y_min)
            pts.append((x, y))

        # Fill area under curve
        if self.fill_color and pts:
            path = self.canv.beginPath()
            path.moveTo(pts[0][0], oy)
            for px, py in pts:
                path.lineTo(px, py)
            path.lineTo(pts[-1][0], oy)
            path.close()
            self.canv.setFillColor(colors.HexColor(self.fill_color))
            self.canv.drawPath(path, fill=1, stroke=0)

        # Line
        self.canv.setStrokeColor(colors.HexColor(self.line_color))
        self.canv.setLineWidth(1.5)
        path = self.canv.beginPath()
        path.moveTo(pts[0][0], pts[0][1])
        for px, py in pts[1:]:
            path.lineTo(px, py)
        self.canv.drawPath(path, fill=0, stroke=1)

        # Dots
        self.canv.setFillColor(colors.HexColor(self.line_color))
        for px, py in pts:
            self.canv.circle(px, py, 2.5, fill=1, stroke=0)

        # X labels
        if self.x_labels:
            self.canv.setFont("Helvetica", 5.5)
            self.canv.setFillColor(colors.HexColor("#757575"))
            for i, lbl in enumerate(self.x_labels[:n]):
                x = ox + plot_w * i / (n - 1)
                short = lbl[:10]
                self.canv.drawCentredString(x, oy - 12, short)

        # Y axis label
        self.canv.setFont("Helvetica", 6.5)
        self.canv.setFillColor(colors.HexColor("#757575"))
        self.canv.saveState()
        self.canv.translate(8, oy + plot_h / 2)
        self.canv.rotate(90)
        self.canv.drawCentredString(0, 0, self.y_label)
        self.canv.restoreState()


class HeatmapFlowable(Flowable):
    """Heatmap matrix: rows x columns with color-coded cells."""
    def __init__(self, row_labels, col_labels, data_matrix, width=420, cell_h=22):
        """data_matrix[row][col] = value (0-10 or 0-1). Colors interpolated."""
        super().__init__()
        self.rows = row_labels
        self.cols = col_labels
        self.data = data_matrix
        self.chart_w = width
        self.cell_h = cell_h

    def wrap(self, availWidth, availHeight):
        h = (len(self.rows) + 1) * self.cell_h + 10
        return self.chart_w, h

    def draw(self):
        n_rows = len(self.rows)
        n_cols = len(self.cols)
        label_w = 90
        cell_w = (self.chart_w - label_w) / max(n_cols, 1)
        y_start = n_rows * self.cell_h

        # Column headers
        self.canv.setFont("Helvetica-Bold", 6.5)
        self.canv.setFillColor(colors.HexColor("#424242"))
        for j, col in enumerate(self.cols):
            x = label_w + j * cell_w + cell_w / 2
            short = col[:12]
            self.canv.drawCentredString(x, y_start + 4, short)

        # Rows
        for i, row_label in enumerate(self.rows):
            y = y_start - (i + 1) * self.cell_h
            # Row label
            self.canv.setFont("Helvetica", 7)
            self.canv.setFillColor(colors.HexColor("#424242"))
            self.canv.drawRightString(label_w - 6, y + self.cell_h / 2 - 3, row_label[:14])
            # Cells
            for j in range(n_cols):
                val = self.data[i][j] if i < len(self.data) and j < len(self.data[i]) else 0
                # Color: red(0) -> yellow(5) -> green(10)
                if val >= 7:
                    hex_c = "#43A047"
                elif val >= 4:
                    hex_c = "#FF8F00"
                else:
                    hex_c = "#E53935"
                alpha = max(0.2, min(1.0, val / 10))
                x = label_w + j * cell_w
                self.canv.setFillColor(colors.HexColor(hex_c))
                self.canv.setStrokeColor(colors.white)
                self.canv.setLineWidth(1)
                self.canv.rect(x, y, cell_w, self.cell_h, fill=1, stroke=1)
                # Value text
                self.canv.setFont("Helvetica-Bold", 7)
                self.canv.setFillColor(colors.white)
                display = f"{val:.0f}" if val == int(val) else f"{val:.1f}"
                self.canv.drawCentredString(x + cell_w / 2, y + self.cell_h / 2 - 3, display)


class FunnelFlowable(Flowable):
    """Funnel visualization for stage-wise drop-off."""
    def __init__(self, stages, width=400, height=220):
        """stages: list of (label, score_0_10, hex_color)"""
        super().__init__()
        self.stages = stages
        self.chart_w = width
        self.chart_h = height

    def wrap(self, availWidth, availHeight):
        return self.chart_w, self.chart_h + 20

    def draw(self):
        n = len(self.stages)
        if not n:
            return
        step_h = self.chart_h / n
        max_w = self.chart_w * 0.8
        cx = self.chart_w / 2
        shrink = 0.15 / max(n, 1) * 3  # Adaptive shrink based on # of stages

        for i, (label, score, hex_color) in enumerate(self.stages):
            pct = score / 10
            top_w = max_w * (1 - i * shrink)
            bot_w = max_w * (1 - (i + 1) * shrink)
            y_top = self.chart_h - i * step_h
            y_bot = self.chart_h - (i + 1) * step_h
            gap = 2  # Small gap between stages

            # Trapezoid with rounded appearance
            base_color = colors.HexColor(hex_color)
            # Slightly lighter version for background
            self.canv.setFillColor(base_color)
            self.canv.setStrokeColor(colors.white)
            self.canv.setLineWidth(2)
            path = self.canv.beginPath()
            path.moveTo(cx - top_w / 2, y_top - gap)
            path.lineTo(cx + top_w / 2, y_top - gap)
            path.lineTo(cx + bot_w / 2, y_bot + gap)
            path.lineTo(cx - bot_w / 2, y_bot + gap)
            path.close()
            self.canv.drawPath(path, fill=1, stroke=1)

            # Label and score
            mid_y = (y_top + y_bot) / 2
            self.canv.setFont("Helvetica-Bold", 9)
            self.canv.setFillColor(colors.white)
            self.canv.drawCentredString(cx, mid_y + 3, f"{label}")
            # Score bar indicator
            bar_w = min(top_w, bot_w) * 0.5
            bar_h = 4
            bar_y = mid_y - 10
            # Background bar
            self.canv.setFillColor(colors.Color(1, 1, 1, 0.3))
            self.canv.rect(cx - bar_w / 2, bar_y, bar_w, bar_h, fill=1, stroke=0)
            # Filled bar
            self.canv.setFillColor(colors.Color(1, 1, 1, 0.8))
            self.canv.rect(cx - bar_w / 2, bar_y, bar_w * pct, bar_h, fill=1, stroke=0)


class GaugeFlowable(Flowable):
    """Half-circle gauge / speedometer for deal momentum."""
    def __init__(self, value, max_val=100, label="Deal Momentum", width=200, height=140):
        super().__init__()
        self.value = value
        self.max_val = max_val
        self.label = label
        self.gauge_w = width
        self.gauge_h = height

    def wrap(self, availWidth, availHeight):
        # Total height = arc area + generous space for value text + label text below
        return self.gauge_w, self.gauge_h + 55

    def draw(self):
        total_h = self.gauge_h + 55
        cx = self.gauge_w / 2
        # Position arc center so everything fits: leave 50px for text below
        r = min(self.gauge_w / 2 - 25, self.gauge_h - 20)
        cy = 52  # Base of the arc — leaves room for value + label below

        # Background arc segments (red/yellow/green) — thick for visibility
        segments = [
            (120, 60, "#E53935"),   # left zone (0-33%) – Hot / red
            (60,  60, "#FF8F00"),   # middle zone (33-67%) – Warm / orange
            (0,   60, "#43A047"),   # right zone (67-100%) – Cold / green
        ]
        for start, sweep, hex_c in segments:
            self.canv.setStrokeColor(colors.HexColor(hex_c))
            self.canv.setLineWidth(18)
            self.canv.arc(cx - r, cy - r, cx + r, cy + r, start, sweep)

        # White separator arcs between segments (thinner than arc to create gap)
        for start_angle in [120, 60]:
            self.canv.setStrokeColor(colors.HexColor("#FFFFFF"))
            self.canv.setLineWidth(20)
            self.canv.arc(cx - r, cy - r, cx + r, cy + r, start_angle, 1.5)

        # Needle
        pct = min(self.value / self.max_val, 1.0) if self.max_val else 0
        angle_deg = 180 - pct * 180
        angle_rad = math.radians(angle_deg)
        needle_len = r - 24
        nx = cx + needle_len * math.cos(angle_rad)
        ny = cy + needle_len * math.sin(angle_rad)
        # Shadow
        self.canv.setStrokeColor(colors.HexColor("#BDBDBD"))
        self.canv.setLineWidth(3)
        self.canv.line(cx + 1, cy - 1, nx + 1, ny - 1)
        # Needle
        self.canv.setStrokeColor(colors.HexColor("#212121"))
        self.canv.setLineWidth(2.5)
        self.canv.line(cx, cy, nx, ny)
        # Center dot
        self.canv.setFillColor(colors.HexColor("#212121"))
        self.canv.circle(cx, cy, 6, fill=1, stroke=0)
        self.canv.setFillColor(colors.HexColor("#FFFFFF"))
        self.canv.circle(cx, cy, 2.5, fill=1, stroke=0)

        # Value text — placed BELOW the arc center with proper spacing
        val_color = "#43A047" if pct >= 0.67 else "#FF8F00" if pct >= 0.33 else "#E53935"
        self.canv.setFont("Helvetica-Bold", 16)
        self.canv.setFillColor(colors.HexColor(val_color))
        self.canv.drawCentredString(cx, cy - 20, f"{self.value}%")

        # Label — below value with extra spacing
        self.canv.setFont("Helvetica", 8)
        self.canv.setFillColor(colors.HexColor("#757575"))
        self.canv.drawCentredString(cx, cy - 36, self.label)

        # Zone labels at arc edges with proper positioning
        self.canv.setFont("Helvetica-Bold", 7.5)
        self.canv.setFillColor(colors.HexColor("#E53935"))
        self.canv.drawString(cx - r - 8, cy + 8, "Hot")
        self.canv.setFillColor(colors.HexColor("#FF8F00"))
        self.canv.drawCentredString(cx, cy + r + 10, "Warm")
        self.canv.setFillColor(colors.HexColor("#43A047"))
        self.canv.drawRightString(cx + r + 8, cy + 8, "Cold")


class QuadrantFlowable(Flowable):
    """2D quadrant chart (e.g., Trust vs Pressure)."""
    def __init__(self, points, width=300, height=300,
                 x_label="Trust-Building →", y_label="Pressure →"):
        """points: list of (label, x_val, y_val) where 0-10 scale."""
        super().__init__()
        self.points = points
        self.chart_w = width
        self.chart_h = height
        self.x_label = x_label
        self.y_label = y_label

    def wrap(self, availWidth, availHeight):
        return self.chart_w + 20, self.chart_h + 20

    def draw(self):
        margin = 35
        pw = self.chart_w - 2 * margin
        ph = self.chart_h - 2 * margin
        ox, oy = margin, margin

        # Quadrant backgrounds
        half_w, half_h = pw / 2, ph / 2
        quads = [
            (ox, oy + half_h, half_w, half_h, "#FFEBEE", "High Pressure\nLow Trust"),
            (ox + half_w, oy + half_h, half_w, half_h, "#FFF8E1", "High Pressure\nHigh Trust"),
            (ox, oy, half_w, half_h, "#E3F2FD", "Low Pressure\nLow Trust"),
            (ox + half_w, oy, half_w, half_h, "#E8F5E9", "Low Pressure\nHigh Trust"),
        ]
        for qx, qy, qw, qh, bg, qlabel in quads:
            self.canv.setFillColor(colors.HexColor(bg))
            self.canv.rect(qx, qy, qw, qh, fill=1, stroke=0)
            self.canv.setFont("Helvetica", 5.5)
            self.canv.setFillColor(colors.HexColor("#BDBDBD"))
            lines = qlabel.split('\n')
            for li, line in enumerate(lines):
                self.canv.drawCentredString(qx + qw / 2, qy + qh / 2 - li * 8 + 4, line)

        # Grid
        self.canv.setStrokeColor(colors.HexColor("#BDBDBD"))
        self.canv.setLineWidth(0.5)
        self.canv.setDash(2, 2)
        self.canv.line(ox + half_w, oy, ox + half_w, oy + ph)
        self.canv.line(ox, oy + half_h, ox + pw, oy + half_h)
        self.canv.setDash()

        # Border
        self.canv.setStrokeColor(colors.HexColor("#9E9E9E"))
        self.canv.setLineWidth(0.5)
        self.canv.rect(ox, oy, pw, ph, fill=0, stroke=1)

        # Points
        for label, xv, yv in self.points:
            px = ox + pw * min(xv / 10, 1.0)
            py = oy + ph * min(yv / 10, 1.0)
            self.canv.setFillColor(colors.HexColor("#1565C0"))
            self.canv.circle(px, py, 5, fill=1, stroke=0)
            self.canv.setFont("Helvetica-Bold", 6)
            self.canv.setFillColor(colors.HexColor("#0D47A1"))
            self.canv.drawString(px + 7, py - 2, label[:20])

        # Ideal zone marker
        ideal_x = ox + pw * 0.75
        ideal_y = oy + ph * 0.25
        self.canv.setStrokeColor(colors.HexColor("#43A047"))
        self.canv.setLineWidth(1)
        self.canv.setDash(3, 3)
        self.canv.circle(ideal_x, ideal_y, 18, fill=0, stroke=1)
        self.canv.setDash()
        self.canv.setFont("Helvetica", 5)
        self.canv.setFillColor(colors.HexColor("#2E7D32"))
        self.canv.drawCentredString(ideal_x, ideal_y - 24, "IDEAL ZONE")

        # Axis labels
        self.canv.setFont("Helvetica-Bold", 7)
        self.canv.setFillColor(colors.HexColor("#424242"))
        self.canv.drawCentredString(ox + pw / 2, oy - 18, self.x_label)
        self.canv.saveState()
        self.canv.translate(ox - 20, oy + ph / 2)
        self.canv.rotate(90)
        self.canv.drawCentredString(0, 0, self.y_label)
        self.canv.restoreState()


class StackedBarFlowable(Flowable):
    """Stacked horizontal bar chart for question quality distribution etc."""
    def __init__(self, categories, width=420, bar_height=24):
        """categories: list of (label, [(segment_label, value, hex_color), ...])"""
        super().__init__()
        self.categories = categories
        self.chart_w = width
        self.bar_h = bar_height

    def wrap(self, availWidth, availHeight):
        h = len(self.categories) * (self.bar_h + 12) + 30
        return self.chart_w, h

    def draw(self):
        label_w = 100
        bar_w = self.chart_w - label_w - 20
        y = len(self.categories) * (self.bar_h + 12)

        # Legend at top
        legend_x = label_w
        all_segments = set()
        for _, segs in self.categories:
            for sl, _, sc in segs:
                all_segments.add((sl, sc))
        self.canv.setFont("Helvetica", 6)
        lx = legend_x
        for sl, sc in sorted(all_segments):
            self.canv.setFillColor(colors.HexColor(sc))
            self.canv.rect(lx, y + 4, 8, 8, fill=1, stroke=0)
            self.canv.setFillColor(colors.HexColor("#424242"))
            self.canv.drawString(lx + 10, y + 4, sl)
            lx += len(sl) * 5 + 24

        for cat_label, segments in self.categories:
            y -= (self.bar_h + 12)
            total = sum(v for _, v, _ in segments) or 1

            # Label
            self.canv.setFont("Helvetica", 7.5)
            self.canv.setFillColor(colors.HexColor("#424242"))
            self.canv.drawRightString(label_w - 6, y + self.bar_h / 2 - 3, cat_label[:16])

            # Stacked bar
            x = label_w
            for seg_label, value, hex_color in segments:
                w = bar_w * (value / total)
                self.canv.setFillColor(colors.HexColor(hex_color))
                self.canv.setStrokeColor(colors.white)
                self.canv.setLineWidth(0.5)
                self.canv.rect(x, y, w, self.bar_h, fill=1, stroke=1)
                # Value inside bar
                if w > 25:
                    self.canv.setFont("Helvetica-Bold", 6)
                    self.canv.setFillColor(colors.white)
                    pct = (value / total) * 100
                    self.canv.drawCentredString(x + w / 2, y + self.bar_h / 2 - 3, f"{pct:.0f}%")
                x += w


# ════════════════════════════════════════════════════════
# MAIN SUMMARIZER CLASS
# ════════════════════════════════════════════════════════

class RealEstateSalesMeetingSummarizer:
    """Combined best-of-both RAG coaching evaluation engine."""

    # Configuration constants for consistent thresholds
    CONFIG = {
        'score_thresholds': {
            'red_max': 4.9,
            'yellow_min': 5.0,
            'yellow_max': 7.4,
            'green_min': 7.5,
        },
        'percentage_thresholds': {
            'red_max': 49,
            'yellow_min': 50,
            'yellow_max': 74,
            'green_min': 75,
        },
        'talk_ratio_ideal': (40, 55),
        'gauge_zones': {
            'red': (0, 33),
            'yellow': (34, 66),
            'green': (67, 100),
        },
    }

    def __init__(self, transcript_file: str = "conversation.json"):
        # Resolve all paths relative to this script's directory (Backend/)
        self._script_dir = os.path.dirname(os.path.abspath(__file__))
        env_path = os.path.join(self._script_dir, '.env')
        load_dotenv(dotenv_path=env_path)

        self.transcript_file = transcript_file
        self.config = configparser.ConfigParser()
        config_path = os.path.join(self._script_dir, 'meeting_config.ini')
        if os.path.exists(config_path):
            self.config.read(config_path)
        self.summaries_folder = os.path.join(self._script_dir, 'Rept Meeting')
        os.makedirs(self.summaries_folder, exist_ok=True)

        self._load_environment_variables()
        self._setup_llm_config()
        self._setup_professional_styles()
        self._setup_rag_system()
        self.viz_logger = VisualizationLogger(report_type="real_estate_sales_coaching")

    # Utility methods for consistent color coding
    def _get_score_color(self, value: float, max_val: float = 10.0) -> str:
        """Return hex color based on score thresholds (red <5, yellow 5-7.4, green >=7.5)."""
        pct = value / max_val if max_val else 0
        if pct >= 0.75:
            return "#43A047"  # green
        elif pct >= 0.50:
            return "#FF8F00"  # orange
        else:
            return "#E53935"  # red

    def _get_percentage_color(self, percentage: float) -> str:
        """Return hex color based on percentage thresholds (red <50, yellow 50-74, green >=75)."""
        if percentage >= 75:
            return "#43A047"
        elif percentage >= 50:
            return "#FF8F00"
        else:
            return "#E53935"

    def _get_zone_from_gauge_value(self, value: float) -> Tuple[str, str]:
        """Return zone name and color based on gauge value (0-100)."""
        if value >= 67:
            return "HOT", "#43A047"
        elif value >= 34:
            return "WARM", "#FF8F00"
        else:
            return "COLD", "#E53935"

    def _get_color_for_status(self, status: str) -> str:
        """Map status strings to colors."""
        status_lower = status.lower()
        if 'confirmed' in status_lower or 'strong' in status_lower or 'high' in status_lower:
            return "#43A047"
        elif 'partially' in status_lower or 'moderate' in status_lower or 'medium' in status_lower:
            return "#FF8F00"
        elif 'not' in status_lower or 'weak' in status_lower or 'low' in status_lower or 'missed' in status_lower:
            return "#E53935"
        else:
            return "#757575"

    def _add_color_legend(self, story, thresholds_type='score'):
        """Add a small legend explaining the color coding used in charts."""
        if thresholds_type == 'score':
            legend_text = (
                "<b>Color legend:</b> "
                "<font color='#43A047'>■ Green = strong (≥7.5/10)</font>, "
                "<font color='#FF8F00'>■ Yellow = moderate (5.0-7.4/10)</font>, "
                "<font color='#E53935'>■ Red = needs improvement (<5.0/10)</font>"
            )
        else:  # percentage
            legend_text = (
                "<b>Color legend:</b> "
                "<font color='#43A047'>■ Green = strong (≥75%)</font>, "
                "<font color='#FF8F00'>■ Yellow = moderate (50-74%)</font>, "
                "<font color='#E53935'>■ Red = needs improvement (<50%)</font>"
            )
        story.append(Paragraph(legend_text, ParagraphStyle('legend', parent=self.ps['body'], fontSize=7.5, textColor=self.C['light'])))
        story.append(Spacer(1, 6))

    # ── Environment & Config ──

    def _load_environment_variables(self):
        self.llm_provider = os.getenv("LLM_PROVIDER").strip().lower()
        self.llm_api_key = os.getenv("LLM_API_KEY")
        self.llm_model = os.getenv("LLM_MODEL")
        self.llm_base_url = os.getenv("LLM_BASE_URL")
        self.llm_temperature = float(os.getenv("LLM_TEMPERATURE"))
        self.llm_max_tokens = int(os.getenv("LLM_MAX_TOKENS"))
        self.llm_timeout = int(os.getenv("LLM_TIMEOUT"))

    def _setup_llm_config(self):
        self.llm_config = LLMConfig(
            provider=self.llm_provider,
            api_key=self.llm_api_key,
            model=self.llm_model,
            base_url=self.llm_base_url,
            temperature=self.llm_temperature,
            max_tokens=self.llm_max_tokens,
            timeout=self.llm_timeout,
        )

    # ── Professional Styles ──

    def _setup_professional_styles(self):
        """Enterprise color palette + compact paragraph style registry."""
        self.C = {
            'primary': colors.HexColor("#1565C0"),
            'primary_dark': colors.HexColor("#0D47A1"),
            'primary_light': colors.HexColor("#42A5F5"),
            'accent': colors.HexColor("#FF8F00"),
            'success': colors.HexColor("#2E7D32"),
            'success_bg': colors.HexColor("#E8F5E9"),
            'danger': colors.HexColor("#C62828"),
            'danger_bg': colors.HexColor("#FFEBEE"),
            'warning': colors.HexColor("#F57F17"),
            'warning_bg': colors.HexColor("#FFF8E1"),
            'dark': colors.HexColor("#263238"),
            'medium': colors.HexColor("#546E7A"),
            'light': colors.HexColor("#78909C"),
            'bg': colors.HexColor("#F5F7FA"),
            'white': colors.white,
            'border': colors.HexColor("#CFD8DC"),
        }

        base = getSampleStyleSheet()

        def ps(name, **kw):
            return ParagraphStyle(name, parent=base['Normal'], **kw)

        self.ps = {
            'cover_title': ps('cover_title', fontName='Helvetica-Bold', fontSize=26,
                              textColor=self.C['primary_dark'], alignment=TA_CENTER, spaceAfter=50, leading=35),
            'cover_sub': ps('cover_sub', fontName='Helvetica', fontSize=12,
                            textColor=self.C['medium'], alignment=TA_CENTER, spaceAfter=30),
            'sec': ps('sec', fontName='Helvetica-Bold', fontSize=14,
                       textColor=self.C['primary_dark'], spaceBefore=14, spaceAfter=6),
            'subsec': ps('subsec', fontName='Helvetica-Bold', fontSize=11,
                          textColor=self.C['primary'], spaceBefore=8, spaceAfter=4),
            'body': ps('body', fontName='Helvetica', fontSize=9.5,
                        textColor=self.C['dark'], leading=13, spaceAfter=3),
            'body_bold': ps('body_bold', fontName='Helvetica-Bold', fontSize=9.5,
                             textColor=self.C['dark'], leading=13, spaceAfter=3),
            'bullet': ps('bullet', fontName='Helvetica', fontSize=9.5,
                          textColor=self.C['dark'], leftIndent=18, leading=13, spaceAfter=2),
            'bullet_success': ps('bullet_success', fontName='Helvetica', fontSize=9.5,
                                  textColor=self.C['success'], leftIndent=18, leading=13, spaceAfter=2),
            'bullet_danger': ps('bullet_danger', fontName='Helvetica', fontSize=9.5,
                                 textColor=self.C['danger'], leftIndent=18, leading=13, spaceAfter=2),
            'bullet_accent': ps('bullet_accent', fontName='Helvetica', fontSize=9.5,
                                 textColor=self.C['accent'], leftIndent=18, leading=13, spaceAfter=2),
            'th': ps('th', fontName='Helvetica-Bold', fontSize=9,
                      textColor=self.C['white'], alignment=TA_CENTER),
            'td': ps('td', fontName='Helvetica', fontSize=9,
                      textColor=self.C['dark'], leading=12),
            'td_c': ps('td_c', fontName='Helvetica', fontSize=9,
                        textColor=self.C['dark'], alignment=TA_CENTER, leading=12),
            'ts': ps('ts', fontName='Helvetica-Bold', fontSize=8.5,
                      textColor=self.C['primary'], spaceAfter=2),
            'quote': ps('quote', fontName='Helvetica-Oblique', fontSize=9,
                         textColor=self.C['medium'], leftIndent=18, spaceAfter=3),
            'score_big': ps('score_big', fontName='Helvetica-Bold', fontSize=28,
                             textColor=self.C['primary_dark'], alignment=TA_CENTER),
            'score_label': ps('score_label', fontName='Helvetica', fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER),
            'tier': ps('tier', fontName='Helvetica-Bold', fontSize=16,
                        textColor=self.C['primary_dark'], alignment=TA_CENTER),
            'conf': ps('conf', fontName='Helvetica', fontSize=7,
                        textColor=self.C['light'], alignment=TA_CENTER),
        }

    # ── RAG System (from Final_backup.py — robust implementation) ──

    def _setup_rag_system(self):
        kb_dir = os.path.join(self._script_dir, "RealState Knowledge Base")
        self.knowledge_base_path = os.path.join(kb_dir, "knowledge_base.docx")
        self.faiss_index_path = os.path.join(kb_dir, "knowledge_base.faiss")
        self.chunks_path = os.path.join(kb_dir, "knowledge_base_chunks.pkl")
        self.embedding_model_name = "all-MiniLM-L6-v2"
        self.embedding_model = None
        self.faiss_index = None
        self.knowledge_chunks: List[KnowledgeChunk] = []

        if not EMBEDDINGS_AVAILABLE or not FAISS_AVAILABLE:
            print("RAG components not available. System will work without knowledge base.")
            return

        try:
            logger.info(f"Loading embedding model: {self.embedding_model_name}...")
            self.embedding_model = SentenceTransformer(self.embedding_model_name)
        except Exception as e:
            logger.error(f"Failed to load embedding model: {e}")
            return

        # Try loading existing index (check both possible pkl filenames)
        if os.path.exists(self.faiss_index_path):
            pkl_path = None
            for candidate in [self.chunks_path, os.path.join(self._script_dir, "knowledge_chunks.pkl")]:
                if os.path.exists(candidate):
                    pkl_path = candidate
                    break
            if pkl_path:
                try:
                    self.faiss_index = faiss.read_index(self.faiss_index_path)
                    with open(pkl_path, 'rb') as f:
                        self.knowledge_chunks = pickle.load(f)
                    logger.info(f"Loaded {len(self.knowledge_chunks)} knowledge chunks from existing index.")
                    return
                except Exception as e:
                    logger.error(f"Failed to load existing index: {e}. Will rebuild...")

        # Build new index from docx
        if not os.path.exists(self.knowledge_base_path) or not DOCX_AVAILABLE:
            logger.warning(f"Knowledge base not found: {self.knowledge_base_path}")
            return

        logger.info(f"Building FAISS index from {self.knowledge_base_path}...")
        self._build_faiss_index()

    def _load_knowledge_base_text(self) -> str:
        """Load knowledge base from .docx — extracts paragraphs AND tables."""
        try:
            doc = Document(self.knowledge_base_path)
            full_text = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            full_text.append(text)
            content = "\n".join(full_text)
            logger.info(f"Loaded {len(content)} characters from knowledge base.")
            return content
        except Exception as e:
            logger.error(f"Error loading knowledge base: {e}")
            return ""

    def _semantic_chunk(self, text: str, chunk_size: int = 500, overlap: int = 100) -> List[str]:
        """Chunk text semantically with overlap for better retrieval."""
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        chunks = []
        current_chunk = []
        current_length = 0

        for para in paragraphs:
            para_length = len(para)
            if para_length > chunk_size:
                sentences = re.split(r'(?<=[.!?])\s+', para)
                for sent in sentences:
                    if current_length + len(sent) > chunk_size and current_chunk:
                        chunks.append(' '.join(current_chunk))
                        overlap_text = ' '.join(current_chunk[-2:]) if len(current_chunk) > 1 else (current_chunk[-1] if current_chunk else "")
                        current_chunk = [overlap_text, sent] if overlap_text else [sent]
                        current_length = len(overlap_text) + len(sent)
                    else:
                        current_chunk.append(sent)
                        current_length += len(sent)
            else:
                if current_length + para_length > chunk_size and current_chunk:
                    chunks.append(' '.join(current_chunk))
                    overlap_text = ' '.join(current_chunk[-2:]) if len(current_chunk) > 1 else (current_chunk[-1] if current_chunk else "")
                    current_chunk = [overlap_text, para] if overlap_text else [para]
                    current_length = len(overlap_text) + para_length
                else:
                    current_chunk.append(para)
                    current_length += para_length

        if current_chunk:
            chunks.append(' '.join(current_chunk))
        return chunks

    def _build_faiss_index(self):
        """Build FAISS index with semantic chunking."""
        kb_content = self._load_knowledge_base_text()
        if not kb_content:
            return

        text_chunks = self._semantic_chunk(kb_content, chunk_size=500, overlap=100)
        logger.info(f"Created {len(text_chunks)} semantic chunks.")

        embeddings = self.embedding_model.encode(text_chunks, show_progress_bar=False, convert_to_numpy=True)
        dimension = embeddings.shape[1]
        index = faiss.IndexFlatL2(dimension)
        index.add(embeddings.astype('float32'))

        self.faiss_index = index
        self.knowledge_chunks = [
            KnowledgeChunk(chunk_id=i, text=chunk, embedding=emb.tolist())
            for i, (chunk, emb) in enumerate(zip(text_chunks, embeddings))
        ]

        faiss.write_index(index, self.faiss_index_path)
        with open(self.chunks_path, 'wb') as f:
            pickle.dump(self.knowledge_chunks, f)
        logger.info(f"FAISS index built with {len(self.knowledge_chunks)} chunks.")

    def _retrieve_relevant_knowledge(self, query: str, top_k: int = 8) -> List[KnowledgeChunk]:
        """Retrieve most relevant knowledge chunks for a query."""
        if not self.faiss_index or not self.knowledge_chunks or not self.embedding_model:
            return []
        query_embedding = self.embedding_model.encode([query], convert_to_numpy=True)
        distances, indices = self.faiss_index.search(query_embedding.astype('float32'), top_k)
        retrieved = []
        for idx in indices[0]:
            if 0 <= idx < len(self.knowledge_chunks):
                retrieved.append(self.knowledge_chunks[idx])
        return retrieved

    # ── Transcript Loading ──

    def _load_transcript(self) -> Optional[TranscriptData]:
        path = Path(self.transcript_file)
        if not path.exists():
            logger.error(f"Transcript file not found: {self.transcript_file}")
            return None

        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        segments = []
        speakers = set()
        total_dur = 0.0

        for i, entry in enumerate(data):
            speaker = entry.get("speaker", "Unknown")
            speakers.add(speaker)
            seg = TranscriptSegment(
                speaker_id=speaker, speaker_name=speaker,
                transcript=entry.get("transcript", ""),
                segment_id=i,
                start=entry.get("start", 0.0),
                end=entry.get("end", 0.0),
                duration=entry.get("duration", 0.0),
            )
            segments.append(seg)
            total_dur += seg.duration

        return TranscriptData(
            transcripts=segments,
            agenda="Real Estate Listing Presentation & Coaching Evaluation",
            summary_info={
                "total_segments": len(segments),
                "unique_speakers": list(speakers),
                "total_duration_seconds": total_dur,
            },
        )

    def _format_transcript_for_analysis(self, data: TranscriptData) -> str:
        lines = []
        for seg in data.transcripts:
            start_m, start_s = int(seg.start // 60), int(seg.start % 60)
            end_m, end_s = int(seg.end // 60), int(seg.end % 60)
            ts = f"[{start_m}:{start_s:02d}-{end_m}:{end_s:02d}]"
            lines.append(f"{ts} {seg.speaker_name}: {seg.transcript}")
        return "\n".join(lines)

    @staticmethod
    def _convert_timestamps_to_mmss(text: str) -> str:
        """Convert any [X.Xs-Y.Ys] or [Xs-Ys] timestamp patterns to [M:SS-M:SS] format."""
        def _seconds_to_mmss(match):
            try:
                start_raw = match.group(1)
                end_raw = match.group(2)
                start_sec = float(start_raw.rstrip('s'))
                end_sec = float(end_raw.rstrip('s'))
                s_m, s_s = int(start_sec // 60), int(start_sec % 60)
                e_m, e_s = int(end_sec // 60), int(end_sec % 60)
                return f"[{s_m}:{s_s:02d}-{e_m}:{e_s:02d}]"
            except (ValueError, AttributeError):
                return match.group(0)
        # Match patterns like [123.5s-145.0s], [0.0s-21.5s], [123.5-145.0], [0.0s - 21.5s]
        text = re.sub(r'\[(\d+\.?\d*s?)\s*-\s*(\d+\.?\d*s?)\]', _seconds_to_mmss, text)
        return text

    # ── LLM Call ──

    def _call_llm(self, prompt: str) -> Optional[str]:
        url = f"{self.llm_config.base_url}/chat/completions"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.llm_config.api_key}",
        }
        payload = {
            "model": self.llm_config.model,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": self.llm_config.temperature,
            "max_tokens": self.llm_config.max_tokens,
        }
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=self.llm_config.timeout)
            resp.raise_for_status()
            return resp.json()["choices"][0]["message"]["content"]
        except Exception as e:
            logger.error(f"LLM call failed: {e}")
            traceback.print_exc()
            return None

    # ════════════════════════════════════════════════════════
    # UNIFIED LLM PROMPT — BRUTALLY HONEST EVALUATION
    # ════════════════════════════════════════════════════════

    def _generate_unified_llm_analysis(self, transcript_text: str, knowledge_context: str) -> Optional[str]:
        prompt = f"""You are a ruthless senior sales performance auditor and executive coach with 25+ years of real estate sales coaching. You do NOT coddle agents. Your job is to protect BOTH the client's financial interests AND the agent's career by delivering unvarnished truth. A mediocre agent who thinks they're good will lose deals and damage clients. Your honesty is an act of professional respect.

YOUR MANDATE: Deliver a BRUTALLY HONEST, forensic-level performance audit. If the agent performed well, acknowledge it with specific evidence — but if they failed, NAME the failure, QUOTE the exact moment, and explain the REVENUE IMPACT. Do NOT sugarcoat. Do NOT use filler phrases like "overall, the agent did a good job." Every sentence must earn its place with evidence or insight. Mediocrity is not acceptable — it costs clients money and agents their reputation.

═══════════════════════════════════════════════
CRITICAL: ANTI-HALLUCINATION & FACTUAL INTEGRITY RULES
═══════════════════════════════════════════════

You MUST follow these anti-hallucination rules without exception:

1. STRICT GROUNDING: Every claim, score, quote, timestamp, and metric you produce MUST be directly supported by evidence found in the transcript below. If you cannot find evidence in the transcript for a specific insight, you MUST state: "Information cannot be determined from the transcript."

2. NO FABRICATION: Under NO circumstances may you:
   - Invent or assume numbers, percentages, or scores that are not supported by transcript evidence
   - Fabricate client intent, emotions, or motivations not explicitly expressed in the transcript
   - Create imaginary property discussions, pricing details, or deal terms not mentioned
   - Generate placeholder metrics or fallback data
   - Assume information that was not discussed (e.g., do not assume budget if never mentioned)

3. PARTIAL INFORMATION: If the transcript covers some topics well but lacks information for others:
   - For sections with sufficient evidence: generate the analysis normally with full detail
   - For sections lacking evidence: clearly state what cannot be determined and WHY
   - Example: "Client budget was not discussed during this conversation. Cannot be determined from the transcript."

4. QUOTE ACCURACY: Every direct quote you attribute to a speaker MUST be an actual phrase from the transcript. Do NOT paraphrase and present it as a direct quote. If you must paraphrase, explicitly mark it as a paraphrase.

5. TIMESTAMP INTEGRITY: Only reference timestamps that correspond to actual moments in the transcript. Do NOT fabricate or estimate timestamps.

6. SCORE JUSTIFICATION: Every numerical score MUST be justified with specific transcript evidence. If evidence is insufficient to score a dimension, state: "Insufficient transcript evidence to score this dimension reliably" and assign a conservative score with AI Confidence: Low.

KNOWLEDGE BASE CONTEXT (use to benchmark against industry best practices):
{knowledge_context if knowledge_context else "No knowledge base available — rely on expert judgment."}

FULL TRANSCRIPT WITH TIMESTAMPS:
{transcript_text}

═══════════════════════════════════════════════
PRODUCE EXACTLY 19 SECTIONS + ACTION ITEMS USING ## HEADERS
═══════════════════════════════════════════════

## 1. Executive Brief
Write a 4-6 sentence ruthless executive summary. State who the agent is, what the meeting was about, and deliver a VERDICT — not a compliment. Was this a competent presentation or a missed-opportunity failure? Include the single most impressive moment AND the single worst failure. End with a one-line prognosis: will this deal close, and at what cost? Be direct — the agent's income depends on hearing the truth.

## 2. Performance Score Dashboard
Rate EACH dimension on a strict 1-10 scale using PIPE-DELIMITED format:
Label|X|10|One-sentence justification with specific evidence|AI Confidence: High/Medium/Low

IMPORTANT: The 5th pipe field MUST contain an AI confidence indicator (High/Medium/Low) reflecting how much transcript evidence supports this score.

Required dimensions:
Communication Effectiveness|?|10|...|AI Confidence: ...
Client Engagement & Rapport|?|10|...|AI Confidence: ...
Objection Handling Skill|?|10|...|AI Confidence: ...
Needs Discovery Depth|?|10|...|AI Confidence: ...
Closing Ability & Deal Progression|?|10|...|AI Confidence: ...
Confidence & Authority|?|10|...|AI Confidence: ...
Value Proposition Clarity|?|10|...|AI Confidence: ...
Risk Management Awareness|?|10|...|AI Confidence: ...
Emotional Intelligence|?|10|...|AI Confidence: ...
Listening Quality|?|10|...|AI Confidence: ...

Then provide:
OVERALL|?|10|Weighted assessment considering all dimensions above|AI Confidence: ...

Be HARSH. A 7/10 means "competent but with clear gaps." An 8+ requires exceptional evidence. A 5 or below means "this actively hurt the deal." Do not inflate scores to be nice — that helps nobody.

## 3. Behavioral Breakdown
Analyze the agent's behavior in granular detail using ### subheadings:
### Communication Patterns
### Body Language Signals (from verbal cues)
### Questioning Technique
### Active Listening Indicators
### Persuasion & Influence Tactics

For EACH, cite specific timestamps and quotes. Identify PATTERNS, not just isolated moments. Be clinical — patterns reveal whether this agent can repeat success or is getting lucky.

## 4. Strengths (Evidence-Anchored)
List 3-5 genuine strengths. For EACH strength, use this exact structure:

STRENGTH: [Name of the strength in 3-5 words]
TIMESTAMP: [exact time range]
EVIDENCE: "Exact quote from transcript proving this strength"
IMPACT: Why this specific strength matters for deal progression and client trust

Do NOT pad this section. If the agent only has 2 real strengths, list 2. Do NOT combine multiple observations into one bullet — each strength must be a separate, clearly structured entry.

## 5. Missed Opportunities & Critical Failures
This is the MOST IMPORTANT section. The agent NEEDS to hear this. For each failure, use this exact structure:

FAILURE TITLE: [Clear 3-5 word name for this failure]
TIMESTAMP: [exact time range]
QUOTE: "What the agent actually said"
WHAT HAPPENED: Describe the failure in 1-2 sentences
WHAT A TOP AGENT WOULD DO: Specific alternative approach with example script
REVENUE IMPACT: How this failure could cost money or trust — be specific with dollar estimates where possible

Provide AT LEAST 3 failures. Be SPECIFIC and CONSTRUCTIVE. Vague criticisms like "could have been better" waste everyone's time.

## 6. Tactical Corrections & Coaching Playbook
This section combines tactical fixes for past mistakes AND coaching drills for future scenarios into ONE unified playbook. For EACH entry (provide 5-7 entries total), use this structure:

SCENARIO: Brief description of the situation or recurring pattern
CLIENT TRIGGER: What the client said/did that created the opening
RISK: What happens if the agent keeps doing this — quantify the cost
ASSERTIVE APPROACH: "Exact word-for-word script using direct, authoritative framing"
CONSULTATIVE APPROACH: "Exact word-for-word script using collaborative, question-based framing"
WHY IT WORKS: Sales psychology principle behind both approaches

Each entry must address a DIFFERENT skill gap. No repetition from other sections.

## 7. 7-Day Intensive Coaching Plan
Provide a comprehensive day-by-day coaching plan for the next 7 days. Each day must include a focus area, a specific drill or exercise, and a measurable outcome. This replaces both short-term focus and long-term roadmap — compress ALL development priorities into one aggressive 7-day sprint.

Format each day as:
Day X: [Focus Area] -- [Specific Drill/Exercise] -- [Measurable Outcome/Success Metric]

Cover all 7 days. Each day must target a DIFFERENT skill gap identified in this evaluation. Include:
- At least 2 role-play exercises across the 7 days
- At least 1 script-writing exercise
- At least 1 objection-handling drill
- At least 1 listening/discovery practice session
Make each drill SPECIFIC and ACTIONABLE — not vague advice like "practice more." Include exact scenarios, word counts, or rep counts where possible.

## 8. Decisions & Commitments Tracker
Split into two clearly labeled subsections:
Decisions Made:
- List every decision made during the meeting with context
Commitments Given:
- List every commitment/promise made by either party

## Action Items
Create a table of specific action items arising from the meeting:
| Task Description | Responsible Person | Deadline | Priority | Status |
|---|---|---|---|---|
| [Specific task] | [Name] | [Date if mentioned, otherwise "No due date mentioned"] | [High/Medium/Low] | Pending |

Include 4-6 concrete action items based on what was discussed and agreed.

## 9. Deal Intelligence Summary
Use PIPE-DELIMITED format for key metrics:
Deal Closing Probability|XX%|Justification
Client Motivation Level|Strong/Moderate/Weak|Evidence
Urgency Level|High/Medium/Low|Evidence
Price Sensitivity|High/Medium/Low|Evidence
Competition Risk|High/Medium/Low|Evidence
Trust Level Established|Strong/Moderate/Weak|Evidence
Objection Resolution Rate|XX%|Evidence

Then add:
Recommended Next Steps:
1. [Specific action]
2. [Specific action]
3. [Specific action]

## 10. Agent Tier Classification
Classify the agent using this tier system:
Agent Tier: Level X - [Title]

Tier Scale:
Level 1: Trainee -- Fundamental gaps in sales methodology; needs intensive coaching
Level 2: Developing -- Shows awareness of techniques but inconsistent execution
Level 3: Competent -- Solid fundamentals with specific areas for growth
Level 4: Advanced -- Strong performer with minor refinements needed
Level 5: Elite -- Exceptional across all dimensions; ready to mentor others

Promotion Criteria (Next Level):
- At least 3 specific, measurable criteria

## 11. Client Emotional Journey & Risk Map
Analyze the client's emotional state throughout the conversation AND provide actionable solutions for each risk identified. For EACH phase:

EARLY PHASE (first third of conversation):
Sentiment: [Neutral/Curious/Skeptical/Engaged/etc.]
Evidence: Quote specific client language
Timestamp: [range]
Risk Identified: What could go wrong based on this sentiment
Solution: Specific action the agent should take to address this risk

MID PHASE (middle third):
Sentiment: [state]
Evidence: How client language changed
Timestamp: [range]
Risk Identified: What danger signs appeared
Solution: What the agent should do differently next time at this stage

LATE PHASE (final third):
Sentiment: [state]
Evidence: Quote specific language shifts
Timestamp: [range]
Risk Identified: End-state risk assessment
Solution: Immediate corrective action to protect the deal

EMOTIONAL TURNING POINTS:
- Identify 2-3 moments where sentiment shifted. For each: TIMESTAMP, TRIGGER, DIRECTION (positive/negative), RECOMMENDED RESPONSE

TRUST ARC: Did trust increase, decrease, or remain flat? Cite evidence and state what the agent must do to strengthen trust going forward.

OVERALL SENTIMENT TRAJECTORY: [e.g., Neutral -> Curious -> Engaged -> Price-Anxious -> Cautiously Interested]

## 12. Communication Balance & Listening Evaluation
Evaluate the agent's listening and communication balance:

TALK RATIO:
Estimated Agent Talk Time: XX%
Estimated Client Talk Time: XX%
Verdict: [Balanced / Agent-dominated / Client-led] — explain WHY this ratio helps or hurts the deal
Recommendation: Specific adjustment for next meeting

INTERRUPTION ANALYSIS:
- Count instances where agent cut short client statements
- For each: TIMESTAMP, what happened, and what the agent should have done instead

REFLECTIVE LISTENING EVALUATION:
- How often did the agent paraphrase, summarize, or validate?
- Cite specific examples with timestamps
- Score: X/10
- Improvement Tip: One specific technique to improve this score

FOLLOW-UP QUESTION QUALITY:
- Rate depth: Surface-level / Moderate / Deep
- Cite best and worst examples with timestamps
- Recommendation: Specific question types the agent should add to their toolkit

CRITICAL LISTENING FAILURES:
- Moments where the agent missed or ignored important client signals
- For each: TIMESTAMP, what the client said, what the agent missed, and what a top agent would have caught

## 13. Five Core Discovery Pillars Assessment
Evaluate how thoroughly the agent explored the 5 essential discovery areas that every top agent MUST cover before any listing presentation can succeed:

Use PIPE-DELIMITED format:
Pillar|Status|Evidence|AI Confidence
Budget/Financial Capacity|Confirmed/Partially Explored/Not Addressed|Specific quote or timestamp|High/Medium/Low
Decision Authority|Confirmed/Partially Explored/Not Addressed|Evidence|High/Medium/Low
Timeline & Urgency|Confirmed/Partially Explored/Not Addressed|Evidence|High/Medium/Low
Emotional Motivations|Explored/Partially Explored/Not Addressed|Evidence|High/Medium/Low
Competitive Alternatives|Explored/Partially Explored/Not Addressed|Evidence|High/Medium/Low

DISCOVERY COMPLETENESS SCORE: X/5 pillars adequately explored

MISSING DISCOVERY QUESTIONS:
- List 3-5 specific questions the agent MUST ask next time, with the exact wording

## 14. Emotional Intelligence Assessment
Evaluate the agent's EQ during the interaction. For EACH dimension, provide the score on its OWN line, then the evidence as separate bullet points below:

EMPATHY SCORE: X/10
- Evidence point 1 with timestamp
- Evidence point 2 with timestamp

ADAPTABILITY SCORE: X/10
- Evidence point 1 with timestamp
- Evidence point 2 with timestamp

SOCIAL AWARENESS SCORE: X/10
- Evidence point 1 with timestamp
- Evidence point 2 with timestamp

EMOTIONAL REGULATION SCORE: X/10
- Evidence point 1 with timestamp
- Evidence point 2 with timestamp

OVERALL EQ SCORE: X/10
AI Confidence: High/Medium/Low

Do NOT overlap score text with evidence. Each score must be on its own line.

## 15. Ethics & Compliance Audit
Evaluate ethical and compliance dimensions:

PRESSURE TACTICS ASSESSMENT:
- Instances of high-pressure tactics with TIMESTAMP and language used
- Risk Level: None/Low/Medium/High

OVER-PROMISING CHECK:
- Claims or guarantees that may be unsupported with TIMESTAMP
- Risk Level: None/Low/Medium/High

REGULATORY LANGUAGE REVIEW:
- Statements creating legal/regulatory exposure
- Fair housing compliance indicators

TRANSPARENCY ASSESSMENT:
- Was the agent transparent about fees, processes, limitations?

OVERALL ETHICS RISK: None/Low/Medium/High
Justification: Brief assessment

ETHICS RECOMMENDATIONS:
- Specific actionable suggestions

## 16. Self-Awareness & Adaptability Profile
Evaluate whether the agent recognizes their own patterns and adapts in real-time. For EACH dimension, provide the findings then a specific improvement recommendation:

SIGNAL RECOGNITION:
- Did the agent recognize client hesitation or disengagement?
- For each signal: TIMESTAMP, description, whether agent noticed it
- Improvement Recommendation: Specific technique to improve signal detection

REAL-TIME ADAPTATION:
- Did the agent adjust approach mid-conversation?
- Cite specific pivot moments or missed adaptation opportunities
- Improvement Recommendation: Specific drill to build adaptive response skills

SELF-CORRECTION INSTANCES:
- Moments where the agent caught and corrected their own approach
- Improvement Recommendation: How to build self-correction into muscle memory

BLIND SPOTS:
- Recurring patterns the agent appears unaware of
- Improvement Recommendation: Specific exercise to expose and address each blind spot

SELF-AWARENESS SCORE: X/10
AI Confidence: High/Medium/Low

## 17. Client Engagement & Tone Analysis
Analyze the agent's tone, energy, and client engagement quality throughout the conversation. This section focuses EXCLUSIVELY on HOW the agent communicates — not WHAT they say.

TONE CONSISTENCY:
- Was the agent's tone professional, warm, authoritative, or inconsistent?
- Identify tone shifts with TIMESTAMPS and describe the trigger
- Score: X/10

ENERGY & ENTHUSIASM:
- Did the agent maintain appropriate energy? Too flat? Too aggressive?
- Cite specific moments where energy helped or hurt the interaction
- Score: X/10

CLIENT ENGAGEMENT QUALITY:
- How effectively did the agent keep the client involved and participating?
- Identify moments of peak engagement vs disengagement with TIMESTAMPS
- Score: X/10

RAPPORT-BUILDING LANGUAGE:
- Specific phrases, mirroring, name usage, personal references that built connection
- Missed opportunities for rapport (with exact moments)
- Score: X/10

OVERALL TONE & ENGAGEMENT SCORE: X/10
AI Confidence: High/Medium/Low

## 18. Negotiation & Persuasion Proficiency
Evaluate the agent's negotiation tactics, persuasion techniques, and deal-advancing behaviors. This section analyzes the agent's ability to move the client toward commitment.

PERSUASION TECHNIQUES USED:
- Identify specific persuasion frameworks (social proof, scarcity, authority, reciprocity, anchoring, framing)
- For each technique: TIMESTAMP, what was said, effectiveness rating (Effective/Partially Effective/Ineffective)

NEGOTIATION POSITIONING:
- Did the agent establish strong positioning or give away leverage?
- Specific moments where negotiation advantage was gained or lost
- Score: X/10

OBJECTION-TO-COMMITMENT CONVERSION:
- Track each objection raised and whether the agent converted it toward commitment
- For each: OBJECTION → AGENT RESPONSE → OUTCOME (Resolved/Partially Resolved/Unresolved)

VALUE FRAMING ABILITY:
- How well did the agent frame their value proposition relative to price/competition?
- Specific examples with TIMESTAMPS
- Score: X/10

CLOSING SIGNALS & TRIAL CLOSES:
- Did the agent use trial closes or test for commitment?
- Missed closing windows with TIMESTAMPS
- Score: X/10

OVERALL NEGOTIATION SCORE: X/10
AI Confidence: High/Medium/Low

## 19. Deal Outcome / Client Decision
Determine the final deal status between the agent and the client based STRICTLY on evidence from the transcript. Do NOT assume or fabricate a decision.

Analyze the transcript for explicit signals such as:
- The deal was successfully closed (e.g., client signed, agreed to list, committed verbally)
- The client refused or declined the deal (e.g., explicit rejection, chose another agent)
- The client postponed the decision (e.g., "we will see", "we will get back to you", "I will think about it", "let us discuss", "reach out later")
- The deal status is unclear or still under discussion

Provide your analysis in the following EXACT format (one field per line):
DEAL_STATUS: [Exactly one of: Deal Closed | Client Declined the Deal | Client Requested Time to Decide | Decision Pending / Not Clearly Stated]
CONFIDENCE: [high | medium | low]
EVIDENCE: [1-3 sentence summary of the client's final stance, citing specific transcript moments]
SUPPORTING_QUOTES:
- "Exact quote from transcript supporting the determination"
- "Another exact quote if available"

RULES FOR THIS SECTION:
1. The determination MUST be based ONLY on actual client statements in the transcript.
2. If the client did not make a clear final decision, use "Decision Pending / Not Clearly Stated".
3. Do NOT interpret silence, politeness, or engagement as agreement.
4. A verbal commitment to list with the agent counts as "Deal Closed" ONLY if the client explicitly stated agreement.
5. Expressions like "we will think about it", "we need to discuss", "I will get back to you" MUST be classified as "Client Requested Time to Decide".
6. If evidence is insufficient, state so clearly with CONFIDENCE: low.

═══════════════════════════════════════════════
ABSOLUTE RULES — VIOLATING ANY RULE INVALIDATES THE ENTIRE OUTPUT
═══════════════════════════════════════════════

1. ZERO REDUNDANCY: If a fact, quote, or observation appears in one section, it MUST NOT appear in any other section. Each section must contain 100% unique content. Sections 11-18 each cover a DIFFERENT analytical dimension — do not repeat findings across them.
2. EVIDENCE-ANCHORED: Every claim must reference a specific timestamp, quote, or observable behavior. No unsupported assertions. ALL TIMESTAMPS must be in M:SS or MM:SS format (minutes:seconds), e.g., [0:30-0:45], [5:15-6:00], [24:30-25:10]. NEVER use raw seconds format.
3. CLEAR SEPARATION: Sections 1-5 = DIAGNOSIS (what happened). Section 6 = COACHING (what to do differently). Sections 7-8 = DEVELOPMENT PLAN & COMMITMENTS. Section 9-10 = DEAL INTELLIGENCE & CLASSIFICATION. Sections 11-18 = ADVANCED INTELLIGENCE (each with a unique analytical lens). Section 19 = DEAL OUTCOME (final client decision). Never mix these purposes.
4. NO FILLER: Remove every sentence that doesn't add NEW information. If a section would be empty, write "No significant findings for this dimension."
5. PIPE FORMAT: Scores in Section 2, metrics in Sections 9 and 13 MUST use the pipe-delimited format specified.
6. HONEST AUTHORITY: Be direct and clinical. Do not soften language to spare feelings — that costs the agent money. Use "significant gap," "below benchmark," "critical development area," "requires immediate attention." Use "exceptional," "expert-level," "textbook execution" ONLY when genuinely earned with evidence.
7. NO EMOJIS, NO HASHTAGS: Professional language only. No markdown headers beyond ## and ###.
8. MINIMUM DEPTH: Section 5 must contain AT LEAST 3 failures. Section 6 must contain AT LEAST 5 playbook entries. Section 4 must contain AT LEAST 3 strengths.
9. SECTION 3 MUST include all 5 ### subheadings listed. Each subheading MUST have its own paragraph with timestamps.
10. SECTION 10 MUST include the full Tier Scale (all 5 levels) and at least 3 Promotion Criteria bullet points.
11. AI CONFIDENCE TRANSPARENCY: Every numerical score must include an AI confidence qualifier.
12. DUAL COACHING STYLES: Section 6 MUST provide both an assertive AND a consultative approach for each entry.
13. ETHICS AWARENESS: Section 15 must objectively assess ethical dimensions.
14. SCORE FORMATTING: In Sections 14, 16, 17, and 18, scores (X/10) must appear on their OWN line, separate from evidence text. Never put a score and evidence paragraph on the same line.
15. SECTION 7 MUST cover all 7 days with a different skill focus each day. Include at least 2 role-play drills across the week.
16. SECTIONS 17-18 must provide unique analysis not covered in earlier sections — focus on HOW the agent communicates (tone) and HOW they negotiate (persuasion).
17. ZERO HALLUCINATION: Every insight, score, quote, and recommendation MUST be grounded in the actual transcript content provided above. If you cannot find supporting evidence in the transcript for any data point, explicitly state "Cannot be determined from the transcript" rather than fabricating information. Never invent quotes, timestamps, numbers, or client intentions not present in the transcript.
18. PARTIAL COVERAGE HONESTY: If the transcript does not contain enough information to fully populate a section, populate what you can with evidence and clearly mark gaps with "Insufficient transcript evidence for this sub-section." Do NOT fill gaps with assumptions or generic advice.
"""
        return self._call_llm(prompt)

    # ════════════════════════════════════════════════════════
    # SECTION PARSING
    # ════════════════════════════════════════════════════════

    def _parse_unified_sections(self, raw_text: str) -> Dict[str, str]:
        sections = {}
        current_key = None
        current_lines = []

        for line in raw_text.split('\n'):
            # Match numbered headers like "## 1. Title" or unnumbered like "## Action Items"
            m = re.match(r'^##\s*(?:(\d+)\.\s*)?(.*)', line)
            if m and m.group(2).strip():
                header_text = m.group(2).strip()
                # Skip lines that look like table separators
                if header_text.startswith('|') or header_text.startswith('---'):
                    if current_key:
                        current_lines.append(line)
                    continue
                if current_key:
                    sections[current_key] = '\n'.join(current_lines).strip()
                num = m.group(1)
                current_key = f"{num}. {header_text}" if num else header_text
                current_lines = []
            elif current_key:
                current_lines.append(line)

        if current_key:
            sections[current_key] = '\n'.join(current_lines).strip()
        return sections

    # ════════════════════════════════════════════════════════
    # ORCHESTRATOR
    # ════════════════════════════════════════════════════════

    def generate_unified_coaching_evaluation(self):
        transcript_data = self._load_transcript()
        if not transcript_data:
            logger.error("Failed to load transcript.")
            return

        # ═══════════════════════════════════════════════════════════
        # TRANSCRIPT VALIDATION — ANTI-HALLUCINATION PRE-CHECK
        # ═══════════════════════════════════════════════════════════
        logger.info("Running transcript validation (anti-hallucination pre-check)...")
        self._validation_result = TranscriptValidator.validate(transcript_data)

        if not self._validation_result.is_valid:
            logger.warning(f"Transcript validation FAILED: {self._validation_result.validation_message}")
            logger.warning(f"Failure reasons: {self._validation_result.failure_reasons}")
            # Generate validation-failure outputs (PDFs + JSONs with messages, no LLM call)
            self._transcript_data = transcript_data
            self._agent_name = self._extract_agent_name(transcript_data)
            self._cached_scores = {}
            self._generate_validation_failure_outputs(transcript_data)
            return
        else:
            logger.info(f"Transcript validation PASSED: {self._validation_result.validation_message}")

        # Extract agent and client names for anonymization
        speakers = transcript_data.summary_info.get('unique_speakers', [])
        if len(speakers) == 1:
            agent_name = speakers[0]
        elif len(speakers) > 1:
            speaker_counts = {}
            for seg in transcript_data.transcripts:
                if seg.speaker_name:
                    speaker_counts[seg.speaker_name] = speaker_counts.get(seg.speaker_name, 0) + 1
            agent_name = max(speaker_counts, key=speaker_counts.get) if speaker_counts else speakers[0]
        else:
            agent_name = "Agent"

        client_names = set()
        for seg in transcript_data.transcripts:
            if seg.speaker_name and seg.speaker_name != agent_name and seg.speaker_name != 'Unknown':
                client_names.add(seg.speaker_name)
        # Also scan agent's dialogue for names
        name_patterns = re.findall(
            r'(?:So|Hey|Now|And|,)\s+([A-Z][a-z]{2,})(?:[,\s])', 
            ' '.join(seg.transcript for seg in transcript_data.transcripts if seg.speaker_name == agent_name)
        )
        common_words = {'The', 'This', 'That', 'When', 'What', 'Where', 'How', 'Why', 'Now',
                        'Here', 'There', 'Well', 'Yes', 'Yeah', 'Cool', 'Right', 'Just',
                        'Let', 'Got', 'But', 'And', 'Because', 'Matter', 'Number', 'Level',
                        'Royal', 'Okay', 'Does', 'Both', 'First', 'Second', 'Third',
                        'Some', 'Like', 'Also', 'Most', 'Very', 'Sure', 'Homes', 'Home'}
        for name in name_patterns:
            if name not in common_words and name != agent_name:
                client_names.add(name)

        speaker_map = {agent_name: 'Agent'}
        for client in client_names:
            speaker_map[client] = 'Client'

        self._agent_name = agent_name  # Store for renderers

        transcript_text = self._format_transcript_for_analysis(transcript_data)

        # RAG retrieval — build a rich query from agenda + transcript preview
        logger.info("Retrieving relevant knowledge from knowledge base...")
        rag_query = f"Sales coaching evaluation for: Real Estate Listing Presentation. {transcript_text[:500]}"
        relevant_chunks = self._retrieve_relevant_knowledge(rag_query, top_k=8)

        if relevant_chunks:
            logger.info(f"Retrieved {len(relevant_chunks)} relevant knowledge chunks.")
            knowledge = "\n\n".join([
                f"[Knowledge Reference {i+1}]\n{chunk.text}"
                for i, chunk in enumerate(relevant_chunks)
            ])
        else:
            logger.info("No knowledge base chunks retrieved — using general expertise.")
            knowledge = ""

        logger.info("Calling LLM for brutally honest evaluation...")
        raw = self._generate_unified_llm_analysis(transcript_text, knowledge)
        if not raw:
            logger.error("LLM returned no response.")
            return

        sections = self._parse_unified_sections(raw)
        if not sections:
            logger.error("Failed to parse LLM response into sections.")
            return

        # Anonymize sections except Executive Brief
        for key, content in sections.items():
            if not key.startswith('1.'):
                for old_name, new_name in speaker_map.items():
                    content = content.replace(old_name, new_name)
                sections[key] = content

        logger.info(f"Parsed {len(sections)} sections. Building 4 PDFs + 8 JSON files...")
        self._transcript_data = transcript_data  # Store for visualization section access

        # Build section number map for content lookup
        self._sec_num_map = {}
        for key in sections:
            m = re.match(r'^(\d+)\.?\s*', key)
            if m:
                self._sec_num_map[int(m.group(1))] = key
            else:
                self._sec_num_map[key] = key

        # Create two separate VisualizationLoggers for the split JSON outputs
        self.coaching_viz_logger = VisualizationLogger(report_type="coaching_insights")
        self.performance_viz_logger = VisualizationLogger(report_type="agent_performance")

        common_metadata = dict(
            agent=getattr(self, '_agent_name', 'Agent'),
            speakers=transcript_data.summary_info.get('unique_speakers', []),
            total_segments=transcript_data.summary_info.get('total_segments', 0),
            total_duration_seconds=transcript_data.summary_info.get('total_duration_seconds', 0),
            generated_date=datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        )
        self.coaching_viz_logger.set_metadata(**common_metadata)
        self.performance_viz_logger.set_metadata(**common_metadata)

        # PDF 1: Coaching Summary
        logger.info("Building PDF 1: coaching_summary.pdf...")
        self._create_coaching_summary_pdf(sections, transcript_data)

        # PDF 2: Agent Profile
        logger.info("Building PDF 2: agent_summary.pdf...")
        self._create_agent_profile_pdf(sections, transcript_data)

        # PDF 3: Coaching Insights Visualization
        logger.info("Building PDF 3: coaching_insights_viz.pdf...")
        self._create_coaching_insights_viz_pdf(sections, transcript_data)

        # PDF 4: Agent Performance Visualization
        logger.info("Building PDF 4: agent_performance_viz.pdf...")
        self._create_agent_performance_viz_pdf(sections, transcript_data)

        # JSON 1: Coaching Insights JSON
        coaching_json_path = os.path.join(self.summaries_folder, "coaching_insights.json")
        self.coaching_viz_logger.save_to_json(coaching_json_path)

        # JSON 2: Agent Performance JSON
        performance_json_path = os.path.join(self.summaries_folder, "agent_performance.json")
        self.performance_viz_logger.save_to_json(performance_json_path)

        # JSON 3: Action Items JSON
        content_actions = self._find_section_content(sections, 0, ['Action', 'Items'])
        action_items_list = self._parse_action_items_to_list(content_actions)
        action_items_json_path = os.path.join(self.summaries_folder, "action_items.json")
        action_items_payload = {
            "generated_at": datetime.now().isoformat(),
            "report_type": "action_items",
            "metadata": {
                "agent": getattr(self, '_agent_name', 'Agent'),
                "total_sections_parsed": len(sections),
                "transcript_file": self.transcript_file,
            },
            "total_action_items": len(action_items_list),
            "action_items": action_items_list,
        }
        try:
            with open(action_items_json_path, 'w', encoding='utf-8') as f:
                json.dump(action_items_payload, f, indent=2, ensure_ascii=False, default=str)
            logger.info(f"Saved action items JSON: {action_items_json_path}")
        except Exception as e:
            logger.error(f"Failed to save action items JSON: {e}")

        # JSON 4: All Visualizations (master combined JSON)
        logger.info("Building JSON 4: all_visualizations.json...")
        self._generate_all_visualizations_json(sections)

        # JSON 5: Coaching Recommendations
        logger.info("Building JSON 5: coaching_recommendations.json...")
        self._generate_coaching_recommendations_json(sections)

        # JSON 6: Agent Tier Calculation
        logger.info("Building JSON 6: agent_tier_calculation.json...")
        self._generate_agent_tier_calculation_json(sections)

        # JSON 7: Agent Performance Aggregate (new)
        print("  Building JSON 7: agent_performance_aggregate.json...")
        self._generate_agent_performance_aggregate_json(sections, transcript_data)

        # JSON 8: Deal Outcome / Client Decision
        logger.info("Building JSON 8: deal_outcome.json...")
        self._generate_deal_outcome_json(sections)

        logger.info("\nAll 4 PDFs + 8 JSON files generated successfully!")
        logger.info("PDF 1: coaching_summary.pdf")
        logger.info("PDF 2: agent_summary.pdf")
        logger.info("PDF 3: coaching_insights_viz.pdf")
        logger.info("PDF 4: agent_performance_viz.pdf")
        logger.info("JSON 1: coaching_insights.json")
        logger.info("JSON 2: agent_performance.json")
        logger.info("JSON 3: action_items.json")
        logger.info("JSON 4: all_visualizations.json")
        logger.info("JSON 5: coaching_recommendations.json")
        logger.info("JSON 6: agent_tier_calculation.json")
        logger.info("JSON 7: agent_performance_aggregate.json")
        logger.info("JSON 8: deal_outcome.json")

    # ════════════════════════════════════════════════════════
    # HELPER: Extract agent name from transcript data
    # ════════════════════════════════════════════════════════

    def _extract_agent_name(self, transcript_data: TranscriptData) -> str:
        """Extract agent name from transcript data (used by validation-failure path too)."""
        speakers = transcript_data.summary_info.get('unique_speakers', [])
        if len(speakers) == 1:
            return speakers[0]
        elif len(speakers) > 1:
            speaker_counts = {}
            for seg in transcript_data.transcripts:
                if seg.speaker_name:
                    speaker_counts[seg.speaker_name] = speaker_counts.get(seg.speaker_name, 0) + 1
            return max(speaker_counts, key=speaker_counts.get) if speaker_counts else speakers[0]
        return "Agent"

    # ════════════════════════════════════════════════════════
    # VALIDATION FAILURE OUTPUTS — NO LLM, NO HALLUCINATION
    # ════════════════════════════════════════════════════════

    def _generate_validation_failure_outputs(self, transcript_data: TranscriptData):
        """Generate all 4 PDFs + 8 JSON files with validation failure messages.
        
        Called when the transcript fails validation (too short or not real estate related).
        Maintains the exact same file structure and layout but replaces content with
        clear validation messages. No LLM call is made.
        """
        validation = self._validation_result
        agent_name = getattr(self, '_agent_name', 'Agent')
        failure_msg = validation.validation_message

        logger.info("Generating validation-failure outputs (no LLM call)...")

        # ── PDF 1: Coaching Summary (validation failure) ──
        self._create_validation_failure_pdf(
            transcript_data,
            pdf_filename="coaching_summary.pdf",
            pdf_title="COACHING SUMMARY REPORT",
            pdf_subtitle="Actionable Performance Coaching & Improvement Plan",
            section_titles=[
                "EXECUTIVE BRIEF",
                "PERFORMANCE SCORE DASHBOARD",
                "STRENGTHS & AREAS FOR IMPROVEMENT",
                "CLIENT INSIGHTS & COMMUNICATION INTELLIGENCE",
                "COACHING RECOMMENDATIONS & TACTICAL PLAYBOOK",
                "7-DAY INTENSIVE COACHING PLAN",
                "DEAL INTELLIGENCE & ACTION PLAN",
                "DEAL OUTCOME / CLIENT DECISION",
            ],
            failure_msg=failure_msg,
        )

        # ── PDF 2: Agent Profile (validation failure) ──
        self._create_validation_failure_pdf(
            transcript_data,
            pdf_filename="agent_summary.pdf",
            pdf_title="AGENT PROFILE & INTELLIGENCE REPORT",
            pdf_subtitle="Advanced Behavioral Intelligence & Compliance Assessment",
            section_titles=[
                "FIVE CORE DISCOVERY PILLARS ASSESSMENT",
                "EMOTIONAL INTELLIGENCE ASSESSMENT",
                "ETHICS & COMPLIANCE AUDIT",
                "SELF-AWARENESS & ADAPTABILITY PROFILE",
                "CLIENT ENGAGEMENT & TONE ANALYSIS",
                "NEGOTIATION & PERSUASION PROFICIENCY",
                "BEHAVIORAL BREAKDOWN & COMMUNICATION PATTERNS",
                "AGENT TIER CLASSIFICATION",
            ],
            failure_msg=failure_msg,
        )

        # ── PDF 3: Coaching Insights Viz (validation failure) ──
        self._create_validation_failure_pdf(
            transcript_data,
            pdf_filename="coaching_insights_viz.pdf",
            pdf_title="COACHING INSIGHTS VISUALIZATION",
            pdf_subtitle="Behavioral Analysis, Communication Quality & Coaching Recommendations",
            section_titles=[
                "COACHING VISUALIZATIONS",
            ],
            failure_msg=failure_msg,
        )

        # ── PDF 4: Agent Performance Viz (validation failure) ──
        self._create_validation_failure_pdf(
            transcript_data,
            pdf_filename="agent_performance_viz.pdf",
            pdf_title="AGENT PERFORMANCE VISUALIZATION",
            pdf_subtitle="Quantitative Performance Metrics, Scoring & Evaluation Charts",
            section_titles=[
                "PERFORMANCE VISUALIZATIONS",
            ],
            failure_msg=failure_msg,
        )

        # ── JSON 1: Coaching Insights (validation failure) ──
        self._save_validation_failure_json(
            "coaching_insights.json",
            report_type="coaching_insights",
            agent_name=agent_name,
            failure_msg=failure_msg,
        )

        # ── JSON 2: Agent Performance (validation failure) ──
        self._save_validation_failure_json(
            "agent_performance.json",
            report_type="agent_performance",
            agent_name=agent_name,
            failure_msg=failure_msg,
        )

        # ── JSON 3: Action Items (validation failure) ──
        action_items_payload = {
            "generated_at": datetime.now().isoformat(),
            "report_type": "action_items",
            "metadata": {
                "agent": agent_name,
                "total_sections_parsed": 0,
                "transcript_file": self.transcript_file,
                "validation_failed": True,
                "validation_message": failure_msg,
                "failure_reasons": validation.failure_reasons,
            },
            "total_action_items": 0,
            "action_items": [],
            "validation_status": "failed",
            "validation_message": failure_msg,
        }
        self._save_json(os.path.join(self.summaries_folder, "action_items.json"), action_items_payload)

        # ── JSON 4: All Visualizations (validation failure) ──
        all_viz_payload = AllVisualizationsOutput(
            generated_at=datetime.now().isoformat(),
            report_type="all_visualizations",
            schema_version="2.0",
            metadata={
                "agent": agent_name,
                "validation_failed": True,
                "validation_message": failure_msg,
                "failure_reasons": validation.failure_reasons,
            },
            total_visualizations=0,
            coaching_insights_count=0,
            agent_performance_count=0,
            visualizations=[],
        )
        self._save_json(
            os.path.join(self.summaries_folder, "all_visualizations.json"),
            all_viz_payload.model_dump(),
        )

        # ── JSON 5: Coaching Recommendations (validation failure) ──
        coaching_rec_payload = CoachingRecommendationsOutput(
            generated_at=datetime.now().isoformat(),
            report_type="coaching_recommendations",
            schema_version="2.0",
            agent_name=agent_name,
            metadata={
                "transcript_file": self.transcript_file,
                "total_sections_parsed": 0,
                "overall_score": 0,
                "validation_failed": True,
                "validation_message": failure_msg,
                "failure_reasons": validation.failure_reasons,
            },
            strengths=[],
            failures=[],
            tactical_playbook=[],
            seven_day_plan=[],
            summary={
                "total_strengths": 0,
                "total_failures": 0,
                "total_playbook_entries": 0,
                "total_plan_days": 0,
                "overall_performance_score": 0,
                "validation_status": "failed",
                "validation_message": failure_msg,
            },
        )
        self._save_json(
            os.path.join(self.summaries_folder, "coaching_recommendations.json"),
            coaching_rec_payload.model_dump(),
        )

        # ── JSON 6: Agent Tier Calculation (validation failure) ──
        tier_payload = AgentTierCalculationOutput(
            generated_at=datetime.now().isoformat(),
            report_type="agent_tier_calculation",
            schema_version="2.0",
            agent_name=agent_name,
            metrics_used={},
            weight_distribution={},
            scoring_formula="not_applicable",
            overall_score=0.0,
            final_agent_tier="Cannot Be Determined",
            tier_justification=failure_msg,
            promotion_criteria=[],
            comparative_metrics=[],
            metadata={
                "transcript_file": self.transcript_file,
                "total_dimensions_scored": 0,
                "validation_failed": True,
                "validation_message": failure_msg,
                "failure_reasons": validation.failure_reasons,
            },
        )
        self._save_json(
            os.path.join(self.summaries_folder, "agent_tier_calculation.json"),
            tier_payload.model_dump(),
        )

        # ── JSON 7: Agent Performance Aggregate (validation failure) ──
        aggregate_payload = AgentPerformanceAggregate(
            generated_at=datetime.now().isoformat(),
            agent_name=agent_name,
            transcript_file=self.transcript_file,
            meeting_date=datetime.now().strftime('%Y-%m-%d'),
            overall_performance_score=0.0,
            overall_deal_probability=None,
            talk_ratio={"agent": None, "client": None},
            performance_dimensions=[],
            discovery_pillars=[],
            emotional_intelligence={},
            tone_scores={},
            negotiation_scores={},
            ethics_risk=None,
            coaching_priority_areas=[],
            conversation_metrics={
                "total_segments": len(transcript_data.transcripts),
                "total_duration_seconds": transcript_data.summary_info.get('total_duration_seconds', 0),
                "total_duration_minutes": round(
                    transcript_data.summary_info.get('total_duration_seconds', 0) / 60, 1
                ),
                "validation_status": "failed",
                "validation_message": failure_msg,
            },
            question_type_counts={},
            objection_type_counts={},
            objection_handling_counts={},
            language_phrase_counts={},
            response_length_distribution={},
            response_delay_distribution={},
            sentiment_progression=[],
            trust_vs_pressure=None,
            deal_outcome=None,
        )
        self._save_json(
            os.path.join(self.summaries_folder, "agent_performance_aggregate.json"),
            aggregate_payload.model_dump(),
        )

        # ── JSON 8: Deal Outcome (validation failure) ──
        deal_outcome_payload = DealOutcomeResult(
            generated_at=datetime.now().isoformat(),
            report_type="deal_outcome",
            schema_version="1.0",
            agent_name=agent_name,
            status="Decision Pending / Not Clearly Stated",
            confidence="low",
            evidence=failure_msg,
            supporting_quotes=[],
            metadata={
                "transcript_file": self.transcript_file,
                "validation_failed": True,
                "validation_message": failure_msg,
                "failure_reasons": validation.failure_reasons,
            },
        )
        self._save_json(
            os.path.join(self.summaries_folder, "deal_outcome.json"),
            deal_outcome_payload.model_dump(),
        )

        logger.info("All 4 PDFs + 8 JSON files generated with validation failure messages.")
        logger.info(f"Validation failure reason: {failure_msg}")

    def _create_validation_failure_pdf(
        self,
        transcript_data: TranscriptData,
        pdf_filename: str,
        pdf_title: str,
        pdf_subtitle: str,
        section_titles: List[str],
        failure_msg: str,
    ):
        """Create a PDF with the standard layout but validation failure messages in every section."""
        pdf_path = os.path.join(self.summaries_folder, pdf_filename)
        doc = SimpleDocTemplate(
            pdf_path, pagesize=A4,
            leftMargin=0.6 * inch, rightMargin=0.6 * inch,
            topMargin=0.7 * inch, bottomMargin=0.5 * inch,
            title=pdf_title,
            author="RAG Coaching Evaluation System",
        )
        story = []

        # Cover page
        self._build_cover_page(story, transcript_data, pdf_title=pdf_title, pdf_subtitle=pdf_subtitle)
        story.append(PageBreak())

        # Validation status banner
        validation = self._validation_result
        banner_lines = [
            Paragraph("<b>TRANSCRIPT VALIDATION FAILED</b>", ParagraphStyle(
                'banner_title', parent=self.ps['sec'], fontSize=16,
                textColor=self.C['danger'], alignment=TA_CENTER, spaceBefore=10, spaceAfter=6)),
            Spacer(1, 6),
            Paragraph(failure_msg, ParagraphStyle(
                'banner_body', parent=self.ps['body'], fontSize=11,
                textColor=self.C['dark'], alignment=TA_CENTER, leading=16)),
            Spacer(1, 8),
        ]

        # Add specific failure reasons
        for reason in validation.failure_reasons:
            banner_lines.append(
                Paragraph(f"<bullet>&bull;</bullet> {reason}", ParagraphStyle(
                    'banner_reason', parent=self.ps['bullet'], fontSize=10,
                    textColor=self.C['medium'], leading=14))
            )
            banner_lines.append(Spacer(1, 4))

        # Validation details
        banner_lines.append(Spacer(1, 6))
        details = (
            f"<b>Total segments:</b> {validation.total_segments} &nbsp; | &nbsp; "
            f"<b>Meaningful exchanges:</b> {validation.meaningful_exchanges} &nbsp; | &nbsp; "
            f"<b>Real estate relevance:</b> {validation.real_estate_relevance_score:.1%}"
        )
        banner_lines.append(Paragraph(details, ParagraphStyle(
            'banner_details', parent=self.ps['body'], fontSize=9,
            textColor=self.C['light'], alignment=TA_CENTER)))

        banner_table = Table([[banner_lines]], colWidths=[6.2 * inch])
        banner_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#FFF3E0")),
            ('BOX', (0, 0), (-1, -1), 2, self.C['danger']),
            ('PADDING', (0, 0), (-1, -1), 18),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(banner_table)
        story.append(Spacer(1, 16))

        # Add each section with the validation message
        for i, title in enumerate(section_titles, 1):
            story.append(CondPageBreak(1.5 * inch))
            story.append(SectionDivider(width=460, color="#0D47A1", thickness=1.5))
            story.append(Spacer(1, 2))
            story.append(Paragraph(f"{i}. {title}", self.ps['sec']))
            story.append(Spacer(1, 8))

            # Validation message box for this section
            section_msg = Table(
                [[Paragraph(
                    f"<i>{failure_msg}</i>",
                    ParagraphStyle('val_msg', parent=self.ps['body'], fontSize=10,
                                   textColor=self.C['medium'], alignment=TA_CENTER, leading=14)
                )]],
                colWidths=[6.0 * inch],
            )
            section_msg.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F5F5")),
                ('BOX', (0, 0), (-1, -1), 0.5, self.C['border']),
                ('PADDING', (0, 0), (-1, -1), 14),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(section_msg)
            story.append(Spacer(1, 12))

        doc.build(story,
                  onFirstPage=self._draw_cover_header_footer,
                  onLaterPages=self._draw_content_header_footer)
        logger.info(f"Saved (validation failure): {pdf_path}")

    def _save_validation_failure_json(self, filename: str, report_type: str, agent_name: str, failure_msg: str):
        """Save a visualization JSON file with validation failure metadata."""
        validation = self._validation_result
        payload = {
            "generated_at": datetime.now().isoformat(),
            "report_type": report_type,
            "schema_version": "2.0",
            "metadata": {
                "agent": agent_name,
                "validation_failed": True,
                "validation_message": failure_msg,
                "failure_reasons": validation.failure_reasons,
                "total_segments": validation.total_segments,
                "meaningful_exchanges": validation.meaningful_exchanges,
                "real_estate_relevance_score": validation.real_estate_relevance_score,
            },
            "total_visualizations": 0,
            "visualizations": [],
            "validation_status": "failed",
            "validation_message": failure_msg,
        }
        self._save_json(os.path.join(self.summaries_folder, filename), payload)

    def _save_json(self, filepath: str, payload: dict):
        """Utility to save a JSON file with error handling."""
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(payload, f, indent=2, ensure_ascii=False, default=str)
            logger.info(f"Saved JSON: {filepath}")
        except Exception as e:
            logger.error(f"Failed to save JSON {filepath}: {e}")

    # ════════════════════════════════════════════════════════
    # HELPER: Find section content by number or keyword
    # ════════════════════════════════════════════════════════

    def _find_section_content(self, sections: Dict[str, str], sec_num: int = 0, title_words: List[str] = None) -> str:
        """Find section content by number or keyword match."""
        sec_num_map = getattr(self, '_sec_num_map', {})
        if sec_num > 0 and sec_num in sec_num_map:
            return sections.get(sec_num_map[sec_num], '')
        if title_words:
            words = [w.lower() for w in title_words if len(w) > 3]
            for key, val in sections.items():
                if any(word in key.lower() for word in words):
                    return val
        return ''

    # ════════════════════════════════════════════════════════
    # HELPER: Parse action items from LLM text into a list of dicts
    # ════════════════════════════════════════════════════════

    def _parse_action_items_to_list(self, content: str) -> List[Dict[str, str]]:
        """Parse action items from LLM output into a JSON-serializable list.
        Returns list of dicts with keys: task, responsible, deadline, priority, status."""
        items = []
        if not content or not content.strip():
            return items

        for line in content.split('\n'):
            line = line.strip()
            if not line or line.startswith('|---') or line.startswith('| Task') or line.startswith('|Task'):
                continue

            if '|' in line:
                parts = [p.strip().replace('**', '') for p in line.split('|') if p.strip()]
                if len(parts) >= 2:
                    responsible = parts[1] if len(parts) > 1 else ''
                    if responsible == 'Agent':
                        responsible = getattr(self, '_agent_name', 'Agent')
                    responsible = ' '.join(dict.fromkeys(responsible.split()))
                    items.append({
                        "task": parts[0],
                        "responsible": responsible,
                        "deadline": parts[2] if len(parts) > 2 else "No due date mentioned",
                        "priority": parts[3] if len(parts) > 3 else "Medium",
                        "status": parts[4] if len(parts) > 4 else "Pending",
                    })
            elif line.startswith('-') or line.startswith('*'):
                text = line.lstrip('-* ').strip()
                if text:
                    items.append({
                        "task": text,
                        "responsible": getattr(self, '_agent_name', 'Agent'),
                        "deadline": "No due date mentioned",
                        "priority": "Medium",
                        "status": "Pending",
                    })
        return items

    # ════════════════════════════════════════════════════════
    # HELPER: Add numbered section to story
    # ════════════════════════════════════════════════════════

    def _add_section_to_story(self, story, counter: int, title: str, renderer, content: str):
        """Add a numbered section header + rendered content to a story."""
        if not content or not content.strip():
            return counter
        counter += 1
        story.append(CondPageBreak(2.0 * inch))
        story.append(SectionDivider(width=460, color="#0D47A1", thickness=1.5))
        story.append(Spacer(1, 2))
        story.append(Paragraph(f"{counter}. {title}", self.ps['sec']))
        story.append(Spacer(1, 4))
        renderer(story, content)
        story.append(Spacer(1, 4))
        return counter

    # ════════════════════════════════════════════════════════
    # PDF 1: COACHING SUMMARY
    # ════════════════════════════════════════════════════════

    def _create_coaching_summary_pdf(self, sections: Dict[str, str], transcript_data: TranscriptData):
        pdf_path = os.path.join(self.summaries_folder, "coaching_summary.pdf")
        doc = SimpleDocTemplate(
            pdf_path, pagesize=A4,
            leftMargin=0.6 * inch, rightMargin=0.6 * inch,
            topMargin=0.7 * inch, bottomMargin=0.5 * inch,
            title="Coaching Summary Report",
            author="RAG Coaching Evaluation System",
            subject="Sales Performance Coaching Summary",
        )
        story = []
        self._build_cover_page(story, transcript_data, pdf_title="COACHING SUMMARY REPORT",
                               pdf_subtitle="Actionable Performance Coaching & Improvement Plan")
        story.append(PageBreak())

        counter = 0

        # 1. Executive Brief (LLM §1)
        content = self._find_section_content(sections, 1)
        counter = self._add_section_to_story(story, counter, 'EXECUTIVE BRIEF', self._render_executive_brief, content)

        # 2. Performance Score Dashboard (LLM §2)
        content = self._find_section_content(sections, 2)
        counter = self._add_section_to_story(story, counter, 'PERFORMANCE SCORE DASHBOARD', self._render_score_dashboard_text, content)

        # 3. Strengths & Areas for Improvement (LLM §4 + §5 combined)
        content_strengths = self._find_section_content(sections, 4)
        content_failures = self._find_section_content(sections, 5)
        if content_strengths or content_failures:
            counter += 1
            story.append(CondPageBreak(2.0 * inch))
            story.append(SectionDivider(width=460, color="#0D47A1", thickness=1.5))
            story.append(Spacer(1, 2))
            story.append(Paragraph(f"{counter}. STRENGTHS & AREAS FOR IMPROVEMENT", self.ps['sec']))
            story.append(Spacer(1, 4))
            if content_strengths:
                story.append(Paragraph("<b>A. Evidence-Anchored Strengths</b>", self.ps['subsec']))
                story.append(Spacer(1, 4))
                self._render_strengths(story, content_strengths)
                story.append(Spacer(1, 8))
            if content_failures:
                story.append(Paragraph("<b>B. Missed Opportunities & Critical Failures</b>", self.ps['subsec']))
                story.append(Spacer(1, 4))
                self._render_objections(story, content_failures)
            story.append(Spacer(1, 4))

        # 4. Client Insights (LLM §11 Client Emotional Journey + §12 Communication Balance as subs)
        content_sentiment = self._find_section_content(sections, 11)
        content_listening = self._find_section_content(sections, 12)
        if content_sentiment or content_listening:
            counter += 1
            story.append(CondPageBreak(2.0 * inch))
            story.append(SectionDivider(width=460, color="#0D47A1", thickness=1.5))
            story.append(Spacer(1, 2))
            story.append(Paragraph(f"{counter}. CLIENT INSIGHTS & COMMUNICATION INTELLIGENCE", self.ps['sec']))
            story.append(Spacer(1, 4))
            if content_sentiment:
                story.append(Paragraph("<b>A. Client Emotional Journey & Risk Map</b>", self.ps['subsec']))
                story.append(Spacer(1, 4))
                self._render_client_sentiment(story, content_sentiment)
                story.append(Spacer(1, 8))
            if content_listening:
                story.append(Paragraph("<b>B. Communication Balance & Listening Evaluation</b>", self.ps['subsec']))
                story.append(Spacer(1, 4))
                self._render_listening_intelligence(story, content_listening)
            story.append(Spacer(1, 4))

        # 5. Coaching Recommendations & Tactical Playbook (LLM §6)
        content = self._find_section_content(sections, 6)
        counter = self._add_section_to_story(story, counter, 'COACHING RECOMMENDATIONS & TACTICAL PLAYBOOK', self._render_tactical_playbook_combined, content)

        # 6. 7-Day Intensive Coaching Plan (LLM §7 — merged day-by-day plan with drills)
        content = self._find_section_content(sections, 7)
        counter = self._add_section_to_story(story, counter, '7-DAY INTENSIVE COACHING PLAN', self._render_7day_coaching_plan, content)

        # 7. Deal Intelligence Summary (LLM §9 as main, with Action Items + Decisions §8 as subs)
        content_deal = self._find_section_content(sections, 9)
        content_actions = self._find_section_content(sections, 0, ['Action', 'Items'])
        content_decisions = self._find_section_content(sections, 8)
        if content_deal or content_actions or content_decisions:
            counter += 1
            story.append(CondPageBreak(2.0 * inch))
            story.append(SectionDivider(width=460, color="#0D47A1", thickness=1.5))
            story.append(Spacer(1, 2))
            story.append(Paragraph(f"{counter}. DEAL INTELLIGENCE & ACTION PLAN", self.ps['sec']))
            story.append(Spacer(1, 4))
            if content_deal:
                self._render_deal_intelligence_text(story, content_deal)
                story.append(Spacer(1, 8))
            if content_actions:
                story.append(Paragraph("<b>A. Action Items</b>", self.ps['subsec']))
                story.append(Spacer(1, 4))
                self._render_action_items(story, content_actions)
                story.append(Spacer(1, 8))
            if content_decisions:
                story.append(Paragraph("<b>B. Decisions & Commitments</b>", self.ps['subsec']))
                story.append(Spacer(1, 4))
                self._render_decisions(story, content_decisions)
            story.append(Spacer(1, 4))

        # 8. Deal Outcome / Client Decision (LLM §19)
        content_deal_outcome = self._find_section_content(sections, 19, ['Deal Outcome', 'Client Decision'])
        deal_outcome_obj = self._parse_deal_outcome(content_deal_outcome)
        deal_outcome_obj.agent_name = getattr(self, '_agent_name', 'Agent')
        self._cached_deal_outcome = deal_outcome_obj  # Cache for JSON generation
        if content_deal_outcome and content_deal_outcome.strip():
            counter += 1
            story.append(CondPageBreak(2.0 * inch))
            story.append(SectionDivider(width=460, color="#0D47A1", thickness=1.5))
            story.append(Spacer(1, 2))
            story.append(Paragraph(f"{counter}. DEAL OUTCOME / CLIENT DECISION", self.ps['sec']))
            story.append(Spacer(1, 4))
            self._render_deal_outcome_section(story, deal_outcome_obj)
            story.append(Spacer(1, 4))

        doc.build(story,
                  onFirstPage=self._draw_cover_header_footer,
                  onLaterPages=self._draw_content_header_footer)
        logger.info(f"Saved: {pdf_path}")

    # ════════════════════════════════════════════════════════
    # PDF 2: AGENT PROFILE
    # ════════════════════════════════════════════════════════

    def _create_agent_profile_pdf(self, sections: Dict[str, str], transcript_data: TranscriptData):
        pdf_path = os.path.join(self.summaries_folder, "agent_summary.pdf")
        doc = SimpleDocTemplate(
            pdf_path, pagesize=A4,
            leftMargin=0.6 * inch, rightMargin=0.6 * inch,
            topMargin=0.7 * inch, bottomMargin=0.5 * inch,
            title="Agent Profile & Intelligence Report",
            author="RAG Coaching Evaluation System",
            subject="Agent Intelligence Profile & Advanced Assessment",
        )
        story = []
        self._build_cover_page(story, transcript_data, pdf_title="AGENT PROFILE & INTELLIGENCE REPORT",
                               pdf_subtitle="Advanced Behavioral Intelligence & Compliance Assessment")
        story.append(PageBreak())

        counter = 0

        # 1. Five Core Discovery Pillars Assessment (LLM §13)
        content = self._find_section_content(sections, 13)
        counter = self._add_section_to_story(story, counter, 'FIVE CORE DISCOVERY PILLARS ASSESSMENT', self._render_discovery_matrix, content)

        # 2. Emotional Intelligence Assessment (LLM §14)
        content = self._find_section_content(sections, 14)
        counter = self._add_section_to_story(story, counter, 'EMOTIONAL INTELLIGENCE ASSESSMENT', self._render_emotional_intelligence, content)

        # 3. Ethics & Compliance Audit (LLM §15)
        content = self._find_section_content(sections, 15)
        counter = self._add_section_to_story(story, counter, 'ETHICS & COMPLIANCE AUDIT', self._render_ethics_compliance, content)

        # 4. Self-Awareness & Adaptability Profile (LLM §16)
        content = self._find_section_content(sections, 16)
        counter = self._add_section_to_story(story, counter, 'SELF-AWARENESS & ADAPTABILITY PROFILE', self._render_agent_self_awareness, content)

        # 5. Client Engagement & Tone Analysis (LLM §17 — NEW)
        content = self._find_section_content(sections, 17)
        counter = self._add_section_to_story(story, counter, 'CLIENT ENGAGEMENT & TONE ANALYSIS', self._render_tone_analysis, content)

        # 6. Negotiation & Persuasion Proficiency (LLM §18 — NEW)
        content = self._find_section_content(sections, 18)
        counter = self._add_section_to_story(story, counter, 'NEGOTIATION & PERSUASION PROFICIENCY', self._render_negotiation_proficiency, content)

        # 7. Behavioral Breakdown (LLM §3 — Communication patterns, body language, questioning)
        content = self._find_section_content(sections, 3)
        counter = self._add_section_to_story(story, counter, 'BEHAVIORAL BREAKDOWN & COMMUNICATION PATTERNS', self._render_behavioral_breakdown, content)

        # 8. Agent Tier Classification (LLM §10)
        content = self._find_section_content(sections, 10)
        counter = self._add_section_to_story(story, counter, 'AGENT TIER CLASSIFICATION', self._render_agent_tier, content)

        doc.build(story,
                  onFirstPage=self._draw_cover_header_footer,
                  onLaterPages=self._draw_content_header_footer)
        logger.info(f"Saved: {pdf_path}")

    # ════════════════════════════════════════════════════════
    # PDF 3: COACHING INSIGHTS VISUALIZATION
    # ════════════════════════════════════════════════════════

    def _create_coaching_insights_viz_pdf(self, sections: Dict[str, str], transcript_data: TranscriptData):
        pdf_path = os.path.join(self.summaries_folder, "coaching_insights_viz.pdf")
        doc = SimpleDocTemplate(
            pdf_path, pagesize=A4,
            leftMargin=0.6 * inch, rightMargin=0.6 * inch,
            topMargin=0.7 * inch, bottomMargin=0.5 * inch,
            title="Coaching Insights Visualization",
            author="RAG Coaching Evaluation System",
            subject="Coaching-Focused Data Visualizations & Behavioral Insights",
        )
        story = []
        self._build_cover_page(story, transcript_data,
                               pdf_title="COACHING INSIGHTS VISUALIZATION",
                               pdf_subtitle="Behavioral Analysis, Communication Quality & Coaching Recommendations")
        story.append(PageBreak())

        # Build COACHING-focused visualizations
        self._build_coaching_insights_viz_section(story, sections)

        doc.build(story,
                  onFirstPage=self._draw_cover_header_footer,
                  onLaterPages=self._draw_content_header_footer)
        logger.info(f"Saved: {pdf_path}")

    # ════════════════════════════════════════════════════════
    # PDF 4: AGENT PERFORMANCE VISUALIZATION
    # ════════════════════════════════════════════════════════

    def _create_agent_performance_viz_pdf(self, sections: Dict[str, str], transcript_data: TranscriptData):
        pdf_path = os.path.join(self.summaries_folder, "agent_performance_viz.pdf")
        doc = SimpleDocTemplate(
            pdf_path, pagesize=A4,
            leftMargin=0.6 * inch, rightMargin=0.6 * inch,
            topMargin=0.7 * inch, bottomMargin=0.5 * inch,
            title="Agent Performance Visualization",
            author="RAG Coaching Evaluation System",
            subject="Agent Performance Metrics, Scoring & Evaluation Charts",
        )
        story = []
        self._build_cover_page(story, transcript_data,
                               pdf_title="AGENT PERFORMANCE VISUALIZATION",
                               pdf_subtitle="Quantitative Performance Metrics, Scoring & Evaluation Charts")
        story.append(PageBreak())

        # Build PERFORMANCE-focused visualizations
        self._build_agent_performance_viz_section(story, sections)

        doc.build(story,
                  onFirstPage=self._draw_cover_header_footer,
                  onLaterPages=self._draw_content_header_footer)
        logger.info(f"Saved: {pdf_path}")

    # ── Header / Footer (matching Final.py enterprise style) ──

    def _draw_cover_header_footer(self, canv, doc):
        canv.saveState()
        w, h = A4
        # Top accent bar
        canv.setStrokeColor(self.C['primary_dark'])
        canv.setLineWidth(4)
        canv.line(0, h - 2, w, h - 2)
        # Footer bar
        canv.setFillColor(self.C['primary_dark'])
        canv.rect(0, 0, w, 0.32 * inch, fill=1, stroke=0)
        canv.setFillColor(colors.white)
        canv.setFont("Helvetica", 7.5)
        canv.drawCentredString(w / 2, 0.12 * inch, f"Page {doc.page}")
        canv.drawString(0.6 * inch, 0.12 * inch, f"(c) {datetime.now().year}")
        canv.drawRightString(w - 0.6 * inch, 0.12 * inch, "CONFIDENTIAL")
        canv.restoreState()

    def _draw_content_header_footer(self, canv, doc):
        canv.saveState()
        w, h = A4
        hdr_h = 0.55 * inch
        # Header bar
        canv.setFillColor(self.C['primary_dark'])
        canv.rect(0, h - hdr_h, w, hdr_h, fill=1, stroke=0)
        canv.setFillColor(colors.white)
        canv.setFont("Helvetica-Bold", 10)
        canv.drawString(0.6 * inch, h - 0.35 * inch, "REAL ESTATE AI COACHING")
        canv.setFont("Helvetica", 7)
        canv.drawString(0.6 * inch, h - 0.48 * inch, "Enterprise Performance Intelligence")
        canv.setFont("Helvetica-Bold", 9)
        canv.drawRightString(w - 0.6 * inch, h - 0.35 * inch, "COACHING REPORT")
        canv.setFont("Helvetica", 7)
        canv.drawRightString(w - 0.6 * inch, h - 0.48 * inch, "RAG-Powered Intelligence")
        # Footer bar
        ftr_h = 0.35 * inch
        canv.setFillColor(self.C['primary'])
        canv.rect(0, 0, w, ftr_h, fill=1, stroke=0)
        canv.setFillColor(colors.white)
        canv.setFont("Helvetica", 7.5)
        canv.drawCentredString(w / 2, 0.12 * inch, f"Page {doc.page}")
        canv.drawString(0.6 * inch, 0.12 * inch, f"(c) {datetime.now().year} | {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        canv.drawRightString(w - 0.6 * inch, 0.12 * inch, "CONFIDENTIAL")
        canv.restoreState()

    # ── Cover Page ──

    def _build_cover_page(self, story, transcript_data: TranscriptData, pdf_title: str = "COACHING EVALUATION REPORT", pdf_subtitle: str = "Enterprise AI-Powered Performance Intelligence"):
        story.append(Spacer(1, 1.0 * inch))
        story.append(Paragraph(pdf_title, self.ps['cover_title']))
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph(pdf_subtitle, self.ps['cover_sub']))
        story.append(Spacer(1, 0.6 * inch))

        info = transcript_data.summary_info
        speakers = info.get("unique_speakers", [])
        dur_sec = info.get("total_duration_seconds", 0)
        dur_min = int(dur_sec // 60)
        dur_s = int(dur_sec % 60)
        duration_str = f"{dur_min} minutes {dur_s} seconds" if dur_min > 0 else f"{dur_s} seconds"
        # Agent name: find the speaker who is NOT a client
        # Exclude names that are generic client labels
        client_labels = {'client', 'seller', 'buyer', 'homeowner', 'prospect', 'unknown'}
        non_client_speakers = [s for s in speakers if s.lower() not in client_labels]

        if len(non_client_speakers) == 1:
            agent_name = non_client_speakers[0]
        elif len(non_client_speakers) > 1:
            # Among non-client speakers, pick the one with the most segments (likely the agent)
            speaker_counts = {}
            for seg in transcript_data.transcripts:
                if seg.speaker_name and seg.speaker_name.lower() not in client_labels:
                    speaker_counts[seg.speaker_name] = speaker_counts.get(seg.speaker_name, 0) + 1
            agent_name = max(speaker_counts, key=speaker_counts.get) if speaker_counts else non_client_speakers[0]
        elif len(speakers) == 1:
            agent_name = speakers[0]
        else:
            # Fallback: pick speaker with most segments overall (excluding 'Client' etc.)
            speaker_counts = {}
            for seg in transcript_data.transcripts:
                if seg.speaker_name:
                    speaker_counts[seg.speaker_name] = speaker_counts.get(seg.speaker_name, 0) + 1
            # Remove known client labels and pick the top remaining
            for label in list(speaker_counts.keys()):
                if label.lower() in client_labels:
                    del speaker_counts[label]
            agent_name = max(speaker_counts, key=speaker_counts.get) if speaker_counts else "N/A"

        # Extract client names dynamically from transcript — detect proper nouns addressed by agent
        client_names = set()
        # Collect all unique speaker names that are not the agent
        for seg in transcript_data.transcripts:
            if seg.speaker_name and seg.speaker_name != agent_name and seg.speaker_name != 'Unknown':
                client_names.add(seg.speaker_name)
        # Also scan agent's dialogue for names (capitalized words after addressing patterns)
        name_patterns = re.findall(
            r'(?:So|Hey|Now|And|,)\s+([A-Z][a-z]{2,})(?:[,\s])', 
            ' '.join(seg.transcript for seg in transcript_data.transcripts if seg.speaker_name == agent_name)
        )
        # Filter out common non-name words
        common_words = {'The', 'This', 'That', 'When', 'What', 'Where', 'How', 'Why', 'Now',
                        'Here', 'There', 'Well', 'Yes', 'Yeah', 'Cool', 'Right', 'Just',
                        'Let', 'Got', 'But', 'And', 'Because', 'Matter', 'Number', 'Level',
                        'Royal', 'Okay', 'Does', 'Both', 'First', 'Second', 'Third',
                        'Some', 'Like', 'Also', 'Most', 'Very', 'Sure', 'Homes', 'Home'}
        for name in name_patterns:
            if name not in common_words and name != agent_name:
                client_names.add(name)
        client_display = ", ".join(sorted(client_names)) if client_names else "Client(s)"

        meeting_id = f"COACH_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        rows = [
            ['Report Details', ''],
            ['Meeting ID', meeting_id],
            ['Date', datetime.now().strftime('%B %d, %Y')],
            ['Agenda', 'Real Estate Listing Presentation & Coaching Evaluation'],
            ['Duration', duration_str],
            ['Agent Evaluated', agent_name],
            ['Report Type', 'Enterprise Coaching Evaluation'],
            ['RAG System', 'Active' if self.faiss_index else 'Inactive'],
        ]
        t = Table(rows, colWidths=[1.8 * inch, 4.6 * inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (1, 0), self.C['primary_dark']),
            ('TEXTCOLOR', (0, 0), (1, 0), self.C['white']),
            ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (1, 0), 11),
            ('ALIGN', (0, 0), (1, 0), 'CENTER'),
            ('PADDING', (0, 0), (1, 0), 10),
            ('BACKGROUND', (0, 1), (0, -1), colors.HexColor("#E3F2FD")),
            ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 1), (0, -1), 9),
            ('ALIGN', (0, 1), (0, -1), 'RIGHT'),
            ('PADDING', (0, 1), (-1, -1), 7),
            ('FONTNAME', (1, 1), (1, -1), 'Helvetica'),
            ('FONTSIZE', (1, 1), (1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#BBDEFB")),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.4 * inch))

        # Confidential notice
        story.append(Table(
            [[Paragraph("CONFIDENTIAL  --  FOR COACHING PURPOSES ONLY", self.ps['conf'])]],
            colWidths=[6.2 * inch],
            style=TableStyle([
                ('BOX', (0, 0), (0, 0), 1.5, self.C['danger']),
                ('PADDING', (0, 0), (0, 0), 10),
            ])
        ))

    # ════════════════════════════════════════════════════════
    # SECTION RENDERERS (A through M)
    # ════════════════════════════════════════════════════════

    # ── A. Executive Brief ──
    def _render_executive_brief(self, story, content: str):
        """Strategic diagnostic in a subtle shaded container."""
        inner = []
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
            if line.startswith('-') or line.startswith('*'):
                text = line.lstrip('-* ')
                inner.append(Paragraph(f"<bullet>&bull;</bullet>{text}", self.ps['bullet']))
            else:
                inner.append(Paragraph(line, self.ps['body']))
            inner.append(Spacer(1, 4))

        if inner:
            t = Table([[inner]], colWidths=[6.2 * inch])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('BOX', (0, 0), (-1, -1), 0.5, self.C['border']),
                ('PADDING', (0, 0), (-1, -1), 12),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(t)
            story.append(Spacer(1, 8))

    # ── B. Score Dashboard (TEXT ONLY — chart goes to visualization section) ──
    def _render_score_dashboard_text(self, story, content: str):
        """Parse pipe-delimited scores and render as table. Charts deferred to visualization section."""
        scores = {}
        overall_score = None
        overall_just = ""
        rows = []

        for line in content.split('\n'):
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            parts = [p.strip() for p in line.split('|')]
            if len(parts) >= 4:
                label = parts[0].replace('**', '').strip()
                try:
                    score_val = int(parts[1])
                except ValueError:
                    continue
                max_val = int(parts[2]) if parts[2].isdigit() else 10
                justification = parts[3] if len(parts) > 3 else ""
                confidence = parts[4].strip() if len(parts) > 4 else ""

                # Validate score range (1-10) and clamp if out of range
                if score_val < 1 or score_val > 10:
                    logger.warning(f"Score {score_val} for {label} out of range (1-10). Clamping.")
                    score_val = max(1, min(10, score_val))

                if label.upper() == 'OVERALL':
                    overall_score = score_val
                    overall_just = justification
                else:
                    scores[label] = score_val
                    rows.append((label, score_val, max_val, justification, confidence))

        # Hero overall score display
        if overall_score is not None:
            pct = overall_score / 10
            border_color = "#43A047" if pct >= 0.75 else "#FF8F00" if pct >= 0.50 else "#E53935"
            score_style = ParagraphStyle('score_hero', parent=self.ps['score_big'],
                                         fontSize=36, leading=42, spaceAfter=0, spaceBefore=0)
            label_style = ParagraphStyle('score_lbl', parent=self.ps['score_label'],
                                         fontSize=9, leading=12, spaceBefore=4, spaceAfter=0)
            just_style = ParagraphStyle('score_just', parent=self.ps['body'],
                                        fontSize=8, leading=11, textColor=self.C['medium'],
                                        alignment=TA_CENTER, spaceBefore=2)
            hero = Table([
                [Paragraph(f"<font color='{border_color}'><b>{overall_score}/10</b></font>", score_style)],
                [Spacer(1, 4)],
                [Paragraph("OVERALL PERFORMANCE SCORE", label_style)],
                [Paragraph(f"<i>{overall_just}</i>", just_style)],
            ], colWidths=[6.2 * inch], rowHeights=[48, 4, 16, None])
            hero.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('LINEBELOW', (0, -1), (-1, -1), 1.5, colors.HexColor(border_color)),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
                ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(hero)
            story.append(Spacer(1, 10))

        # Individual scores as table rows
        if rows:
            td_wrap = ParagraphStyle('td_wrap', parent=self.ps['td'], wordWrap='LTR', fontSize=8, leading=11)
            conf_style = ParagraphStyle('td_conf', parent=self.ps['td'], wordWrap='LTR', fontSize=7, leading=10,
                                        textColor=self.C['light'])
            table_data = [[
                Paragraph("<b>Dimension</b>", self.ps['th']),
                Paragraph("<b>Score</b>", self.ps['th']),
                Paragraph("<b>Justification</b>", self.ps['th']),
                Paragraph("<b>AI Confidence</b>", self.ps['th']),
            ]]
            for row_data in rows:
                label, score_val, max_val, just = row_data[0], row_data[1], row_data[2], row_data[3]
                confidence = row_data[4] if len(row_data) > 4 else ""
                pct = score_val / max_val
                sc = "#43A047" if pct >= 0.75 else "#FF8F00" if pct >= 0.50 else "#E53935"
                # Color-code confidence level
                conf_color = "#2E7D32" if 'high' in confidence.lower() else "#F57F17" if 'medium' in confidence.lower() else "#C62828" if 'low' in confidence.lower() else "#757575"
                table_data.append([
                    Paragraph(label, self.ps['td']),
                    Paragraph(f"<b><font color='{sc}'>{score_val}/{max_val}</font></b>", self.ps['td_c']),
                    Paragraph(f"<i>{just}</i>", td_wrap),
                    Paragraph(f"<font color='{conf_color}'><i>{confidence}</i></font>", conf_style),
                ])
            t = Table(table_data, colWidths=[1.4 * inch, 0.5 * inch, 3.3 * inch, 1.0 * inch])
            style_cmds = [
                ('BACKGROUND', (0, 0), (-1, 0), self.C['primary_dark']),
                ('TEXTCOLOR', (0, 0), (-1, 0), self.C['white']),
                ('GRID', (0, 0), (-1, -1), 0.5, self.C['border']),
                ('PADDING', (0, 0), (-1, -1), 5),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
            for i in range(1, len(table_data)):
                if i % 2 == 0:
                    style_cmds.append(('BACKGROUND', (0, i), (-1, i), colors.HexColor("#F5F7FA")))
            t.setStyle(TableStyle(style_cmds))
            story.append(t)

        # Store scores for visualization section
        self._cached_scores = scores

    # ── C. Behavioral Breakdown ──
    def _render_behavioral_breakdown(self, story, content: str):
        content = self._convert_timestamps_to_mmss(content)
        subsections = re.split(r'###\s*(.*)', content)
        # subsections: ['', 'Title1', 'Content1', 'Title2', 'Content2', ...]
        i = 1
        while i < len(subsections) - 1:
            sub_title = subsections[i].strip()
            sub_content = subsections[i + 1].strip()
            story.append(Paragraph(f"<b>{sub_title}</b>", self.ps['subsec']))
            for line in sub_content.split('\n'):
                line = line.strip()
                if not line:
                    continue
                # Highlight timestamps
                line = re.sub(r'\[(\d+[\.\d]*s?\s*-\s*\d+[\.\d]*s?)\]',
                              r"<font color='#1565C0'><b>[\1]</b></font>", line)
                line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
                if line.startswith('-') or line.startswith('*'):
                    text = line.lstrip('-* ')
                    story.append(Paragraph(f"<bullet>&bull;</bullet>{text}", self.ps['bullet']))
                else:
                    story.append(Paragraph(line, self.ps['body']))
                story.append(Spacer(1, 2))
            i += 2

        # Fallback: if no ### subheadings found, look for **bold headers** or numbered patterns
        if len(subsections) <= 1:
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                # Detect bold header lines (e.g. "**Communication Patterns**" alone on a line)
                bold_match = re.match(r'^\*\*(.*?)\*\*\s*$', line)
                if bold_match:
                    story.append(Spacer(1, 4))
                    story.append(Paragraph(f"<b>{bold_match.group(1)}</b>", self.ps['subsec']))
                    continue
                line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
                line = re.sub(r'\[(\d+[\.\d]*s?\s*-\s*\d+[\.\d]*s?)\]',
                              r"<font color='#1565C0'><b>[\1]</b></font>", line)
                if line.startswith('-') or line.startswith('*'):
                    text = line.lstrip('-* ')
                    story.append(Paragraph(f"<bullet>&bull;</bullet>{text}", self.ps['bullet']))
                else:
                    story.append(Paragraph(line, self.ps['body']))
                story.append(Spacer(1, 2))

    # ── D. Strengths ──
    def _render_strengths(self, story, content: str):
        """Parse structured STRENGTH/TIMESTAMP/EVIDENCE/IMPACT entries into a single professional table."""
        content = self._convert_timestamps_to_mmss(content)
        td_s = ParagraphStyle('td_strength', parent=self.ps['td'], fontSize=8, leading=11,
                              textColor=self.C['dark'], wordWrap='LTR')
        td_ev = ParagraphStyle('td_evidence', parent=self.ps['td'], fontSize=7.5, leading=10,
                               textColor=self.C['medium'], fontName='Helvetica-Oblique', wordWrap='LTR')
        td_imp = ParagraphStyle('td_impact', parent=self.ps['td'], fontSize=7.5, leading=10,
                                textColor=colors.HexColor('#2E7D32'), wordWrap='LTR')

        # Parse structured entries
        entries = []
        current_entry = {}
        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                if current_entry:
                    entries.append(current_entry)
                    current_entry = {}
                continue
            matched_field = False
            for label in ['STRENGTH', 'TIMESTAMP', 'EVIDENCE', 'IMPACT', 'WHY IT MATTERS']:
                if stripped.upper().startswith(label + ':'):
                    value = stripped[len(label) + 1:].strip().strip('"')
                    if label == 'STRENGTH' and current_entry:
                        entries.append(current_entry)
                        current_entry = {}
                    current_entry[label.upper()] = value
                    matched_field = True
                    break
            if not matched_field:
                # Treat as bullet/extra content
                cleaned = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
                if cleaned.startswith('-') or cleaned.startswith('*'):
                    cleaned = cleaned.lstrip('-* ')
                if not current_entry:
                    current_entry['STRENGTH'] = cleaned
                elif 'EVIDENCE' not in current_entry:
                    current_entry.setdefault('STRENGTH', '')
                    current_entry['STRENGTH'] += ' ' + cleaned
        if current_entry:
            entries.append(current_entry)

        # Build proper table with columns: #, Strength, Timestamp, Evidence, Impact
        if entries and any('STRENGTH' in e for e in entries):
            rows = [[
                Paragraph("<b>#</b>", self.ps['th']),
                Paragraph("<b>Strength</b>", self.ps['th']),
                Paragraph("<b>When</b>", self.ps['th']),
                Paragraph("<b>Evidence</b>", self.ps['th']),
                Paragraph("<b>Impact</b>", self.ps['th']),
            ]]
            entry_num = 0
            for entry in entries:
                if not entry.get('STRENGTH'):
                    continue
                entry_num += 1
                name = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', entry.get('STRENGTH', ''))
                ts = entry.get('TIMESTAMP', '--')
                evidence = entry.get('EVIDENCE', '--').strip('"')
                impact = entry.get('IMPACT', entry.get('WHY IT MATTERS', '--'))
                rows.append([
                    Paragraph(f"<b>{entry_num}</b>", self.ps['td_c']),
                    Paragraph(f"<b>{name}</b>", td_s),
                    Paragraph(f"<font color='#1565C0'>{ts}</font>", self.ps['td_c']),
                    Paragraph(f'<i>"{evidence}"</i>' if evidence != '--' else '--', td_ev),
                    Paragraph(impact, td_imp),
                ])
            t = Table(rows, colWidths=[0.35 * inch, 1.55 * inch, 0.85 * inch, 2.05 * inch, 1.4 * inch])
            style_cmds = [
                ('BACKGROUND', (0, 0), (-1, 0), self.C['success']),
                ('TEXTCOLOR', (0, 0), (-1, 0), self.C['white']),
                ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#C8E6C9')),
                ('TOPPADDING', (0, 0), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ALIGN', (0, 0), (0, -1), 'CENTER'),
                ('ALIGN', (2, 0), (2, -1), 'CENTER'),
            ]
            for i in range(1, len(rows)):
                if i % 2 == 0:
                    style_cmds.append(('BACKGROUND', (0, i), (-1, i), self.C['success_bg']))
            t.setStyle(TableStyle(style_cmds))
            story.append(t)
        else:
            # Fallback: parse as flat bullets into a simple 2-column table
            bullet_items = []
            for line in content.split('\n'):
                stripped = line.strip()
                if not stripped:
                    continue
                stripped = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', stripped)
                if stripped.startswith('-') or stripped.startswith('*'):
                    bullet_items.append(stripped.lstrip('-* '))
                elif not any(stripped.upper().startswith(l + ':') for l in ['STRENGTH','TIMESTAMP','EVIDENCE','IMPACT']):
                    bullet_items.append(stripped)
            if bullet_items:
                rows = [[Paragraph("<b>#</b>", self.ps['th']),
                         Paragraph("<b>Identified Strength</b>", self.ps['th'])]]
                for i, item in enumerate(bullet_items, 1):
                    rows.append([
                        Paragraph(f"<b>{i}</b>", self.ps['td_c']),
                        Paragraph(item, td_s),
                    ])
                t = Table(rows, colWidths=[0.4 * inch, 5.8 * inch])
                style_cmds = [
                    ('BACKGROUND', (0, 0), (-1, 0), self.C['success']),
                    ('TEXTCOLOR', (0, 0), (-1, 0), self.C['white']),
                    ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#C8E6C9')),
                    ('TOPPADDING', (0, 0), (-1, -1), 5),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('ALIGN', (0, 1), (0, -1), 'CENTER'),
                ]
                for i in range(1, len(rows)):
                    if i % 2 == 0:
                        style_cmds.append(('BACKGROUND', (0, i), (-1, i), self.C['success_bg']))
                t.setStyle(TableStyle(style_cmds))
                story.append(t)

    # ── E. Missed Opportunities & Critical Failures ──
    def _render_objections(self, story, content: str):
        """Render failures as professional structured cards with clear field labels."""
        content = self._convert_timestamps_to_mmss(content)
        entries = re.split(r'\n\s*\n', content)
        entry_num = 0
        for entry in entries:
            entry = entry.strip()
            if not entry:
                continue
            fields = {}
            extra_lines = []
            for line in entry.split('\n'):
                line = line.strip()
                if not line:
                    continue
                matched = False
                for label in ['FAILURE TITLE', 'TIMESTAMP', 'QUOTE', 'WHAT HAPPENED',
                              'WHAT A TOP AGENT WOULD DO', 'WHAT WAS MISSED',
                              'REVENUE IMPACT', 'DEAL IMPACT', 'BUSINESS IMPACT']:
                    if line.upper().startswith(label + ':'):
                        value = line[len(label) + 1:].strip()
                        fields[label.upper()] = value
                        matched = True
                        break
                if not matched:
                    extra_lines.append(line)

            if not fields and not extra_lines:
                continue
            entry_num += 1

            # Build structured card
            title = fields.get('FAILURE TITLE', f'Issue {entry_num}')
            title = re.sub(r'\*\*(.*?)\*\*', r'\1', title)

            rows = []
            # Title row spanning full width
            rows.append([Paragraph(
                f"<b>Failure {entry_num}: {title}</b>",
                ParagraphStyle('ft', parent=self.ps['subsec'],
                               textColor=self.C['danger'], fontSize=11)),
                Paragraph('', self.ps['td'])])

            if 'TIMESTAMP' in fields:
                rows.append([
                    Paragraph("<b>When</b>", self.ps['td']),
                    Paragraph(f"<font color='#1565C0'><b>{fields['TIMESTAMP']}</b></font>", self.ps['td'])
                ])
            if 'QUOTE' in fields:
                q = fields['QUOTE'].strip('"').strip("'")
                rows.append([
                    Paragraph("<b>Agent Said</b>", self.ps['td']),
                    Paragraph(f'<i>"{q}"</i>', ParagraphStyle(
                        'fq', parent=self.ps['td'], textColor=self.C['medium'],
                        fontName='Helvetica-Oblique'))
                ])
            if 'WHAT HAPPENED' in fields:
                rows.append([
                    Paragraph("<b>What Went Wrong</b>", self.ps['td']),
                    Paragraph(fields['WHAT HAPPENED'], self.ps['td'])
                ])
            missed = fields.get('WHAT A TOP AGENT WOULD DO', fields.get('WHAT WAS MISSED', ''))
            if missed:
                rows.append([
                    Paragraph("<b>Correct Approach</b>", self.ps['td']),
                    Paragraph(f"<font color='#2E7D32'>{missed}</font>", self.ps['td'])
                ])
            impact = fields.get('REVENUE IMPACT', fields.get('DEAL IMPACT', fields.get('BUSINESS IMPACT', '')))
            if impact:
                rows.append([
                    Paragraph("<b>Revenue Impact</b>", self.ps['td']),
                    Paragraph(f"<font color='#C62828'><b>{impact}</b></font>", self.ps['td'])
                ])
            for el in extra_lines:
                el = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', el)
                if el.startswith('-') or el.startswith('*'):
                    rows.append([Paragraph('', self.ps['td']),
                                 Paragraph(f"<bullet>&bull;</bullet>{el.lstrip('-* ')}", self.ps['bullet'])])
                else:
                    rows.append([Paragraph('', self.ps['td']), Paragraph(el, self.ps['td'])])

            if rows:
                t = Table(rows, colWidths=[1.3 * inch, 4.9 * inch])
                style_cmds = [
                    ('BACKGROUND', (0, 0), (-1, 0), self.C['danger_bg']),
                    ('BACKGROUND', (0, 1), (0, -1), colors.HexColor('#FFF5F5')),
                    ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#FFCDD2')),
                    ('PADDING', (0, 0), (-1, -1), 6),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('SPAN', (0, 0), (1, 0)),
                ]
                t.setStyle(TableStyle(style_cmds))
                story.append(t)
                story.append(Spacer(1, 4))

    # ── F. Tactical Corrections & Coaching Playbook (Combined) ──
    def _render_tactical_playbook_combined(self, story, content: str):
        """Unified playbook: tactical corrections + coaching drills in structured cards."""
        content = self._convert_timestamps_to_mmss(content)
        entries = re.split(r'\n\s*\n', content)
        entry_num = 0

        for entry in entries:
            entry = entry.strip()
            if not entry:
                continue

            # Parse structured fields
            fields = {}
            extra_lines = []
            for line in entry.split('\n'):
                line = line.strip()
                if not line:
                    continue
                matched = False
                for label in ['SCENARIO', 'CLIENT TRIGGER', 'RISK', 'ASSERTIVE APPROACH',
                              'CONSULTATIVE APPROACH', 'ASSERTIVE FRAME', 'CONSULTATIVE FRAME',
                              'CORRECT FRAME', 'WHY IT WORKS', 'ETHICS-SAFE NOTE']:
                    if line.upper().startswith(label + ':'):
                        fields[label.upper()] = line[len(label) + 1:].strip().strip('"')
                        matched = True
                        break
                if not matched:
                    extra_lines.append(line)

            if fields:
                entry_num += 1
                rows = []
                scenario = fields.get('SCENARIO', '')
                if scenario:
                    rows.append([
                        Paragraph(f"<b>Playbook {entry_num}: {scenario}</b>",
                                  ParagraphStyle('pb_t', parent=self.ps['subsec'],
                                                 textColor=self.C['primary_dark'], fontSize=10)),
                        Paragraph('', self.ps['td'])
                    ])
                if 'CLIENT TRIGGER' in fields:
                    rows.append([Paragraph("<b>Client Trigger</b>", self.ps['td']),
                                 Paragraph(f'<i>"{fields["CLIENT TRIGGER"]}"</i>', self.ps['td'])])
                if 'RISK' in fields:
                    rows.append([Paragraph("<b>Risk</b>", self.ps['td']),
                                 Paragraph(f"<font color='#C62828'>{fields['RISK']}</font>", self.ps['td'])])
                assertive = fields.get('ASSERTIVE APPROACH', fields.get('ASSERTIVE FRAME', fields.get('CORRECT FRAME', '')))
                if assertive:
                    rows.append([Paragraph("<b>Assertive Script</b>", self.ps['td']),
                                 Paragraph(f"<font color='#2E7D32'><i>\"{assertive}\"</i></font>", self.ps['td'])])
                consultative = fields.get('CONSULTATIVE APPROACH', fields.get('CONSULTATIVE FRAME', ''))
                if consultative:
                    rows.append([Paragraph("<b>Consultative Script</b>", self.ps['td']),
                                 Paragraph(f"<font color='#1565C0'><i>\"{consultative}\"</i></font>", self.ps['td'])])
                if 'WHY IT WORKS' in fields:
                    rows.append([Paragraph("<b>Why It Works</b>", self.ps['td']),
                                 Paragraph(f"<i>{fields['WHY IT WORKS']}</i>", self.ps['td'])])
                if 'ETHICS-SAFE NOTE' in fields:
                    rows.append([Paragraph("<b>Ethics Note</b>", self.ps['td']),
                                 Paragraph(f"<font color='#F57F17'><i>{fields['ETHICS-SAFE NOTE']}</i></font>", self.ps['td'])])

                if rows:
                    t = Table(rows, colWidths=[1.3 * inch, 4.9 * inch])
                    style_cmds = [
                        ('GRID', (0, 0), (-1, -1), 0.5, self.C['border']),
                        ('PADDING', (0, 0), (-1, -1), 6),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#E3F2FD')),
                    ]
                    if scenario:
                        style_cmds.append(('SPAN', (0, 0), (1, 0)))
                        style_cmds.append(('BACKGROUND', (0, 0), (1, 0), colors.HexColor('#E8EAF6')))
                    if entry_num % 2 == 0:
                        style_cmds.append(('BACKGROUND', (1, 1), (1, -1), colors.HexColor('#FAFAFA')))
                    t.setStyle(TableStyle(style_cmds))
                    story.append(t)
                    story.append(Spacer(1, 6))
            else:
                # Fallback: render as formatted text
                text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', entry)
                story.append(Paragraph(text, self.ps['body']))
                story.append(Spacer(1, 3))

    # ── H. 7-Day Focus ──
    def _render_7day_focus(self, story, content: str):
        content = self._convert_timestamps_to_mmss(content)
        story.append(Paragraph("<b>For the next 7 days, focus on ONLY these areas:</b>",
                               ParagraphStyle('focus_hdr', parent=self.ps['body_bold'],
                                              textColor=self.C['danger'], fontSize=10)))
        story.append(Spacer(1, 6))
        for line in content.split('\n'):
            line = line.strip()
            m = re.match(r'^(\d+)\.\s*(.*)', line)
            if m:
                num, text = m.group(1), m.group(2)
                text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
                t = Table([[
                    Paragraph(f"<b>{num}</b>", ParagraphStyle('fn', parent=self.ps['td_c'],
                              fontSize=14, textColor=self.C['white'])),
                    Paragraph(text, self.ps['body'])
                ]], colWidths=[0.4 * inch, 5.8 * inch])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), self.C['primary_dark']),
                    ('BACKGROUND', (1, 0), (1, 0), colors.HexColor("#F5F7FA")),
                    ('BOX', (0, 0), (-1, -1), 0.5, self.C['border']),
                    ('PADDING', (0, 0), (-1, -1), 8),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                story.append(t)
                story.append(Spacer(1, 4))

    # ── I. 7-Week Development Roadmap ──
    def _render_7week_roadmap(self, story, content: str):
        content = self._convert_timestamps_to_mmss(content)
        rows = [[
            Paragraph("<b>Week</b>", self.ps['th']),
            Paragraph("<b>Focus Area</b>", self.ps['th']),
            Paragraph("<b>Drill / Exercise</b>", self.ps['th']),
        ]]
        for line in content.split('\n'):
            line = line.strip()
            m = re.match(r'Week\s*(\d+)\s*:\s*(.*?)(?:\s*--\s*|\s*[-:]\s*)(.*)', line, re.IGNORECASE)
            if m:
                wk, focus, drill = m.group(1), m.group(2).strip(), m.group(3).strip()
                rows.append([
                    Paragraph(f"<b>{wk}</b>", self.ps['td_c']),
                    Paragraph(focus, self.ps['td']),
                    Paragraph(drill, self.ps['td']),
                ])

        if len(rows) > 1:
            t = Table(rows, colWidths=[0.5 * inch, 2.5 * inch, 3.2 * inch])
            style_cmds = [
                ('BACKGROUND', (0, 0), (-1, 0), self.C['primary_dark']),
                ('TEXTCOLOR', (0, 0), (-1, 0), self.C['white']),
                ('GRID', (0, 0), (-1, -1), 0.5, self.C['border']),
                ('PADDING', (0, 0), (-1, -1), 6),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
            for i in range(1, len(rows)):
                if i % 2 == 0:
                    style_cmds.append(('BACKGROUND', (0, i), (-1, i), colors.HexColor("#F5F7FA")))
            t.setStyle(TableStyle(style_cmds))
            story.append(t)
        else:
            story.append(Paragraph(content, self.ps['body']))
        story.append(Spacer(1, 8))

    # ── J. Decisions & Commitments ──
    def _render_decisions(self, story, content: str):
        current = None
        for line in content.split('\n'):
            stripped = line.strip()
            if 'Decisions Made:' in stripped or 'Decisions:' in stripped:
                story.append(Paragraph("<b>Decisions Made</b>", self.ps['subsec']))
                current = 'decisions'
            elif 'Commitments Given:' in stripped or 'Commitments:' in stripped:
                story.append(Spacer(1, 6))
                story.append(Paragraph("<b>Commitments Given</b>", self.ps['subsec']))
                current = 'commitments'
            elif stripped.startswith('-') or stripped.startswith('*'):
                text = stripped.lstrip('-* ').strip()
                text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
                style = self.ps['bullet_success'] if current == 'decisions' else self.ps['bullet_accent']
                story.append(Paragraph(f"<bullet>&bull;</bullet>{text}", style))
                story.append(Spacer(1, 2))
            elif stripped:
                story.append(Paragraph(re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', stripped), self.ps['body']))

    # ── Action Items ──
    def _render_action_items(self, story, content: str):
        """Parse action items table from LLM output."""
        td_wrap = ParagraphStyle('td_act', parent=self.ps['td'], wordWrap='LTR', fontSize=8, leading=11)
        rows = [[
            Paragraph("<b>Task</b>", self.ps['th']),
            Paragraph("<b>Responsible</b>", self.ps['th']),
            Paragraph("<b>Deadline</b>", self.ps['th']),
            Paragraph("<b>Priority</b>", self.ps['th']),
            Paragraph("<b>Status</b>", self.ps['th']),
        ]]
        for line in content.split('\n'):
            line = line.strip()
            if not line or line.startswith('|---') or line.startswith('| Task') or line.startswith('|Task'):
                continue
            # Parse pipe-delimited table rows
            if '|' in line:
                parts = [p.strip() for p in line.split('|') if p.strip()]
                if len(parts) >= 4:
                    task = parts[0].replace('**', '')
                    responsible = parts[1] if len(parts) > 1 else ''
                    if responsible == 'Agent':
                        responsible = getattr(self, '_agent_name', 'Agent')
                    responsible = ' '.join(dict.fromkeys(responsible.split()))
                    deadline = parts[2] if len(parts) > 2 else ''
                    priority = parts[3] if len(parts) > 3 else 'Medium'
                    status = parts[4] if len(parts) > 4 else 'Pending'
                    rows.append([
                        Paragraph(task, td_wrap),
                        Paragraph(responsible, self.ps['td_c']),
                        Paragraph(deadline, td_wrap),
                        Paragraph(priority, self.ps['td_c']),
                        Paragraph(status, self.ps['td_c']),
                    ])
            elif line.startswith('-') or line.startswith('*'):
                text = line.lstrip('-* ').strip()
                rows.append([
                    Paragraph(text, td_wrap),
                    Paragraph('', self.ps['td_c']),
                    Paragraph('', td_wrap),
                    Paragraph('Medium', self.ps['td_c']),
                    Paragraph('Pending', self.ps['td_c']),
                ])

        if len(rows) > 1:
            t = Table(rows, colWidths=[2.2 * inch, 1.0 * inch, 1.2 * inch, 0.8 * inch, 0.8 * inch])
            style_cmds = [
                ('BACKGROUND', (0, 0), (-1, 0), self.C['primary_dark']),
                ('TEXTCOLOR', (0, 0), (-1, 0), self.C['white']),
                ('GRID', (0, 0), (-1, -1), 0.5, self.C['border']),
                ('PADDING', (0, 0), (-1, -1), 5),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
            for i in range(1, len(rows)):
                if i % 2 == 0:
                    style_cmds.append(('BACKGROUND', (0, i), (-1, i), colors.HexColor("#F5F7FA")))
            t.setStyle(TableStyle(style_cmds))
            story.append(t)
        else:
            # Fallback: render as text
            for line in content.split('\n'):
                line = line.strip()
                if line:
                    line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
                    story.append(Paragraph(line, self.ps['body']))
                    story.append(Spacer(1, 2))

    # ── K. Deal Intelligence (TEXT ONLY — chart goes to visualization section) ──
    def _render_deal_intelligence_text(self, story, content: str):
        """Parse pipe-delimited metrics and next steps. Hero + cards rendered here; chart deferred."""
        content = self._convert_timestamps_to_mmss(content)
        metrics = []
        next_steps = []
        in_next_steps = False
        deal_prob = None

        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            if 'Recommended Next Steps' in stripped:
                in_next_steps = True
                continue
            if in_next_steps:
                m = re.match(r'^\d+\.\s*(.*)', stripped)
                if m:
                    next_steps.append(m.group(1))
                elif stripped.startswith('-'):
                    next_steps.append(stripped.lstrip('- '))
                continue

            parts = [p.strip() for p in stripped.split('|')]
            if len(parts) >= 3:
                label = parts[0].replace('**', '')
                value = parts[1]
                just = parts[2] if len(parts) > 2 else ""
                confidence = parts[3].strip() if len(parts) > 3 else ""
                metrics.append((label, value, just, confidence))
                if 'probability' in label.lower():
                    deal_prob = value

        # Store for visualization
        self._cached_deal_metrics = metrics

        # Deal probability hero
        if deal_prob:
            try:
                prob_num = int(deal_prob.replace('%', '').strip())
            except ValueError:
                prob_num = 50
            color = "#43A047" if prob_num >= 70 else "#FF8F00" if prob_num >= 50 else "#E53935"
            prob_style = ParagraphStyle('dp_hero', parent=self.ps['score_big'],
                                        fontSize=40, leading=46, spaceAfter=0, spaceBefore=0)
            lbl_style = ParagraphStyle('dp_lbl', parent=self.ps['score_label'],
                                       fontSize=9, leading=12, spaceBefore=4, spaceAfter=0)
            hero = Table([
                [Paragraph(f"<font color='{color}'><b>{deal_prob}</b></font>", prob_style)],
                [Spacer(1, 4)],
                [Paragraph("OVERALL DEAL PROBABILITY", lbl_style)],
                [Paragraph("<i>AI-estimated confidence score based on observed signals</i>",
                          ParagraphStyle('dp_conf', parent=self.ps['conf'], fontSize=7,
                                         textColor=self.C['light']))],
            ], colWidths=[6.2 * inch], rowHeights=[52, 4, 18, 14])
            hero.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('LINEBELOW', (0, -1), (-1, -1), 1.5, colors.HexColor(color)),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
                ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(hero)
            story.append(Spacer(1, 10))

        # Metric cards in 2-column grid
        non_prob = []
        for m_item in metrics:
            if 'probability' not in m_item[0].lower():
                non_prob.append(m_item)
        if non_prob:
            row_cells = []
            for m_item in non_prob:
                label = m_item[0]
                value = m_item[1]
                just = m_item[2] if len(m_item) > 2 else ""
                confidence = m_item[3] if len(m_item) > 3 else ""
                val_lower = value.lower()
                if val_lower in ['strong', 'high'] and 'risk' not in label.lower():
                    bg, tc = "#E8F5E9", "#2E7D32"
                elif val_lower in ['weak', 'low'] and 'risk' not in label.lower():
                    bg, tc = "#FFEBEE", "#C62828"
                elif val_lower in ['high'] and 'risk' in label.lower():
                    bg, tc = "#FFEBEE", "#C62828"
                elif val_lower in ['low'] and 'risk' in label.lower():
                    bg, tc = "#E8F5E9", "#2E7D32"
                else:
                    bg, tc = "#FFF8E1", "#F57F17"

                cell_content = [
                    Paragraph(f"<b>{value}</b>", ParagraphStyle('mv', parent=self.ps['body_bold'],
                              fontSize=14, textColor=colors.HexColor(tc), alignment=TA_CENTER)),
                    Spacer(1, 2),
                    Paragraph(label, ParagraphStyle('ml', parent=self.ps['body'],
                              fontSize=8, textColor=self.C['light'], alignment=TA_CENTER)),
                ]
                if just:
                    cell_content.append(Spacer(1, 2))
                    cell_content.append(Paragraph(f"<i>{just}</i>", ParagraphStyle(
                        'mj', parent=self.ps['body'], fontSize=7, textColor=self.C['light'], alignment=TA_CENTER)))
                if confidence:
                    cell_content.append(Spacer(1, 1))
                    cell_content.append(Paragraph(f"<i>{confidence}</i>", ParagraphStyle(
                        'mc', parent=self.ps['body'], fontSize=6, textColor=self.C['light'], alignment=TA_CENTER)))

                card = Table([[cell_content]], colWidths=[2.9 * inch])
                card.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(bg)),
                    ('PADDING', (0, 0), (-1, -1), 8),
                ]))
                row_cells.append(card)

            for i in range(0, len(row_cells), 2):
                pair = row_cells[i:i + 2]
                if len(pair) == 1:
                    pair.append(Spacer(1, 1))
                grid = Table([pair], colWidths=[3.1 * inch, 3.1 * inch])
                grid.setStyle(TableStyle([('PADDING', (0, 0), (-1, -1), 4)]))
                story.append(grid)
                story.append(Spacer(1, 4))

        # Next Steps
        if next_steps:
            story.append(Spacer(1, 6))
            story.append(Paragraph("<b>Recommended Next Steps</b>", self.ps['subsec']))
            for i, step in enumerate(next_steps, 1):
                story.append(Paragraph(f"<b>{i}.</b> {step}", self.ps['body']))
                story.append(Spacer(1, 2))

    # ── K2. Deal Outcome / Client Decision ──

    def _parse_deal_outcome(self, content: str) -> DealOutcomeResult:
        """Parse Section 19 deal outcome into a structured DealOutcomeResult."""
        result = DealOutcomeResult(generated_at=datetime.now().isoformat())

        if not content or not content.strip():
            return result

        status = ""
        confidence = ""
        evidence = ""
        quotes = []
        in_quotes = False

        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            upper = stripped.upper()

            if upper.startswith('DEAL_STATUS') or upper.startswith('DEAL STATUS'):
                status = re.sub(r'^DEAL[_ ]?STATUS\s*[:\-]\s*', '', stripped, flags=re.IGNORECASE).strip()
                in_quotes = False
            elif upper.startswith('CONFIDENCE'):
                confidence = re.sub(r'^CONFIDENCE\s*[:\-]\s*', '', stripped, flags=re.IGNORECASE).strip().lower()
                in_quotes = False
            elif upper.startswith('EVIDENCE'):
                evidence = re.sub(r'^EVIDENCE\s*[:\-]\s*', '', stripped, flags=re.IGNORECASE).strip()
                in_quotes = False
            elif upper.startswith('SUPPORTING_QUOTES') or upper.startswith('SUPPORTING QUOTES'):
                in_quotes = True
            elif in_quotes and (stripped.startswith('-') or stripped.startswith('"') or stripped.startswith('\u201c')):
                quote = stripped.lstrip('- ').strip().strip('""\u201c\u201d\'')
                if quote:
                    quotes.append(quote)

        # Normalize status to canonical values
        status_lower = status.lower()
        if any(kw in status_lower for kw in ['closed', 'signed', 'agreed', 'committed']):
            result.status = "Deal Closed"
        elif any(kw in status_lower for kw in ['declined', 'refused', 'rejected', 'chose another']):
            result.status = "Client Declined the Deal"
        elif any(kw in status_lower for kw in ['requested time', 'think about', 'get back', 'will see',
                                                  'postpone', 'discuss', 'reach out', 'time to decide']):
            result.status = "Client Requested Time to Decide"
        elif status:
            result.status = status  # Use LLM's exact wording if it doesn't match patterns
        # else keep default: "Decision Pending / Not Clearly Stated"

        if confidence in ('high', 'medium', 'low'):
            result.confidence = confidence
        elif confidence:
            result.confidence = confidence

        if evidence:
            result.evidence = evidence
        result.supporting_quotes = quotes

        return result

    def _render_deal_outcome_section(self, story, deal_outcome: DealOutcomeResult):
        """Render the Deal Outcome / Client Decision section in the PDF."""

        # Status color mapping
        status = deal_outcome.status
        status_lower = status.lower()
        if 'closed' in status_lower:
            status_color = "#2E7D32"
            status_bg = "#E8F5E9"
            status_icon = "\u2713"  # checkmark
        elif 'declined' in status_lower or 'refused' in status_lower:
            status_color = "#C62828"
            status_bg = "#FFEBEE"
            status_icon = "\u2717"  # cross
        elif 'requested time' in status_lower or 'time to decide' in status_lower:
            status_color = "#F57F17"
            status_bg = "#FFF8E1"
            status_icon = "\u23F3"  # hourglass
        else:
            status_color = "#546E7A"
            status_bg = "#ECEFF1"
            status_icon = "\u2014"  # em-dash

        # Confidence color
        conf = deal_outcome.confidence.lower()
        if conf == 'high':
            conf_color = "#2E7D32"
            conf_label = "HIGH"
        elif conf == 'medium':
            conf_color = "#F57F17"
            conf_label = "MEDIUM"
        else:
            conf_color = "#E53935"
            conf_label = "LOW"

        # Hero status display
        status_style = ParagraphStyle('deal_status_hero', parent=self.ps['score_big'],
                                       fontSize=18, leading=24, spaceAfter=0, spaceBefore=0,
                                       textColor=colors.HexColor(status_color))
        conf_style = ParagraphStyle('deal_conf', parent=self.ps['conf'], fontSize=8,
                                     textColor=colors.HexColor(conf_color))

        hero = Table([
            [Paragraph(f"<b>{status}</b>", status_style)],
            [Spacer(1, 4)],
            [Paragraph(f"Confidence: <b>{conf_label}</b>", conf_style)],
        ], colWidths=[6.2 * inch], rowHeights=[30, 4, 16])
        hero.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(status_bg)),
            ('LINEBELOW', (0, -1), (-1, -1), 2, colors.HexColor(status_color)),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(hero)
        story.append(Spacer(1, 10))

        # Evidence summary
        if deal_outcome.evidence and deal_outcome.evidence != "Insufficient transcript evidence to determine deal outcome.":
            story.append(Paragraph("<b>Evidence Summary</b>", self.ps['subsec']))
            story.append(Spacer(1, 4))
            evidence_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', deal_outcome.evidence)
            story.append(Paragraph(evidence_text, self.ps['body']))
            story.append(Spacer(1, 6))

        # Supporting quotes
        if deal_outcome.supporting_quotes:
            story.append(Paragraph("<b>Supporting Quotes from Transcript</b>", self.ps['subsec']))
            story.append(Spacer(1, 4))
            for quote in deal_outcome.supporting_quotes:
                story.append(Paragraph(f'<i>"{quote}"</i>', self.ps['quote']))
                story.append(Spacer(1, 3))

    # ── L. Comparative Performance Context ──
    def _render_comparative(self, story, content: str):
        td_wrap = ParagraphStyle('td_comp', parent=self.ps['td'], wordWrap='LTR', fontSize=8.5, leading=11.5)
        rows = [[
            Paragraph("<b>Metric</b>", self.ps['th']),
            Paragraph("<b>Baseline</b>", self.ps['th']),
            Paragraph("<b>Trend / Insight</b>", self.ps['th']),
        ]]
        note_text = ""
        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.lower().startswith('note:'):
                note_text = stripped[5:].strip()
                continue
            parts = [p.strip() for p in stripped.split('|')]
            if len(parts) >= 2:
                label = parts[0].replace('**', '').strip()
                value = parts[1].strip()
                trend = parts[2].strip() if len(parts) >= 3 else '--'

                # Enrich baseline value with context if it's a bare number
                display_value = value
                if re.match(r'^\d+\.?\d*$', value):
                    num = float(value)
                    label_lower = label.lower()
                    if num <= 10:
                        # Likely a /10 score
                        display_value = f"{value}/10"
                    elif num <= 100 and any(w in label_lower for w in ['rate', 'ratio', 'percent', '%', 'resolution', 'clarity']):
                        display_value = f"{value}%"
                    else:
                        display_value = value
                elif re.match(r'^\d+%$', value):
                    display_value = value  # already has %
                elif '/' in value:
                    display_value = value  # already has /N format

                # Color-code the baseline value using central method
                try:
                    num_val = float(re.sub(r'[^\d.]', '', value))
                    if num_val <= 10:
                        val_color = self._get_score_color(num_val, 10)
                    else:
                        val_color = self._get_percentage_color(num_val)
                except (ValueError, ZeroDivisionError):
                    val_color = "#757575"

                # Determine trend color indicator
                trend_lower = trend.lower()
                if any(w in trend_lower for w in ['below', 'weak', 'poor', 'low', 'needs', 'decline']):
                    trend_color = "#E53935"
                elif any(w in trend_lower for w in ['above', 'strong', 'good', 'high', 'exceed']):
                    trend_color = "#2E7D32"
                elif any(w in trend_lower for w in ['moderate', 'average', 'room']):
                    trend_color = "#F57F17"
                else:
                    trend_color = "#757575"

                rows.append([
                    Paragraph(f"<b>{label}</b>", td_wrap),
                    Paragraph(f"<b><font color='{val_color}'>{display_value}</font></b>", self.ps['td_c']),
                    Paragraph(f"<font color='{trend_color}'>{trend}</font>", td_wrap),
                ])

        if len(rows) > 1:
            t = Table(rows, colWidths=[2.0 * inch, 1.2 * inch, 3.0 * inch])
            style_cmds = [
                ('BACKGROUND', (0, 0), (-1, 0), self.C['primary_dark']),
                ('TEXTCOLOR', (0, 0), (-1, 0), self.C['white']),
                ('GRID', (0, 0), (-1, -1), 0.4, self.C['border']),
                ('TOPPADDING', (0, 0), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]
            for i in range(1, len(rows)):
                if i % 2 == 0:
                    style_cmds.append(('BACKGROUND', (0, i), (-1, i), colors.HexColor("#F5F7FA")))
            t.setStyle(TableStyle(style_cmds))
            story.append(t)
            # Add explanatory note about the scale
            story.append(Spacer(1, 4))
            story.append(Paragraph(
                "<i>Note: Baseline scores shown as X/10 represent performance on a 10-point scale. "
                "Percentages (%) indicate rate-based metrics. Color coding: "
                "<font color='#2E7D32'>Green</font> = strong, "
                "<font color='#F57F17'>Amber</font> = moderate, "
                "<font color='#C62828'>Red</font> = needs improvement.</i>",
                ParagraphStyle('comp_note', parent=self.ps['body'], textColor=self.C['light'], fontSize=7.5)))
        else:
            # Fallback: render content as text
            for line in content.split('\n'):
                line = line.strip()
                if line:
                    line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
                    story.append(Paragraph(line, self.ps['body']))
                    story.append(Spacer(1, 2))

        if note_text:
            story.append(Spacer(1, 6))
            story.append(Paragraph(f"<i>Note: {note_text}</i>", ParagraphStyle(
                'note', parent=self.ps['body'], textColor=self.C['light'], fontSize=8)))
        story.append(Spacer(1, 6))

    # ── M. Agent Tier Classification ──
    def _render_agent_tier(self, story, content: str):
        tier_line = ""
        tier_scale = []
        promotion = []
        in_scale = False
        in_promo = False
        level_num = 3  # default

        for line in content.split('\n'):
            stripped = line.strip()
            if stripped.lower().startswith('agent tier:'):
                tier_line = stripped.split(':', 1)[1].strip()
            elif stripped.lower().startswith('tier scale:'):
                in_scale = True
                in_promo = False
            elif stripped.lower().startswith('promotion criteria'):
                in_promo = True
                in_scale = False
            elif in_scale and stripped.startswith('Level'):
                tier_scale.append(stripped)
            elif in_promo and (stripped.startswith('-') or stripped.startswith('*')):
                promotion.append(stripped.lstrip('-* ').strip())
            elif in_scale and not stripped:
                in_scale = False
            elif in_promo and not stripped:
                in_promo = False

        # Tier badge
        if tier_line:
            level_match = re.search(r'Level\s*(\d)', tier_line)
            level_num = int(level_match.group(1)) if level_match else 3
            level_colors = {
                1: ("#FFEBEE", "#C62828"),
                2: ("#FFF8E1", "#F57F17"),
                3: ("#E3F2FD", "#1565C0"),
                4: ("#E8F5E9", "#2E7D32"),
                5: ("#F3E5F5", "#6A1B9A"),
            }
            bg, tc = level_colors.get(level_num, ("#E3F2FD", "#1565C0"))

            badge = Table([
                [Paragraph(tier_line, ParagraphStyle('tb', parent=self.ps['tier'],
                           textColor=colors.HexColor(tc), fontSize=18, leading=24))],
            ], colWidths=[6.4 * inch])
            badge.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(bg)),
                ('BOX', (0, 0), (-1, -1), 1.5, colors.HexColor(tc)),
                ('TOPPADDING', (0, 0), (-1, -1), 16),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 16),
                ('LEFTPADDING', (0, 0), (-1, -1), 20),
                ('RIGHTPADDING', (0, 0), (-1, -1), 20),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(badge)
            story.append(Spacer(1, 10))

        # Tier scale table — use fallback if LLM didn't provide them
        if not tier_scale:
            tier_scale = [
                "Level 1: Trainee -- Fundamental gaps in sales methodology; needs intensive coaching",
                "Level 2: Developing -- Shows awareness of techniques but inconsistent execution",
                "Level 3: Competent -- Solid fundamentals with specific areas for growth",
                "Level 4: Advanced -- Strong performer with minor refinements needed",
                "Level 5: Elite -- Exceptional across all dimensions; ready to mentor others",
            ]
        if tier_scale:
            story.append(Paragraph("<b>Tier Scale</b>", self.ps['subsec']))
            trows = [[Paragraph("<b>Level</b>", self.ps['th']),
                      Paragraph("<b>Title</b>", self.ps['th']),
                      Paragraph("<b>Description</b>", self.ps['th'])]]
            for ts_line in tier_scale:
                m = re.match(r'Level\s*(\d+)\s*:\s*(.*?)(?:\s*--\s*|\s*[-:]\s*)(.*)', ts_line)
                if m:
                    trows.append([
                        Paragraph(m.group(1), self.ps['td_c']),
                        Paragraph(f"<b>{m.group(2).strip()}</b>", self.ps['td']),
                        Paragraph(m.group(3).strip(), self.ps['td']),
                    ])
            if len(trows) > 1:
                t = Table(trows, colWidths=[0.6 * inch, 1.5 * inch, 4.1 * inch])
                style_cmds = [
                    ('BACKGROUND', (0, 0), (-1, 0), self.C['primary_dark']),
                    ('TEXTCOLOR', (0, 0), (-1, 0), self.C['white']),
                    ('GRID', (0, 0), (-1, -1), 0.5, self.C['border']),
                    ('PADDING', (0, 0), (-1, -1), 5),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]
                for i, ts_line in enumerate(tier_scale, 1):
                    lm = re.match(r'Level\s*(\d+)', ts_line)
                    if lm and int(lm.group(1)) == level_num:
                        style_cmds.append(('BACKGROUND', (0, i), (-1, i), colors.HexColor("#E3F2FD")))
                t.setStyle(TableStyle(style_cmds))
                story.append(t)
                story.append(Spacer(1, 8))

        # Promotion criteria
        if promotion:
            story.append(Paragraph("<b>Promotion Criteria (Next Level)</b>", self.ps['subsec']))
            for crit in promotion:
                crit = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', crit)
                t = Table([[Paragraph(f"<bullet>&bull;</bullet>{crit}", self.ps['bullet_accent'])]],
                          colWidths=[6.2 * inch])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), self.C['warning_bg']),
                    ('PADDING', (0, 0), (-1, -1), 5),
                ]))
                story.append(t)
                story.append(Spacer(1, 3))

    # ════════════════════════════════════════════════════════
    # NEW INTELLIGENCE LAYER RENDERERS (N through S)
    # ════════════════════════════════════════════════════════

    # ── N. Client Sentiment Timeline ──
    def _render_client_sentiment(self, story, content: str):
        """Render client emotional journey with structured risk/solution cards per phase."""
        content = self._convert_timestamps_to_mmss(content)
        # Intro
        story.append(Paragraph(
            "This section maps the client's emotional state throughout the meeting. "
            "Each phase identifies the dominant sentiment, concrete risks to the deal, "
            "and the recommended corrective action the agent should have taken.",
            ParagraphStyle('sent_intro', parent=self.ps['body'], fontSize=9,
                           textColor=self.C['medium'], spaceAfter=8)))

        phases = []
        turning_points = []
        trust_arc = ""
        trajectory = ""
        current_phase = None
        current_lines = []
        in_turning = False
        in_trust = False

        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            phase_match = re.match(r'^(EARLY|MID|LATE)\s*PHASE.*:', stripped, re.IGNORECASE)
            if phase_match:
                if current_phase and current_lines:
                    phases.append((current_phase, '\n'.join(current_lines)))
                current_phase = phase_match.group(1).upper()
                current_lines = []
                in_turning = False
                in_trust = False
                continue
            if 'EMOTIONAL TURNING POINT' in stripped.upper():
                if current_phase and current_lines:
                    phases.append((current_phase, '\n'.join(current_lines)))
                    current_phase = None
                    current_lines = []
                in_turning = True
                in_trust = False
                continue
            if 'TRUST ARC' in stripped.upper():
                if current_phase and current_lines:
                    phases.append((current_phase, '\n'.join(current_lines)))
                    current_phase = None
                    current_lines = []
                in_turning = False
                in_trust = True
                val = stripped.split(':', 1)
                if len(val) > 1 and val[1].strip():
                    trust_arc = val[1].strip()
                continue
            if 'OVERALL SENTIMENT TRAJECTORY' in stripped.upper():
                in_turning = False
                in_trust = False
                val = stripped.split(':', 1)
                if len(val) > 1:
                    trajectory = val[1].strip()
                continue
            if in_turning:
                turning_points.append(stripped.lstrip('-* ').strip() if stripped.startswith(('-', '*')) else stripped)
            elif in_trust:
                trust_arc += " " + stripped
            elif current_phase:
                current_lines.append(stripped)

        if current_phase and current_lines:
            phases.append((current_phase, '\n'.join(current_lines)))

        # Phase cards — structured table rows with risk + solution
        phase_colors = {
            'EARLY': ('#E3F2FD', '#1565C0'),
            'MID': ('#FFF8E1', '#F57F17'),
            'LATE': ('#F3E5F5', '#6A1B9A'),
        }
        for phase_name, phase_content in phases:
            bg, tc = phase_colors.get(phase_name, ('#F5F7FA', '#424242'))
            # Parse fields from phase content
            fields = {}
            extra = []
            for pl in phase_content.split('\n'):
                pl = pl.strip()
                if not pl:
                    continue
                matched = False
                for lbl in ['SENTIMENT', 'EVIDENCE', 'TIMESTAMP', 'RISK', 'SOLUTION',
                            'DEAL RISK', 'RECOMMENDED ACTION', 'CLIENT SIGNAL']:
                    if pl.upper().startswith(lbl + ':'):
                        fields[lbl.upper()] = pl[len(lbl) + 1:].strip()
                        matched = True
                        break
                if not matched:
                    extra.append(pl)

            rows = []
            # Phase title row
            rows.append([
                Paragraph(f"<b>{phase_name} PHASE</b>",
                          ParagraphStyle('ph_' + phase_name, parent=self.ps['subsec'],
                                         textColor=colors.HexColor(tc), fontSize=10)),
                Paragraph('', self.ps['td'])
            ])
            if 'SENTIMENT' in fields:
                rows.append([Paragraph("<b>Sentiment</b>", self.ps['td']),
                             Paragraph(fields['SENTIMENT'], self.ps['td'])])
            if 'TIMESTAMP' in fields:
                rows.append([Paragraph("<b>When</b>", self.ps['td']),
                             Paragraph(f"<font color='#1565C0'><b>{fields['TIMESTAMP']}</b></font>", self.ps['td'])])
            if 'EVIDENCE' in fields or 'CLIENT SIGNAL' in fields:
                ev = fields.get('EVIDENCE', fields.get('CLIENT SIGNAL', ''))
                rows.append([Paragraph("<b>Evidence</b>", self.ps['td']),
                             Paragraph(f"<i>{ev}</i>", self.ps['td'])])
            risk = fields.get('RISK', fields.get('DEAL RISK', ''))
            if risk:
                rows.append([Paragraph("<b>Deal Risk</b>", self.ps['td']),
                             Paragraph(f"<font color='#C62828'><b>{risk}</b></font>", self.ps['td'])])
            solution = fields.get('SOLUTION', fields.get('RECOMMENDED ACTION', ''))
            if solution:
                rows.append([Paragraph("<b>Solution</b>", self.ps['td']),
                             Paragraph(f"<font color='#2E7D32'>{solution}</font>", self.ps['td'])])
            for el in extra:
                el = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', el)
                rows.append([Paragraph('', self.ps['td']), Paragraph(el, self.ps['td'])])

            if rows:
                t = Table(rows, colWidths=[1.3 * inch, 4.9 * inch])
                style_cmds = [
                    ('SPAN', (0, 0), (1, 0)),
                    ('BACKGROUND', (0, 0), (1, 0), colors.HexColor(bg)),
                    ('BACKGROUND', (0, 1), (0, -1), colors.HexColor(bg)),
                    ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor(tc)),
                    ('PADDING', (0, 0), (-1, -1), 6),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]
                t.setStyle(TableStyle(style_cmds))
                story.append(t)
                story.append(Spacer(1, 5))

        # Turning Points
        if turning_points:
            story.append(Paragraph("<b>Emotional Turning Points</b>", self.ps['subsec']))
            for tp in turning_points:
                tp = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', tp)
                tp = re.sub(r'\[(\d+[\.\d]*s?\s*-\s*\d+[\.\d]*s?)\]',
                            r"<font color='#1565C0'><b>[\1]</b></font>", tp)
                story.append(Paragraph(f"<bullet>&bull;</bullet>{tp}", self.ps['bullet_accent']))
                story.append(Spacer(1, 2))

        # Trust Arc
        if trust_arc:
            story.append(Spacer(1, 4))
            t = Table([[Paragraph(f"<b>Trust Arc:</b> {trust_arc.strip()}", self.ps['body'])]],
                      colWidths=[6.2 * inch])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#E8F5E9')),
                ('PADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(t)

        # Trajectory
        if trajectory:
            story.append(Spacer(1, 6))
            story.append(Paragraph(f"<b>Overall Sentiment Trajectory:</b> <i>{trajectory}</i>",
                         ParagraphStyle('traj', parent=self.ps['body_bold'], fontSize=10,
                                        textColor=self.C['primary_dark'])))

        # Cache for visualization section
        self._cached_sentiment_phases = phases
        self._cached_turning_points = turning_points
        self._cached_trust_arc = trust_arc

    # ── N. Communication Balance & Listening Evaluation ──
    def _render_listening_intelligence(self, story, content: str):
        """Render listening evaluation with verdicts and recommendations per area."""
        content = self._convert_timestamps_to_mmss(content)
        subsections = {}
        current_key = None
        current_lines = []

        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            # Broader matching: look for known subsection headers
            header_match = re.match(
                r'^(TALK RATIO|INTERRUPTION ANALYSIS|REFLECTIVE LISTENING.*|'
                r'FOLLOW-UP QUESTION.*|ACTIVE LISTENING GAPS|CRITICAL LISTENING FAILURES?|'
                r'LISTENING FAILURE).*:', stripped, re.IGNORECASE)
            if header_match:
                if current_key and current_lines:
                    subsections[current_key] = '\n'.join(current_lines)
                raw_key = header_match.group(1).upper().strip()
                # Normalize to standard keys
                if 'TALK' in raw_key:
                    current_key = 'TALK RATIO'
                elif 'INTERRUPT' in raw_key:
                    current_key = 'INTERRUPTION ANALYSIS'
                elif 'REFLECT' in raw_key:
                    current_key = 'REFLECTIVE LISTENING'
                elif 'FOLLOW' in raw_key:
                    current_key = 'FOLLOW-UP QUESTION'
                elif 'CRITICAL' in raw_key or 'FAILURE' in raw_key:
                    current_key = 'CRITICAL LISTENING FAILURES'
                else:
                    current_key = raw_key
                remainder = stripped[header_match.end():].strip()
                current_lines = [remainder] if remainder else []
                continue
            # Also detect ** bold headers ** as subsection markers
            bold_header = re.match(r'^\*\*(TALK RATIO|INTERRUPTION|REFLECTIVE|FOLLOW.UP|CRITICAL|LISTENING).*?\*\*', stripped, re.IGNORECASE)
            if bold_header:
                if current_key and current_lines:
                    subsections[current_key] = '\n'.join(current_lines)
                raw_key = bold_header.group(1).upper().strip()
                if 'TALK' in raw_key:
                    current_key = 'TALK RATIO'
                elif 'INTERRUPT' in raw_key:
                    current_key = 'INTERRUPTION ANALYSIS'
                elif 'REFLECT' in raw_key:
                    current_key = 'REFLECTIVE LISTENING'
                elif 'FOLLOW' in raw_key:
                    current_key = 'FOLLOW-UP QUESTION'
                elif 'CRITICAL' in raw_key or 'LISTENING' in raw_key:
                    current_key = 'CRITICAL LISTENING FAILURES'
                else:
                    current_key = raw_key
                current_lines = []
                continue
            if current_key:
                current_lines.append(stripped)

        if current_key and current_lines:
            subsections[current_key] = '\n'.join(current_lines)

        # FALLBACK: If no subsections were detected at all, parse entire content generically
        if not subsections:
            # Try to extract talk ratio from anywhere in content
            agent_match = re.search(r'(?:Agent|Estimated Agent).*?(\d+)\s*%', content, re.IGNORECASE)
            client_match = re.search(r'(?:Client|Estimated Client).*?(\d+)\s*%', content, re.IGNORECASE)
            if agent_match and client_match:
                subsections['TALK RATIO'] = f"Estimated Agent Talk Time: {agent_match.group(1)}%\nEstimated Client Talk Time: {client_match.group(1)}%"
            # Put everything else as generic content
            remaining_lines = []
            for line in content.split('\n'):
                stripped = line.strip()
                if stripped and not re.match(r'Estimated (Agent|Client)', stripped, re.IGNORECASE):
                    remaining_lines.append(stripped)
            if remaining_lines:
                subsections.setdefault('GENERAL', '\n'.join(remaining_lines))

        # Talk Ratio visual — NO AI-estimated note
        talk_content = subsections.get('TALK RATIO', '')
        if talk_content:
            story.append(Paragraph("<b>Talk Ratio Analysis</b>", self.ps['subsec']))
            agent_pct = re.search(r'Agent.*?(\d+)%', talk_content)
            client_pct = re.search(r'Client.*?(\d+)%', talk_content)
            if agent_pct and client_pct:
                a_val = int(agent_pct.group(1))
                c_val = int(client_pct.group(1))
                a_color = '#E53935' if a_val > 65 else '#FF8F00' if a_val > 55 else '#43A047'
                c_color = '#43A047' if c_val > 35 else '#FF8F00'
                a_verdict = "Agent-dominated — too much talking" if a_val > 65 else "Balanced" if a_val <= 55 else "Slightly high"
                c_verdict = "Good engagement" if c_val > 35 else "Low participation — not enough probing"
                ratio_data = [
                    [Paragraph("<b>Speaker</b>", self.ps['th']),
                     Paragraph("<b>Talk Time</b>", self.ps['th']),
                     Paragraph("<b>Verdict</b>", self.ps['th']),
                     Paragraph("<b>Recommendation</b>", self.ps['th'])],
                    [Paragraph("Agent", self.ps['td']),
                     Paragraph(f"<b><font color='{a_color}'>{a_val}%</font></b>", self.ps['td_c']),
                     Paragraph(a_verdict, self.ps['td']),
                     Paragraph("Ask more open-ended questions; aim for 40-50% talk time" if a_val > 55
                               else "Maintain current approach", self.ps['td'])],
                    [Paragraph("Client", self.ps['td']),
                     Paragraph(f"<b><font color='{c_color}'>{c_val}%</font></b>", self.ps['td_c']),
                     Paragraph(c_verdict, self.ps['td']),
                     Paragraph("Use silence and probing to get client talking more" if c_val <= 35
                               else "Keep encouraging engagement", self.ps['td'])],
                ]
                rt = Table(ratio_data, colWidths=[1.0 * inch, 1.0 * inch, 2.0 * inch, 2.2 * inch])
                rt.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), self.C['primary_dark']),
                    ('TEXTCOLOR', (0, 0), (-1, 0), self.C['white']),
                    ('GRID', (0, 0), (-1, -1), 0.5, self.C['border']),
                    ('PADDING', (0, 0), (-1, -1), 6),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                story.append(rt)
            else:
                for tl in talk_content.split('\n'):
                    tl = tl.strip()
                    if tl:
                        tl = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', tl)
                        story.append(Paragraph(tl, self.ps['body']))
                        story.append(Spacer(1, 2))
            story.append(Spacer(1, 6))

        # Other subsections — structured with verdict color coding
        sub_config = [
            ('INTERRUPTION ANALYSIS', 'Interruption Analysis', '#FFEBEE', '#C62828'),
            ('REFLECTIVE LISTENING', 'Reflective Listening', '#E8F5E9', '#2E7D32'),
            ('FOLLOW-UP QUESTION', 'Follow-Up Question Quality', '#E3F2FD', '#1565C0'),
            ('ACTIVE LISTENING GAPS', 'Active Listening Gaps', '#FFF8E1', '#F57F17'),
            ('CRITICAL LISTENING FAILURES', 'Critical Listening Failures', '#FFEBEE', '#C62828'),
            ('GENERAL', 'Detailed Analysis', '#F5F7FA', '#424242'),
        ]
        for key, display, bg, tc in sub_config:
            sub_content = subsections.get(key, '')
            if not sub_content:
                continue
            inner = [
                Paragraph(f"<b>{display}</b>",
                          ParagraphStyle('li_' + key, parent=self.ps['subsec'],
                                         textColor=colors.HexColor(tc))),
                Spacer(1, 3),
            ]
            for sl in sub_content.split('\n'):
                sl = sl.strip()
                if not sl:
                    continue
                sl = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', sl)
                sl = re.sub(r'\[(\d+[\.\d]*s?\s*-\s*\d+[\.\d]*s?)\]',
                            r"<font color='#1565C0'><b>[\1]</b></font>", sl)
                if sl.startswith('-') or sl.startswith('*'):
                    inner.append(Paragraph(f"<bullet>&bull;</bullet>{sl.lstrip('-* ')}",
                                           self.ps['bullet']))
                elif re.match(r'^Score:', sl, re.IGNORECASE):
                    inner.append(Paragraph(f"<b>{sl}</b>", self.ps['body_bold']))
                elif re.match(r'^Verdict:', sl, re.IGNORECASE):
                    inner.append(Paragraph(f"<font color='{tc}'><b>{sl}</b></font>", self.ps['body_bold']))
                elif re.match(r'^Recommendation:', sl, re.IGNORECASE):
                    inner.append(Paragraph(f"<font color='#2E7D32'>{sl}</font>", self.ps['body']))
                else:
                    inner.append(Paragraph(sl, self.ps['body']))
                inner.append(Spacer(1, 2))

            card = Table([[inner]], colWidths=[6.2 * inch])
            card.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(bg)),
                ('PADDING', (0, 0), (-1, -1), 8),
                ('LINEBELOW', (0, 0), (-1, -1), 0.5, colors.HexColor(tc)),
            ]))
            story.append(card)
            story.append(Spacer(1, 5))

        # Cache talk ratio for visualizations
        talk_content = subsections.get('TALK RATIO', '')
        _a_match = re.search(r'Agent.*?(\d+)%', talk_content) if talk_content else None
        _c_match = re.search(r'Client.*?(\d+)%', talk_content) if talk_content else None
        self._cached_talk_ratio = {
            'agent': int(_a_match.group(1)) if _a_match else None,
            'client': int(_c_match.group(1)) if _c_match else None,
        }
        self._cached_listening_subsections = subsections

    # ── O. Five Core Discovery Pillars ──
    def _render_discovery_matrix(self, story, content: str):
        """Render the 5 essential discovery pillars every top agent must cover."""
        content = self._convert_timestamps_to_mmss(content)
        # Intro explanation
        story.append(Paragraph(
            "Every successful listing presentation requires thorough discovery across five essential pillars. "
            "Failing to explore any of these areas leaves money on the table and creates deal risk. "
            "Below is the agent's coverage assessment:",
            ParagraphStyle('disc_intro', parent=self.ps['body'], fontSize=9,
                           textColor=self.C['medium'], spaceAfter=8)))

        pillars = []
        score_line = ""
        missing_qs = []
        in_missing = False

        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            if 'DISCOVERY COMPLETENESS SCORE' in stripped.upper():
                score_line = stripped.split(':', 1)[-1].strip()
                in_missing = False
                continue
            if 'MISSING DISCOVERY QUESTION' in stripped.upper():
                in_missing = True
                continue
            if in_missing:
                if stripped.startswith('-') or stripped.startswith('*'):
                    missing_qs.append(stripped.lstrip('-* ').strip())
                elif re.match(r'^\d+\.', stripped):
                    missing_qs.append(re.sub(r'^\d+\.\s*', '', stripped).strip())
                continue

            parts = [p.strip() for p in stripped.split('|')]
            if len(parts) >= 3:
                pillar = parts[0].replace('**', '').strip()
                status = parts[1].strip()
                evidence = parts[2].strip() if len(parts) > 2 else ''
                confidence = parts[3].strip() if len(parts) > 3 else ''
                if pillar and not pillar.lower().startswith('pillar'):
                    pillars.append((pillar, status, evidence, confidence))

        # Matrix table
        if pillars:
            td_wrap = ParagraphStyle('td_disc', parent=self.ps['td'], wordWrap='LTR',
                                     fontSize=8, leading=11)
            rows = [[
                Paragraph("<b>Discovery Pillar</b>", self.ps['th']),
                Paragraph("<b>Status</b>", self.ps['th']),
                Paragraph("<b>Evidence</b>", self.ps['th']),
                Paragraph("<b>Confidence</b>", self.ps['th']),
            ]]
            for pillar, status, evidence, confidence in pillars:
                status_color = self._get_color_for_status(status)
                rows.append([
                    Paragraph(f"<b>{pillar}</b>", self.ps['td']),
                    Paragraph(f"<font color='{status_color}'><b>{status}</b></font>", self.ps['td_c']),
                    Paragraph(f"<i>{evidence}</i>", td_wrap),
                    Paragraph(f"<i>{confidence}</i>",
                             ParagraphStyle('dc', parent=self.ps['td'], fontSize=7,
                                            textColor=self.C['light'])),
                ])

            t = Table(rows, colWidths=[1.4 * inch, 1.2 * inch, 2.6 * inch, 1.0 * inch])
            style_cmds = [
                ('BACKGROUND', (0, 0), (-1, 0), self.C['primary_dark']),
                ('TEXTCOLOR', (0, 0), (-1, 0), self.C['white']),
                ('GRID', (0, 0), (-1, -1), 0.5, self.C['border']),
                ('PADDING', (0, 0), (-1, -1), 5),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
            for i, (_, status, _, _) in enumerate(pillars, 1):
                status_color = self._get_color_for_status(status)
                # Lighten background
                if status_color == "#43A047":
                    style_cmds.append(('BACKGROUND', (1, i), (1, i), colors.HexColor('#E8F5E9')))
                elif status_color == "#FF8F00":
                    style_cmds.append(('BACKGROUND', (1, i), (1, i), colors.HexColor('#FFF8E1')))
                elif status_color == "#E53935":
                    style_cmds.append(('BACKGROUND', (1, i), (1, i), colors.HexColor('#FFEBEE')))
            t.setStyle(TableStyle(style_cmds))
            story.append(t)
            story.append(Spacer(1, 6))

        # Score
        if score_line:
            score_match = re.search(r'(\d+)', score_line)
            s_val = int(score_match.group(1)) if score_match else 0
            s_color = '#43A047' if s_val >= 4 else '#FF8F00' if s_val >= 3 else '#E53935'
            story.append(Paragraph(
                f"<b>Pillar Coverage Score:</b> <font color='{s_color}'><b>{score_line}</b></font>",
                ParagraphStyle('dcs', parent=self.ps['body_bold'], fontSize=11,
                               textColor=self.C['primary_dark'])))
            story.append(Spacer(1, 6))

        # Missing questions
        if missing_qs:
            story.append(Paragraph("<b>Questions the Agent Must Ask Next Time</b>", self.ps['subsec']))
            for i, q in enumerate(missing_qs, 1):
                q = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', q)
                story.append(Paragraph(f"<b>{i}.</b> {q}", self.ps['bullet_danger']))
                story.append(Spacer(1, 2))

        # Cache for visualization
        self._cached_discovery_pillars = pillars

    # ── Q. Emotional Intelligence Assessment ──
    def _render_emotional_intelligence(self, story, content: str):
        """Render EQ assessment with score cards for each dimension."""
        content = self._convert_timestamps_to_mmss(content)
        eq_scores = {}
        overall_eq = ""
        ai_conf = ""
        current_dim = None
        current_lines = []

        dimensions = ['EMPATHY', 'ADAPTABILITY', 'SOCIAL AWARENESS', 'EMOTIONAL REGULATION']

        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue

            dim_match = None
            for dim in dimensions:
                if stripped.upper().startswith(dim + ' SCORE'):
                    dim_match = dim
                    break

            if dim_match:
                if current_dim and current_lines:
                    eq_scores[current_dim] = '\n'.join(current_lines)
                current_dim = dim_match
                score_part = stripped.split(':', 1)
                current_lines = [score_part[1].strip()] if len(score_part) > 1 and score_part[1].strip() else []
                continue

            if 'OVERALL EQ SCORE' in stripped.upper():
                if current_dim and current_lines:
                    eq_scores[current_dim] = '\n'.join(current_lines)
                    current_dim = None
                    current_lines = []
                val = stripped.split(':', 1)
                if len(val) > 1:
                    overall_eq = val[1].strip()
                continue

            if 'AI CONFIDENCE' in stripped.upper() and not current_dim:
                val = stripped.split(':', 1)
                if len(val) > 1:
                    ai_conf = val[1].strip()
                continue

            if current_dim:
                current_lines.append(stripped)

        if current_dim and current_lines:
            eq_scores[current_dim] = '\n'.join(current_lines)

        # Overall EQ hero
        if overall_eq:
            score_match = re.search(r'(\d+)\s*/\s*(\d+)', overall_eq)
            if score_match:
                s_val = int(score_match.group(1))
                pct = s_val / 10
                color = "#43A047" if pct >= 0.75 else "#FF8F00" if pct >= 0.50 else "#E53935"
                hero = Table([
                    [Paragraph(f"<font color='{color}'><b>{overall_eq}</b></font>",
                     ParagraphStyle('eq_hero', parent=self.ps['score_big'], fontSize=32, leading=38))],
                    [Spacer(1, 1)],
                    [Paragraph("EMOTIONAL INTELLIGENCE SCORE", self.ps['score_label'])],
                    [Paragraph(f"<i>{ai_conf}</i>" if ai_conf else "",
                     ParagraphStyle('eq_conf', parent=self.ps['conf'], fontSize=7))],
                ], colWidths=[6.2 * inch], rowHeights=[44, 6, 16, 14])
                hero.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F5F7FA')),
                    ('LINEBELOW', (0, -1), (-1, -1), 1.5, colors.HexColor(color)),
                    ('PADDING', (0, 0), (-1, -1), 8),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                story.append(hero)
                story.append(Spacer(1, 10))

        # Dimension cards
        dim_display = {
            'EMPATHY': ('Empathy', '#E3F2FD', '#1565C0'),
            'ADAPTABILITY': ('Adaptability', '#E8F5E9', '#2E7D32'),
            'SOCIAL AWARENESS': ('Social Awareness', '#FFF8E1', '#F57F17'),
            'EMOTIONAL REGULATION': ('Emotional Regulation', '#F3E5F5', '#6A1B9A'),
        }
        for dim, dim_content in eq_scores.items():
            display_name, bg, tc = dim_display.get(dim, (dim.title(), '#F5F7FA', '#424242'))
            inner = [Paragraph(f"<b>{display_name}</b>",
                     ParagraphStyle('eq_' + dim, parent=self.ps['subsec'],
                                    textColor=colors.HexColor(tc)))]
            for dl in dim_content.split('\n'):
                dl = dl.strip()
                if not dl:
                    continue
                dl = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', dl)
                dl = re.sub(r'\[(\d+[\.\d]*s?\s*-\s*\d+[\.\d]*s?)\]',
                            r"<font color='#1565C0'><b>[\1]</b></font>", dl)
                if dl.startswith('-') or dl.startswith('*'):
                    inner.append(Paragraph(f"<bullet>&bull;</bullet>{dl.lstrip('-* ')}",
                                           self.ps['bullet']))
                else:
                    inner.append(Paragraph(dl, self.ps['body']))
                inner.append(Spacer(1, 2))

            card = Table([[inner]], colWidths=[6.2 * inch])
            card.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(bg)),
                ('PADDING', (0, 0), (-1, -1), 8),
                ('LINEBELOW', (0, 0), (-1, -1), 0.5, colors.HexColor(tc)),
            ]))
            story.append(card)
            story.append(Spacer(1, 6))

        # Cache EQ dimension scores for visualization
        self._cached_eq_dimensions = {dim: eq_scores.get(dim, '') for dim in dimensions}
        self._cached_eq_overall = overall_eq

    # ── R. Ethics & Compliance Audit ──
    def _render_ethics_compliance(self, story, content: str):
        """Render ethics/compliance audit with risk indicators."""
        content = self._convert_timestamps_to_mmss(content)
        subsections = {}
        current_key = None
        current_lines = []

        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            header_match = re.match(
                r'^(PRESSURE TACTICS|OVER-PROMISING|REGULATORY|TRANSPARENCY|OVERALL ETHICS RISK|'
                r'ETHICS RECOMMENDATION).*:', stripped, re.IGNORECASE)
            if header_match:
                if current_key and current_lines:
                    subsections[current_key] = '\n'.join(current_lines)
                current_key = header_match.group(1).upper()
                remainder = stripped[header_match.end():].strip()
                current_lines = [remainder] if remainder else []
                continue
            if current_key:
                current_lines.append(stripped)

        if current_key and current_lines:
            subsections[current_key] = '\n'.join(current_lines)

        # Overall risk badge at top
        risk_content = subsections.get('OVERALL ETHICS RISK', '')
        if risk_content:
            risk_color = self._get_color_for_status(risk_content)
            risk_bg = "#FFEBEE" if risk_color == "#E53935" else "#FFF8E1" if risk_color == "#FF8F00" else "#E8F5E9"
            badge = Table([
                [Paragraph(
                    f"<b>Overall Ethics Risk: {risk_content.split(chr(10))[0].strip()}</b>",
                    ParagraphStyle('er', parent=self.ps['body_bold'], fontSize=12,
                                   textColor=colors.HexColor(risk_color), alignment=TA_CENTER))],
            ], colWidths=[6.2 * inch])
            badge.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(risk_bg)),
                ('PADDING', (0, 0), (-1, -1), 12),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ]))
            story.append(badge)
            story.append(Spacer(1, 10))

        # Cache ethics data for visualization
        self._cached_ethics_subsections = subsections

        # Sub-audit sections
        audit_sections = [
            ('PRESSURE TACTICS', 'Pressure Tactics Assessment', '#FFEBEE', '#C62828'),
            ('OVER-PROMISING', 'Over-Promising Check', '#FFF8E1', '#F57F17'),
            ('REGULATORY', 'Regulatory Language Review', '#E3F2FD', '#1565C0'),
            ('TRANSPARENCY', 'Transparency Assessment', '#E8F5E9', '#2E7D32'),
        ]
        for key, display, bg, tc in audit_sections:
            sub_content = subsections.get(key, '')
            if not sub_content:
                continue
            inner = [Paragraph(f"<b>{display}</b>",
                     ParagraphStyle('ea_' + key, parent=self.ps['subsec'],
                                    textColor=colors.HexColor(tc)))]
            for al in sub_content.split('\n'):
                al = al.strip()
                if not al:
                    continue
                al = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', al)
                al = re.sub(r'\[(\d+[\.\d]*s?\s*-\s*\d+[\.\d]*s?)\]',
                            r"<font color='#1565C0'><b>[\1]</b></font>", al)
                for rl in ['High', 'Medium', 'Low', 'None']:
                    if f'Risk Level: {rl}' in al:
                        rl_color = '#C62828' if rl == 'High' else '#F57F17' if rl == 'Medium' else '#2E7D32'
                        al = al.replace(
                            f'Risk Level: {rl}',
                            f"<b>Risk Level: <font color='{rl_color}'>{rl}</font></b>")
                if al.startswith('-') or al.startswith('*'):
                    inner.append(Paragraph(f"<bullet>&bull;</bullet>{al.lstrip('-* ')}",
                                           self.ps['bullet']))
                else:
                    inner.append(Paragraph(al, self.ps['body']))
                inner.append(Spacer(1, 2))

            card = Table([[inner]], colWidths=[6.2 * inch])
            card.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(bg)),
                ('PADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(card)
            story.append(Spacer(1, 6))

        # Ethics Recommendations
        rec_content = subsections.get('ETHICS RECOMMENDATION', '')
        if rec_content:
            story.append(Paragraph("<b>Ethics Recommendations</b>", self.ps['subsec']))
            for rl in rec_content.split('\n'):
                rl = rl.strip()
                if not rl:
                    continue
                rl = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', rl)
                if rl.startswith('-') or rl.startswith('*'):
                    story.append(Paragraph(f"<bullet>&bull;</bullet>{rl.lstrip('-* ')}",
                                           self.ps['bullet_success']))
                else:
                    story.append(Paragraph(rl, self.ps['body']))
                story.append(Spacer(1, 2))

    # ── S. Agent Self-Awareness Assessment ──
    def _render_agent_self_awareness(self, story, content: str):
        """Render agent self-awareness evaluation with categorized findings."""
        content = self._convert_timestamps_to_mmss(content)
        subsections = {}
        current_key = None
        current_lines = []
        sa_score = ""
        ai_conf = ""

        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            header_match = re.match(
                r'^(SIGNAL RECOGNITION|REAL-TIME ADAPTATION|SELF-CORRECTION|BLIND SPOTS).*:',
                stripped, re.IGNORECASE)
            if header_match:
                if current_key and current_lines:
                    subsections[current_key] = '\n'.join(current_lines)
                current_key = header_match.group(1).upper()
                remainder = stripped[header_match.end():].strip()
                current_lines = [remainder] if remainder else []
                continue
            if 'SELF-AWARENESS SCORE' in stripped.upper():
                if current_key and current_lines:
                    subsections[current_key] = '\n'.join(current_lines)
                    current_key = None
                    current_lines = []
                val = stripped.split(':', 1)
                if len(val) > 1:
                    sa_score = val[1].strip()
                continue
            if 'AI CONFIDENCE' in stripped.upper() and not current_key:
                val = stripped.split(':', 1)
                if len(val) > 1:
                    ai_conf = val[1].strip()
                continue
            if current_key:
                current_lines.append(stripped)

        if current_key and current_lines:
            subsections[current_key] = '\n'.join(current_lines)

        # Score hero — fixed layout to prevent text overlapping
        if sa_score:
            score_match = re.search(r'(\d+)\s*/\s*(\d+)', sa_score)
            if score_match:
                s_val = int(score_match.group(1))
                pct = s_val / 10
                color = "#43A047" if pct >= 0.75 else "#FF8F00" if pct >= 0.50 else "#E53935"
                hero = Table([
                    [Paragraph(f"<font color='{color}'><b>{sa_score}</b></font>",
                     ParagraphStyle('sa_hero', parent=self.ps['score_big'], fontSize=28, leading=34))],
                    [Spacer(1, 1)],  # spacer row to prevent overlap
                    [Paragraph("SELF-AWARENESS SCORE", self.ps['score_label'])],
                    [Paragraph(f"<i>{ai_conf}</i>" if ai_conf else "",
                     ParagraphStyle('sa_conf', parent=self.ps['conf'], fontSize=7))],
                ], colWidths=[6.2 * inch], rowHeights=[40, 6, 16, 14])
                hero.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F5F7FA')),
                    ('LINEBELOW', (0, -1), (-1, -1), 1.5, colors.HexColor(color)),
                    ('PADDING', (0, 0), (-1, -1), 8),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ]))
                story.append(hero)
                story.append(Spacer(1, 10))

        # Subsection cards
        sa_display = {
            'SIGNAL RECOGNITION': (
                'Signal Recognition', '#E3F2FD', '#1565C0',
                'Did the agent recognize client hesitation, confusion, or disengagement?'),
            'REAL-TIME ADAPTATION': (
                'Real-Time Adaptation', '#E8F5E9', '#2E7D32',
                'Did the agent adjust approach mid-conversation based on client cues?'),
            'SELF-CORRECTION': (
                'Self-Correction Instances', '#FFF8E1', '#F57F17',
                'Moments where the agent caught and corrected their own approach.'),
            'BLIND SPOTS': (
                'Blind Spots', '#FFEBEE', '#C62828',
                'Recurring patterns the agent appears unaware of.'),
        }
        for key, (display, bg, tc, desc) in sa_display.items():
            sub_content = subsections.get(key, '')
            if not sub_content:
                continue
            inner = [
                Paragraph(f"<b>{display}</b>",
                         ParagraphStyle('sa_' + key, parent=self.ps['subsec'],
                                        textColor=colors.HexColor(tc))),
                Paragraph(f"<i>{desc}</i>",
                         ParagraphStyle('sa_d_' + key, parent=self.ps['body'], fontSize=7.5,
                                        textColor=self.C['light'])),
                Spacer(1, 4),
            ]
            for sl in sub_content.split('\n'):
                sl = sl.strip()
                if not sl:
                    continue
                sl = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', sl)
                sl = re.sub(r'\[(\d+[\.\d]*s?\s*-\s*\d+[\.\d]*s?)\]',
                            r"<font color='#1565C0'><b>[\1]</b></font>", sl)
                if sl.startswith('-') or sl.startswith('*'):
                    inner.append(Paragraph(f"<bullet>&bull;</bullet>{sl.lstrip('-* ')}",
                                           self.ps['bullet']))
                elif re.match(r'^(Improvement|Recommendation|Action):', sl, re.IGNORECASE):
                    inner.append(Paragraph(f"<font color='#2E7D32'><b>{sl}</b></font>", self.ps['body_bold']))
                else:
                    inner.append(Paragraph(sl, self.ps['body']))
                inner.append(Spacer(1, 2))

            card = Table([[inner]], colWidths=[6.2 * inch])
            card.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(bg)),
                ('PADDING', (0, 0), (-1, -1), 8),
                ('LINEBELOW', (0, 0), (-1, -1), 0.5, colors.HexColor(tc)),
            ]))
            story.append(card)
            story.append(Spacer(1, 6))

    # ── T. 7-Day Intensive Coaching Plan (merged day-by-day with drills) ──
    def _render_7day_coaching_plan(self, story, content: str):
        """Render the comprehensive 7-day coaching plan with drills and exercises."""
        content = self._convert_timestamps_to_mmss(content)
        story.append(Paragraph(
            "<b>Complete 7-Day Coaching Sprint — One Skill Gap Per Day</b>",
            ParagraphStyle('plan_hdr', parent=self.ps['body_bold'],
                           textColor=self.C['danger'], fontSize=10)))
        story.append(Spacer(1, 6))

        rows = [[
            Paragraph("<b>Day</b>", self.ps['th']),
            Paragraph("<b>Focus Area</b>", self.ps['th']),
            Paragraph("<b>Drill / Exercise</b>", self.ps['th']),
            Paragraph("<b>Success Metric</b>", self.ps['th']),
        ]]
        # Parse "Day X: Focus -- Drill -- Metric" format
        for line in content.split('\n'):
            line = line.strip()
            m = re.match(r'Day\s*(\d+)\s*:\s*(.*?)(?:\s*--\s*)(.*?)(?:\s*--\s*)(.*)', line, re.IGNORECASE)
            if m:
                day, focus, drill, metric = m.group(1), m.group(2).strip(), m.group(3).strip(), m.group(4).strip()
                rows.append([
                    Paragraph(f"<b>{day}</b>", self.ps['td_c']),
                    Paragraph(focus, self.ps['td']),
                    Paragraph(drill, self.ps['td']),
                    Paragraph(metric, self.ps['td']),
                ])
            else:
                # Fallback: try "Day X: Focus -- Drill" (no metric)
                m2 = re.match(r'Day\s*(\d+)\s*:\s*(.*?)(?:\s*--\s*|-\s*)(.*)', line, re.IGNORECASE)
                if m2:
                    day, focus, drill = m2.group(1), m2.group(2).strip(), m2.group(3).strip()
                    rows.append([
                        Paragraph(f"<b>{day}</b>", self.ps['td_c']),
                        Paragraph(focus, self.ps['td']),
                        Paragraph(drill, self.ps['td']),
                        Paragraph("", self.ps['td']),
                    ])

        if len(rows) > 1:
            t = Table(rows, colWidths=[0.4 * inch, 1.6 * inch, 2.2 * inch, 2.0 * inch])
            style_cmds = [
                ('BACKGROUND', (0, 0), (-1, 0), self.C['primary_dark']),
                ('TEXTCOLOR', (0, 0), (-1, 0), self.C['white']),
                ('GRID', (0, 0), (-1, -1), 0.5, self.C['border']),
                ('PADDING', (0, 0), (-1, -1), 6),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
            for i in range(1, len(rows)):
                if i % 2 == 0:
                    style_cmds.append(('BACKGROUND', (0, i), (-1, i), colors.HexColor("#F5F7FA")))
            t.setStyle(TableStyle(style_cmds))
            story.append(t)
        else:
            # Fallback: render as numbered list if table parsing fails
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    continue
                m = re.match(r'^(\d+)\.\s*(.*)', line)
                if m:
                    num, text = m.group(1), m.group(2)
                    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
                    t = Table([[
                        Paragraph(f"<b>{num}</b>", ParagraphStyle('fn', parent=self.ps['td_c'],
                                  fontSize=14, textColor=self.C['white'])),
                        Paragraph(text, self.ps['body'])
                    ]], colWidths=[0.4 * inch, 5.8 * inch])
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (0, 0), self.C['primary_dark']),
                        ('BACKGROUND', (1, 0), (1, 0), colors.HexColor("#F5F7FA")),
                        ('BOX', (0, 0), (-1, -1), 0.5, self.C['border']),
                        ('PADDING', (0, 0), (-1, -1), 8),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 4))
                else:
                    line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
                    if line.startswith('-') or line.startswith('*'):
                        story.append(Paragraph(f"<bullet>&bull;</bullet>{line.lstrip('-* ')}", self.ps['bullet']))
                    else:
                        story.append(Paragraph(line, self.ps['body']))
                    story.append(Spacer(1, 3))
        story.append(Spacer(1, 8))

    # ── U. Client Engagement & Tone Analysis (NEW §17) ──
    def _render_tone_analysis(self, story, content: str):
        """Render tone, energy, and client engagement quality analysis."""
        content = self._convert_timestamps_to_mmss(content)
        subsections = {}
        current_key = None
        current_lines = []
        tone_scores = {}
        overall_score = ""
        ai_conf = ""

        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            header_match = re.match(
                r'^(TONE CONSISTENCY|ENERGY|ENTHUSIASM|CLIENT ENGAGEMENT|RAPPORT.BUILDING|'
                r'OVERALL TONE).*:', stripped, re.IGNORECASE)
            if header_match:
                if current_key and current_lines:
                    subsections[current_key] = '\n'.join(current_lines)
                current_key = header_match.group(1).upper().replace('&', 'AND')
                remainder = stripped[header_match.end():].strip()
                current_lines = [remainder] if remainder else []
                continue
            # Capture scores
            score_match = re.match(r'Score:\s*(\d+)\s*/\s*(\d+)', stripped, re.IGNORECASE)
            if score_match and current_key:
                tone_scores[current_key] = int(score_match.group(1))
                current_lines.append(stripped)
                continue
            if 'OVERALL TONE' in stripped.upper() and 'SCORE' in stripped.upper():
                val = stripped.split(':', 1)
                if len(val) > 1:
                    overall_score = val[1].strip()
                    sc_match = re.search(r'(\d+)\s*/\s*(\d+)', overall_score)
                    if sc_match:
                        tone_scores['OVERALL'] = int(sc_match.group(1))
                continue
            if 'AI CONFIDENCE' in stripped.upper() and not current_key:
                val = stripped.split(':', 1)
                if len(val) > 1:
                    ai_conf = val[1].strip()
                continue
            if current_key:
                current_lines.append(stripped)

        if current_key and current_lines:
            subsections[current_key] = '\n'.join(current_lines)

        # Cache tone scores for visualization
        self._cached_tone_scores = tone_scores

        # Score hero
        if overall_score:
            score_match = re.search(r'(\d+)\s*/\s*(\d+)', overall_score)
            if score_match:
                s_val = int(score_match.group(1))
                pct = s_val / 10
                color = "#43A047" if pct >= 0.75 else "#FF8F00" if pct >= 0.50 else "#E53935"
                hero = Table([
                    [Paragraph(f"<font color='{color}'><b>{overall_score}</b></font>",
                     ParagraphStyle('tone_hero', parent=self.ps['score_big'], fontSize=28, leading=34))],
                    [Spacer(1, 1)],
                    [Paragraph("OVERALL TONE & ENGAGEMENT SCORE", self.ps['score_label'])],
                    [Paragraph(f"<i>{ai_conf}</i>" if ai_conf else "",
                     ParagraphStyle('tone_conf', parent=self.ps['conf'], fontSize=7))],
                ], colWidths=[6.2 * inch], rowHeights=[40, 6, 16, 14])
                hero.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F5F7FA')),
                    ('LINEBELOW', (0, -1), (-1, -1), 1.5, colors.HexColor(color)),
                    ('PADDING', (0, 0), (-1, -1), 8),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ]))
                story.append(hero)
                story.append(Spacer(1, 10))

        # Subsection cards
        tone_display = {
            'TONE CONSISTENCY': ('Tone Consistency', '#E3F2FD', '#1565C0',
                                 'Was the agent professional, warm, authoritative, or inconsistent?'),
            'ENERGY': ('Energy & Enthusiasm', '#E8F5E9', '#2E7D32',
                       'Did the agent maintain appropriate energy throughout?'),
            'ENTHUSIASM': ('Energy & Enthusiasm', '#E8F5E9', '#2E7D32',
                           'Did the agent maintain appropriate energy throughout?'),
            'CLIENT ENGAGEMENT': ('Client Engagement Quality', '#FFF8E1', '#F57F17',
                                  'How effectively did the agent keep the client involved?'),
            'RAPPORT-BUILDING': ('Rapport-Building Language', '#F3E5F5', '#7B1FA2',
                                 'Specific phrases and techniques that built connection.'),
        }
        for key, (display, bg, tc, desc) in tone_display.items():
            sub_content = subsections.get(key, '')
            if not sub_content:
                continue
            inner = [
                Paragraph(f"<b>{display}</b>",
                         ParagraphStyle('tn_' + key[:8], parent=self.ps['subsec'],
                                        textColor=colors.HexColor(tc))),
                Paragraph(f"<i>{desc}</i>",
                         ParagraphStyle('tn_d_' + key[:8], parent=self.ps['body'], fontSize=7.5,
                                        textColor=self.C['light'])),
                Spacer(1, 4),
            ]
            for sl in sub_content.split('\n'):
                sl = sl.strip()
                if not sl:
                    continue
                sl = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', sl)
                sl = re.sub(r'\[(\d+[\.\d]*s?\s*-\s*\d+[\.\d]*s?)\]',
                            r"<font color='#1565C0'><b>[\1]</b></font>", sl)
                if sl.startswith('-') or sl.startswith('*'):
                    inner.append(Paragraph(f"<bullet>&bull;</bullet>{sl.lstrip('-* ')}",
                                           self.ps['bullet']))
                elif re.match(r'^Score:', sl, re.IGNORECASE):
                    inner.append(Paragraph(f"<font color='{tc}'><b>{sl}</b></font>", self.ps['body_bold']))
                else:
                    inner.append(Paragraph(sl, self.ps['body']))
                inner.append(Spacer(1, 2))

            card = Table([[inner]], colWidths=[6.2 * inch])
            card.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(bg)),
                ('PADDING', (0, 0), (-1, -1), 8),
                ('LINEBELOW', (0, 0), (-1, -1), 0.5, colors.HexColor(tc)),
            ]))
            story.append(card)
            story.append(Spacer(1, 6))

    # ── V. Negotiation & Persuasion Proficiency (NEW §18) ──
    def _render_negotiation_proficiency(self, story, content: str):
        """Render negotiation tactics, persuasion techniques, and deal-advancing analysis."""
        content = self._convert_timestamps_to_mmss(content)
        subsections = {}
        current_key = None
        current_lines = []
        neg_scores = {}
        overall_score = ""
        ai_conf = ""

        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            header_match = re.match(
                r'^(PERSUASION TECHNIQUES|NEGOTIATION POSITIONING|OBJECTION.TO.COMMITMENT|'
                r'VALUE FRAMING|CLOSING SIGNALS|OVERALL NEGOTIATION).*:', stripped, re.IGNORECASE)
            if header_match:
                if current_key and current_lines:
                    subsections[current_key] = '\n'.join(current_lines)
                current_key = header_match.group(1).upper().replace('-', '_')
                remainder = stripped[header_match.end():].strip()
                current_lines = [remainder] if remainder else []
                continue
            score_match = re.match(r'Score:\s*(\d+)\s*/\s*(\d+)', stripped, re.IGNORECASE)
            if score_match and current_key:
                neg_scores[current_key] = int(score_match.group(1))
                current_lines.append(stripped)
                continue
            if 'OVERALL NEGOTIATION' in stripped.upper() and 'SCORE' in stripped.upper():
                val = stripped.split(':', 1)
                if len(val) > 1:
                    overall_score = val[1].strip()
                    sc_match = re.search(r'(\d+)\s*/\s*(\d+)', overall_score)
                    if sc_match:
                        neg_scores['OVERALL'] = int(sc_match.group(1))
                continue
            if 'AI CONFIDENCE' in stripped.upper() and not current_key:
                val = stripped.split(':', 1)
                if len(val) > 1:
                    ai_conf = val[1].strip()
                continue
            if current_key:
                current_lines.append(stripped)

        if current_key and current_lines:
            subsections[current_key] = '\n'.join(current_lines)

        # Cache negotiation scores for visualization
        self._cached_negotiation_scores = neg_scores

        # Score hero
        if overall_score:
            score_match = re.search(r'(\d+)\s*/\s*(\d+)', overall_score)
            if score_match:
                s_val = int(score_match.group(1))
                pct = s_val / 10
                color = "#43A047" if pct >= 0.75 else "#FF8F00" if pct >= 0.50 else "#E53935"
                hero = Table([
                    [Paragraph(f"<font color='{color}'><b>{overall_score}</b></font>",
                     ParagraphStyle('neg_hero', parent=self.ps['score_big'], fontSize=28, leading=34))],
                    [Spacer(1, 1)],
                    [Paragraph("OVERALL NEGOTIATION & PERSUASION SCORE", self.ps['score_label'])],
                    [Paragraph(f"<i>{ai_conf}</i>" if ai_conf else "",
                     ParagraphStyle('neg_conf', parent=self.ps['conf'], fontSize=7))],
                ], colWidths=[6.2 * inch], rowHeights=[40, 6, 16, 14])
                hero.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F5F7FA')),
                    ('LINEBELOW', (0, -1), (-1, -1), 1.5, colors.HexColor(color)),
                    ('PADDING', (0, 0), (-1, -1), 8),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ]))
                story.append(hero)
                story.append(Spacer(1, 10))

        # Subsection cards
        neg_display = {
            'PERSUASION TECHNIQUES': ('Persuasion Techniques Used', '#E3F2FD', '#1565C0',
                                      'Social proof, scarcity, authority, reciprocity, anchoring, framing.'),
            'NEGOTIATION POSITIONING': ('Negotiation Positioning', '#E8F5E9', '#2E7D32',
                                        'Did the agent establish strong positioning or give away leverage?'),
            'OBJECTION_TO_COMMITMENT': ('Objection-to-Commitment Conversion', '#FFF8E1', '#F57F17',
                                        'Tracking each objection and whether it was converted toward commitment.'),
            'VALUE FRAMING': ('Value Framing Ability', '#F3E5F5', '#7B1FA2',
                              'How well did the agent frame value relative to price/competition?'),
            'CLOSING SIGNALS': ('Closing Signals & Trial Closes', '#FFEBEE', '#C62828',
                                'Did the agent use trial closes or test for commitment?'),
        }
        for key, (display, bg, tc, desc) in neg_display.items():
            sub_content = subsections.get(key, '')
            if not sub_content:
                continue
            inner = [
                Paragraph(f"<b>{display}</b>",
                         ParagraphStyle('ng_' + key[:8], parent=self.ps['subsec'],
                                        textColor=colors.HexColor(tc))),
                Paragraph(f"<i>{desc}</i>",
                         ParagraphStyle('ng_d_' + key[:8], parent=self.ps['body'], fontSize=7.5,
                                        textColor=self.C['light'])),
                Spacer(1, 4),
            ]
            for sl in sub_content.split('\n'):
                sl = sl.strip()
                if not sl:
                    continue
                sl = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', sl)
                sl = re.sub(r'\[(\d+[\.\d]*s?\s*-\s*\d+[\.\d]*s?)\]',
                            r"<font color='#1565C0'><b>[\1]</b></font>", sl)
                # Highlight effectiveness ratings
                for rating in ['Effective', 'Partially Effective', 'Ineffective',
                               'Resolved', 'Partially Resolved', 'Unresolved']:
                    if rating in sl:
                        r_color = '#43A047' if rating in ['Effective', 'Resolved'] else \
                                  '#FF8F00' if 'Partial' in rating else '#E53935'
                        sl = sl.replace(rating, f"<font color='{r_color}'><b>{rating}</b></font>")
                if sl.startswith('-') or sl.startswith('*'):
                    inner.append(Paragraph(f"<bullet>&bull;</bullet>{sl.lstrip('-* ')}",
                                           self.ps['bullet']))
                elif re.match(r'^Score:', sl, re.IGNORECASE):
                    inner.append(Paragraph(f"<font color='{tc}'><b>{sl}</b></font>", self.ps['body_bold']))
                elif any(marker in sl.upper() for marker in ['OBJECTION', '→', 'OUTCOME']):
                    inner.append(Paragraph(f"<b>{sl}</b>", self.ps['body_bold']))
                else:
                    inner.append(Paragraph(sl, self.ps['body']))
                inner.append(Spacer(1, 2))

            card = Table([[inner]], colWidths=[6.2 * inch])
            card.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(bg)),
                ('PADDING', (0, 0), (-1, -1), 8),
                ('LINEBELOW', (0, 0), (-1, -1), 0.5, colors.HexColor(tc)),
            ]))
            story.append(card)
            story.append(Spacer(1, 6))

    # ════════════════════════════════════════════════════════
    # COACHING INSIGHTS VISUALIZATION SECTION (PDF 3)
    # ════════════════════════════════════════════════════════

    def _build_coaching_insights_viz_section(self, story, sections: Dict[str, str]):
        """Coaching-focused visualizations — 12 charts:
        1. Conversation Stage Distribution (Donut)
        2. Agent vs Client Talking Ratio (Donut)
        3. Question Type Analysis (Stacked Bar)
        4. Active Listening Indicators (Radar)
        5. Objection Type Breakdown (HBar)
        6. Objection Handling Effectiveness (Stacked Bar)
        7. Sentiment Progression (Line)
        8. Behavioral Skill Radar (Radar)
        9. Sales Technique Usage (HBar)
        10. Coaching Priority Areas (HBar)
        11. Confidence vs Uncertainty Language (Stacked Bar) — redistributed
        12. Discovery Completeness Matrix (HBar) — redistributed
        """
        scores = getattr(self, '_cached_scores', {})
        talk_ratio = getattr(self, '_cached_talk_ratio', {})
        sentiment_phases = getattr(self, '_cached_sentiment_phases', [])
        discovery_pillars = getattr(self, '_cached_discovery_pillars', [])
        ethics_subs = getattr(self, '_cached_ethics_subsections', {})
        listening_subs = getattr(self, '_cached_listening_subsections', {})
        eq_dims = getattr(self, '_cached_eq_dimensions', {})
        tone_scores = getattr(self, '_cached_tone_scores', {})
        neg_scores = getattr(self, '_cached_negotiation_scores', {})
        transcript_data_ref = getattr(self, '_transcript_data', None)

        # ── Pre-compute transcript-level data ──
        agent_segments = []
        client_segments = []
        all_segments = []
        if transcript_data_ref and transcript_data_ref.transcripts:
            for seg in transcript_data_ref.transcripts:
                all_segments.append(seg)
                if seg.speaker_name and seg.speaker_name != 'Client':
                    agent_segments.append(seg)
                else:
                    client_segments.append(seg)
        agent_text_all = ' '.join(s.transcript for s in agent_segments).lower()
        client_text_all = ' '.join(s.transcript for s in client_segments).lower()
        agent_word_count = len(agent_text_all.split())
        client_word_count = len(client_text_all.split())

        story.append(SectionDivider(width=460, color="#0D47A1", thickness=2))
        story.append(Spacer(1, 2))
        story.append(Paragraph("COACHING INSIGHTS — VISUAL ANALYTICS", self.ps['sec']))
        story.append(Spacer(1, 4))
        story.append(Paragraph(
            "These visualizations focus on coaching-relevant insights: conversation behavior analysis, "
            "communication quality, objection handling, discovery gaps, and improvement recommendations. "
            "All data is derived from the actual transcript analysis.",
            self.ps['body']))
        story.append(Spacer(1, 10))

        # ════════════════════════════════════════════════════════
        # CHART 1 — Conversation Stage Distribution (Donut)
        # ════════════════════════════════════════════════════════
        story.append(Paragraph("<b>1. Conversation Stage Distribution</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        # Classify each segment into a conversation stage based on position + content keywords
        stage_counts = {}
        if all_segments:
            total_segs = len(all_segments)
            stage_keywords = {
                'Opening / Rapport': ['hello', 'hi ', 'nice to meet', 'how are you', 'thank you for',
                                      'appreciate', 'good morning', 'good afternoon', 'welcome'],
                'Needs Discovery': ['what are you looking', 'what do you need', 'tell me about',
                                    'what\'s important', 'budget', 'timeline', 'motivation',
                                    'why are you', 'how long', 'what brings', 'goals'],
                'Solution Presentation': ['i recommend', 'here\'s what', 'our approach', 'let me show',
                                          'the value', 'what we offer', 'the benefit', 'solution',
                                          'marketing plan', 'strategy', 'comparable', 'market analysis'],
                'Objection Handling': ['concern', 'worry', 'but what about', 'i\'m not sure',
                                       'that\'s too', 'what if', 'i don\'t think', 'hesit',
                                       'the problem', 'my concern', 'i understand your concern'],
                'Closing / Next Steps': ['next step', 'moving forward', 'sign', 'contract', 'agreement',
                                         'schedule', 'follow up', 'i\'ll send', 'let\'s set up',
                                         'action item', 'commitment', 'we\'ll do'],
            }
            for i, seg in enumerate(all_segments):
                text_lower = seg.transcript.lower()
                position_frac = i / max(total_segs - 1, 1)
                best_stage = None
                best_hits = 0
                for stage_name, keywords in stage_keywords.items():
                    hits = sum(1 for kw in keywords if kw in text_lower)
                    if hits > best_hits:
                        best_hits = hits
                        best_stage = stage_name
                if best_stage is None:
                    # Fallback: classify by position
                    if position_frac <= 0.15:
                        best_stage = 'Opening / Rapport'
                    elif position_frac <= 0.40:
                        best_stage = 'Needs Discovery'
                    elif position_frac <= 0.65:
                        best_stage = 'Solution Presentation'
                    elif position_frac <= 0.85:
                        best_stage = 'Objection Handling'
                    else:
                        best_stage = 'Closing / Next Steps'
                stage_counts[best_stage] = stage_counts.get(best_stage, 0) + 1

        if stage_counts:
            stage_colors = {
                'Opening / Rapport': '#43A047',
                'Needs Discovery': '#1565C0',
                'Solution Presentation': '#FF8F00',
                'Objection Handling': '#7B1FA2',
                'Closing / Next Steps': '#E53935',
            }
            donut_data = [(stage, count, stage_colors.get(stage, '#757575'))
                          for stage, count in stage_counts.items() if count > 0]
            if donut_data:
                story.append(Paragraph(
                    "Distribution of conversation segments across sales stages, classified by content keywords and position. "
                    "<font color='#43A047'>■</font> Opening / Rapport, <font color='#1565C0'>■</font> Needs Discovery, "
                    "<font color='#FF8F00'>■</font> Solution Presentation, <font color='#7B1FA2'>■</font> Objection Handling, "
                    "<font color='#E53935'>■</font> Closing / Next Steps.",
                    ParagraphStyle('ci_stage_d', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
                story.append(Spacer(1, 4))
                donut = DonutChartFlowable(donut_data, width=220, height=220)
                d_table = Table([[donut]], colWidths=[6.2 * inch])
                d_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                    ('PADDING', (0, 0), (-1, -1), 10),
                ]))
                story.append(KeepTogether(d_table))
                story.append(Spacer(1, 12))

                self.coaching_viz_logger.log_chart(
                    chart_id="conversation_stage_distribution",
                    title="Conversation Stage Distribution",
                    chart_type="donut",
                    labels=[d[0] for d in donut_data],
                    datasets=[{"label": "Segments", "values": [d[1] for d in donut_data],
                               "colors": [d[2] for d in donut_data]}],
                    description="Distribution of conversation segments across sales stages.",
                    metadata={"total_segments": sum(d[1] for d in donut_data)},
                )
            else:
                story.append(Paragraph(
                    "<i>Insufficient data from the transcript to generate this chart. (No conversation segments available for stage classification.)</i>",
                    ParagraphStyle('ci_no_stage', parent=self.ps['body'], fontSize=8,
                                   textColor=self.C['light'], alignment=TA_CENTER)))
                story.append(Spacer(1, 12))
        else:
            story.append(Paragraph(
                "<i>Insufficient data from the transcript to generate this chart. (No transcript data available for conversation stage analysis.)</i>",
                ParagraphStyle('ci_no_stage2', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 2 — Agent vs Client Talking Ratio (Donut)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(2.5 * inch))
        story.append(Paragraph("<b>2. Agent vs Client Talking Ratio</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        agent_pct = talk_ratio.get('agent')
        client_pct = talk_ratio.get('client')
        # Also calculate from raw transcript if cached ratio unavailable
        if agent_pct is None and agent_word_count + client_word_count > 0:
            total_words = agent_word_count + client_word_count
            agent_pct = round(agent_word_count / total_words * 100) if total_words else None
            client_pct = 100 - agent_pct if agent_pct is not None else None
        if agent_pct is not None and client_pct is not None:
            story.append(Paragraph(
                "Ideal zone: 40-55% agent talk. Excessive agent talk suggests pushy behavior; too little suggests weak guidance. "
                "<font color='#43A047'>■</font> Ideal, <font color='#FF8F00'>■</font> Slightly off, <font color='#E53935'>■</font> Needs adjustment.",
                ParagraphStyle('ci_talk_d', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            donut = DonutChartFlowable([
                ("Agent", agent_pct, "#1565C0"),
                ("Client", client_pct, "#43A047")], width=220, height=220)
            donut_table = Table([[donut]], colWidths=[6.2 * inch])
            donut_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 10),
            ]))
            story.append(KeepTogether(donut_table))

            if 40 <= agent_pct <= 55:
                zone_text = "Agent is in the IDEAL talk-time zone."
                zone_color = "#2E7D32"
            elif agent_pct > 65:
                zone_text = "Agent-dominated conversation — reduce talking, ask more questions."
                zone_color = "#C62828"
            elif agent_pct < 40:
                zone_text = "Agent too passive — needs to guide the conversation more firmly."
                zone_color = "#F57F17"
            else:
                zone_text = "Slightly outside ideal range — minor adjustment needed."
                zone_color = "#F57F17"
            story.append(Paragraph(f"<font color='{zone_color}'><b>{zone_text}</b></font>",
                         ParagraphStyle('ci_zone2', parent=self.ps['body'], fontSize=8, alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

            # Add note about derivation
            note = "(Derived from transcript analysis – word counts of agent and client turns.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ci_note', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.coaching_viz_logger.log_chart(
                chart_id="agent_vs_client_talking_ratio",
                title="Agent vs Client Talking Ratio",
                chart_type="donut",
                labels=["Agent", "Client"],
                datasets=[{"label": "Talk Time %", "values": [agent_pct, client_pct],
                           "colors": ["#1565C0", "#43A047"]}],
                description=f"Ideal zone: {self.CONFIG['talk_ratio_ideal'][0]}-{self.CONFIG['talk_ratio_ideal'][1]}% agent talk. Shows conversation dominance balance. {note}",
                metadata={"ideal_range": [40, 55], "zone_verdict": zone_text, "metric_source": "transcript_analysis"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient data from the transcript to generate this chart. (Talk-time ratio could not be determined from the transcript analysis.)</i>",
                ParagraphStyle('ci_no_donut2', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 3 — Question Type Analysis (Stacked Bar)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(2.5 * inch))
        story.append(Paragraph("<b>3. Question Type Analysis</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        agent_questions = []
        for seg in agent_segments:
            sentences = re.split(r'[.!?]+', seg.transcript)
            for s in sentences:
                s = s.strip()
                if '?' in s and s and any(w in s.lower() for w in ['what', 'how', 'why', 'when', 'where', 'who', 'which', 'do', 'does', 'did', 'would', 'could', 'can', 'is', 'are', 'was', 'were', 'have', 'has']):
                    agent_questions.append(s)

        if agent_questions:
            open_ended = sum(1 for q in agent_questions if any(w in q.lower() for w in ['what', 'how', 'why', 'tell me', 'describe', 'explain', 'walk me through']))
            closed = sum(1 for q in agent_questions if any(w in q.lower() for w in ['is ', 'are ', 'do ', 'does ', 'did ', 'can ', 'was ', 'were ', 'have you', 'has ']))
            leading = sum(1 for q in agent_questions if any(w in q.lower() for w in ['don\'t you', 'wouldn\'t you', 'isn\'t it', 'right?', 'agree?', 'correct?', 'fair?']))
            assumptive = sum(1 for q in agent_questions if any(w in q.lower() for w in ['when you', 'once we', 'after we', 'so you\'re']))
            q_total = open_ended + closed + leading + assumptive
            if q_total > 0:
                story.append(Paragraph(
                    f"Classified {len(agent_questions)} questions from the agent's actual dialogue. "
                    "<font color='#43A047'>■</font> Open-Ended, <font color='#1565C0'>■</font> Closed, "
                    "<font color='#FF8F00'>■</font> Leading, <font color='#E53935'>■</font> Assumptive.",
                    ParagraphStyle('ci_qt_d3', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
                story.append(Spacer(1, 4))
                q_data = [
                    ("All Questions", [
                        ("Open-Ended", open_ended, "#43A047"),
                        ("Closed", closed, "#1565C0"),
                        ("Leading", leading, "#FF8F00"),
                        ("Assumptive", assumptive, "#E53935"),
                    ]),
                ]
                q_chart = StackedBarFlowable(q_data, width=420, bar_height=28)
                qc_table = Table([[q_chart]], colWidths=[6.2 * inch])
                qc_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                    ('PADDING', (0, 0), (-1, -1), 8),
                ]))
                story.append(KeepTogether(qc_table))
                story.append(Spacer(1, 12))

                note = "(Derived from transcript analysis – counts of question types detected in agent dialogue.)"
                story.append(Paragraph(f"<i>{note}</i>",
                             ParagraphStyle('ci_note_q', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
                story.append(Spacer(1, 2))

                self.coaching_viz_logger.log_chart(
                    chart_id="question_type_analysis",
                    title="Question Type Analysis",
                    chart_type="stackedBar",
                    labels=["All Questions"],
                    datasets=[
                        {"label": "Open-Ended", "values": [open_ended], "colors": ["#43A047"]},
                        {"label": "Closed", "values": [closed], "colors": ["#1565C0"]},
                        {"label": "Leading", "values": [leading], "colors": ["#FF8F00"]},
                        {"label": "Assumptive", "values": [assumptive], "colors": ["#E53935"]},
                    ],
                    description=f"Classified {len(agent_questions)} questions from agent dialogue.",
                    metadata={"total_questions": len(agent_questions),
                              "open_ended": open_ended, "closed": closed,
                              "leading": leading, "assumptive": assumptive,
                              "metric_source": "transcript_analysis"},
                )
            else:
                story.append(Paragraph(
                    "<i>Insufficient data from the transcript to generate this chart. (No classifiable questions detected in agent dialogue.)</i>",
                    ParagraphStyle('ci_no_q3', parent=self.ps['body'], fontSize=8,
                                   textColor=self.C['light'], alignment=TA_CENTER)))
                story.append(Spacer(1, 12))
        else:
            story.append(Paragraph(
                "<i>Insufficient data from the transcript to generate this chart. (No questions detected in agent dialogue.)</i>",
                ParagraphStyle('ci_no_q4', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 4 — Active Listening Indicators (Radar)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(3.0 * inch))
        story.append(Paragraph("<b>4. Active Listening Indicators</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        listening_scores = {}
        if listening_subs:
            for key, content in listening_subs.items():
                if 'SCORE' in key.upper() or 'RATING' in key.upper():
                    sc_m = re.search(r'(\d+)\s*/\s*(\d+)', str(content))
                    if sc_m:
                        listening_scores[key.title()] = int(sc_m.group(1))
                elif isinstance(content, (int, float)):
                    listening_scores[key.title()] = int(content)

        if not listening_scores and agent_segments:
            agent_text_listen = ' '.join(s.transcript for s in agent_segments).lower()
            paraphrasing = sum(agent_text_listen.count(p) for p in ['i understand', 'what you\'re saying', 'in other words', 'so you mean', 'let me make sure'])
            validation = sum(agent_text_listen.count(p) for p in ['that makes sense', 'i can see why', 'i hear you', 'i appreciate', 'thank you for sharing'])
            clarification = sum(agent_text_listen.count(p) for p in ['can you tell me more', 'help me understand', 'what do you mean by', 'could you elaborate'])
            empathy_count = sum(agent_text_listen.count(p) for p in ['i can imagine', 'that must be', 'i understand how', 'that sounds', 'i feel'])
            follow_up = sum(agent_text_listen.count(p) for p in ['building on that', 'related to what you said', 'earlier you mentioned', 'going back to'])
            listening_scores = {
                'Paraphrasing': min(10, paraphrasing * 2),
                'Validation': min(10, validation * 2),
                'Clarification': min(10, clarification * 2),
                'Empathy': min(10, empathy_count * 2),
                'Follow-Up': min(10, follow_up * 2),
            }
            # Add note that these are derived from transcript
            note = "Note: Scores are derived from transcript analysis – counts of paraphrasing, validation, etc."
        else:
            note = ""

        if listening_scores and len(listening_scores) >= 3:
            story.append(Paragraph(
                "Multi-dimensional assessment of active listening behaviors demonstrated in the conversation. "
                "<font color='#43A047'>■</font> Green = strong (≥7.5), <font color='#FF8F00'>■</font> Yellow = moderate (5.0-7.4), <font color='#E53935'>■</font> Red = needs improvement (<5.0).",
                ParagraphStyle('ci_listen_d4', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            listen_radar = RadarChartFlowable(listening_scores, width=300, height=300)
            lr_table = Table([[listen_radar]], colWidths=[6.2 * inch])
            lr_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 10),
            ]))
            story.append(KeepTogether(lr_table))
            story.append(Spacer(1, 12))

            if note:
                story.append(Paragraph(f"<i>{note}</i>",
                             ParagraphStyle('ci_note_listen', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
                story.append(Spacer(1, 2))

            self.coaching_viz_logger.log_chart(
                chart_id="active_listening_indicators",
                title="Active Listening Indicators",
                chart_type="radar",
                labels=list(listening_scores.keys()),
                datasets=[{"label": "Listening Score", "values": list(listening_scores.values()),
                           "colors": ["#1565C0"], "max": 10}],
                description="Multi-dimensional assessment of active listening behaviors.",
                options={"max_value": 10},
                metadata={"metric_source": "transcript_analysis" if note else "llm_evaluation"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the transcript to generate this chart. (No active listening indicators detected in the transcript.)</i>",
                ParagraphStyle('ci_no_listen4', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 5 — Objection Type Breakdown (HBar)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(2.5 * inch))
        story.append(Paragraph("<b>5. Objection Type Breakdown</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        objection_counts = {}
        objection_patterns = {
            'Price Concern': ['too expensive', 'cost', 'price is', 'budget', 'afford', 'money', 'commission'],
            'Timing Concern': ['not ready', 'timing', 'too soon', 'wait', 'later', 'time'],
            'Competition': ['other agent', 'another company', 'competitor', 'elsewhere', 'friend recommended'],
            'Trust / Skepticism': ['not sure', 'doubt', 'skeptical', 'worried', 'concerned', 'hesitant'],
            'Need Clarification': ['don\'t understand', 'confused', 'explain', 'clarify', 'what do you mean'],
            'Stall Tactic': ['think about it', 'let me talk to', 'discuss with', 'get back to you'],
        }
        if client_segments:
            client_text = ' '.join(s.transcript for s in client_segments).lower()
            for obj_type, patterns in objection_patterns.items():
                count = sum(client_text.count(p) for p in patterns)
                if count > 0:
                    objection_counts[obj_type] = count

        if objection_counts:
            obj_bar_data = [(obj, count, "#E53935") for obj, count in objection_counts.items()]
            obj_bar_data.sort(key=lambda x: x[1], reverse=True)
            story.append(Paragraph(
                "Frequency of different objection types raised by the client during the conversation. "
                "<font color='#43A047'>■</font> Strong, <font color='#FF8F00'>■</font> Moderate, <font color='#E53935'>■</font> Needs improvement.",
                ParagraphStyle('ci_obj_d5', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            obj_chart = HBarChartFlowable(obj_bar_data, width=440, bar_height=20,
                                          max_val=max(c for _, c, _ in obj_bar_data))
            oc_table = Table([[obj_chart]], colWidths=[6.2 * inch])
            oc_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(KeepTogether(oc_table))
            story.append(Spacer(1, 12))

            note = "(Derived from transcript analysis – counts of objection keywords in client dialogue.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ci_note_obj', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.coaching_viz_logger.log_chart(
                chart_id="objection_type_breakdown",
                title="Objection Type Breakdown",
                chart_type="horizontalBar",
                labels=[d[0] for d in obj_bar_data],
                datasets=[{"label": "Frequency", "values": [d[1] for d in obj_bar_data],
                           "colors": [d[2] for d in obj_bar_data]}],
                description="Frequency of different objection types raised by the client.",
                metadata={"total_objections": sum(d[1] for d in obj_bar_data), "metric_source": "transcript_analysis"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the transcript to generate this chart. (No objections detected in client dialogue.)</i>",
                ParagraphStyle('ci_no_obj5', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 6 — Objection Handling Effectiveness (Stacked Bar)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(2.5 * inch))
        story.append(Paragraph("<b>6. Objection Handling Effectiveness</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        handling_data = {}
        if objection_counts and agent_segments:
            for obj_type in objection_counts.keys():
                addressed = 0
                total_mentions = 0
                for seg in agent_segments:
                    seg_lower = seg.transcript.lower()
                    if any(p in seg_lower for p in objection_patterns.get(obj_type, [])):
                        total_mentions += 1
                        if any(eff in seg_lower for eff in ['i understand', 'let me address', 'here\'s why',
                                                            'what if we', 'consider this', 'i can offer',
                                                            'the reason', 'that\'s a great question',
                                                            'let me explain', 'the benefit']):
                            addressed += 1
                if total_mentions > 0:
                    handling_data[obj_type] = {'addressed': addressed, 'unaddressed': total_mentions - addressed}

        if handling_data:
            eff_categories = []
            for obj_type, counts in handling_data.items():
                eff_categories.append((obj_type, [
                    ("Addressed", counts['addressed'], "#43A047"),
                    ("Unaddressed", counts['unaddressed'], "#E53935"),
                ]))
            story.append(Paragraph(
                "How effectively the agent addressed each type of objection. "
                "<font color='#43A047'>■</font> Addressed, <font color='#E53935'>■</font> Unaddressed.",
                ParagraphStyle('ci_eff_d6', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            eff_chart = StackedBarFlowable(eff_categories, width=420, bar_height=24)
            ef_table = Table([[eff_chart]], colWidths=[6.2 * inch])
            ef_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(KeepTogether(ef_table))
            story.append(Spacer(1, 12))

            note = "(Derived from transcript analysis – based on agent's responses following client objection mentions.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ci_note_eff', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.coaching_viz_logger.log_chart(
                chart_id="objection_handling_effectiveness",
                title="Objection Handling Effectiveness",
                chart_type="stackedBar",
                labels=[obj for obj in handling_data.keys()],
                datasets=[
                    {"label": "Addressed", "values": [v['addressed'] for v in handling_data.values()],
                     "colors": ["#43A047"]},
                    {"label": "Unaddressed", "values": [v['unaddressed'] for v in handling_data.values()],
                     "colors": ["#E53935"]},
                ],
                description="How effectively the agent addressed each type of objection.",
                metadata={"handling_rates": {k: v for k, v in handling_data.items()}, "metric_source": "transcript_analysis"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the transcript to generate this chart. (No objection handling patterns detected.)</i>",
                ParagraphStyle('ci_no_eff6', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 7 — Sentiment Progression (Line)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(2.5 * inch))
        story.append(Paragraph("<b>7. Sentiment Progression</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        sentiment_data = []
        sentiment_labels = []
        if sentiment_phases:
            for phase_name, phase_content in sentiment_phases:
                sent_match = re.search(r'(\d+)\s*/\s*10', phase_content)
                if sent_match:
                    sentiment_data.append(int(sent_match.group(1)))
                else:
                    content_lower = phase_content.lower()
                    if any(w in content_lower for w in ['skeptic', 'hesitant', 'anxious', 'distrust']):
                        sentiment_data.append(3)
                    elif any(w in content_lower for w in ['neutral', 'cautious', 'guarded']):
                        sentiment_data.append(5)
                    elif any(w in content_lower for w in ['engaged', 'interested', 'warming']):
                        sentiment_data.append(7)
                    elif any(w in content_lower for w in ['confident', 'trust', 'committed', 'positive']):
                        sentiment_data.append(8)
                    else:
                        continue
                sentiment_labels.append(f"{phase_name} Phase")

        if sentiment_data and len(sentiment_data) >= 2:
            story.append(Paragraph(
                "Shows where trust increased or dropped across meeting phases. "
                "<font color='#43A047'>■</font> Green = positive, <font color='#FF8F00'>■</font> Yellow = neutral, <font color='#E53935'>■</font> Red = negative.",
                ParagraphStyle('ci_sent_d7', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            line_chart = LineChartFlowable(
                sentiment_data, width=420, height=160,
                y_label="Sentiment", x_labels=sentiment_labels,
                line_color="#1565C0", fill_color="#E3F2FD",
                y_range=(0, 10))
            lc_table = Table([[line_chart]], colWidths=[6.2 * inch])
            lc_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 10),
            ]))
            story.append(KeepTogether(lc_table))
            story.append(Spacer(1, 12))

            note = "(Derived from transcript analysis – based on LLM-evaluated sentiment per phase.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ci_note_sent', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.coaching_viz_logger.log_chart(
                chart_id="sentiment_progression",
                title="Sentiment Progression",
                chart_type="line",
                labels=sentiment_labels,
                datasets=[{"label": "Sentiment Score", "values": sentiment_data,
                           "colors": ["#1565C0"], "fill": True, "fillColor": "#E3F2FD"}],
                description="Shows where trust increased or dropped across meeting phases.",
                options={"y_range": [0, 10], "y_label": "Sentiment"},
                metadata={"metric_source": "llm_evaluation"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the transcript to generate this chart. (No sentiment progression data available from the transcript.)</i>",
                ParagraphStyle('ci_no_sent7', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 8 — Behavioral Skill Radar (Radar)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(3.0 * inch))
        story.append(Paragraph("<b>8. Behavioral Skill Radar</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        behavioral_scores = {}
        if scores:
            skill_mapping = {
                'Rapport Building': ['rapport', 'engagement', 'client engagement'],
                'Questioning': ['needs discovery', 'questioning', 'discovery'],
                'Active Listening': ['listening', 'listening quality'],
                'Objection Handling': ['objection', 'handling'],
                'Value Communication': ['value proposition', 'communication'],
                'Closing Skills': ['closing', 'deal progression'],
                'Confidence': ['confidence', 'authority'],
                'Adaptability': ['emotional intelligence', 'adaptability'],
            }
            for skill, keywords in skill_mapping.items():
                for score_key, score_val in scores.items():
                    if any(kw in score_key.lower() for kw in keywords):
                        behavioral_scores[skill] = score_val
                        break

        if behavioral_scores and len(behavioral_scores) >= 3:
            story.append(Paragraph(
                "Multi-dimensional view of behavioral skills demonstrated during the conversation. "
                "<font color='#43A047'>■</font> Green = strong (≥7.5), <font color='#FF8F00'>■</font> Yellow = moderate (5.0-7.4), <font color='#E53935'>■</font> Red = needs improvement (<5.0).",
                ParagraphStyle('ci_behav_d8', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            behave_radar = RadarChartFlowable(behavioral_scores, width=300, height=300)
            br_table = Table([[behave_radar]], colWidths=[6.2 * inch])
            br_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 10),
            ]))
            story.append(KeepTogether(br_table))
            story.append(Spacer(1, 12))

            note = "(Derived from LLM-evaluated performance scores.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ci_note_behav', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.coaching_viz_logger.log_chart(
                chart_id="behavioral_skill_radar",
                title="Behavioral Skill Radar",
                chart_type="radar",
                labels=list(behavioral_scores.keys()),
                datasets=[{"label": "Skill Score", "values": list(behavioral_scores.values()),
                           "colors": ["#FF8F00"], "max": 10}],
                description="Multi-dimensional view of behavioral skills demonstrated.",
                options={"max_value": 10},
                metadata={"metric_source": "llm_evaluation"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the transcript to generate this chart. (No behavioral skill scores available from transcript analysis.)</i>",
                ParagraphStyle('ci_no_behav8', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 9 — Sales Technique Usage (HBar)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(2.5 * inch))
        story.append(Paragraph("<b>9. Sales Technique Usage</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        technique_counts = {}
        if agent_segments:
            agent_text_tech = ' '.join(s.transcript for s in agent_segments).lower()
            techniques = {
                'Social Proof': ['other clients', 'similar situation', 'many people', 'our clients have',
                                 'case study', 'track record', 'experience with'],
                'Scarcity': ['limited time', 'only available', 'running out', 'don\'t miss',
                             'time sensitive', 'won\'t last'],
                'Authority': ['as a professional', 'my experience', 'industry standard', 'expertise',
                              'certified', 'years in the business'],
                'Reciprocity': ['i\'ll do this for you', 'in return', 'because you',
                                'since you mentioned', 'as a courtesy'],
                'Anchoring': ['starting at', 'normally costs', 'regular price', 'discounted from',
                              'valued at', 'compared to'],
                'Framing': ['think of it as', 'consider this', 'the real value', 'what matters most',
                            'the way i see it', 'perspective'],
            }
            for tech, patterns in techniques.items():
                count = sum(agent_text_tech.count(p) for p in patterns)
                if count > 0:
                    technique_counts[tech] = count

        if technique_counts:
            tech_bar_data = [(tech, count, "#7B1FA2") for tech, count in technique_counts.items()]
            tech_bar_data.sort(key=lambda x: x[1], reverse=True)
            story.append(Paragraph(
                "Frequency of different sales/persuasion techniques detected in agent dialogue. "
                "<font color='#43A047'>■</font> Strong usage, <font color='#FF8F00'>■</font> Moderate, <font color='#E53935'>■</font> Low.",
                ParagraphStyle('ci_tech_d9', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            tech_chart = HBarChartFlowable(tech_bar_data, width=440, bar_height=20,
                                           max_val=max(c for _, c, _ in tech_bar_data))
            tc_table = Table([[tech_chart]], colWidths=[6.2 * inch])
            tc_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(KeepTogether(tc_table))
            story.append(Spacer(1, 12))

            note = "(Derived from transcript analysis – keyword matching for sales techniques.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ci_note_tech', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.coaching_viz_logger.log_chart(
                chart_id="sales_technique_usage",
                title="Sales Technique Usage",
                chart_type="horizontalBar",
                labels=[d[0] for d in tech_bar_data],
                datasets=[{"label": "Usage Frequency", "values": [d[1] for d in tech_bar_data],
                           "colors": [d[2] for d in tech_bar_data]}],
                description="Frequency of different sales/persuasion techniques detected in agent dialogue.",
                metadata={"total_techniques": sum(d[1] for d in tech_bar_data), "metric_source": "transcript_analysis"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the transcript to generate this chart. (No sales techniques detected in agent dialogue.)</i>",
                ParagraphStyle('ci_no_tech9', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 10 — Coaching Priority Areas (HBar)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(2.5 * inch))
        story.append(Paragraph("<b>10. Coaching Priority Areas</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        if scores:
            sorted_scores = sorted(scores.items(), key=lambda x: x[1])
            prio_bar_data = []
            for label, val in sorted_scores[:5]:
                pct = val / 10
                col = self._get_score_color(val)
                prio_bar_data.append((label, val, col))
            story.append(Paragraph(
                "The five lowest-scoring dimensions requiring immediate coaching attention. "
                "<font color='#43A047'>■</font> Green = strong (≥7.5), <font color='#FF8F00'>■</font> Yellow = moderate (5.0-7.4), <font color='#E53935'>■</font> Red = needs improvement (<5.0).",
                ParagraphStyle('ci_prio_d10', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            prio_chart = HBarChartFlowable(prio_bar_data, width=440, bar_height=18, max_val=10)
            pc_table = Table([[prio_chart]], colWidths=[6.2 * inch])
            pc_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(KeepTogether(pc_table))
            story.append(Spacer(1, 12))

            note = "(Derived from LLM-evaluated performance scores.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ci_note_prio', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.coaching_viz_logger.log_chart(
                chart_id="coaching_priority_areas",
                title="Coaching Priority Areas",
                chart_type="horizontalBar",
                labels=[d[0] for d in prio_bar_data],
                datasets=[{"label": "Score", "values": [d[1] for d in prio_bar_data],
                           "colors": [d[2] for d in prio_bar_data]}],
                description="The five lowest-scoring dimensions requiring immediate coaching attention.",
                options={"max_value": 10},
                metadata={"metric_source": "llm_evaluation"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the transcript to generate this chart. (No performance scores available for coaching priorities.)</i>",
                ParagraphStyle('ci_no_prio10', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 11 — Confidence vs Uncertainty Language (Stacked Bar)
        # (Redistributed from full_report_archive)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(2.5 * inch))
        story.append(Paragraph("<b>11. Confidence vs Uncertainty Language</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        transcript_data_ref = getattr(self, '_transcript_data', None)
        agent_segs_text = []
        if transcript_data_ref and transcript_data_ref.transcripts:
            agent_segs_text = [seg.transcript for seg in transcript_data_ref.transcripts
                               if seg.speaker_name and seg.speaker_name != 'Client']
        agent_text_lang = ' '.join(agent_segs_text).lower()

        confident_phrases = ['i will', 'we will', 'absolutely', 'definitely', 'i guarantee',
                             'without a doubt', 'i assure', 'i promise', 'my goal is',
                             'what i want', 'here\'s what', 'the fact is', 'i believe',
                             'i recommend', 'my recommendation', 'i\'m going to',
                             'we\'re going to', 'let me walk you through', 'let me explain']
        hedging_phrases = ['i think', 'maybe', 'perhaps', 'probably', 'might', 'could be',
                           'i guess', 'sort of', 'kind of', 'it depends', 'not sure',
                           'i don\'t know', 'possibly', 'hopefully', 'we\'ll see']
        filler_phrases = ['you know', 'i mean', 'like', 'um', 'uh', 'basically',
                          'obviously', 'literally', 'actually', 'right right right']

        confident_count = sum(agent_text_lang.count(p) for p in confident_phrases)
        hedging_count = sum(agent_text_lang.count(p) for p in hedging_phrases)
        filler_count = sum(agent_text_lang.count(p) for p in filler_phrases)
        total_lang = confident_count + hedging_count + filler_count

        if total_lang > 0:
            story.append(Paragraph(
                "Counts of confident vs hedging phrases from the agent's actual dialogue in this transcript. "
                "<font color='#43A047'>■</font> Confident, <font color='#FF8F00'>■</font> Hedging, <font color='#E53935'>■</font> Filler.",
                ParagraphStyle('ci_conf_d11', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            conf_data = [("Language Type", [
                ("Confident", confident_count, "#43A047"),
                ("Hedging", hedging_count, "#FF8F00"),
                ("Filler Words", filler_count, "#E53935"),
            ])]
            conf_chart = StackedBarFlowable(conf_data, width=420, bar_height=28)
            cc_table = Table([[conf_chart]], colWidths=[6.2 * inch])
            cc_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(KeepTogether(cc_table))
            story.append(Spacer(1, 4))
            story.append(Paragraph(
                f"<i>Detected {confident_count} confident, {hedging_count} hedging, and {filler_count} filler phrase instances.</i>",
                ParagraphStyle('ci_lang_note', parent=self.ps['body'], fontSize=7.5,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

            note = "(Derived from transcript analysis – counts of specific phrases.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ci_note_lang', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.coaching_viz_logger.log_chart(
                chart_id="confidence_vs_uncertainty_language",
                title="Confidence vs Uncertainty Language",
                chart_type="stackedBar",
                labels=["Language Type"],
                datasets=[
                    {"label": "Confident", "values": [confident_count], "colors": ["#43A047"]},
                    {"label": "Hedging", "values": [hedging_count], "colors": ["#FF8F00"]},
                    {"label": "Filler Words", "values": [filler_count], "colors": ["#E53935"]},
                ],
                description="Counts of confident vs hedging phrases from the agent's actual dialogue.",
                metadata={"total_phrases": total_lang, "confident_count": confident_count,
                          "hedging_count": hedging_count, "filler_count": filler_count,
                          "metric_source": "transcript_analysis"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the transcript to generate this chart. (No language patterns detected in agent dialogue.)</i>",
                ParagraphStyle('ci_no_lang11', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 12 — Discovery Completeness Matrix (HBar)
        # (Redistributed from full_report_archive)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(2.0 * inch))
        story.append(Paragraph("<b>12. Discovery Completeness Matrix</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        if discovery_pillars:
            story.append(Paragraph(
                "Objective diagnostic: which pillars were covered, partially explored, or completely missed. "
                "<font color='#43A047'>■</font> Green = Confirmed/Strong, <font color='#FF8F00'>■</font> Yellow = Partially explored/Moderate, <font color='#E53935'>■</font> Red = Not addressed/Weak.",
                ParagraphStyle('ci_disc_d12', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            disc_bar_data = []
            for pillar, status, evidence, confidence in discovery_pillars:
                status_lower = status.lower()
                if 'confirmed' in status_lower:
                    val, col = 9, "#43A047"
                elif 'partially' in status_lower:
                    val, col = 5, "#FF8F00"
                else:
                    val, col = 2, "#E53935"
                disc_bar_data.append((pillar, val, col))
            disc_chart = HBarChartFlowable(disc_bar_data, width=440, bar_height=20, max_val=10)
            dc_table = Table([[disc_chart]], colWidths=[6.2 * inch])
            dc_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(KeepTogether(dc_table))
            story.append(Spacer(1, 12))

            note = "(Derived from LLM-evaluated discovery pillars.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ci_note_disc', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.coaching_viz_logger.log_chart(
                chart_id="discovery_completeness_matrix",
                title="Discovery Completeness Matrix",
                chart_type="horizontalBar",
                labels=[d[0] for d in disc_bar_data],
                datasets=[{"label": "Coverage Score", "values": [d[1] for d in disc_bar_data],
                           "colors": [d[2] for d in disc_bar_data]}],
                description="Objective diagnostic: which pillars were covered, partially explored, or completely missed.",
                options={"max_value": 10},
                metadata={"metric_source": "llm_evaluation"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the transcript to generate this chart. (No discovery pillar data available from transcript analysis.)</i>",
                ParagraphStyle('ci_no_disc12', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

    # ════════════════════════════════════════════════════════
    # AGENT PERFORMANCE VISUALIZATION SECTION (PDF 4)
    # ════════════════════════════════════════════════════════

    def _build_agent_performance_viz_section(self, story, sections: Dict[str, str]):
        """Agent performance-focused visualizations — 18 charts:
        Part 2 — Agent Performance (8 charts):
        1. Agent Performance Score (Gauge)
        2. Client Engagement Level (Gauge)
        3. Response Length Distribution (HBar histogram)
        4. Response Delay Analysis (HBar)
        5. Conversation Outcome Indicators (HBar)
        6. Trust Indicators (Radar)
        7. Sentiment Distribution (Donut)
        8. Conversation Efficiency (Gauge)
        Part 3 — Advanced Analytics (5 charts):
        9. Trust vs Pressure Quadrant (Scatter)
        10. Conversation Topic Heatmap (Heatmap)
        11. Dialogue Flow Diagram (Table-based flow)
        12. Sales Funnel Drop-Off (Funnel)
        13. Manager-Facing Summary Dashboard (ScoreCard)
        Part 4 — Redistributed from Archive (5 charts):
        14. Pressure & Risk Indicator (HBar)
        15. Agent Skill Profile Radar (Radar)
        16. EQ Radar (Radar)
        17. Tone & Engagement Scores (HBar)
        18. Negotiation Proficiency Scores (HBar)
        """
        scores = getattr(self, '_cached_scores', {})
        deal_metrics = getattr(self, '_cached_deal_metrics', [])
        talk_ratio = getattr(self, '_cached_talk_ratio', {})
        sentiment_phases = getattr(self, '_cached_sentiment_phases', [])
        ethics_subs = getattr(self, '_cached_ethics_subsections', {})
        listening_subs = getattr(self, '_cached_listening_subsections', {})
        eq_dims = getattr(self, '_cached_eq_dimensions', {})
        tone_scores = getattr(self, '_cached_tone_scores', {})
        neg_scores = getattr(self, '_cached_negotiation_scores', {})
        transcript_data_ref = getattr(self, '_transcript_data', None)

        # Pre-compute transcript-level data
        agent_segments = []
        client_segments = []
        all_segments = []
        if transcript_data_ref and transcript_data_ref.transcripts:
            for seg in transcript_data_ref.transcripts:
                all_segments.append(seg)
                if seg.speaker_name and seg.speaker_name != 'Client':
                    agent_segments.append(seg)
                else:
                    client_segments.append(seg)

        agent_pct = talk_ratio.get('agent', 0)
        client_pct = talk_ratio.get('client', 0)

        story.append(SectionDivider(width=460, color="#0D47A1", thickness=2))
        story.append(Spacer(1, 2))
        story.append(Paragraph("AGENT PERFORMANCE — VISUAL ANALYTICS", self.ps['sec']))
        story.append(Spacer(1, 4))
        story.append(Paragraph(
            "These visualizations focus on quantitative agent performance metrics: scoring dashboards, "
            "engagement gauges, response patterns, trust indicators, and executive summaries. "
            "All data is derived from the actual transcript analysis.",
            self.ps['body']))
        story.append(Spacer(1, 10))

        # ════════════════════════════════════════════════════════
        # CHART 1 — Agent Performance Score (Gauge)
        # ════════════════════════════════════════════════════════
        story.append(Paragraph("<b>1. Agent Performance Score</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        if scores:
            overall = sum(scores.values()) / len(scores)
            overall_pct = int(min(100, max(0, overall * 10)))
            gauge = GaugeFlowable(overall_pct, max_val=100, label="Overall Performance", width=280, height=170)
            perf_zone, zone_color = self._get_zone_from_gauge_value(overall_pct)
            legend_content = [
                Paragraph(f"<b><font color='{zone_color}' size='16'>{overall:.1f}/10</font></b>",
                         ParagraphStyle('ap_g1_val', parent=self.ps['body'], alignment=TA_CENTER)),
                Spacer(1, 4),
                Paragraph(f"<b><font color='{zone_color}'>{perf_zone}</font></b>",
                         ParagraphStyle('ap_g1_zone', parent=self.ps['body_bold'], fontSize=11, alignment=TA_CENTER)),
                Spacer(1, 8),
                Paragraph("<font color='#E53935'>■</font> Needs Improvement: 0-49%", self.ps['body']),
                Paragraph("<font color='#FF8F00'>■</font> Adequate: 50-74%", self.ps['body']),
                Paragraph("<font color='#43A047'>■</font> Excellent: 75-100%", self.ps['body']),
            ]
            g1_table = Table([[gauge, legend_content]], colWidths=[4.2 * inch, 2.7 * inch])
            g1_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (0, 0), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 12),
                ('BOX', (0, 0), (-1, -1), 0.5, self.C['border']),
            ]))
            story.append(KeepTogether(g1_table))
            story.append(Spacer(1, 14))

            note = "(Calculated from LLM-evaluated performance scores.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note1', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="agent_performance_score_gauge",
                title="Agent Performance Score",
                chart_type="gauge",
                labels=["Needs Improvement", "Adequate", "Excellent"],
                datasets=[{"label": "Overall Performance", "values": [overall_pct]}],
                description="Overall agent performance as a percentage gauge.",
                metadata={"raw_score": round(overall, 1), "zone": perf_zone, "zone_color": zone_color,
                          "thresholds": {"needs_improvement": [0, 49], "adequate": [50, 74], "excellent": [75, 100]},
                          "metric_source": "llm_evaluation"},
                options={"max_value": 100},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the LLM analysis to generate this chart. (No performance scores available.)</i>",
                ParagraphStyle('ap_no_scores1', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 2 — Client Engagement Level (Gauge)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(4.2 * inch))
        story.append(Paragraph("<b>2. Client Engagement Level</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        engagement_val = None
        if scores:
            eng_keys = [v for k, v in scores.items()
                        if any(w in k.lower() for w in ['engagement', 'rapport', 'client'])]
            if eng_keys:
                engagement_val = int(min(100, max(0, (sum(eng_keys) / len(eng_keys)) * 10)))
        # Also check tone scores
        if engagement_val is None and tone_scores:
            tone_vals = [v for k, v in tone_scores.items() if k != 'OVERALL' and isinstance(v, (int, float))]
            if tone_vals:
                engagement_val = int(min(100, max(0, (sum(tone_vals) / len(tone_vals)) * 10)))

        if engagement_val is not None:
            gauge2 = GaugeFlowable(engagement_val, max_val=100, label="Client Engagement", width=280, height=170)
            eng_zone, eng_color = self._get_zone_from_gauge_value(engagement_val)
            legend2 = [
                Paragraph(f"<b><font color='{eng_color}' size='16'>{engagement_val}%</font></b>",
                         ParagraphStyle('ap_g2_val', parent=self.ps['body'], alignment=TA_CENTER)),
                Spacer(1, 4),
                Paragraph(f"<b><font color='{eng_color}'>{eng_zone} ENGAGEMENT</font></b>",
                         ParagraphStyle('ap_g2_zone', parent=self.ps['body_bold'], fontSize=11, alignment=TA_CENTER)),
                Spacer(1, 8),
                Paragraph("<font color='#E53935'>■</font> Low: 0-33%", self.ps['body']),
                Paragraph("<font color='#FF8F00'>■</font> Moderate: 34-66%", self.ps['body']),
                Paragraph("<font color='#43A047'>■</font> High: 67-100%", self.ps['body']),
            ]
            g2_table = Table([[gauge2, legend2]], colWidths=[4.2 * inch, 2.7 * inch])
            g2_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (0, 0), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 12),
                ('BOX', (0, 0), (-1, -1), 0.5, self.C['border']),
            ]))
            story.append(KeepTogether(g2_table))
            story.append(Spacer(1, 14))

            note = "(Calculated from performance scores and meeting duration.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note2', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="client_engagement_level_gauge",
                title="Client Engagement Level",
                chart_type="gauge",
                labels=["Low", "Moderate", "High"],
                datasets=[{"label": "Client Engagement", "values": [engagement_val]}],
                description="How engaged the client was during the conversation.",
                metadata={"zone": eng_zone, "zone_color": eng_color, "metric_source": "llm_evaluation"},
                options={"max_value": 100},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the LLM analysis to generate this chart. (Client engagement level could not be determined.)</i>",
                ParagraphStyle('ap_no_eng2', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 3 — Response Length Distribution (HBar histogram)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(2.5 * inch))
        story.append(Paragraph("<b>3. Response Length Distribution</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        if agent_segments:
            # Bucket agent responses by word count
            buckets = {'1-20 words': 0, '21-50 words': 0, '51-100 words': 0,
                       '101-200 words': 0, '200+ words': 0}
            for seg in agent_segments:
                wc = len(seg.transcript.split())
                if wc <= 20:
                    buckets['1-20 words'] += 1
                elif wc <= 50:
                    buckets['21-50 words'] += 1
                elif wc <= 100:
                    buckets['51-100 words'] += 1
                elif wc <= 200:
                    buckets['101-200 words'] += 1
                else:
                    buckets['200+ words'] += 1

            hist_data = [(label, count, "#1565C0") for label, count in buckets.items() if count > 0]
            if hist_data:
                story.append(Paragraph(
                    "Distribution of agent response lengths (word count per turn). "
                    "<font color='#43A047'>■</font> Ideal length, <font color='#FF8F00'>■</font> Moderate, <font color='#E53935'>■</font> Needs adjustment.",
                    ParagraphStyle('ap_hist_d3', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
                story.append(Spacer(1, 4))
                hist_chart = HBarChartFlowable(hist_data, width=440, bar_height=20,
                                               max_val=max(c for _, c, _ in hist_data))
                hc_table = Table([[hist_chart]], colWidths=[6.2 * inch])
                hc_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                    ('PADDING', (0, 0), (-1, -1), 8),
                ]))
                story.append(KeepTogether(hc_table))
                story.append(Spacer(1, 12))

                note = "(Derived from transcript analysis – word counts per agent turn.)"
                story.append(Paragraph(f"<i>{note}</i>",
                             ParagraphStyle('ap_note3', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
                story.append(Spacer(1, 2))

                self.performance_viz_logger.log_chart(
                    chart_id="response_length_distribution",
                    title="Response Length Distribution",
                    chart_type="histogram",
                    labels=[d[0] for d in hist_data],
                    datasets=[{"label": "Count", "values": [d[1] for d in hist_data],
                               "colors": [d[2] for d in hist_data]}],
                    description="Distribution of agent response lengths (word count per turn).",
                    metadata={"total_responses": len(agent_segments), "metric_source": "transcript_analysis"},
                )
            else:
                story.append(Paragraph(
                    "<i>Insufficient information from the transcript to generate this chart. (No agent responses found.)</i>",
                    ParagraphStyle('ap_no_hist3', parent=self.ps['body'], fontSize=8,
                                   textColor=self.C['light'], alignment=TA_CENTER)))
                story.append(Spacer(1, 12))
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the transcript to generate this chart. (No agent segments available for response length analysis.)</i>",
                ParagraphStyle('ap_no_hist3b', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 4 — Response Delay Analysis (HBar)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(2.5 * inch))
        story.append(Paragraph("<b>4. Response Delay Analysis</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        delay_data = []
        if all_segments and len(all_segments) >= 2:
            delays = []
            for i in range(1, len(all_segments)):
                prev = all_segments[i - 1]
                curr = all_segments[i]
                if curr.speaker_name and curr.speaker_name != 'Client' and prev.speaker_name == 'Client':
                    delay = (curr.start or 0) - (prev.end or 0)
                    if delay >= 0:
                        delays.append(delay)
            if delays:
                buckets_delay = {'0-2s': 0, '2-5s': 0, '5-10s': 0, '10-20s': 0, '20s+': 0}
                for d in delays:
                    if d <= 2:
                        buckets_delay['0-2s'] += 1
                    elif d <= 5:
                        buckets_delay['2-5s'] += 1
                    elif d <= 10:
                        buckets_delay['5-10s'] += 1
                    elif d <= 20:
                        buckets_delay['10-20s'] += 1
                    else:
                        buckets_delay['20s+'] += 1
                delay_data = [(label, count, "#FF8F00") for label, count in buckets_delay.items() if count > 0]

        if delay_data:
            story.append(Paragraph(
                "Agent response times after client finishes speaking (seconds between turns). "
                "<font color='#43A047'>■</font> Quick, <font color='#FF8F00'>■</font> Moderate, <font color='#E53935'>■</font> Slow.",
                ParagraphStyle('ap_delay_d4', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            delay_chart = HBarChartFlowable(delay_data, width=440, bar_height=20,
                                            max_val=max(c for _, c, _ in delay_data))
            dl_table = Table([[delay_chart]], colWidths=[6.2 * inch])
            dl_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(KeepTogether(dl_table))
            story.append(Spacer(1, 12))

            avg_delay = sum(delays) / len(delays) if delays else 0
            note = "(Derived from transcript analysis – timestamp differences between turns.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note4', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="response_delay_analysis",
                title="Response Delay Analysis",
                chart_type="horizontalBar",
                labels=[d[0] for d in delay_data],
                datasets=[{"label": "Count", "values": [d[1] for d in delay_data],
                           "colors": [d[2] for d in delay_data]}],
                description="Agent response times after client finishes speaking.",
                metadata={"avg_delay_seconds": round(avg_delay, 2), "total_transitions": len(delays), "metric_source": "transcript_analysis"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the transcript to generate this chart. (No timestamp data available for response delay analysis.)</i>",
                ParagraphStyle('ap_no_delay4', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 5 — Conversation Outcome Indicators (HBar)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(2.5 * inch))
        story.append(Paragraph("<b>5. Conversation Outcome Indicators</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        outcome_data = []
        if deal_metrics:
            for dm in deal_metrics:
                label = dm[0]
                value_str = dm[1].strip()
                # Try to parse numeric value
                num_val = None
                if '%' in value_str:
                    try:
                        num_val = int(value_str.replace('%', '').strip())
                    except ValueError:
                        pass
                elif value_str.lower() in ['strong', 'high']:
                    num_val = 8
                elif value_str.lower() in ['moderate', 'medium']:
                    num_val = 5
                elif value_str.lower() in ['weak', 'low']:
                    num_val = 3
                if num_val is not None:
                    col = self._get_percentage_color(num_val) if num_val <= 100 else self._get_score_color(num_val)
                    outcome_data.append((label, num_val, col))

        if outcome_data:
            story.append(Paragraph(
                "Key deal intelligence indicators scored from the LLM analysis. "
                "<font color='#43A047'>■</font> Green = strong, <font color='#FF8F00'>■</font> Yellow = moderate, <font color='#E53935'>■</font> Red = weak.",
                ParagraphStyle('ap_out_d5', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            out_chart = HBarChartFlowable(outcome_data, width=440, bar_height=18, max_val=10)
            ot_table = Table([[out_chart]], colWidths=[6.2 * inch])
            ot_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(KeepTogether(ot_table))
            story.append(Spacer(1, 12))

            note = "(Derived from LLM-evaluated deal intelligence metrics.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note5', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="conversation_outcome_indicators",
                title="Conversation Outcome Indicators",
                chart_type="horizontalBar",
                labels=[d[0] for d in outcome_data],
                datasets=[{"label": "Score", "values": [d[1] for d in outcome_data],
                           "colors": [d[2] for d in outcome_data]}],
                description="Key deal intelligence indicators scored from the LLM analysis.",
                options={"max_value": 10},
                metadata={"metric_source": "llm_evaluation"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the LLM analysis to generate this chart. (No deal intelligence metrics available.)</i>",
                ParagraphStyle('ap_no_out5', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 6 — Trust Indicators (Radar)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(3.0 * inch))
        story.append(Paragraph("<b>6. Trust Indicators</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        trust_scores = {}
        if scores:
            trust_mapping = {
                'Rapport': ['rapport', 'engagement'],
                'Empathy': ['emotional intelligence', 'empathy'],
                'Transparency': ['communication', 'value proposition'],
                'Listening': ['listening', 'listening quality'],
                'Reliability': ['confidence', 'authority'],
                'Follow-Through': ['closing', 'deal progression'],
            }
            for dim, keywords in trust_mapping.items():
                for score_key, score_val in scores.items():
                    if any(kw in score_key.lower() for kw in keywords):
                        trust_scores[dim] = score_val
                        break

        if trust_scores and len(trust_scores) >= 3:
            story.append(Paragraph(
                "Trust-building dimensions derived from performance scores. "
                "<font color='#43A047'>■</font> Green = strong (≥7.5), <font color='#FF8F00'>■</font> Yellow = moderate (5.0-7.4), <font color='#E53935'>■</font> Red = needs improvement (<5.0).",
                ParagraphStyle('ap_trust_d6', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            trust_radar = RadarChartFlowable(trust_scores, width=300, height=300)
            tr_table = Table([[trust_radar]], colWidths=[6.2 * inch])
            tr_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 10),
            ]))
            story.append(KeepTogether(tr_table))
            story.append(Spacer(1, 12))

            note = "(Derived from LLM-evaluated performance scores.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note6', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="trust_indicators_radar",
                title="Trust Indicators",
                chart_type="radar",
                labels=list(trust_scores.keys()),
                datasets=[{"label": "Trust Score", "values": list(trust_scores.values()),
                           "colors": ["#43A047"], "max": 10}],
                description="Trust-building dimensions derived from performance scores.",
                options={"max_value": 10},
                metadata={"metric_source": "llm_evaluation"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the LLM analysis to generate this chart. (No trust indicator data available.)</i>",
                ParagraphStyle('ap_no_trust6', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 7 — Sentiment Distribution (Donut)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(2.5 * inch))
        story.append(Paragraph("<b>7. Sentiment Distribution</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        sent_counts = {'Positive': 0, 'Neutral': 0, 'Negative': 0}
        if client_segments:
            positive_words = ['great', 'excellent', 'perfect', 'love', 'amazing', 'wonderful', 'fantastic',
                              'impressed', 'happy', 'excited', 'agree', 'yes', 'absolutely', 'definitely']
            negative_words = ['no', 'not sure', 'worried', 'concerned', 'expensive', 'doubt', 'hesitant',
                              'confused', 'disappointed', 'frustrated', 'problem', 'issue', 'don\'t think']
            for seg in client_segments:
                text_lower = seg.transcript.lower()
                pos_hits = sum(1 for w in positive_words if w in text_lower)
                neg_hits = sum(1 for w in negative_words if w in text_lower)
                if pos_hits > neg_hits:
                    sent_counts['Positive'] += 1
                elif neg_hits > pos_hits:
                    sent_counts['Negative'] += 1
                else:
                    sent_counts['Neutral'] += 1

        total_sent = sum(sent_counts.values())
        if total_sent > 0:
            sent_donut_data = [
                ("Positive", sent_counts['Positive'], "#43A047"),
                ("Neutral", sent_counts['Neutral'], "#FF8F00"),
                ("Negative", sent_counts['Negative'], "#E53935"),
            ]
            sent_donut_data = [(l, v, c) for l, v, c in sent_donut_data if v > 0]
            story.append(Paragraph(
                "Distribution of client sentiment across conversation turns. "
                "<font color='#43A047'>■</font> Positive, <font color='#FF8F00'>■</font> Neutral, <font color='#E53935'>■</font> Negative.",
                ParagraphStyle('ap_sent_d7', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            sent_donut = DonutChartFlowable(sent_donut_data, width=220, height=220)
            sd_table = Table([[sent_donut]], colWidths=[6.2 * inch])
            sd_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 10),
            ]))
            story.append(KeepTogether(sd_table))
            story.append(Spacer(1, 12))

            note = "(Derived from transcript analysis – keyword-based sentiment classification.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note7', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="sentiment_distribution_donut",
                title="Sentiment Distribution",
                chart_type="donut",
                labels=[d[0] for d in sent_donut_data],
                datasets=[{"label": "Segments", "values": [d[1] for d in sent_donut_data],
                           "colors": [d[2] for d in sent_donut_data]}],
                description="Distribution of client sentiment across conversation turns.",
                metadata={"total_segments": total_sent, "counts": sent_counts, "metric_source": "transcript_analysis"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the transcript to generate this chart. (No client segments available for sentiment analysis.)</i>",
                ParagraphStyle('ap_no_sent7', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 8 — Conversation Efficiency (Gauge)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(4.2 * inch))
        story.append(Paragraph("<b>8. Conversation Efficiency</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        efficiency_val = None
        if all_segments:
            # Guard division by zero
            if all_segments:
                total_duration = max((seg.end or 0) for seg in all_segments) - min((seg.start or 0) for seg in all_segments)
                total_duration_mins = max(total_duration / 60, 1) if total_duration > 0 else 1
                # Efficiency = weighted blend of outcomes achieved vs time spent
                if scores:
                    avg_score = sum(scores.values()) / len(scores) if scores else 0
                    # Normalize: shorter meeting with higher score = more efficient
                    # Ideal meeting: 20-40 min. Penalize if too short (< 15) or too long (> 60)
                    time_factor = 1.0
                    if total_duration_mins < 15:
                        time_factor = 0.7  # Too short, probably incomplete
                    elif total_duration_mins > 60:
                        time_factor = max(0.5, 1.0 - (total_duration_mins - 60) / 120)
                    efficiency_val = int(min(100, max(0, avg_score * 10 * time_factor)))

        if efficiency_val is not None:
            gauge3 = GaugeFlowable(efficiency_val, max_val=100, label="Conversation Efficiency", width=280, height=170)
            eff_zone, eff_color = self._get_zone_from_gauge_value(efficiency_val)
            legend3 = [
                Paragraph(f"<b><font color='{eff_color}' size='16'>{efficiency_val}%</font></b>",
                         ParagraphStyle('ap_g3_val', parent=self.ps['body'], alignment=TA_CENTER)),
                Spacer(1, 4),
                Paragraph(f"<b><font color='{eff_color}'>{eff_zone} EFFICIENCY</font></b>",
                         ParagraphStyle('ap_g3_zone', parent=self.ps['body_bold'], fontSize=11, alignment=TA_CENTER)),
                Spacer(1, 8),
                Paragraph(f"Duration: {total_duration_mins:.0f} min", self.ps['body']),
                Paragraph(f"Segments: {len(all_segments)}", self.ps['body']),
            ]
            g3_table = Table([[gauge3, legend3]], colWidths=[4.2 * inch, 2.7 * inch])
            g3_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (0, 0), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 12),
                ('BOX', (0, 0), (-1, -1), 0.5, self.C['border']),
            ]))
            story.append(KeepTogether(g3_table))
            story.append(Spacer(1, 14))

            note = "(Calculated from performance scores and meeting duration.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note8', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="conversation_efficiency_gauge",
                title="Conversation Efficiency",
                chart_type="gauge",
                labels=["Low", "Moderate", "High"],
                datasets=[{"label": "Efficiency", "values": [efficiency_val]}],
                description="Overall conversation efficiency (score vs time).",
                metadata={"zone": eff_zone, "duration_min": round(total_duration_mins, 1), "metric_source": "llm_evaluation"},
                options={"max_value": 100},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the transcript to generate this chart. (Cannot calculate conversation efficiency.)</i>",
                ParagraphStyle('ap_no_eff8', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # PART 3 — ADVANCED ANALYTICS
        # ════════════════════════════════════════════════════════
        story.append(SectionDivider(width=460, color="#0D47A1", thickness=1.5))
        story.append(Spacer(1, 2))
        story.append(Paragraph("ADVANCED ANALYTICS", self.ps['sec']))
        story.append(Spacer(1, 10))

        # ════════════════════════════════════════════════════════
        # CHART 9 — Trust vs Pressure Quadrant (Scatter)
        # ════════════════════════════════════════════════════════
        story.append(Paragraph("<b>9. Trust vs Pressure Assessment</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        trust_val = None
        pressure_val = None
        if scores:
            trust_keys = [v for k, v in scores.items()
                          if any(w in k.lower() for w in ['rapport', 'empathy', 'trust', 'listening', 'emotional'])]
            if trust_keys:
                trust_val = sum(trust_keys) / len(trust_keys)
            pressure_sub = ethics_subs.get('PRESSURE TACTICS', '').lower()
            if 'high' in pressure_sub:
                pressure_val = 8
            elif 'medium' in pressure_sub or 'moderate' in pressure_sub:
                pressure_val = 5
            elif 'low' in pressure_sub or 'none' in pressure_sub:
                pressure_val = 2

        if trust_val is not None and pressure_val is not None:
            quadrant = QuadrantFlowable(
                [("This Agent", trust_val, pressure_val)],
                width=280, height=280,
                x_label="Trust-Building →",
                y_label="Pressure →")
            quad_table = Table([[quadrant]], colWidths=[6.2 * inch])
            quad_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#FAFAFA")),
                ('PADDING', (0, 0), (-1, -1), 10),
            ]))
            story.append(KeepTogether(quad_table))
            story.append(Spacer(1, 12))

            note = "(Trust derived from performance scores; pressure from ethics assessment.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note9', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="trust_vs_pressure_quadrant",
                title="Trust vs Pressure Quadrant",
                chart_type="scatter",
                labels=["This Agent"],
                datasets=[{"label": "Agent Position",
                           "values": [{"x": round(trust_val, 1), "y": round(pressure_val, 1)}],
                           "colors": ["#1565C0"]}],
                description="Best agents live in HIGH TRUST / LOW PRESSURE (bottom-right).",
                metadata={"quadrants": [
                    {"name": "High Pressure / Low Trust", "x_range": [0, 5], "y_range": [5, 10]},
                    {"name": "High Pressure / High Trust", "x_range": [5, 10], "y_range": [5, 10]},
                    {"name": "Low Pressure / Low Trust", "x_range": [0, 5], "y_range": [0, 5]},
                    {"name": "Low Pressure / High Trust", "x_range": [5, 10], "y_range": [0, 5]},
                ], "ideal_zone": {"x": 7.5, "y": 2.5}, "metric_source": "llm_evaluation"},
                options={"x_label": "Trust-Building", "y_label": "Pressure", "axis_range": [0, 10]},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the LLM analysis to generate this chart. (Need both trust-related scores and ethics data for quadrant plot.)</i>",
                ParagraphStyle('ap_no_quad9', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 10 — Conversation Topic Heatmap (Heatmap)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(3.0 * inch))
        story.append(Paragraph("<b>10. Conversation Topic Heatmap</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        topic_heatmap_data = []
        topic_rows = ['Property', 'Pricing', 'Timeline', 'Concerns', 'Next Steps']
        topic_cols = ['Opening', 'Mid', 'Close']
        if all_segments:
            third = len(all_segments) // 3
            phase_groups = [all_segments[:third], all_segments[third:2*third], all_segments[2*third:]]
            topic_keywords = {
                'Property': ['property', 'house', 'home', 'bedroom', 'bathroom', 'square', 'location', 'neighborhood'],
                'Pricing': ['price', 'cost', 'commission', 'fee', 'market value', 'comparable', 'worth'],
                'Timeline': ['when', 'timeline', 'schedule', 'ready', 'move', 'list', 'sell by'],
                'Concerns': ['worry', 'concern', 'afraid', 'nervous', 'risk', 'what if', 'hesitant'],
                'Next Steps': ['next step', 'follow up', 'schedule', 'contract', 'sign', 'meeting'],
            }
            for row_topic in topic_rows:
                row_data = []
                kws = topic_keywords.get(row_topic, [])
                for phase_segs in phase_groups:
                    phase_text = ' '.join(s.transcript.lower() for s in phase_segs)
                    hits = sum(phase_text.count(kw) for kw in kws)
                    row_data.append(min(10, hits))
                topic_heatmap_data.append(row_data)

        if topic_heatmap_data and any(any(v > 0 for v in row) for row in topic_heatmap_data):
            story.append(Paragraph(
                "Topic intensity across conversation phases (Opening / Mid / Closing). "
                "<font color='#43A047'>■</font> High intensity, <font color='#FF8F00'>■</font> Moderate, <font color='#E53935'>■</font> Low.",
                ParagraphStyle('ap_hm_d10', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            hm = HeatmapFlowable(topic_rows, topic_cols, topic_heatmap_data, width=420, cell_h=24)
            hm_table = Table([[hm]], colWidths=[6.2 * inch])
            hm_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(KeepTogether(hm_table))
            story.append(Spacer(1, 12))

            note = "(Derived from transcript analysis – keyword frequency per phase.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note10', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="conversation_topic_heatmap",
                title="Conversation Topic Heatmap",
                chart_type="heatmap",
                labels=topic_cols,
                datasets=[{"label": row, "values": topic_heatmap_data[i]}
                          for i, row in enumerate(topic_rows)],
                description="Topic intensity across conversation phases.",
                metadata={"rows": topic_rows, "cols": topic_cols, "metric_source": "transcript_analysis"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the transcript to generate this chart. (No topic data available for heatmap.)</i>",
                ParagraphStyle('ap_no_hm10', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 11 — Dialogue Flow Diagram (Table-based)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(2.5 * inch))
        story.append(Paragraph("<b>11. Dialogue Flow Summary</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        if all_segments and len(all_segments) >= 4:
            # Build a simplified flow: show speaker transitions with segment counts
            flow_transitions = {}
            for i in range(1, len(all_segments)):
                prev_speaker = "Agent" if all_segments[i-1].speaker_name and all_segments[i-1].speaker_name != 'Client' else "Client"
                curr_speaker = "Agent" if all_segments[i].speaker_name and all_segments[i].speaker_name != 'Client' else "Client"
                key = f"{prev_speaker} → {curr_speaker}"
                flow_transitions[key] = flow_transitions.get(key, 0) + 1

            flow_rows = [[Paragraph("<b>Transition</b>", self.ps['th']),
                          Paragraph("<b>Count</b>", self.ps['th']),
                          Paragraph("<b>% of Total</b>", self.ps['th'])]]
            total_transitions = sum(flow_transitions.values())
            # Guard division by zero
            if total_transitions > 0:
                for trans, count in sorted(flow_transitions.items(), key=lambda x: x[1], reverse=True):
                    pct = (count / total_transitions) * 100
                    flow_rows.append([
                        Paragraph(trans, self.ps['td']),
                        Paragraph(f"<b>{count}</b>", self.ps['td_c']),
                        Paragraph(f"{pct:.1f}%", self.ps['td_c']),
                    ])
            else:
                flow_rows.append([Paragraph("No transitions", self.ps['td']), Paragraph("", self.ps['td']), Paragraph("", self.ps['td'])])
            flow_t = Table(flow_rows, colWidths=[2.5 * inch, 1.5 * inch, 2.2 * inch])
            flow_t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#0D47A1")),
                ('TEXTCOLOR', (0, 0), (-1, 0), self.C['white']),
                ('GRID', (0, 0), (-1, -1), 0.5, self.C['border']),
                ('PADDING', (0, 0), (-1, -1), 5),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(KeepTogether(flow_t))
            story.append(Spacer(1, 12))

            note = "(Derived from transcript analysis – speaker transitions.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note11', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="dialogue_flow_summary",
                title="Dialogue Flow Summary",
                chart_type="table",
                labels=["Transition", "Count", "Percentage"],
                datasets=[{"label": "Speaker Transitions",
                           "values": [{"transition": k, "count": v, "pct": round((v / total_transitions) * 100, 1)}
                                      for k, v in flow_transitions.items()]}],
                description="Speaker transition patterns showing dialogue flow.",
                metadata={"total_transitions": total_transitions, "total_segments": len(all_segments), "metric_source": "transcript_analysis"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the transcript to generate this chart. (Not enough segments for dialogue flow analysis.)</i>",
                ParagraphStyle('ap_no_flow11', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 12 — Sales Funnel Drop-Off (Funnel)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(3.0 * inch))
        story.append(Paragraph("<b>12. Sales Funnel Analysis</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        funnel_stages = []
        stage_keys = ['Rapport', 'Discovery', 'Solution Fit', 'Objection Handling', 'Closing']
        stage_score_map = {
            'Rapport': ['rapport', 'engagement', 'client engagement'],
            'Discovery': ['discovery', 'needs discovery', 'needs'],
            'Solution Fit': ['value proposition', 'solution', 'communication'],
            'Objection Handling': ['objection', 'handling'],
            'Closing': ['closing', 'deal progression', 'close'],
        }
        stage_colors = ['#43A047', '#1565C0', '#FF8F00', '#7B1FA2', '#E53935']
        for i, stage in enumerate(stage_keys):
            keywords = stage_score_map.get(stage, [stage.lower()])
            for s_label, s_val in scores.items():
                if any(kw in s_label.lower() for kw in keywords):
                    funnel_stages.append((f"{stage} ({s_val:.0f}/10)", s_val, stage_colors[i]))
                    break

        if funnel_stages:
            story.append(Paragraph(
                "Shows where momentum was lost across conversation stages. Each stage scored 0-10.",
                ParagraphStyle('ap_funnel_d12', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            funnel = FunnelFlowable(funnel_stages, width=400, height=250)
            f_table = Table([[funnel]], colWidths=[6.2 * inch])
            f_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 12),
                ('BOX', (0, 0), (-1, -1), 0.5, self.C['border']),
            ]))
            story.append(KeepTogether(f_table))
            story.append(Spacer(1, 12))

            note = "(Derived from LLM-evaluated performance scores mapped to funnel stages.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note12', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="sales_funnel_dropoff",
                title="Sales Funnel Drop-Off",
                chart_type="funnel",
                labels=[s[0] for s in funnel_stages],
                datasets=[{"label": "Stage Score", "values": [s[1] for s in funnel_stages],
                           "colors": [s[2] for s in funnel_stages]}],
                description="Shows where momentum was lost across conversation stages.",
                options={"max_value": 10},
                metadata={"metric_source": "llm_evaluation"},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the LLM analysis to generate this chart. (No matching scores found for funnel stages.)</i>",
                ParagraphStyle('ap_no_funnel12', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 13 — Manager-Facing Summary Dashboard (ScoreCard)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(3.0 * inch))
        story.append(Paragraph("<b>13. Manager-Facing Summary Dashboard</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        if scores:
            overall = sum(scores.values()) / len(scores)
            pct = overall / 10
            sc = self._get_score_color(overall)

            # Deal momentum
            momentum_val = None
            for dm in deal_metrics:
                if 'probability' in dm[0].lower():
                    try:
                        momentum_val = int(dm[1].replace('%', '').strip())
                    except (ValueError, AttributeError):
                        pass
                    break
            if momentum_val is None:
                momentum_val = int(min(100, max(0, overall * 10)))

            momentum_zone, mz_color = self._get_zone_from_gauge_value(momentum_val)

            dash_row1 = Table([
                [Paragraph(f"<b>{overall:.1f}</b>",
                           ParagraphStyle('ap_ds_num2', fontName='Helvetica-Bold', fontSize=28,
                                          leading=34, textColor=colors.HexColor(sc), alignment=TA_CENTER)),
                 Paragraph(f"<b>{momentum_val}%</b>",
                           ParagraphStyle('ap_dm_num2', fontName='Helvetica-Bold', fontSize=28,
                                          leading=34, textColor=colors.HexColor(mz_color), alignment=TA_CENTER))],
                [Paragraph("out of 10",
                           ParagraphStyle('ap_ds_sub2', fontName='Helvetica', fontSize=8,
                                          leading=10, textColor=self.C['light'], alignment=TA_CENTER)),
                 Paragraph(f"{momentum_zone} Zone",
                           ParagraphStyle('ap_dm_sub2', fontName='Helvetica-Bold', fontSize=9,
                                          leading=11, textColor=colors.HexColor(mz_color), alignment=TA_CENTER))],
                [Paragraph("OVERALL PERFORMANCE",
                           ParagraphStyle('ap_ds_cat2', fontName='Helvetica', fontSize=8,
                                          leading=10, textColor=self.C['light'], alignment=TA_CENTER,
                                          spaceBefore=4)),
                 Paragraph("DEAL MOMENTUM",
                           ParagraphStyle('ap_dm_cat2', fontName='Helvetica', fontSize=8,
                                          leading=10, textColor=self.C['light'], alignment=TA_CENTER,
                                          spaceBefore=4))],
            ], colWidths=[3.1 * inch, 3.1 * inch],
               rowHeights=[40, 16, 20])
            dash_row1.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
                ('VALIGN', (0, 1), (-1, -1), 'TOP'),
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor("#F5F7FA")),
                ('BACKGROUND', (1, 0), (1, -1), colors.HexColor("#F5F7FA")),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
                ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 1), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                ('BOX', (0, 0), (0, -1), 1, colors.HexColor(sc)),
                ('BOX', (1, 0), (1, -1), 1, colors.HexColor(mz_color)),
            ]))
            story.append(KeepTogether(dash_row1))
            story.append(Spacer(1, 8))

            # Key metrics
            risk_level = ethics_subs.get('OVERALL ETHICS RISK', '')
            risk_first = risk_level.split('\n')[0].strip() if risk_level else 'N/A'
            if not risk_first:
                risk_first = 'N/A'
            risk_color = self._get_color_for_status(risk_first)
            if agent_pct and client_pct:
                talk_color = "#C62828" if agent_pct > 65 else "#2E7D32" if 40 <= agent_pct <= 55 else "#F57F17"
                talk_display = f"{agent_pct}% / {client_pct}%"
            else:
                talk_color = "#757575"
                talk_display = "N/A"

            # Sentiment trend (if available)
            if sentiment_phases:
                # Simplified: get first and last sentiment from phases if possible
                first_sent = None
                last_sent = None
                for phase_name, phase_content in sentiment_phases:
                    sent_match = re.search(r'(\d+)\s*/\s*10', phase_content)
                    if sent_match:
                        val = int(sent_match.group(1))
                        if first_sent is None:
                            first_sent = val
                        last_sent = val
                if first_sent is not None and last_sent is not None:
                    sent_direction = "Improving" if last_sent > first_sent else "Declining" if last_sent < first_sent else "Stable"
                else:
                    sent_direction = "N/A"
            else:
                sent_direction = "N/A"
            sent_color = "#2E7D32" if sent_direction == "Improving" else "#C62828" if sent_direction == "Declining" else "#F57F17" if sent_direction == "Stable" else "#757575"

            top_weakness = sorted(scores.items(), key=lambda x: x[1])[0][0] if scores else 'N/A'

            dash_row2 = Table([
                [Paragraph("<b>Talk Ratio<br/>(Agent/Client)</b>", self.ps['th']),
                 Paragraph("<b>Ethics Risk</b>", self.ps['th']),
                 Paragraph("<b>Sentiment Trend</b>", self.ps['th']),
                 Paragraph("<b>Top Weakness</b>", self.ps['th'])],
                [Paragraph(f"<font color='{talk_color}'><b>{talk_display}</b></font>", self.ps['td_c']),
                 Paragraph(f"<font color='{risk_color}'><b>{risk_first}</b></font>", self.ps['td_c']),
                 Paragraph(f"<font color='{sent_color}'><b>{sent_direction}</b></font>", self.ps['td_c']),
                 Paragraph(f"<font color='#C62828'><b>{top_weakness}</b></font>", self.ps['td_c'])],
            ], colWidths=[1.55 * inch, 1.55 * inch, 1.55 * inch, 1.55 * inch])
            dash_row2.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), self.C['primary_dark']),
                ('TEXTCOLOR', (0, 0), (-1, 0), self.C['white']),
                ('GRID', (0, 0), (-1, -1), 0.5, self.C['border']),
                ('PADDING', (0, 0), (-1, -1), 6),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(KeepTogether(dash_row2))
            story.append(Spacer(1, 8))

            note = "(All metrics derived from LLM evaluation and transcript analysis.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note13', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="manager_dashboard_summary",
                title="Manager-Facing Summary Dashboard",
                chart_type="scoreCard",
                labels=["Overall Performance", "Deal Momentum", "Talk Ratio", "Ethics Risk", "Sentiment Trend", "Top Weakness"],
                datasets=[{
                    "label": "Dashboard Metrics",
                    "values": [
                        {"metric": "Overall Performance", "value": round(overall, 1), "max": 10, "color": sc},
                        {"metric": "Deal Momentum", "value": momentum_val, "max": 100,
                         "zone": momentum_zone, "color": mz_color},
                        {"metric": "Talk Ratio (Agent/Client)", "value": talk_display, "color": talk_color},
                        {"metric": "Ethics Risk", "value": risk_first, "color": risk_color},
                        {"metric": "Sentiment Trend", "value": sent_direction, "color": sent_color},
                        {"metric": "Top Weakness", "value": top_weakness, "color": "#C62828"},
                    ],
                }],
                description="Executive summary dashboard with key performance indicators.",
                metadata={"metric_source": "mixed"},
            )

            # Deal Intelligence Signal Map
            if deal_metrics:
                story.append(Spacer(1, 8))
                story.append(Paragraph("<b>Deal Intelligence Signal Map</b>", self.ps['subsec']))
                story.append(Spacer(1, 4))
                for metric_item in deal_metrics:
                    label = metric_item[0]
                    value = metric_item[1]
                    just = metric_item[2] if len(metric_item) > 2 else ''
                    val_lower = value.lower().strip()
                    ind_color = self._get_color_for_status(value)
                    row = Table([[
                        Paragraph(label, self.ps['td']),
                        Paragraph(f"<b><font color='{ind_color}'>{value}</font></b>", self.ps['td_c']),
                        Paragraph(f"<i>{just}</i>" if just else "", self.ps['td']),
                    ]], colWidths=[1.8 * inch, 1.0 * inch, 3.4 * inch])
                    row.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                        ('PADDING', (0, 0), (-1, -1), 5),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ]))
                    story.append(row)
                    story.append(Spacer(1, 3))

                note_deal = "(Derived from LLM-evaluated deal intelligence metrics.)"
                story.append(Paragraph(f"<i>{note_deal}</i>",
                             ParagraphStyle('ap_note13b', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
                story.append(Spacer(1, 2))

                self.performance_viz_logger.log_chart(
                    chart_id="deal_intelligence_signal_map",
                    title="Deal Intelligence Signal Map",
                    chart_type="table",
                    labels=["Metric", "Value", "Justification"],
                    datasets=[{
                        "label": "Deal Signals",
                        "values": [
                            {"metric": m[0], "value": m[1],
                             "justification": m[2] if len(m) > 2 else "",
                             "confidence": m[3].strip() if len(m) > 3 else ""}
                            for m in deal_metrics
                        ],
                    }],
                    description="AI-evaluated deal intelligence signals with confidence ratings.",
                    metadata={"metric_source": "llm_evaluation"},
                )

        # ════════════════════════════════════════════════════════
        # CHART 14 — Pressure & Risk Indicator (HBar)
        # (Redistributed from full_report_archive)
        # ════════════════════════════════════════════════════════
        story.append(CondPageBreak(2.5 * inch))
        story.append(Paragraph("<b>14. Pressure & Risk Indicator</b>", self.ps['subsec']))
        story.append(Spacer(1, 6))

        ethics_subs = getattr(self, '_cached_ethics_subsections', {})
        risk_categories = ['PRESSURE TACTICS', 'OVER-PROMISING', 'REGULATORY', 'TRANSPARENCY']
        risk_bar_data = []
        for key in risk_categories:
            sub = ethics_subs.get(key, '')
            if not sub:
                continue
            sub_lower = sub.lower()
            if 'high' in sub_lower or 'flag' in sub_lower or 'violation' in sub_lower:
                val, col = 8, "#E53935"
            elif 'medium' in sub_lower or 'moderate' in sub_lower:
                val, col = 5, "#FF8F00"
            elif 'low' in sub_lower or 'none' in sub_lower or 'compliant' in sub_lower:
                val, col = 2, "#43A047"
            else:
                continue
            risk_bar_data.append((key.title(), val, col))

        if risk_bar_data:
            story.append(Paragraph(
                "Enterprise-critical: flags urgency pressure, over-promising, discount pushing, scarcity language. "
                "<font color='#43A047'>■</font> Low risk, <font color='#FF8F00'>■</font> Moderate, <font color='#E53935'>■</font> High.",
                ParagraphStyle('ap_risk_d14', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            risk_chart = HBarChartFlowable(risk_bar_data, width=420, bar_height=22, max_val=10)
            rc_table = Table([[risk_chart]], colWidths=[6.2 * inch])
            rc_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(KeepTogether(rc_table))
            story.append(Spacer(1, 12))

            note = "(Derived from LLM-evaluated ethics & compliance section.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note14', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="pressure_risk_indicator",
                title="Pressure & Risk Indicator",
                chart_type="horizontalBar",
                labels=[d[0] for d in risk_bar_data],
                datasets=[{"label": "Risk Level", "values": [d[1] for d in risk_bar_data],
                           "colors": [d[2] for d in risk_bar_data]}],
                description="Enterprise-critical: flags urgency pressure, over-promising, discount pushing, scarcity language.",
                metadata={"scale_note": "Higher = more risk detected", "metric_source": "llm_evaluation"},
                options={"max_value": 10},
            )
        else:
            story.append(Paragraph(
                "<i>Insufficient information from the LLM analysis to generate this chart. (No ethics/risk data available.)</i>",
                ParagraphStyle('ap_no_risk14', parent=self.ps['body'], fontSize=8,
                               textColor=self.C['light'], alignment=TA_CENTER)))
            story.append(Spacer(1, 12))

        # ════════════════════════════════════════════════════════
        # CHART 15 — Agent Skill Profile Radar (Radar)
        # (Redistributed from full_report_archive)
        # ════════════════════════════════════════════════════════
        if scores and len(scores) >= 3:
            story.append(CondPageBreak(3.0 * inch))
            story.append(Paragraph("<b>15. Agent Skill Profile Radar</b>", self.ps['subsec']))
            story.append(Spacer(1, 6))
            story.append(Paragraph(
                "Multi-dimensional view — used for long-term growth tracking across all evaluated dimensions. "
                "<font color='#43A047'>■</font> Green = strong (≥7.5), <font color='#FF8F00'>■</font> Yellow = moderate (5.0-7.4), <font color='#E53935'>■</font> Red = needs improvement (<5.0).",
                ParagraphStyle('ap_skill_d15', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            chart = RadarChartFlowable(scores, width=300, height=300)
            chart_table = Table([[chart]], colWidths=[6.2 * inch])
            chart_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('PADDING', (0, 0), (-1, -1), 10),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
            ]))
            story.append(KeepTogether(chart_table))
            story.append(Spacer(1, 12))

            note = "(Derived from LLM-evaluated performance scores.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note15', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="agent_skill_profile_radar",
                title="Agent Skill Profile Radar",
                chart_type="radar",
                labels=list(scores.keys()),
                datasets=[{"label": "Skill Score", "values": list(scores.values()),
                           "colors": ["#1565C0"], "max": 10}],
                description="Multi-dimensional skill profile for long-term growth tracking.",
                options={"max_value": 10},
                metadata={"metric_source": "llm_evaluation"},
            )

        # ════════════════════════════════════════════════════════
        # CHART 16 — EQ Radar (Radar)
        # (Redistributed from full_report_archive)
        # ════════════════════════════════════════════════════════
        eq_dims = getattr(self, '_cached_eq_dimensions', {})
        eq_scores_viz = {}
        for dim_name, dim_content in eq_dims.items():
            sc = re.search(r'(\d+)\s*/\s*(\d+)', str(dim_content))
            if sc:
                eq_scores_viz[dim_name.title()] = int(sc.group(1))

        if eq_scores_viz and len(eq_scores_viz) >= 3:
            story.append(CondPageBreak(3.0 * inch))
            story.append(Paragraph("<b>16. Emotional Intelligence Dimensions</b>", self.ps['subsec']))
            story.append(Spacer(1, 6))
            story.append(Paragraph(
                "Plots each EQ dimension scored by the LLM. Managers can identify emotional skill gaps at a glance.",
                ParagraphStyle('ap_eq_d16', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            eq_radar = RadarChartFlowable(eq_scores_viz, width=280, height=280)
            eq_table = Table([[eq_radar]], colWidths=[6.2 * inch])
            eq_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 10),
            ]))
            story.append(KeepTogether(eq_table))
            story.append(Spacer(1, 12))

            note = "(Derived from LLM-evaluated emotional intelligence section.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note16', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="emotional_intelligence_radar",
                title="EQ Radar — Emotional Intelligence Dimensions",
                chart_type="radar",
                labels=list(eq_scores_viz.keys()),
                datasets=[{"label": "EQ Score", "values": list(eq_scores_viz.values()),
                           "colors": ["#7B1FA2"], "max": 10}],
                description="Plots each EQ dimension to identify emotional skill gaps.",
                options={"max_value": 10},
                metadata={"metric_source": "llm_evaluation"},
            )

        # ════════════════════════════════════════════════════════
        # CHART 17 — Tone & Engagement Scores (HBar)
        # (Redistributed from full_report_archive)
        # ════════════════════════════════════════════════════════
        tone_scores = getattr(self, '_cached_tone_scores', {})
        tone_viz = {k: v for k, v in tone_scores.items() if k != 'OVERALL' and isinstance(v, (int, float))}

        if tone_viz:
            story.append(CondPageBreak(2.5 * inch))
            story.append(Paragraph("<b>17. Tone & Engagement Scores</b>", self.ps['subsec']))
            story.append(Spacer(1, 6))
            story.append(Paragraph(
                "How well the agent managed tone, energy, and client rapport throughout the conversation. "
                "<font color='#43A047'>■</font> Green = strong (≥7.5), <font color='#FF8F00'>■</font> Yellow = moderate (5.0-7.4), <font color='#E53935'>■</font> Red = needs improvement (<5.0).",
                ParagraphStyle('ap_tone_d17', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            tone_bar_data = []
            for label, val in tone_viz.items():
                pct = val / 10
                col = self._get_score_color(val)
                tone_bar_data.append((label.title().replace('_', ' '), val, col))
            tone_chart = HBarChartFlowable(tone_bar_data, width=440, bar_height=22, max_val=10)
            tc_table = Table([[tone_chart]], colWidths=[6.2 * inch])
            tc_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(KeepTogether(tc_table))
            story.append(Spacer(1, 12))

            note = "(Derived from LLM-evaluated tone analysis section.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note17', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="tone_engagement_scores",
                title="Client Engagement & Tone Score Breakdown",
                chart_type="horizontalBar",
                labels=[d[0] for d in tone_bar_data],
                datasets=[{"label": "Tone Score", "values": [d[1] for d in tone_bar_data],
                           "colors": [d[2] for d in tone_bar_data]}],
                description="How well the agent managed tone, energy, and client rapport.",
                options={"max_value": 10},
                metadata={"metric_source": "llm_evaluation"},
            )

        # ════════════════════════════════════════════════════════
        # CHART 18 — Negotiation Proficiency Scores (HBar)
        # (Redistributed from full_report_archive)
        # ════════════════════════════════════════════════════════
        neg_scores = getattr(self, '_cached_negotiation_scores', {})
        neg_viz = {k: v for k, v in neg_scores.items() if k != 'OVERALL' and isinstance(v, (int, float))}

        if neg_viz:
            story.append(CondPageBreak(2.5 * inch))
            story.append(Paragraph("<b>18. Negotiation Proficiency Scores</b>", self.ps['subsec']))
            story.append(Spacer(1, 6))
            story.append(Paragraph(
                "How effectively the agent negotiated, persuaded, and advanced the deal. "
                "<font color='#43A047'>■</font> Green = strong (≥7.5), <font color='#FF8F00'>■</font> Yellow = moderate (5.0-7.4), <font color='#E53935'>■</font> Red = needs improvement (<5.0).",
                ParagraphStyle('ap_neg_d18', parent=self.ps['body'], fontSize=7.5, textColor=self.C['medium'])))
            story.append(Spacer(1, 4))
            neg_bar_data = []
            for label, val in neg_viz.items():
                pct = val / 10
                col = self._get_score_color(val)
                neg_bar_data.append((label.title().replace('_', ' '), val, col))
            neg_chart = HBarChartFlowable(neg_bar_data, width=440, bar_height=22, max_val=10)
            nc_table = Table([[neg_chart]], colWidths=[6.2 * inch])
            nc_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
                ('PADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(KeepTogether(nc_table))
            story.append(Spacer(1, 12))

            note = "(Derived from LLM-evaluated negotiation proficiency section.)"
            story.append(Paragraph(f"<i>{note}</i>",
                         ParagraphStyle('ap_note18', parent=self.ps['body'], fontSize=7, textColor=self.C['light'])))
            story.append(Spacer(1, 2))

            self.performance_viz_logger.log_chart(
                chart_id="negotiation_proficiency_scores",
                title="Negotiation Proficiency Score Breakdown",
                chart_type="horizontalBar",
                labels=[d[0] for d in neg_bar_data],
                datasets=[{"label": "Negotiation Score", "values": [d[1] for d in neg_bar_data],
                           "colors": [d[2] for d in neg_bar_data]}],
                description="How effectively the agent negotiated, persuaded, and advanced the deal.",
                options={"max_value": 10},
                metadata={"metric_source": "llm_evaluation"},
            )

    # ════════════════════════════════════════════════════════
    # NEW HELPER METHODS FOR AGGREGATE JSON
    # ════════════════════════════════════════════════════════

    def _compute_talk_ratio(self, transcript_data: TranscriptData) -> Dict[str, Optional[float]]:
        """Return agent/client talk ratio based on word count."""
        agent_words = client_words = 0
        for seg in transcript_data.transcripts:
            if seg.speaker_name and seg.speaker_name != 'Client':
                agent_words += len(seg.transcript.split())
            else:
                client_words += len(seg.transcript.split())
        total = agent_words + client_words
        if total == 0:
            return {"agent": None, "client": None}
        return {
            "agent": round(agent_words / total * 100, 1),
            "client": round(client_words / total * 100, 1)
        }

    def _compute_question_type_counts(self, agent_segments: List[TranscriptSegment]) -> Dict[str, int]:
        """Classify agent questions into types."""
        counts = {"open_ended": 0, "closed": 0, "leading": 0, "assumptive": 0}
        question_texts = []
        for seg in agent_segments:
            sentences = re.split(r'[.!?]+', seg.transcript)
            for s in sentences:
                s = s.strip()
                if '?' in s and any(w in s.lower() for w in ['what','how','why','when','where','who','which','do','does','did','would','could','can','is','are','was','were','have','has']):
                    question_texts.append(s)
        for q in question_texts:
            ql = q.lower()
            if any(w in ql for w in ['what','how','why','tell me','describe','explain']):
                counts["open_ended"] += 1
            elif any(w in ql for w in ['is ','are ','do ','does ','did ','can ','was ','were ','have you','has ']):
                counts["closed"] += 1
            elif any(w in ql for w in ["don't you","wouldn't you","isn't it","right?","agree?","correct?"]):
                counts["leading"] += 1
            elif any(w in ql for w in ["when you","once we","after we","so you're"]):
                counts["assumptive"] += 1
        return counts

    def _compute_objection_type_counts(self, client_segments: List[TranscriptSegment]) -> Dict[str, int]:
        """Count objection types from client dialogue."""
        counts = {}
        objection_patterns = {
            'Price Concern': ['too expensive','cost','price is','budget','afford','money','commission'],
            'Timing Concern': ['not ready','timing','too soon','wait','later','time'],
            'Competition': ['other agent','another company','competitor','elsewhere','friend recommended'],
            'Trust / Skepticism': ['not sure','doubt','skeptical','worried','concerned','hesitant'],
            'Need Clarification': ["don't understand",'confused','explain','clarify','what do you mean'],
            'Stall Tactic': ['think about it','let me talk to','discuss with','get back to you']
        }
        client_text = ' '.join(s.transcript.lower() for s in client_segments)
        for obj_type, patterns in objection_patterns.items():
            cnt = sum(client_text.count(p) for p in patterns)
            if cnt > 0:
                counts[obj_type] = cnt
        return counts

    def _compute_objection_handling_counts(self, agent_segments: List[TranscriptSegment], client_segments: List[TranscriptSegment]) -> Dict[str, Dict[str, int]]:
        """For each objection type, count addressed vs unaddressed."""
        handling = {}
        objection_patterns = {
            'Price Concern': ['too expensive','cost','price is','budget','afford','money','commission'],
            'Timing Concern': ['not ready','timing','too soon','wait','later','time'],
            'Competition': ['other agent','another company','competitor','elsewhere','friend recommended'],
            'Trust / Skepticism': ['not sure','doubt','skeptical','worried','concerned','hesitant'],
            'Need Clarification': ["don't understand",'confused','explain','clarify','what do you mean'],
            'Stall Tactic': ['think about it','let me talk to','discuss with','get back to you']
        }
        client_text = ' '.join(s.transcript.lower() for s in client_segments)
        for obj_type, patterns in objection_patterns.items():
            total_mentions = sum(client_text.count(p) for p in patterns)
            if total_mentions > 0:
                # approximate addressed by counting agent segments containing resolution phrases
                resolution_phrases = ['i understand','let me address',"here's why",'what if we','consider this','i can offer','the reason',"that's a great question",'let me explain','the benefit']
                addressed = sum(1 for seg in agent_segments if any(r in seg.transcript.lower() for r in resolution_phrases))
                unaddressed = total_mentions - addressed
                handling[obj_type] = {"addressed": addressed, "unaddressed": max(0, unaddressed)}
        return handling

    def _compute_language_phrase_counts(self, agent_segments: List[TranscriptSegment]) -> Dict[str, int]:
        """Count confident, hedging, filler phrases."""
        agent_text = ' '.join(s.transcript.lower() for s in agent_segments)
        confident = ['i will','we will','absolutely','definitely','i guarantee','without a doubt','i assure','i promise','my goal is','what i want',"here's what",'the fact is','i believe','i recommend','my recommendation',"i'm going to","we're going to",'let me walk you through','let me explain']
        hedging = ['i think','maybe','perhaps','probably','might','could be','i guess','sort of','kind of','it depends','not sure',"i don't know",'possibly','hopefully',"we'll see"]
        filler = ['you know','i mean','like','um','uh','basically','obviously','literally','actually','right right right']
        return {
            "confident": sum(agent_text.count(p) for p in confident),
            "hedging": sum(agent_text.count(p) for p in hedging),
            "filler": sum(agent_text.count(p) for p in filler)
        }

    def _compute_response_length_distribution(self, agent_segments: List[TranscriptSegment]) -> Dict[str, int]:
        """Bucket agent response lengths."""
        buckets = {'1-20 words': 0, '21-50 words': 0, '51-100 words': 0, '101-200 words': 0, '200+ words': 0}
        for seg in agent_segments:
            wc = len(seg.transcript.split())
            if wc <= 20:
                buckets['1-20 words'] += 1
            elif wc <= 50:
                buckets['21-50 words'] += 1
            elif wc <= 100:
                buckets['51-100 words'] += 1
            elif wc <= 200:
                buckets['101-200 words'] += 1
            else:
                buckets['200+ words'] += 1
        return buckets

    def _compute_response_delay_distribution(self, all_segments: List[TranscriptSegment]) -> Dict[str, int]:
        """Bucket response delays (agent after client)."""
        buckets = {'0-2s':0, '2-5s':0, '5-10s':0, '10-20s':0, '20s+':0}
        for i in range(1, len(all_segments)):
            prev, curr = all_segments[i-1], all_segments[i]
            if curr.speaker_name and curr.speaker_name != 'Client' and prev.speaker_name == 'Client':
                delay = (curr.start or 0) - (prev.end or 0)
                if delay >= 0:
                    if delay <= 2:
                        buckets['0-2s'] += 1
                    elif delay <= 5:
                        buckets['2-5s'] += 1
                    elif delay <= 10:
                        buckets['5-10s'] += 1
                    elif delay <= 20:
                        buckets['10-20s'] += 1
                    else:
                        buckets['20s+'] += 1
        return buckets

    def _compute_sentiment_progression(self, sentiment_phases: List[Tuple[str, str]]) -> List[Dict[str, Any]]:
        """Convert sentiment phases to list of dicts with phase and score."""
        result = []
        for phase, content in sentiment_phases:
            score = None
            # extract numeric score
            m = re.search(r'(\d+)\s*/\s*10', content)
            if m:
                score = int(m.group(1))
            else:
                # fallback keyword mapping
                cl = content.lower()
                if any(w in cl for w in ['skeptic','hesitant','anxious','distrust']):
                    score = 3
                elif any(w in cl for w in ['neutral','cautious','guarded']):
                    score = 5
                elif any(w in cl for w in ['engaged','interested','warming']):
                    score = 7
                elif any(w in cl for w in ['confident','trust','committed','positive']):
                    score = 8
            if score is not None:
                result.append({"phase": phase, "score": score})
        return result

    def _compute_trust_vs_pressure(self, scores: Dict[str, float], ethics_subs: Dict[str, str]) -> Optional[Dict[str, float]]:
        """Return trust and pressure values if available."""
        trust_keys = [v for k,v in scores.items() if any(w in k.lower() for w in ['rapport','empathy','trust','listening','emotional'])]
        trust = sum(trust_keys)/len(trust_keys) if trust_keys else None
        pressure_sub = ethics_subs.get('PRESSURE TACTICS', '').lower()
        pressure = None
        if 'high' in pressure_sub:
            pressure = 8
        elif 'medium' in pressure_sub or 'moderate' in pressure_sub:
            pressure = 5
        elif 'low' in pressure_sub or 'none' in pressure_sub:
            pressure = 2
        if trust is not None and pressure is not None:
            return {"trust": round(trust,1), "pressure": pressure}
        return None

    # ════════════════════════════════════════════════════════
    # NEW METHOD: AGENT PERFORMANCE AGGREGATE JSON (JSON 7)
    # ════════════════════════════════════════════════════════

    def _generate_agent_performance_aggregate_json(self, sections: Dict[str, str], transcript_data: TranscriptData):
        """Create a comprehensive JSON with all numerical values and calculations."""
        agent_name = getattr(self, '_agent_name', 'Agent')
        scores = getattr(self, '_cached_scores', {})
        deal_metrics = getattr(self, '_cached_deal_metrics', [])
        talk_ratio_cached = getattr(self, '_cached_talk_ratio', {})
        sentiment_phases = getattr(self, '_cached_sentiment_phases', [])
        discovery_pillars = getattr(self, '_cached_discovery_pillars', [])
        ethics_subs = getattr(self, '_cached_ethics_subsections', {})
        eq_dims = getattr(self, '_cached_eq_dimensions', {})
        tone_scores = getattr(self, '_cached_tone_scores', {})
        neg_scores = getattr(self, '_cached_negotiation_scores', {})

        # Split agent/client segments
        agent_segments = []
        client_segments = []
        all_segments = transcript_data.transcripts
        for seg in all_segments:
            if seg.speaker_name and seg.speaker_name != 'Client':
                agent_segments.append(seg)
            else:
                client_segments.append(seg)

        # Compute overall deal probability
        overall_deal_prob = None
        for dm in deal_metrics:
            if 'probability' in dm[0].lower():
                try:
                    overall_deal_prob = float(dm[1].replace('%', '').strip())
                except:
                    pass
                break

        # Prepare performance dimensions list
        perf_dims = []
        for label, val in scores.items():
            perf_dims.append({
                "label": label,
                "score": val,
                "max": 10,
                "justification": "",
                "confidence": ""
            })

        # Discovery pillars list
        disc_list = []
        for pillar, status, evidence, confidence in discovery_pillars:
            disc_list.append({
                "pillar": pillar,
                "status": status,
                "evidence": evidence,
                "confidence": confidence
            })

        # Emotional intelligence
        eq_dict = {}
        for dim, content in eq_dims.items():
            sc = re.search(r'(\d+)\s*/\s*(\d+)', str(content))
            if sc:
                eq_dict[dim.lower()] = int(sc.group(1))
        overall_eq = getattr(self, '_cached_eq_overall', '')
        if overall_eq:
            m = re.search(r'(\d+)\s*/\s*(\d+)', overall_eq)
            if m:
                eq_dict['overall'] = int(m.group(1))

        # Tone scores (excluding overall)
        tone_dict = {k.lower(): v for k, v in tone_scores.items() if k != 'OVERALL' and isinstance(v, (int, float))}

        # Negotiation scores
        neg_dict = {k.lower(): v for k, v in neg_scores.items() if k != 'OVERALL' and isinstance(v, (int, float))}

        # Ethics risk
        ethics_risk = ethics_subs.get('OVERALL ETHICS RISK', '')
        ethics_risk = ethics_risk.split('\n')[0].strip() if ethics_risk else None

        # Coaching priority areas (lowest 5)
        sorted_scores = sorted(scores.items(), key=lambda x: x[1])
        priority_areas = [{"skill": label, "score": val} for label, val in sorted_scores[:5]]

        # Talk ratio (prefer cached, else compute)
        talk_ratio = talk_ratio_cached
        if talk_ratio.get('agent') is None:
            talk_ratio = self._compute_talk_ratio(transcript_data)

        # Question types
        question_counts = self._compute_question_type_counts(agent_segments)

        # Objection types
        objection_counts = self._compute_objection_type_counts(client_segments)

        # Objection handling
        handling_counts = self._compute_objection_handling_counts(agent_segments, client_segments)

        # Language phrases
        phrase_counts = self._compute_language_phrase_counts(agent_segments)

        # Response length distribution
        length_dist = self._compute_response_length_distribution(agent_segments)

        # Response delay distribution
        delay_dist = self._compute_response_delay_distribution(all_segments)

        # Sentiment progression
        sentiment_prog = self._compute_sentiment_progression(sentiment_phases)

        # Trust vs pressure
        trust_pressure = self._compute_trust_vs_pressure(scores, ethics_subs)

        # Conversation metrics
        total_duration = transcript_data.summary_info.get('total_duration_seconds', 0)
        conv_metrics = {
            "total_segments": len(all_segments),
            "total_duration_seconds": total_duration,
            "total_duration_minutes": round(total_duration / 60, 1),
            "agent_segments": len(agent_segments),
            "client_segments": len(client_segments)
        }

        # Overall performance score
        overall = round(sum(scores.values()) / len(scores), 2) if scores else 0.0

        # Build the aggregate object
        aggregate = AgentPerformanceAggregate(
            generated_at=datetime.now().isoformat(),
            agent_name=agent_name,
            transcript_file=self.transcript_file,
            meeting_date=datetime.now().strftime('%Y-%m-%d'),
            overall_performance_score=overall,
            overall_deal_probability=overall_deal_prob,
            talk_ratio=talk_ratio,
            performance_dimensions=perf_dims,
            discovery_pillars=disc_list,
            emotional_intelligence=eq_dict,
            tone_scores=tone_dict,
            negotiation_scores=neg_dict,
            ethics_risk=ethics_risk,
            coaching_priority_areas=priority_areas,
            conversation_metrics=conv_metrics,
            question_type_counts=question_counts,
            objection_type_counts=objection_counts,
            objection_handling_counts=handling_counts,
            language_phrase_counts=phrase_counts,
            response_length_distribution=length_dist,
            response_delay_distribution=delay_dist,
            sentiment_progression=sentiment_prog,
            trust_vs_pressure=trust_pressure,
            deal_outcome=getattr(self, '_cached_deal_outcome', DealOutcomeResult()).model_dump() if hasattr(self, '_cached_deal_outcome') and self._cached_deal_outcome else None
        )

        # Save to JSON
        filepath = os.path.join(self.summaries_folder, "agent_performance_aggregate.json")
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(aggregate.model_dump(), f, indent=2, ensure_ascii=False, default=str)
            logger.info(f"Saved agent performance aggregate JSON: {filepath}")
            print(f"    Saved agent performance aggregate JSON: {filepath}")
        except Exception as e:
            logger.error(f"Failed to save agent_performance_aggregate JSON: {e}")
            print(f"    Failed to save agent_performance_aggregate JSON: {e}")

    # ════════════════════════════════════════════════════════
    # JSON 8: DEAL OUTCOME / CLIENT DECISION
    # ════════════════════════════════════════════════════════

    def _generate_deal_outcome_json(self, sections: Dict[str, str]):
        """Generate deal_outcome.json — final deal status extracted from LLM Section 19."""
        deal_outcome = getattr(self, '_cached_deal_outcome', None)
        if not deal_outcome:
            # Parse from sections if not cached (fallback)
            content = self._find_section_content(sections, 19, ['Deal Outcome', 'Client Decision'])
            deal_outcome = self._parse_deal_outcome(content)
            deal_outcome.agent_name = getattr(self, '_agent_name', 'Agent')

        deal_outcome.generated_at = datetime.now().isoformat()
        deal_outcome.metadata = {
            "transcript_file": self.transcript_file,
            "generated_date": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        }

        filepath = os.path.join(self.summaries_folder, "deal_outcome.json")
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(deal_outcome.model_dump(), f, indent=2, ensure_ascii=False, default=str)
            logger.info(f"Saved deal outcome JSON: {filepath}")
        except Exception as e:
            logger.error(f"Failed to save deal_outcome JSON: {e}")

    # ════════════════════════════════════════════════════════
    # JSON 4: ALL VISUALIZATIONS (Master Combined JSON)
    # ════════════════════════════════════════════════════════

    def _generate_all_visualizations_json(self, sections: Dict[str, str]):
        """Generate all_visualizations.json — master JSON combining coaching + performance viz data
        with enhanced Pydantic-compatible schema."""
        all_viz_entries = []

        # Merge coaching insights visualizations
        for viz in self.coaching_viz_logger.get_all():
            entry = VisualizationEntry(
                visualization_id=viz.get("id", ""),
                visualization_name=viz.get("title", ""),
                category="coaching_insights",
                related_pdf="coaching_insights_viz.pdf",
                chart_type=viz.get("type", ""),
                description=viz.get("description", ""),
                labels=viz.get("labels", []),
                datasets=viz.get("datasets", []),
                metric_source=viz.get("metadata", {}).get("metric_source", "transcript_analysis"),
                explanation=viz.get("description", ""),
                visualization_priority=self._get_viz_priority(viz.get("id", "")),
                display_style=viz.get("options", {}),
                metadata=viz.get("metadata", {}),
                timestamp=viz.get("timestamp", ""),
            )
            all_viz_entries.append(entry)

        # Merge agent performance visualizations
        for viz in self.performance_viz_logger.get_all():
            entry = VisualizationEntry(
                visualization_id=viz.get("id", ""),
                visualization_name=viz.get("title", ""),
                category="agent_performance",
                related_pdf="agent_performance_viz.pdf",
                chart_type=viz.get("type", ""),
                description=viz.get("description", ""),
                labels=viz.get("labels", []),
                datasets=viz.get("datasets", []),
                metric_source=viz.get("metadata", {}).get("metric_source", "transcript_analysis"),
                explanation=viz.get("description", ""),
                visualization_priority=self._get_viz_priority(viz.get("id", "")),
                display_style=viz.get("options", {}),
                metadata=viz.get("metadata", {}),
                timestamp=viz.get("timestamp", ""),
            )
            all_viz_entries.append(entry)

        # Build master output using Pydantic model
        output = AllVisualizationsOutput(
            generated_at=datetime.now().isoformat(),
            report_type="all_visualizations",
            schema_version="2.0",
            metadata={
                "agent": getattr(self, '_agent_name', 'Agent'),
                "transcript_file": self.transcript_file,
                "generated_date": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            },
            total_visualizations=len(all_viz_entries),
            coaching_insights_count=sum(1 for e in all_viz_entries if e.category == "coaching_insights"),
            agent_performance_count=sum(1 for e in all_viz_entries if e.category == "agent_performance"),
            visualizations=all_viz_entries,
        )

        filepath = os.path.join(self.summaries_folder, "all_visualizations.json")
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(output.model_dump(), f, indent=2, ensure_ascii=False, default=str)
            logger.info(f"Saved master visualization JSON: {filepath}")
        except Exception as e:
            logger.error(f"Failed to save all_visualizations JSON: {e}")

    def _get_viz_priority(self, chart_id: str) -> int:
        """Return a priority level (1=highest, 10=lowest) for a visualization."""
        high_priority = ['agent_performance_score_gauge', 'coaching_priority_areas',
                         'behavioral_skill_radar', 'trust_vs_pressure_quadrant',
                         'sales_funnel_dropoff', 'manager_dashboard_summary']
        medium_priority = ['conversation_stage_distribution', 'agent_vs_client_talking_ratio',
                           'active_listening_indicators', 'sentiment_progression',
                           'objection_type_breakdown', 'confidence_vs_uncertainty_language',
                           'agent_skill_profile_radar', 'emotional_intelligence_radar']
        if chart_id in high_priority:
            return 1
        elif chart_id in medium_priority:
            return 3
        return 5

    # ════════════════════════════════════════════════════════
    # JSON 5: COACHING RECOMMENDATIONS
    # ════════════════════════════════════════════════════════

    def _generate_coaching_recommendations_json(self, sections: Dict[str, str]):
        """Generate coaching_recommendations.json — structured extraction of
        strengths (§4), failures (§5), tactical playbook (§6), 7-day plan (§7)."""

        agent_name = getattr(self, '_agent_name', 'Agent')

        # Parse Strengths (Section 4)
        strengths_content = self._find_section_content(sections, 4, ['Strength'])
        strengths_list = self._parse_strengths_for_json(strengths_content)

        # Parse Failures (Section 5)
        failures_content = self._find_section_content(sections, 5, ['Failure', 'Weakness'])
        failures_list = self._parse_failures_for_json(failures_content)

        # Parse Tactical Playbook (Section 6)
        playbook_content = self._find_section_content(sections, 6, ['Playbook', 'Tactical', 'Coaching'])
        playbook_list = self._parse_playbook_for_json(playbook_content)

        # Parse 7-Day Plan (Section 7)
        plan_content = self._find_section_content(sections, 7, ['Day', 'Plan', 'Development'])
        plan_list = self._parse_7day_plan_for_json(plan_content)

        scores = getattr(self, '_cached_scores', {})
        overall = round(sum(scores.values()) / len(scores), 1) if scores else 0

        output = CoachingRecommendationsOutput(
            generated_at=datetime.now().isoformat(),
            report_type="coaching_recommendations",
            schema_version="2.0",
            agent_name=agent_name,
            metadata={
                "transcript_file": self.transcript_file,
                "total_sections_parsed": len(sections),
                "overall_score": overall,
            },
            strengths=strengths_list,
            failures=failures_list,
            tactical_playbook=playbook_list,
            seven_day_plan=plan_list,
            summary={
                "total_strengths": len(strengths_list),
                "total_failures": len(failures_list),
                "total_playbook_entries": len(playbook_list),
                "total_plan_days": len(plan_list),
                "overall_performance_score": overall,
            },
        )

        filepath = os.path.join(self.summaries_folder, "coaching_recommendations.json")
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(output.model_dump(), f, indent=2, ensure_ascii=False, default=str)
            logger.info(f"Saved coaching recommendations JSON: {filepath}")
        except Exception as e:
            logger.error(f"Failed to save coaching_recommendations JSON: {e}")

    def _parse_strengths_for_json(self, content: str) -> List['StrengthEntry']:
        """Parse Section 4 strengths into StrengthEntry list."""
        entries = []
        if not content:
            return entries
        # Parse STRENGTH / TIMESTAMP / EVIDENCE / IMPACT blocks
        blocks = re.split(r'(?:^|\n)(?:STRENGTH|###?\s*STRENGTH)\s*(?:\d+)?[:\s]*', content, flags=re.IGNORECASE)
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            strength_title = ""
            timestamp = ""
            evidence = ""
            impact = ""
            lines = block.split('\n')
            for line in lines:
                line_stripped = line.strip()
                if not line_stripped:
                    continue
                if line_stripped.upper().startswith('TIMESTAMP'):
                    timestamp = re.sub(r'^TIMESTAMP\s*[:\-]\s*', '', line_stripped, flags=re.IGNORECASE).strip()
                elif line_stripped.upper().startswith('EVIDENCE'):
                    evidence = re.sub(r'^EVIDENCE\s*[:\-]\s*', '', line_stripped, flags=re.IGNORECASE).strip()
                elif line_stripped.upper().startswith('IMPACT'):
                    impact = re.sub(r'^IMPACT\s*[:\-]\s*', '', line_stripped, flags=re.IGNORECASE).strip()
                elif not strength_title:
                    strength_title = line_stripped.strip('- *')
            if strength_title:
                entries.append(StrengthEntry(
                    strength=strength_title, timestamp=timestamp,
                    evidence=evidence, impact=impact))
        # Fallback: treat numbered/bulleted items as strengths
        if not entries:
            for line in content.split('\n'):
                line = line.strip()
                m = re.match(r'^[\-\*\d\.]+\s*(.+)', line)
                if m and len(m.group(1)) > 10:
                    entries.append(StrengthEntry(strength=m.group(1).strip()))
        return entries

    def _parse_failures_for_json(self, content: str) -> List['FailureEntry']:
        """Parse Section 5 failures into FailureEntry list."""
        entries = []
        if not content:
            return entries
        blocks = re.split(r'(?:^|\n)(?:FAILURE\s*TITLE|###?\s*FAILURE)\s*(?:\d+)?[:\s]*', content, flags=re.IGNORECASE)
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            title = ""
            timestamp = ""
            quote = ""
            what_happened = ""
            top_agent = ""
            revenue = ""
            lines = block.split('\n')
            for line in lines:
                ls = line.strip()
                if not ls:
                    continue
                upper = ls.upper()
                if upper.startswith('TIMESTAMP'):
                    timestamp = re.sub(r'^TIMESTAMP\s*[:\-]\s*', '', ls, flags=re.IGNORECASE).strip()
                elif upper.startswith('QUOTE'):
                    quote = re.sub(r'^QUOTE\s*[:\-]\s*', '', ls, flags=re.IGNORECASE).strip()
                elif upper.startswith('WHAT HAPPENED'):
                    what_happened = re.sub(r'^WHAT HAPPENED\s*[:\-]\s*', '', ls, flags=re.IGNORECASE).strip()
                elif upper.startswith('WHAT A TOP AGENT'):
                    top_agent = re.sub(r'^WHAT A TOP AGENT WOULD DO\s*[:\-]\s*', '', ls, flags=re.IGNORECASE).strip()
                elif upper.startswith('REVENUE IMPACT'):
                    revenue = re.sub(r'^REVENUE IMPACT\s*[:\-]\s*', '', ls, flags=re.IGNORECASE).strip()
                elif not title:
                    title = ls.strip('- *')
            if title:
                entries.append(FailureEntry(
                    failure_title=title, timestamp=timestamp, quote=quote,
                    what_happened=what_happened, what_top_agent_would_do=top_agent,
                    revenue_impact=revenue))
        if not entries:
            for line in content.split('\n'):
                line = line.strip()
                m = re.match(r'^[\-\*\d\.]+\s*(.+)', line)
                if m and len(m.group(1)) > 10:
                    entries.append(FailureEntry(failure_title=m.group(1).strip()))
        return entries

    def _parse_playbook_for_json(self, content: str) -> List['PlaybookEntry']:
        """Parse Section 6 tactical playbook into PlaybookEntry list."""
        entries = []
        if not content:
            return entries
        blocks = re.split(r'(?:^|\n)(?:SCENARIO|###?\s*SCENARIO|ENTRY)\s*(?:\d+)?[:\s]*', content, flags=re.IGNORECASE)
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            scenario = ""
            client_trigger = ""
            risk = ""
            assertive = ""
            consultative = ""
            why_works = ""
            lines = block.split('\n')
            for line in lines:
                ls = line.strip()
                if not ls:
                    continue
                upper = ls.upper()
                if upper.startswith('CLIENT TRIGGER'):
                    client_trigger = re.sub(r'^CLIENT TRIGGER\s*[:\-]\s*', '', ls, flags=re.IGNORECASE).strip()
                elif upper.startswith('RISK'):
                    risk = re.sub(r'^RISK\s*[:\-]\s*', '', ls, flags=re.IGNORECASE).strip()
                elif upper.startswith('ASSERTIVE'):
                    assertive = re.sub(r'^ASSERTIVE\s*(?:APPROACH)?\s*[:\-]\s*', '', ls, flags=re.IGNORECASE).strip()
                elif upper.startswith('CONSULTATIVE'):
                    consultative = re.sub(r'^CONSULTATIVE\s*(?:APPROACH)?\s*[:\-]\s*', '', ls, flags=re.IGNORECASE).strip()
                elif upper.startswith('WHY IT WORKS') or upper.startswith('WHY THIS WORKS'):
                    why_works = re.sub(r'^WHY (?:IT|THIS) WORKS\s*[:\-]\s*', '', ls, flags=re.IGNORECASE).strip()
                elif not scenario:
                    scenario = ls.strip('- *')
            if scenario:
                entries.append(PlaybookEntry(
                    scenario=scenario, client_trigger=client_trigger, risk=risk,
                    assertive_approach=assertive, consultative_approach=consultative,
                    why_it_works=why_works))
        if not entries:
            for line in content.split('\n'):
                line = line.strip()
                m = re.match(r'^[\-\*\d\.]+\s*(.+)', line)
                if m and len(m.group(1)) > 10:
                    entries.append(PlaybookEntry(scenario=m.group(1).strip()))
        return entries

    def _parse_7day_plan_for_json(self, content: str) -> List['CoachingDayPlan']:
        """Parse Section 7 seven-day coaching plan into CoachingDayPlan list."""
        entries = []
        if not content:
            return entries
        # Try "Day X: Focus -- Drill -- Metric" format
        day_blocks = re.split(r'(?:^|\n)(?:Day\s+(\d+))', content, flags=re.IGNORECASE)
        i = 1
        while i < len(day_blocks) - 1:
            day_num_str = day_blocks[i].strip()
            day_content = day_blocks[i + 1].strip() if i + 1 < len(day_blocks) else ""
            try:
                day_num = int(day_num_str)
            except ValueError:
                i += 2
                continue
            focus = ""
            drill = ""
            metric = ""
            # Parse pipe-delimited or line-based
            if '|' in day_content:
                parts = [p.strip() for p in day_content.split('|')]
                if len(parts) >= 1:
                    focus = parts[0].strip(':- ')
                if len(parts) >= 2:
                    drill = parts[1].strip()
                if len(parts) >= 3:
                    metric = parts[2].strip()
            elif '--' in day_content:
                parts = [p.strip() for p in day_content.split('--')]
                if len(parts) >= 1:
                    focus = parts[0].strip(':- ')
                if len(parts) >= 2:
                    drill = parts[1].strip()
                if len(parts) >= 3:
                    metric = parts[2].strip()
            else:
                lines = day_content.split('\n')
                for line in lines:
                    ls = line.strip()
                    if not ls:
                        continue
                    upper = ls.upper()
                    if upper.startswith('FOCUS'):
                        focus = re.sub(r'^FOCUS\s*[:\-]\s*', '', ls, flags=re.IGNORECASE).strip()
                    elif upper.startswith('DRILL'):
                        drill = re.sub(r'^DRILL\s*[:\-]\s*', '', ls, flags=re.IGNORECASE).strip()
                    elif upper.startswith('METRIC') or upper.startswith('SUCCESS'):
                        metric = re.sub(r'^(?:METRIC|SUCCESS\s*METRIC)\s*[:\-]\s*', '', ls, flags=re.IGNORECASE).strip()
                    elif not focus:
                        focus = ls.strip(':- ')
            if focus:
                entries.append(CoachingDayPlan(day=day_num, focus=focus, drill=drill, metric=metric))
            i += 2
        return entries

    # ════════════════════════════════════════════════════════
    # JSON 6: AGENT TIER CALCULATION
    # ════════════════════════════════════════════════════════

    def _generate_agent_tier_calculation_json(self, sections: Dict[str, str]):
        """Generate agent_tier_calculation.json — tier methodology extracted from
        Section 10 (Agent Tier) and Section 2 (Performance Scores)."""

        agent_name = getattr(self, '_agent_name', 'Agent')
        scores = getattr(self, '_cached_scores', {})

        # Calculate overall score and tier
        overall = round(sum(scores.values()) / len(scores), 2) if scores else 0.0

        # Determine tier
        if overall >= 9.0:
            tier = "Elite Performer"
        elif overall >= 7.0:
            tier = "Strong Performer"
        elif overall >= 5.0:
            tier = "Developing Agent"
        elif overall >= 3.0:
            tier = "Needs Improvement"
        else:
            tier = "Critical Development"

        # Parse tier justification from Section 10
        tier_content = self._find_section_content(sections, 10, ['Tier', 'Classification', 'Comparative'])
        tier_justification = ""
        promotion_criteria = []
        comparative_metrics = []

        if tier_content:
            lines = tier_content.split('\n')
            in_promotion = False
            for line in lines:
                ls = line.strip()
                if not ls:
                    continue
                upper = ls.upper()
                if 'PROMOTION' in upper or 'CRITERIA' in upper:
                    in_promotion = True
                    continue
                if in_promotion and ls.startswith('-'):
                    promotion_criteria.append(ls.strip('- '))
                elif '|' in ls:
                    parts = [p.strip() for p in ls.split('|')]
                    if len(parts) >= 2:
                        comparative_metrics.append({
                            "metric": parts[0],
                            "value": parts[1],
                            "trend": parts[2] if len(parts) > 2 else "",
                        })
                elif not tier_justification and len(ls) > 20 and not ls.startswith('#'):
                    tier_justification = ls

        # Weight distribution based on section categories
        weight_distribution = {
            "Communication & Rapport": 20.0,
            "Discovery & Questioning": 15.0,
            "Objection Handling": 15.0,
            "Value Proposition & Closing": 15.0,
            "Emotional Intelligence": 10.0,
            "Ethics & Compliance": 10.0,
            "Tone & Engagement": 7.5,
            "Negotiation & Persuasion": 7.5,
        }

        output = AgentTierCalculationOutput(
            generated_at=datetime.now().isoformat(),
            report_type="agent_tier_calculation",
            schema_version="2.0",
            agent_name=agent_name,
            metrics_used={k: float(v) for k, v in scores.items()},
            weight_distribution=weight_distribution,
            scoring_formula="weighted_average_of_all_llm_evaluated_dimensions",
            overall_score=overall,
            final_agent_tier=tier,
            tier_justification=tier_justification or f"Agent scored {overall}/10 overall, placing in the '{tier}' tier.",
            promotion_criteria=promotion_criteria if promotion_criteria else [
                "Maintain consistent scores above 7.0 across all dimensions",
                "Demonstrate improved objection handling in next 3 evaluations",
                "Reduce filler language and increase confident phrasing",
            ],
            comparative_metrics=comparative_metrics,
            metadata={
                "transcript_file": self.transcript_file,
                "total_dimensions_scored": len(scores),
                "highest_score": max(scores.values()) if scores else 0,
                "lowest_score": min(scores.values()) if scores else 0,
                "score_range": round(max(scores.values()) - min(scores.values()), 1) if scores else 0,
            },
        )

        filepath = os.path.join(self.summaries_folder, "agent_tier_calculation.json")
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(output.model_dump(), f, indent=2, ensure_ascii=False, default=str)
            logger.info(f"Saved agent tier calculation JSON: {filepath}")
        except Exception as e:
            logger.error(f"Failed to save agent_tier_calculation JSON: {e}")

    # ════════════════════════════════════════════════════════
    # LEGACY COMPAT METHODS (kept for compatibility)
    # ════════════════════════════════════════════════════════

    def create_professional_header_footer(self, canvas, doc):
        self._draw_content_header_footer(canvas, doc)

    def create_first_page_header_footer(self, canvas, doc):
        self._draw_cover_header_footer(canvas, doc)

    def create_unified_first_page_header_footer(self, canvas, doc):
        self._draw_cover_header_footer(canvas, doc)

    def create_unified_header_footer(self, canvas, doc):
        self._draw_content_header_footer(canvas, doc)


# ════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════

def main():
    print("=" * 60)
    print("  REAL ESTATE COACHING EVALUATION SYSTEM v5.0")
    print("  Enterprise Intelligence Engine — 4-PDF + 8-JSON Output")
    print("  PDF 1: Coaching Summary (text-heavy coaching report)")
    print("  PDF 2: Agent Profile (performance & compliance)")
    print("  PDF 3: Coaching Insights Visualizations")
    print("  PDF 4: Agent Performance Visualizations")
    print("  JSON 1: coaching_insights.json (coaching charts)")
    print("  JSON 2: agent_performance.json (performance charts)")
    print("  JSON 3: action_items.json (all action items)")
    print("  JSON 4: all_visualizations.json (master combined)")
    print("  JSON 5: coaching_recommendations.json (structured recs)")
    print("  JSON 6: agent_tier_calculation.json (tier methodology)")
    print("  JSON 7: agent_performance_aggregate.json (aggregated dashboard data)")
    print("  JSON 8: deal_outcome.json (deal status / client decision)")
    print("=" * 60)

    default_file = "conversation.json"
    transcript_file = input(f"  Enter JSON transcript path (Enter for '{default_file}'): ").strip().strip('"')
    if not transcript_file:
        transcript_file = default_file

    evaluator = RealEstateSalesMeetingSummarizer(transcript_file=transcript_file)
    logger.info(f"Generating brutally honest evaluation (4 PDFs + 8 JSONs)...")
    evaluator.generate_unified_coaching_evaluation()


if __name__ == "__main__":
    main()